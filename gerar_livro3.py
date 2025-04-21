from openai import OpenAI
from docx import Document
from dotenv import load_dotenv
import os
import tiktoken # Useful for splitting by tokens

# === CARREGA VARIÁVEIS DE AMBIENTE DO .env ===
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")

# === CONFIGURAÇÕES ===
INPUT_TXT = "rascunho.txt"
TEMPLATE_DOCX = "Estrutura.docx"
OUTPUT_DOCX = "Livro_Final_Formatado.docx"

# Define a maximum size for each chunk (in tokens) to send to the API
# You need to leave room for the prompt itself and the expected output.
# A common practice is to split into chunks significantly smaller than the total input context window.
# Let's say the input context is 128k. A chunk size of 8k to 16k tokens might be reasonable for the text content,
# leaving space for the prompt and response. Let's pick a value.
MAX_CHUNK_TOKENS = 10000 # Example chunk size in tokens

# === SETUP OPENAI ===
client = OpenAI(api_key=API_KEY)

# === PASSO 1 – Lê o texto bruto ===
try:
    with open(INPUT_TXT, "r", encoding="utf-8") as f:
        texto_original = f.read()
except FileNotFoundError:
    print(f"Erro: O arquivo de rascunho '{INPUT_TXT}' não foi encontrado.")
    exit()

# === PASSO 2 – Split the text into chunks ===
print("✂️ Fragmentando o texto...")
# This is a simplified splitting method. A more advanced method would
# try to split at natural breaks (sentences, paragraphs, chapters) and
# potentially include some overlap between chunks to maintain context.
# Using tiktoken is good for measuring token count precisely.
# tokenizer = tiktoken.encoding_for_model("gpt-4-turbo") # Or relevant model
# tokens = tokenizer.encode(texto_original)
# num_tokens = len(tokens)

# For simplicity here, let's just split by characters, but be aware this isn't token-precise
# A rough estimate: 1 token is about 4 characters for English text. Portuguese might vary.
# Let's estimate a character-based split based on token limit: 10000 tokens * ~3.5 chars/token = ~35000 characters
CHUNK_SIZE_CHARS = MAX_CHUNK_TOKENS * 3 # Adjust based on testing with your text/language

text_chunks = []
current_chunk = ""
for paragraph in texto_original.split("\n\n"): # Split by paragraphs as a basic unit
    # Estimate tokens if you don't use tiktoken directly
    # num_par_tokens = len(tokenizer.encode(paragraph))

    # A simple character-based split
    if len(current_chunk) + len(paragraph) + 2 < CHUNK_SIZE_CHARS: # +2 for potential \n\n
        current_chunk += paragraph + "\n\n"
    else:
        text_chunks.append(current_chunk.strip())
        current_chunk = paragraph + "\n\n"

# Add the last chunk
if current_chunk.strip():
    text_chunks.append(current_chunk.strip())

print(f"✅ Texto dividido em {len(text_chunks)} fragmentos.")


# === PASSO 3 – Carrega o template ou cria um novo documento ===
try:
    doc = Document(TEMPLATE_DOCX)
    # === PASSO 4 – Limpa corpo do documento (opcional) ===
    if doc._body:
        doc._body.clear_content()
except FileNotFoundError:
    print(f"Template '{TEMPLATE_DOCX}' não encontrado. Criando um novo documento.")
    doc = Document()


# === PASSO 5 – Process and insert each chunk ===
print("🔄 Processando fragmentos com ChatGPT e inserindo no documento...")

for i, chunk in enumerate(text_chunks):
    print(f"  Processando fragmento {i+1}/{len(text_chunks)}...")

    # === Build the prompt for this chunk ===
    # You might need to adjust the prompt slightly for chunks,
    # potentially mentioning it's part of a larger text.
    chunk_prompt = f"""
    Você é um editor literário. Corrija e formate o texto a seguir como parte de um livro.

    Regras:
    - Corrija erros gramaticais, ortográficos e de concordância
    - **NÃO** adicione títulos de capítulo se eles não estiverem explicitamente no texto do fragmento.
    - Use parágrafos justificados com espaçamento apropriado.
    - Se houver um marcador de quebra de página (===QUEBRA_DE_PAGINA===) dentro deste fragmento, mantenha-o.
    - Mantenha estilo literário fluido e bem escrito.
    - **NÃO** adicione texto introdutório ou conclusivo que não faça parte do rascunho.
    - O resultado deve ser APENAS o texto formatado.

    Texto do fragmento:
    \"\"\"{chunk}\"\"\"
    """

    # === Call the API for this chunk ===
    try:
        response = client.chat.completions.create(
            model="gpt-4-turbo", # or gpt-4o
            messages=[{"role": "user", "content": chunk_prompt}],
            temperature=0.7,
            max_tokens=MAX_CHUNK_TOKENS + 1000 # Allow slightly more tokens for output than input chunk size if formatting adds content
        )
        formatted_chunk_text = response.choices[0].message.content
        # print(f"    ✅ Fragmento {i+1} formatado.") # Uncomment for verbose output

    except Exception as e:
        print(f"    ❌ Erro ao processar fragmento {i+1}: {e}")
        # Decide how to handle errors: skip, retry, save raw chunk?
        formatted_chunk_text = chunk # Fallback: use the unformatted chunk if API fails

    # === Insert formatted chunk text into the document ===
    # Split the formatted text by the page break marker if it's present in the response
    parts = formatted_chunk_text.split("===QUEBRA_DE_PAGINA===")

    for j, part in enumerate(parts):
        part = part.strip()
        if not part:
            continue # Skip empty parts

        # If it's not the very first part of the very first chunk, add a page break before a new section
        # This logic needs refinement based on how you want chapters/sections to break
        # A common pattern is to add a page break *before* the content that follows a marker
        # or before the start of a new chunk *if* that chunk represents a new chapter/section.
        if i > 0 or j > 0: # If it's not the first part of the first chunk
             # You'll need smarter logic here. If the API included the marker,
             # splitting by it and adding a break *after* the content before it makes sense.
             # If splitting chunks manually, you might add a break *before* a chunk if it starts a new chapter.

             # Simplified approach: Add a page break *before* the content of a new part *if* it follows a marker
             # or before the content of any chunk after the first one.
             if j > 0: # Add break if this part followed a marker within the chunk
                  doc.add_page_break()
             elif i > 0 and j == 0: # Add break if this is the first part of a new chunk (and not the first chunk overall)
                  # This assumes each new chunk *might* be a new section/chapter.
                  # You might need a different strategy if chapters can span chunks.
                   doc.add_page_break() # Potential page break before a new chunk

        # Add paragraphs from the part
        paragraphs_in_part = part.split("\n\n")
        for par in paragraphs_in_part:
             if par.strip():
                 p = doc.add_paragraph(par.strip())
                 # Add logic here to identify and center titles within the chunk's content
                 if par.strip().startswith("Capítulo ") or par.strip().startswith("CHAPTER "):
                     p.alignment = 1 # 1 for center alignment

# === PASSO 6 – Salva o novo documento ===
doc.save(OUTPUT_DOCX)
print(f"📘 Livro gerado com sucesso: {OUTPUT_DOCX}")
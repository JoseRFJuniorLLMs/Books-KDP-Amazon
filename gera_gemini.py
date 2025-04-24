# --- Using Google's Gemini API (gemini-1.5-pro) ---

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.shared import RGBColor

from dotenv import load_dotenv
import os
import re
import logging
from tqdm import tqdm
import time
import shutil
import traceback # Para log de erros detalhado
import glob # Para encontrar arquivos .txt
# import fcntl # Potentially for file locking if needed in concurrent environments, but likely overkill here

# === SETUP LOGGING ===
log_dir = "logs"
if not os.path.exists(log_dir): os.makedirs(log_dir)
log_filepath = os.path.join(log_dir, "book_processor_multi_author_mem.log")
# ADDED: Central log for processed files
PROCESSED_LOG_FILE = os.path.join(log_dir, "processed_books.log")

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(module)s:%(lineno)d - %(funcName)s - %(message)s',
    handlers=[ logging.FileHandler(log_filepath, encoding='utf-8'), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# === CARREGA VARIÁVEIS DE AMBIENTE ===
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# === CONFIGURAÇÕES ===

# -- Diretórios Base --
BASE_INPUT_TXT_DIR = "txt"
BASE_OUTPUT_DOCX_DIR = "docx"
BASE_OUTPUT_TXT_DIR = "txt"

# -- Nomes de Arquivos Base --
TEMPLATE_DOCX = "Estrutura.docx" # Template OBRIGATÓRIO

# -- Nomes Base dos Arquivos de Saída --
FINAL_DOCX_BASENAME = "Livro_Final_Formatado_Sem_Notas.docx"
# MODIFIED: Updated name as requested
FINAL_NUMBERED_TXT_BASENAME = "Livro_Final_Com_Notas_Numeros.txt"
NOTES_TXT_FILE_BASENAME = "notas_rodape.txt"
# REMOVED: PROCESSED_MARKER_FILENAME = ".processed_ok"

# -- Configurações da API e Processamento --
MODEL_NAME = "gemini-1.5-pro"
MAX_CHUNK_TOKENS = 1500
MAX_OUTPUT_TOKENS = 8192
TEMPERATURE = 0.5

# -- Estilos e Padrões --
NORMAL_STYLE_NAME = "Normal"
CHAPTER_PATTERNS = [
    r'^\s*Capítulo \w+', r'^\s*CAPÍTULO \w+', r'^\s*Capítulo \d+',
    r'^\s*CHAPTER \w+', r'^\s*Chapter \d+', r'^\s*LIVRO \w+', r'^\s*PARTE \w+',
]
PAGE_BREAK_MARKER = "===QUEBRA_DE_PAGINA==="
AI_FAILURE_MARKER = "*** FALHA NA IA - TEXTO ORIGINAL ABAIXO ***"
FORMATTING_ERROR_MARKER = "*** ERRO DE FORMATAÇÃO - TEXTO ORIGINAL ABAIXO ***"

# --- Validação API Key ---
if not GOOGLE_API_KEY:
    logger.error("GOOGLE_API_KEY não encontrada nas variáveis de ambiente.")
    exit(1)

# --- Setup Gemini Client ---
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    safety_settings_lenient = {
        'HATE': 'BLOCK_NONE', 'HARASSMENT': 'BLOCK_NONE',
        'SEXUAL' : 'BLOCK_NONE', 'DANGEROUS' : 'BLOCK_NONE'
    }
    generation_config = genai.GenerationConfig(
                    temperature=TEMPERATURE,
                    max_output_tokens=MAX_OUTPUT_TOKENS
                )
    gemini_model = genai.GenerativeModel(
        MODEL_NAME,
        safety_settings=safety_settings_lenient,
        generation_config=generation_config
    )
    logger.info(f"Modelo Gemini '{MODEL_NAME}' inicializado.")
except Exception as e:
    logger.error(f"Falha ao inicializar modelo Gemini ({MODEL_NAME}): {e}")
    logger.error(traceback.format_exc())
    exit(1)

# --- Funções Auxiliares (Chunking, API Call - inalteradas) ---
def count_tokens_approx(text):
    # ... (código inalterado)
    if not text: return 0
    return len(text) // 3

def create_chunks(text, max_tokens, author_name="N/A", book_name="N/A"):
    """Divide o texto em chunks, subdividindo parágrafos grandes."""
    log_prefix = f"[{author_name}/{book_name}]" # Add book name to log
    logger.info(f"{log_prefix} Iniciando criação de chunks. Máx tokens (aprox): {max_tokens}")
    chunks = []
    current_chunk = ""
    current_chunk_tokens = 0
    paragraphs = text.split("\n\n")
    paragraphs_stripped = [p.strip() for p in paragraphs if p.strip()]

    logger.info(f"{log_prefix} Texto dividido inicialmente em {len(paragraphs_stripped)} blocos não vazios ('\\n\\n').")

    for i, paragraph_text in enumerate(paragraphs):
        if not paragraph_text.strip():
            if chunks and chunks[-1].strip():
                 if not chunks[-1].endswith("\n\n"):
                      chunks[-1] += "\n\n"
            continue
        paragraph_tokens = count_tokens_approx(paragraph_text)
        tokens_with_separator = paragraph_tokens + (count_tokens_approx("\n\n") if current_chunk else 0)

        if current_chunk and (current_chunk_tokens + tokens_with_separator > max_tokens):
            chunks.append(current_chunk)
            logger.debug(f"{log_prefix} Chunk {len(chunks)} salvo (limite atingido). Tokens: {current_chunk_tokens}.")
            current_chunk = paragraph_text
            current_chunk_tokens = paragraph_tokens
        elif current_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if current_chunk else ""
            current_chunk += separator + paragraph_text
            current_chunk_tokens = count_tokens_approx(current_chunk)

        if not chunks or (chunks and chunks[-1] != current_chunk) :
            if paragraph_tokens > max_tokens:
                logger.warning(f"{log_prefix} Parágrafo {i+1} ({paragraph_tokens} tk) excede limite {max_tokens}. Iniciando SUBDIVISÃO.")
                prefix_to_save = ""
                if current_chunk != paragraph_text and current_chunk.strip():
                    if current_chunk.endswith("\n\n" + paragraph_text):
                        prefix_to_save = current_chunk[:-len("\n\n" + paragraph_text)]
                    elif current_chunk == paragraph_text :
                         pass
                    else:
                         logger.warning(f"{log_prefix} Lógica de prefixo de subdivisão inesperada.")
                         prefix_to_save = current_chunk
                    if prefix_to_save.strip():
                         chunks.append(prefix_to_save)
                         logger.debug(f"{log_prefix} Chunk {len(chunks)} salvo (prefixo antes da subdivisão). Tokens: {count_tokens_approx(prefix_to_save)}.")

                sub_chunks_added_count = 0
                sentences = re.split(r'(?<=[.!?])\s+', paragraph_text)
                if len(sentences) <= 1 :
                    sentences = paragraph_text.split('\n')
                current_sub_chunk = ""
                current_sub_chunk_tokens = 0
                for sentence_num, sentence in enumerate(sentences):
                    sentence_clean = sentence.strip()
                    if not sentence_clean: continue
                    sentence_tokens = count_tokens_approx(sentence)
                    tokens_with_sub_separator = sentence_tokens + (count_tokens_approx("\n") if current_sub_chunk else 0)
                    if current_sub_chunk and (current_sub_chunk_tokens + tokens_with_sub_separator > max_tokens):
                        chunks.append(current_sub_chunk)
                        sub_chunks_added_count += 1
                        logger.debug(f"{log_prefix} Sub-chunk {len(chunks)} salvo (Parág. {i+1}). Tokens: {current_sub_chunk_tokens}.")
                        current_sub_chunk = sentence
                        current_sub_chunk_tokens = sentence_tokens
                    elif sentence_tokens > max_tokens:
                        if current_sub_chunk:
                            chunks.append(current_sub_chunk)
                            sub_chunks_added_count += 1
                            logger.debug(f"{log_prefix} Sub-chunk {len(chunks)} salvo (antes sentença longa, Parág. {i+1}). Tokens: {current_sub_chunk_tokens}.")
                        chunks.append(sentence)
                        sub_chunks_added_count += 1
                        logger.warning(f"{log_prefix}  -> Sentença/Linha {sentence_num+1} ({sentence_tokens} tk) excede limite. Adicionada como sub-chunk individual (PODE FALHAR NA API).")
                        current_sub_chunk = ""
                        current_sub_chunk_tokens = 0
                    else:
                        sub_separator = "\n" if current_sub_chunk else ""
                        current_sub_chunk += sub_separator + sentence
                        current_sub_chunk_tokens = count_tokens_approx(current_sub_chunk)
                if current_sub_chunk:
                    chunks.append(current_sub_chunk)
                    sub_chunks_added_count += 1
                    logger.debug(f"{log_prefix} Último sub-chunk {len(chunks)} salvo (Parág. {i+1}). Tokens: {current_sub_chunk_tokens}.")
                if sub_chunks_added_count == 0:
                     logger.warning(f"{log_prefix} Parágrafo {i+1} excedeu limite, mas não foi subdividido. Adicionando original como chunk (PODE FALHAR NA API).")
                     chunks.append(paragraph_text)
                current_chunk = ""
                current_chunk_tokens = 0
            elif not current_chunk:
                 current_chunk = paragraph_text
                 current_chunk_tokens = paragraph_tokens
    if current_chunk:
        chunks.append(current_chunk)
        logger.debug(f"{log_prefix} Chunk final {len(chunks)} salvo. Tokens: {current_chunk_tokens}.")

    merged_chunks = []
    temp_chunk = ""
    temp_chunk_tokens = 0
    for i, chunk in enumerate(chunks):
        chunk_tokens = count_tokens_approx(chunk)
        tokens_with_separator = chunk_tokens + (count_tokens_approx("\n\n") if temp_chunk else 0)
        if temp_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if temp_chunk else ""
            temp_chunk += separator + chunk
            temp_chunk_tokens = count_tokens_approx(temp_chunk)
        else:
            merged_chunks.append(temp_chunk)
            logger.debug(f"{log_prefix} Merged chunk {len(merged_chunks)} salvo. Tokens: {temp_chunk_tokens}.")
            temp_chunk = chunk
            temp_chunk_tokens = chunk_tokens
    if temp_chunk:
        merged_chunks.append(temp_chunk)
        logger.debug(f"{log_prefix} Merged chunk final {len(merged_chunks)} salvo. Tokens: {temp_chunk_tokens}.")

    logger.info(f"{log_prefix} ✅ Chunking concluído. {len(merged_chunks)} chunks finais (após merge).")
    return merged_chunks

def _call_gemini_api(model, prompt_text, chunk_for_log, author_name="N/A", book_name="N/A"):
    """Função interna para chamar a API Gemini com retries."""
    # ... (código inalterado, talvez adicionar book_name ao log_prefix se desejado) ...
    log_prefix = f"[{author_name}/{book_name}]"
    max_retries = 5
    base_wait_time = 5
    log_chunk_preview = chunk_for_log[:150].replace('\n', '\\n') + ('...' if len(chunk_for_log) > 150 else '')

    for attempt in range(max_retries):
        logger.info(f"{log_prefix} Chamando API (Tentativa {attempt + 1}/{max_retries}). Chunk (início): '{log_chunk_preview}'")
        try:
            response = model.generate_content(prompt_text)
            finish_reason = "UNKNOWN"; safety_ratings = "UNKNOWN"; block_reason = "N/A"; result_text = None
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                 if hasattr(response.prompt_feedback, 'block_reason') and response.prompt_feedback.block_reason:
                      block_reason = response.prompt_feedback.block_reason.name
                      logger.error(f"{log_prefix} API BLOQUEOU O PROMPT (Tentativa {attempt + 1}). Razão: {block_reason}. Chunk: '{log_chunk_preview}'")
                      return None
            if not response.candidates:
                 logger.error(f"{log_prefix} API retornou SEM CANDIDATOS (Tentativa {attempt + 1}). Resposta: {response}. Chunk: '{log_chunk_preview}'")
            else:
                 try:
                    candidate = response.candidates[0]
                    finish_reason = candidate.finish_reason.name if hasattr(candidate, 'finish_reason') and candidate.finish_reason else "FINISH_REASON_UNKNOWN"
                    safety_ratings = [(r.category.name, r.probability.name) for r in candidate.safety_ratings] if candidate.safety_ratings else "N/A"
                    logger.debug(f"{log_prefix} API Call OK (Tentativa {attempt + 1}). Finish: {finish_reason}. Safety: {safety_ratings}")
                    if finish_reason == "STOP": pass
                    elif finish_reason == "MAX_TOKENS": logger.warning(f"{log_prefix} API TRUNCOU resposta devido a MAX_OUTPUT_TOKENS.")
                    elif finish_reason == "SAFETY": logger.warning(f"{log_prefix} API interrompeu resposta devido a SAFETY. Ratings: {safety_ratings}.")
                    elif finish_reason == "RECITATION": logger.warning(f"{log_prefix} API interrompeu resposta devido a RECITATION.")
                    elif finish_reason == "OTHER": logger.warning(f"{log_prefix} API interrompeu resposta por OUTRA RAZÃO.")
                    else: logger.warning(f"{log_prefix} API retornou com finish_reason inesperado: {finish_reason}.")

                    result_text = ""
                    if hasattr(candidate, 'content') and hasattr(candidate.content, 'parts'):
                        text_parts = [part.text for part in candidate.content.parts if hasattr(part, 'text')]
                        if text_parts:
                            result_text = "".join(text_parts).strip()
                        else:
                             logger.warning(f"{log_prefix} Resposta API tem 'parts' mas não foi possível extrair texto (Tentativa {attempt+1}). Parts: {candidate.content.parts}")
                    elif hasattr(response, 'text') and response.text:
                        result_text = response.text.strip()
                        logger.debug(f"{log_prefix} Texto extraído via response.text (fallback).")

                    if result_text:
                        logger.debug(f"{log_prefix} Texto API recebido (100 chars): '{result_text[:100].replace('\n', '\\n')}...'")
                        if len(result_text) < len(chunk_for_log) * 0.1 and len(chunk_for_log) > 100:
                            logger.warning(f"{log_prefix} Resposta da API parece muito curta. Input len: {len(chunk_for_log)}, Output len: {len(result_text)}.")
                        return result_text
                    else:
                         logger.warning(f"{log_prefix} Resposta API não continha texto utilizável (Tentativa {attempt+1}), mesmo com candidato. Finish Reason: {finish_reason}.")
                 except Exception as e_details:
                    logger.error(f"{log_prefix} Erro ao extrair detalhes/texto da resposta API (Tentativa {attempt+1}): {e_details} - Resposta Crua: {response}")
                    logger.error(traceback.format_exc())
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                logger.info(f"{log_prefix} Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                 logger.error(f"{log_prefix} Falha final na API após {max_retries} tentativas para o chunk: '{log_chunk_preview}'")
                 return None
        except Exception as e:
            logger.warning(f"{log_prefix} Erro durante a chamada da API ({model.model_name}) (Tentativa {attempt + 1}/{max_retries}): {e}")
            logger.error(traceback.format_exc())
            if "RESOURCE_EXHAUSTED" in str(e) or "429" in str(e):
                 logger.warning(f"{log_prefix} Erro de cota (RESOURCE_EXHAUSTED / 429). Aumentando espera.")
                 base_wait_time = 15
            elif "Internal error encountered." in str(e) or "500" in str(e):
                 logger.warning(f"{log_prefix} Erro interno do servidor (500). Tentando novamente.")
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                logger.info(f"{log_prefix} Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                logger.error(f"{log_prefix} Falha final na API após {max_retries} tentativas (erro na chamada) para o chunk: '{log_chunk_preview}'")
                return None
    logger.error(f"{log_prefix} Loop de tentativas da API concluído sem sucesso explícito para o chunk: '{log_chunk_preview}'")
    return None

def format_with_ai_correction_only(model, chunk, author_name, book_name, is_first_chunk=False):
    """Chama a API Gemini focando APENAS na correção de OCR/gramática."""
    # ... (prompt inalterado, mas passa author_name para ele) ...
    log_prefix = f"[{author_name}/{book_name}]"
    # ... (prompt generation logic is the same, uses author_name) ...
    context_start = "Você está formatando o início de um livro." if is_first_chunk else "Você está continuando a formatação de um texto de livro existente."
    ocr_errors_examples = "..." # Conteúdo inalterado
    chunk_prompt = f"""
{context_start} Você é um editor literário proficiente em português do Brasil. Sua tarefa é CORRIGIR e FORMATAR o fragmento de texto a seguir, que pertence a um livro do autor {author_name}.
... (restante do prompt inalterado) ...
"""
    logger.debug(f"{log_prefix} Enviando chunk para CORREÇÃO (API: {model.model_name}). Tam Aprox: {count_tokens_approx(chunk)} tk")
    return _call_gemini_api(model, chunk_prompt, chunk, author_name, book_name)


def format_with_ai_footnote_only(model, chunk, author_name, book_name):
    """Chama a API Gemini focando APENAS na identificação de notas."""
     # ... (prompt inalterado, mas passa author_name para ele) ...
    log_prefix = f"[{author_name}/{book_name}]"
    # ... (prompt generation logic is the same, uses author_name) ...
    chunk_prompt = f"""
Você é um assistente de edição trabalhando no texto do autor {author_name}. Sua tarefa é analisar o fragmento de texto A SEGUIR, que JÁ FOI CORRIGIDO no passo anterior, e APENAS inserir marcadores para potenciais notas de rodapé onde estritamente necessário.
... (restante do prompt inalterado) ...
"""
    logger.debug(f"{log_prefix} Enviando chunk para IDENTIFICAÇÃO DE NOTAS (API: {model.model_name}). Tam Aprox: {count_tokens_approx(chunk)} tk")
    return _call_gemini_api(model, chunk_prompt, chunk, author_name, book_name)


# --- FUNÇÕES DE PROCESSAMENTO DOS PASSOS ---

def apply_formatting_pass1(doc, formatted_chunk_text, normal_style_name, chapter_patterns, corrected_text_list, author_name, book_name):
    """
    Aplica formatação ao DOCX (Passo 1 - sem notas) e coleta texto para retornar.
    MODIFIED: Includes book_name in logs.
    """
    log_prefix = f"[{author_name}/{book_name}]"
    if not formatted_chunk_text or not formatted_chunk_text.strip():
        logger.warning(f"{log_prefix} Chunk formatado vazio ou apenas espaços recebido (Passo 1). Pulando.")
        return

    # Coleta texto (inalterado)
    plain_text_for_list = formatted_chunk_text.replace(PAGE_BREAK_MARKER, "\n\n").strip()
    if plain_text_for_list:
        corrected_text_list.append(plain_text_for_list)
    else:
        if formatted_chunk_text.strip() == PAGE_BREAK_MARKER:
             logger.debug(f"{log_prefix} Chunk continha apenas marcador de página, não adicionado à lista de texto.")
        else:
             logger.warning(f"{log_prefix} Texto formatado resultou em vazio. Original: '{formatted_chunk_text[:50]}...'")

    # Aplica formatação DOCX
    normal_style = None
    try:
        if normal_style_name in doc.styles:
             normal_style = doc.styles[normal_style_name]
             # logger.debug(f"{log_prefix} Estilo '{normal_style_name}' encontrado no template.") # Too verbose
        else:
             # Warning is logged once in run_correction_pass if style is missing
             pass
    except Exception as e_style:
         logger.error(f"{log_prefix} Erro ao acessar estilo '{normal_style_name}': {e_style}.")

    chapter_regex = re.compile('|'.join(chapter_patterns), re.IGNORECASE)
    parts = formatted_chunk_text.split(PAGE_BREAK_MARKER)
    content_present_before = any(p.text.strip() for p in doc.paragraphs)

    for part_index, part in enumerate(parts):
        part_clean = part.strip()

        # Add page break logic (inalterado)
        if part_index > 0:
             last_para_is_page_break = False
             if doc.paragraphs:
                 last_p = doc.paragraphs[-1]
                 if not last_p.text.strip() and any(run.text and '\f' in run.text for run in last_p.runs):
                     last_para_is_page_break = True
             if not last_para_is_page_break:
                 logger.debug(f"{log_prefix} Adicionando quebra de página ao DOCX (antes da parte {part_index + 1}).")
                 doc.add_page_break()
             else:
                 logger.debug(f"{log_prefix} Quebra de página omitida (duplicada).")
        # ... (rest of page break logic) ...

        if not part_clean:
            if part_index > 0 : content_present_before = True
            continue

        paragraphs_in_part = part_clean.split("\n\n")
        for paragraph_text in paragraphs_in_part:
            paragraph_text_clean = paragraph_text.strip()
            if not paragraph_text_clean:
                if doc.paragraphs and doc.paragraphs[-1].text.strip():
                     p = doc.add_paragraph()
                     if normal_style: p.style = normal_style
                continue

            is_ai_failure_marker = paragraph_text_clean.startswith(AI_FAILURE_MARKER)
            is_formatting_error_marker = paragraph_text_clean.startswith(FORMATTING_ERROR_MARKER)
            is_chapter = not is_ai_failure_marker and not is_formatting_error_marker and chapter_regex.match(paragraph_text_clean) is not None

            p = doc.add_paragraph()
            run = p.add_run(paragraph_text_clean)
            content_present_before = True

            # Aplica estilos/formatação
            try:
                if is_chapter:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.name = 'French Script MT'
                    run.font.size = Pt(48)
                    run.bold = False
                    logger.debug(f"{log_prefix} Aplicada formatação específica de capítulo (French Script MT 48pt).")
                elif is_ai_failure_marker or is_formatting_error_marker:
                    if normal_style: p.style = normal_style
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    logger.debug(f"{log_prefix} Aplicada formatação de marcador de erro.")
                else: # Parágrafo normal
                    if normal_style:
                        p.style = normal_style
                    else:
                        pass # Rely on document default if Normal style missing
            except Exception as e_apply_style:
                 logger.error(f"{log_prefix} Erro ao aplicar estilo/formatação: {e_apply_style}. Texto: '{paragraph_text_clean[:50]}...'")
                 pass


def run_correction_pass(model, input_txt_path, template_docx_path, output_docx_path, author_name, book_name):
    """
    Executa o Passo 1: Corrige texto e gera DOCX base.
    MODIFIED: Includes book_name in logs and function calls.
              Checks for Normal style presence once.
    Retorna: (bool, str | None) -> (success_status, corrected_text_content)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 1: Correção ---")
    logger.info(f"{log_prefix} Lendo texto original de: {input_txt_path}")
    try:
        with open(input_txt_path, "r", encoding="utf-8") as f: texto_original = f.read()
        logger.info(f"{log_prefix} Entrada '{os.path.basename(input_txt_path)}' carregada ({len(texto_original)} chars).")
    except FileNotFoundError:
        logger.error(f"{log_prefix} Fatal: Arquivo de entrada '{input_txt_path}' não encontrado."); return (False, None)
    except Exception as e:
        logger.error(f"{log_prefix} Fatal ao ler entrada '{input_txt_path}': {e}")
        logger.error(traceback.format_exc()); return (False, None)

    output_docx_dir = os.path.dirname(output_docx_path)
    if not os.path.exists(output_docx_dir):
        logger.info(f"{log_prefix} Criando diretório de saída DOCX: {output_docx_dir}")
        os.makedirs(output_docx_dir, exist_ok=True)

    logger.info(f"{log_prefix} Dividindo texto original em chunks...")
    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS, author_name, book_name) # Pass book_name
    if not text_chunks:
        logger.error(f"{log_prefix} Nenhum chunk gerado. Abortando Passo 1."); return (False, None)
    logger.info(f"{log_prefix} Texto dividido em {len(text_chunks)} chunks.")

    doc = None
    normal_style_exists = False # Flag to check style once
    logger.info(f"{log_prefix} Preparando documento DOCX para: {output_docx_path}")
    try:
        # Backup logic (inalterado)
        if os.path.exists(output_docx_path):
             backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
             backup_docx_path = os.path.join(output_docx_dir, f"backup_{os.path.splitext(os.path.basename(output_docx_path))[0]}_{backup_timestamp}.docx")
             try:
                 shutil.copy2(output_docx_path, backup_docx_path)
                 logger.info(f"{log_prefix} Backup do DOCX anterior criado: {backup_docx_path}")
             except Exception as e_bkp:
                 logger.warning(f"{log_prefix} Falha ao criar backup de '{output_docx_path}': {e_bkp}")

        # Load MANDATORY template
        if os.path.exists(template_docx_path):
            try:
                doc = Document(template_docx_path)
                logger.info(f"{log_prefix} Template '{template_docx_path}' carregado. Limpando conteúdo...")
                # Clear content logic (inalterado)
                for para in doc.paragraphs: p_element = para._element; p_element.getparent().remove(p_element)
                for table in doc.tables: tbl_element = table._element; tbl_element.getparent().remove(tbl_element)
                logger.info(f"{log_prefix} Conteúdo principal do template limpo.")

                # Check for essential 'Normal' style ONCE
                if NORMAL_STYLE_NAME in doc.styles:
                    normal_style_exists = True
                    logger.info(f"{log_prefix} Estilo '{NORMAL_STYLE_NAME}' encontrado no template.")
                else:
                    logger.warning(f"{log_prefix} Estilo '{NORMAL_STYLE_NAME}' NÃO encontrado no template '{TEMPLATE_DOCX}'. A formatação de parágrafos normais dependerá do padrão do documento.")

            except Exception as e_load_template:
                 logger.error(f"{log_prefix} FALHA ao carregar/limpar template OBRIGATÓRIO '{template_docx_path}': {e_load_template}. Abortando.")
                 return (False, None)
        else:
            logger.error(f"{log_prefix} Template OBRIGATÓRIO '{template_docx_path}' não encontrado. Abortando.")
            return (False, None)

        # REMOVED: Default style creation block

    except Exception as e_doc:
        logger.error(f"{log_prefix} Erro crítico ao preparar DOCX: {e_doc}")
        logger.error(traceback.format_exc()); return (False, None)

    # Process chunks
    logger.info(f"{log_prefix} Iniciando chamadas à API para CORREÇÃO de {len(text_chunks)} chunks...")
    corrected_text_list_pass1 = []
    processed_chunks_count = 0
    failed_chunks_count = 0
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P1: Corrigindo", unit="chunk")

    for i, chunk in progress_bar:
        progress_bar.set_description(f"{log_prefix} P1: Corrigindo Chunk {i+1}/{len(text_chunks)}")
        # Pass book_name to API call helper
        corrected_chunk = format_with_ai_correction_only(model, chunk, author_name, book_name, is_first_chunk=(i == 0))

        if corrected_chunk and corrected_chunk.strip():
            try:
                # Pass book_name to formatting helper
                apply_formatting_pass1(doc, corrected_chunk, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
                processed_chunks_count += 1
            except Exception as format_err:
                logger.error(f"{log_prefix} Erro na apply_formatting_pass1 (Chunk {i+1}): {format_err}.")
                logger.error(traceback.format_exc())
                failed_chunks_count += 1
                try:
                    fallback_text = f"{FORMATTING_ERROR_MARKER}\n\n{chunk}"
                    # Pass book_name to formatting helper
                    apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
                except Exception as fallback_format_err:
                    logger.critical(f"{log_prefix} Falha CRÍTICA ao aplicar fallback de erro FORMATAÇÃO (Chunk {i+1}): {fallback_format_err}.")
        else:
            logger.warning(f"{log_prefix} Chunk {i+1} falhou na CORREÇÃO (API). Usando fallback com marcador.")
            failed_chunks_count += 1
            try:
                fallback_text = f"{AI_FAILURE_MARKER}\n\n{chunk}"
                # Pass book_name to formatting helper
                apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
            except Exception as fallback_format_err:
                 logger.critical(f"{log_prefix} Falha CRÍTICA ao aplicar fallback de falha API (Chunk {i+1}): {fallback_format_err}.")

        # Save progress periodically (inalterado)
        if (processed_chunks_count + failed_chunks_count) > 0 and \
           (((processed_chunks_count + failed_chunks_count) % 10 == 0) or ((i + 1) == len(text_chunks))):
            temp_save_path = f"{output_docx_path}.temp_save"
            try:
                # logger.debug(f"{log_prefix} Salvando progresso DOCX (chunk {i+1})...") # Verbose
                doc.save(temp_save_path)
                shutil.move(temp_save_path, output_docx_path)
                logger.info(f"{log_prefix} Progresso DOCX (Passo 1) salvo ({processed_chunks_count + failed_chunks_count} chunks processados).")
            except Exception as e_save:
                 logger.error(f"{log_prefix} Erro ao salvar progresso DOCX (Chunk {i+1}) para '{output_docx_path}': {e_save}")

    # Save final DOCX (inalterado)
    try:
        logger.info(f"{log_prefix} Salvando DOCX final (Passo 1) em: {output_docx_path}")
        doc.save(output_docx_path)
    except Exception as e_final_save:
        logger.error(f"{log_prefix} Erro no salvamento final do DOCX (Passo 1): {e_final_save}")
        logger.error(traceback.format_exc())

    # Join corrected text (inalterado)
    full_corrected_text = "\n\n".join(corrected_text_list_pass1)
    logger.info(f"{log_prefix} Acumulado texto corrigido ({len(full_corrected_text)} chars).")

    logger.info(f"{log_prefix} --- Passo 1 concluído. Chunks OK: {processed_chunks_count}, Falhas/Fallback: {failed_chunks_count} ---")
    return (True, full_corrected_text)


def run_footnote_id_pass(model, corrected_text_content, author_name, book_name):
    """
    Executa o Passo 2: Identifica notas no texto já corrigido.
    MODIFIED: Includes book_name in logs and function calls.
    Retorna: (bool, str | None) -> (success_status, marked_text_content)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 2: Identificação de Notas ---")
    if not corrected_text_content:
        logger.error(f"{log_prefix} Texto corrigido de entrada vazio. Abortando Passo 2.")
        return (False, None)
    logger.info(f"{log_prefix} Recebido texto corrigido para Passo 2 ({len(corrected_text_content)} chars).")

    logger.info(f"{log_prefix} Dividindo texto corrigido em chunks para notas...")
    text_chunks = create_chunks(corrected_text_content, MAX_CHUNK_TOKENS, author_name, book_name) # Pass book_name
    if not text_chunks:
        logger.error(f"{log_prefix} Nenhum chunk gerado. Abortando Passo 2."); return (False, None)
    logger.info(f"{log_prefix} Texto corrigido dividido em {len(text_chunks)} chunks.")

    logger.info(f"{log_prefix} Iniciando chamadas à API para IDENTIFICAÇÃO DE NOTAS em {len(text_chunks)} chunks...")
    marked_text_list_pass2 = []
    processed_chunks_count = 0
    failed_chunks_count = 0
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P2: Notas", unit="chunk")

    for i, chunk in progress_bar:
        progress_bar.set_description(f"{log_prefix} P2: Notas Chunk {i+1}/{len(text_chunks)}")
        # Pass book_name to API call helper
        marked_chunk = format_with_ai_footnote_only(model, chunk, author_name, book_name)

        if marked_chunk:
            marked_text_list_pass2.append(marked_chunk)
            processed_chunks_count += 1
            if "[NOTA_" in marked_chunk: pass # logger.debug(f"{log_prefix} Chunk {i+1}: marcadores de nota encontrados.")
            else: pass # logger.debug(f"{log_prefix} Chunk {i+1}: NENHUM marcador de nota adicionado.")
        else:
            logger.warning(f"{log_prefix} Chunk {i+1} falhou na IDENTIFICAÇÃO DE NOTAS (API). Usando texto original do chunk.")
            marked_text_list_pass2.append(chunk)
            failed_chunks_count += 1

    full_marked_text = "\n\n".join(marked_text_list_pass2)
    logger.info(f"{log_prefix} Acumulado texto com marcadores ({len(full_marked_text)} chars).")

    logger.info(f"{log_prefix} --- Passo 2 concluído. Chunks OK: {processed_chunks_count}, Falhas/Fallback: {failed_chunks_count} ---")
    return (True, full_marked_text)


def run_final_txt_generation(marked_text_content, output_notes_path, output_numbered_txt_path, author_name, book_name):
    """
    Executa o Passo 3: Processa marcadores [NOTA_...] para gerar TXT final numerado [N] e arquivo de notas.
    MODIFIED: Includes book_name in logs.
    Retorna: bool -> success_status
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 3: Geração Final TXT (Notas e Numerado) ---")
    if not marked_text_content:
        logger.error(f"{log_prefix} Texto marcado de entrada vazio. Abortando Passo 3.")
        return False
    logger.info(f"{log_prefix} Recebido texto com marcadores para Passo 3 ({len(marked_text_content)} chars).")

    # Ensure output dirs exist (might be redundant but safe)
    output_notes_dir = os.path.dirname(output_notes_path)
    output_numbered_dir = os.path.dirname(output_numbered_txt_path)
    if not os.path.exists(output_notes_dir): os.makedirs(output_notes_dir, exist_ok=True)
    if output_numbered_dir != output_notes_dir and not os.path.exists(output_numbered_dir):
        os.makedirs(output_numbered_dir, exist_ok=True)


    footnote_counter = 1
    notes_found = []
    footnote_pattern = re.compile(
        r'(\[NOTA_(?:IDIOMA|CITACAO|NOME|TERMO):[^\]]+?\])\s*(\[CONTEUDO_NOTA:([^\]]*?)\])',
        re.IGNORECASE
    )

    def replace_marker_and_collect_note(match):
        nonlocal footnote_counter
        original_marker = match.group(1)
        content_marker = match.group(2)
        content = match.group(3).strip()
        if not content:
             logger.warning(f"{log_prefix} Encontrado marcador [CONTEUDO_NOTA:] vazio após {original_marker}. Ignorando.")
             return ""
        notes_found.append(f"{footnote_counter}. {content}")
        replacement = f"[{footnote_counter}]"
        # logger.debug(f"{log_prefix} Nota {footnote_counter}: '{content}'") # Verbose
        footnote_counter += 1
        return replacement

    logger.info(f"{log_prefix} Processando marcadores e gerando arquivos finais...")
    try:
        final_numbered_text = footnote_pattern.sub(replace_marker_and_collect_note, marked_text_content)

        # Save notes file
        logger.info(f"{log_prefix} Salvando arquivo de notas em: {output_notes_path}")
        with open(output_notes_path, "w", encoding="utf-8") as f_notes:
            f_notes.write(f"Notas de Rodapé Geradas para {author_name} - {book_name}\n") # Add book name
            f_notes.write("=" * (30 + len(author_name) + len(book_name)) + "\n\n")
            if notes_found:
                f_notes.write("\n".join(notes_found))
                f_notes.write("\n")
                logger.info(f"{log_prefix} {len(notes_found)} notas salvas.")
            else:
                f_notes.write("(Nenhuma nota de rodapé foi identificada ou extraída com sucesso)\n")
                logger.info(f"{log_prefix} Nenhuma nota de rodapé identificada/salva.")

        # Save final numbered TXT
        logger.info(f"{log_prefix} Salvando TXT final com números [{footnote_counter-1}] em: {output_numbered_txt_path}")
        with open(output_numbered_txt_path, "w", encoding="utf-8") as f_numbered:
            f_numbered.write(final_numbered_text)
        logger.info(f"{log_prefix} TXT final com números salvo ({len(final_numbered_text)} chars).")

    except Exception as e_final_gen:
        logger.error(f"{log_prefix} Erro durante a geração final dos arquivos TXT (Passo 3): {e_final_gen}")
        logger.error(traceback.format_exc())
        return False

    logger.info(f"{log_prefix} --- Passo 3 concluído. ---")
    return True

# --- Funções para Gerenciar Arquivos Processados ---

def load_processed_files(filepath):
    """Lê o arquivo de log central e retorna um set com identificadores de arquivos processados."""
    processed = set()
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                processed.add(line.strip())
        logger.info(f"Carregados {len(processed)} registros de arquivos processados de '{filepath}'.")
    except FileNotFoundError:
        logger.info(f"Arquivo de log de processados '{filepath}' não encontrado. Iniciando do zero.")
    except Exception as e:
        logger.error(f"Erro ao carregar log de processados '{filepath}': {e}")
    return processed

def log_processed_file(filepath, file_identifier):
    """Adiciona um identificador de arquivo ao log central de processados."""
    try:
        with open(filepath, 'a', encoding='utf-8') as f:
            # Basic file locking for append (might be needed in concurrent scenarios)
            # fcntl.flock(f, fcntl.LOCK_EX)
            f.write(f"{file_identifier}\n")
            # fcntl.flock(f, fcntl.LOCK_UN)
        logger.debug(f"Registrado '{file_identifier}' como processado em '{filepath}'.")
    except Exception as e:
        logger.error(f"Erro ao registrar '{file_identifier}' no log de processados '{filepath}': {e}")


# --- FUNÇÃO PRINCIPAL (main) ---
def main():
    start_time_main = time.time()
    logger.info("========================================================")
    logger.info(f"Iniciando Processador Multi-Autor (Resumo por Arquivo) - {time.strftime('%Y-%m-%d %H:%M:%S')}")
    logger.info(f"Diretório de Entrada TXT: {BASE_INPUT_TXT_DIR}")
    logger.info(f"Diretório de Saída DOCX: {BASE_OUTPUT_DOCX_DIR}")
    logger.info(f"Diretório de Saída TXT (Notas/Final): {BASE_OUTPUT_TXT_DIR}")
    logger.info(f"Template DOCX OBRIGATÓRIO: {TEMPLATE_DOCX}")
    logger.info(f"Log de Arquivos Processados: {PROCESSED_LOG_FILE}")
    logger.info("========================================================")

    # Carrega a lista de arquivos já processados
    processed_files_set = load_processed_files(PROCESSED_LOG_FILE)

    if not os.path.isdir(BASE_INPUT_TXT_DIR):
        logger.error(f"Diretório de entrada base '{BASE_INPUT_TXT_DIR}' não encontrado! Abortando.")
        return

    try:
        author_folders = [f for f in os.listdir(BASE_INPUT_TXT_DIR) if os.path.isdir(os.path.join(BASE_INPUT_TXT_DIR, f))]
    except Exception as e:
        logger.error(f"Erro ao listar diretórios em '{BASE_INPUT_TXT_DIR}': {e}")
        return

    if not author_folders:
        logger.warning(f"Nenhuma subpasta de autor encontrada em '{BASE_INPUT_TXT_DIR}'. Saindo.")
        return

    logger.info(f"Autores encontrados: {len(author_folders)} -> {', '.join(author_folders)}")

    total_books_processed = 0
    total_books_skipped = 0
    total_books_failed = 0

    # Loop por AUTOR
    for author_name in author_folders:
        author_input_dir = os.path.join(BASE_INPUT_TXT_DIR, author_name)
        author_output_docx_dir = os.path.join(BASE_OUTPUT_DOCX_DIR, author_name)
        author_output_txt_dir = os.path.join(BASE_OUTPUT_TXT_DIR, author_name) # Para notas e txt final

        logger.info(f"--- Verificando Autor: {author_name} em '{author_input_dir}' ---")

        # Encontra TODOS os arquivos .txt na pasta do autor
        try:
            input_txt_files = glob.glob(os.path.join(author_input_dir, "*.txt"))
            # Filtra arquivos que parecem ser de saída do próprio script
            input_txt_files = [f for f in input_txt_files if not (
                                 os.path.basename(f).endswith(FINAL_NUMBERED_TXT_BASENAME) or
                                 os.path.basename(f).endswith(NOTES_TXT_FILE_BASENAME) or
                                 # Adicionar outros padrões de exclusão se necessário
                                 os.path.basename(f).startswith("backup_")
                             )]
        except Exception as e:
            logger.error(f"[{author_name}] Erro ao listar arquivos .txt em '{author_input_dir}': {e}")
            continue # Pula para o próximo autor

        if not input_txt_files:
            logger.warning(f"[{author_name}] Nenhum arquivo .txt de entrada (livro) encontrado em '{author_input_dir}'.")
            continue

        logger.info(f"[{author_name}] Encontrados {len(input_txt_files)} arquivos .txt potenciais para processar.")

        # Loop por LIVRO (arquivo .txt) dentro da pasta do autor
        for input_txt_path in input_txt_files:
            book_filename = os.path.basename(input_txt_path)
            # Cria um identificador único para o log de processados
            file_identifier = f"{author_name}/{book_filename}"
            log_prefix_book = f"[{file_identifier}]" # Log prefix específico do livro

            logger.info(f"{log_prefix_book} Verificando...")

            # ** Verifica se este ARQUIVO já foi processado **
            if file_identifier in processed_files_set:
                logger.info(f"{log_prefix_book} Já processado anteriormente (encontrado em '{PROCESSED_LOG_FILE}'). Pulando.")
                total_books_skipped += 1
                continue # Pula para o próximo livro

            # Se não foi processado, continua
            logger.info(f"{log_prefix_book} Iniciando processamento...")
            book_start_time = time.time()

            # Deriva o nome base do livro (removendo extensão)
            base_book_name = os.path.splitext(book_filename)[0]

            # Constrói nomes de arquivos de saída específicos para este livro
            output_docx_path = os.path.join(author_output_docx_dir, f"{base_book_name}_{FINAL_DOCX_BASENAME}")
            output_notes_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{NOTES_TXT_FILE_BASENAME}")
            output_numbered_txt_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{FINAL_NUMBERED_TXT_BASENAME}")

            # Cria diretórios de saída se não existirem (pode ser redundante mas seguro)
            try:
                if not os.path.exists(author_output_docx_dir): os.makedirs(author_output_docx_dir)
                if not os.path.exists(author_output_txt_dir): os.makedirs(author_output_txt_dir)
            except Exception as e_mkdir:
                 logger.error(f"{log_prefix_book} Erro ao criar diretórios de saída: {e_mkdir}. Pulando este livro.")
                 total_books_failed += 1
                 continue # Pula para o próximo livro

            # --- Executa a Sequência de Passos para o LIVRO ---
            all_steps_successful_for_book = True
            corrected_text_content = None
            marked_text_content = None

            # === PASSO 1: CORREÇÃO ===
            pass1_success, corrected_text_content = run_correction_pass(
                gemini_model, input_txt_path, TEMPLATE_DOCX, output_docx_path,
                author_name, base_book_name # Passa nome base do livro
            )
            if not pass1_success:
                logger.error(f"{log_prefix_book} Passo 1 (Correção) FALHOU.")
                all_steps_successful_for_book = False
            else:
                # === PASSO 2: IDENTIFICAÇÃO DE NOTAS ===
                pass2_success, marked_text_content = run_footnote_id_pass(
                    gemini_model, corrected_text_content,
                    author_name, base_book_name # Passa nome base do livro
                )
                if not pass2_success:
                    logger.error(f"{log_prefix_book} Passo 2 (Identificação de Notas) FALHOU.")
                    all_steps_successful_for_book = False
                else:
                    # === PASSO 3: GERAÇÃO FINAL TXT ===
                    pass3_success = run_final_txt_generation(
                        marked_text_content, output_notes_path, output_numbered_txt_path,
                        author_name, base_book_name # Passa nome base do livro
                    )
                    if not pass3_success:
                        logger.error(f"{log_prefix_book} Passo 3 (Geração Final TXT) FALHOU.")
                        all_steps_successful_for_book = False

            # --- Conclusão do Processamento do Livro ---
            book_end_time = time.time()
            book_total_time = book_end_time - book_start_time

            if all_steps_successful_for_book:
                logger.info(f"✅ {log_prefix_book} Processamento concluído com SUCESSO em {book_total_time:.2f} seg.")
                # Registra o sucesso no arquivo central
                log_processed_file(PROCESSED_LOG_FILE, file_identifier)
                processed_files_set.add(file_identifier) # Atualiza o set em memória
                total_books_processed += 1
            else:
                logger.warning(f"⚠️ {log_prefix_book} Processamento concluído com FALHAS em {book_total_time:.2f} seg. O arquivo NÃO será marcado como concluído e será reprocessado na próxima execução.")
                total_books_failed += 1
                # Opcional: Remover arquivos de saída parciais? (Pode ser útil mantê-los para depuração)
                # try:
                #     if os.path.exists(output_docx_path): os.remove(output_docx_path)
                #     if os.path.exists(output_notes_path): os.remove(output_notes_path)
                #     if os.path.exists(output_numbered_txt_path): os.remove(output_numbered_txt_path)
                #     logger.info(f"{log_prefix_book} Arquivos de saída parciais removidos devido à falha.")
                # except Exception as e_clean:
                #     logger.warning(f"{log_prefix_book} Erro ao tentar remover arquivos parciais: {e_clean}")

            logger.info(f"{log_prefix_book} --- Fim do processamento ---")
            # Pausa curta opcional entre livros para evitar sobrecarga rápida da API?
            # time.sleep(1)

        # Fim do loop de livros para o autor atual
        logger.info(f"--- Concluída verificação do Autor: {author_name} ---")

    # --- Fim do Loop Principal ---
    end_time_main = time.time()
    total_time_main = end_time_main - start_time_main
    logger.info("========================================================")
    logger.info("🏁 Processamento Multi-Autor/Multi-Livro Concluído!")
    logger.info(f"Tempo total geral: {total_time_main:.2f} seg ({total_time_main/60:.2f} min).")
    logger.info(f"Livros Processados com Sucesso nesta execução: {total_books_processed}")
    logger.info(f"Livros Pulados (já processados anteriormente): {total_books_skipped}")
    logger.info(f"Livros com Falha (serão reprocessados): {total_books_failed}")
    logger.info(f"Log detalhado salvo em: {log_filepath}")
    logger.info(f"Registro de livros concluídos: {PROCESSED_LOG_FILE}")
    logger.info("Verifique os diretórios de saída para os resultados:")
    logger.info(f"  - DOCX: {BASE_OUTPUT_DOCX_DIR}/<nome_do_autor>/<nome_livro>_Livro_Final_Formatado_Sem_Notas.docx")
    logger.info(f"  - TXTs: {BASE_OUTPUT_TXT_DIR}/<nome_do_autor>/<nome_livro>_Livro_Final_Com_Notas_Numeros.txt")
    logger.info(f"  - TXTs: {BASE_OUTPUT_TXT_DIR}/<nome_do_autor>/<nome_livro>_notas_rodape.txt")
    logger.info("========================================================")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("\nProcesso interrompido manualmente (Ctrl+C). Arquivos incompletos não serão marcados como processados.")
    except Exception as e_main:
        logger.critical(f"Erro fatal inesperado durante a execução de main: {e_main}")
        logger.critical(traceback.format_exc())
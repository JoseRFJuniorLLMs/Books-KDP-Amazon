from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from dotenv import load_dotenv
import os
import tiktoken
import re
import logging
from tqdm import tqdm
import time
import shutil

# === SETUP LOGGING ===
# Ensure logs directory exists
log_dir = "logs"
if not os.path.exists(log_dir):
    os.makedirs(log_dir)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(log_dir, "book_processor.log")),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# === CARREGA VARIÁVEIS DE AMBIENTE DO .env ===
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    logger.error("API key not found. Please set OPENAI_API_KEY in your .env file.")
    exit(1)

# === CONFIGURAÇÕES ===
INPUT_TXT = "rascunho.txt"
TEMPLATE_DOCX = "Estrutura.docx"
OUTPUT_DOCX = "Livro_Final_Formatado.docx"
# Create a timestamped backup filename
backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
BACKUP_DOCX = f"backup_{os.path.splitext(OUTPUT_DOCX)[0]}_{backup_timestamp}.docx"


# Configuration constants
MODEL_NAME = "gpt-4-turbo"  # or "gpt-4o" - Ensure this model name is correct and matches your OpenAI access
# MAX_CHUNK_TOKENS: This is the maximum number of tokens for the INPUT chunk sent to the API
# GPT-4-turbo context window is 128k tokens. Setting input chunk size significantly lower is safer.
MAX_CHUNK_TOKENS = 10000 # Reverted to 10000 as the input context can handle this.

# MAX_OUTPUT_TOKENS: This is the maximum number of tokens the model will GENERATE in its response for each chunk.
# This value MUST NOT exceed the model's completion limit (4096 for gpt-4-turbo).
MAX_OUTPUT_TOKENS = 4096 # CORRECTED: Set to the model's completion limit

TEMPERATURE = 0.7
CHAPTER_PATTERNS = [
    r'^Capítulo \w+',
    r'^CAPÍTULO \w+',
    r'^Capítulo \d+',
    r'^CHAPTER \w+',
    r'^Chapter \d+'
]

# Setup token counter
try:
    tokenizer = tiktoken.encoding_for_model(MODEL_NAME)
    logger.info(f"Tiktoken loaded for model '{MODEL_NAME}'.")
except Exception as e:
    logger.warning(f"Could not load tokenizer for {MODEL_NAME}: {e}. Using approximate character count instead.")
    tokenizer = None

# === SETUP OPENAI ===
try:
    client = OpenAI(api_key=API_KEY)
    # Optional: Test authentication
    # client.models.list()
    # logger.info("OpenAI client initialized and authenticated successfully.")
except Exception as e:
    logger.error(f"Failed to initialize OpenAI client: {e}")
    exit(1)


def count_tokens(text):
    """Count tokens in text, either with tiktoken or approximately with characters."""
    if not text:
        return 0
    if tokenizer:
        return len(tokenizer.encode(text))
    else:
        # Fallback: rough approximation (1 token ≈ 3.5-4 chars for Latin scripts)
        return len(text) // 4


def create_chunks(text, max_tokens):
    """
    Split text into chunks respecting paragraph and potential chapter boundaries.
    Returns a list of chunks, each containing text content <= max_tokens.
    Attempts to avoid splitting mid-paragraph or mid-sentence if possible.
    """
    logger.info(f"Starting chunk creation with max_tokens per chunk: {max_tokens}")
    chunks = []
    current_chunk = ""
    current_chunk_tokens = 0

    # Compile the chapter patterns for efficiency
    chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS), re.IGNORECASE)

    paragraphs = text.split("\n\n")
    logger.info(f"Split text into {len(paragraphs)} paragraphs.")

    for i, paragraph in enumerate(paragraphs):
        paragraph = paragraph.strip()
        if not paragraph:
            continue

        paragraph_tokens = count_tokens(paragraph)

        # Check if adding this paragraph would exceed the limit
        # Add 2 tokens for the potential "\n\n" if it's not the first paragraph in the chunk
        tokens_with_separator = paragraph_tokens + (2 if current_chunk else 0)

        # If adding this paragraph pushes the current chunk over the limit, save current chunk and start a new one
        # Ensure the paragraph itself isn't already too big
        if current_chunk_tokens + tokens_with_separator > max_tokens and current_chunk_tokens > 0:
            chunks.append(current_chunk.strip())
            logger.debug(f"Saved chunk {len(chunks)}. Tokens: {current_chunk_tokens}")
            current_chunk = ""
            current_chunk_tokens = 0

        # Handle cases where a single paragraph is larger than max_tokens
        if paragraph_tokens > max_tokens:
            logger.warning(
                f"Paragraph {i+1} exceeds chunk token limit ({paragraph_tokens} > {max_tokens}). Attempting splitting.")
            # Attempt to split the oversized paragraph into smaller pieces (e.g., by sentences)
            # This is a fallback for extremely long paragraphs; ideally, input text should be structured
            sentences = re.split(r'(?<=[.!?])\s+', paragraph)
            temp_chunk = ""
            temp_tokens = 0
            for sentence in sentences:
                sentence_tokens = count_tokens(sentence)
                if temp_tokens + sentence_tokens <= max_tokens:
                    temp_chunk += sentence + " "
                    temp_tokens += sentence_tokens
                else:
                    if temp_chunk.strip():
                        chunks.append(temp_chunk.strip())
                        logger.debug(f"Saved split sentence chunk {len(chunks)}. Tokens: {temp_tokens}")
                    temp_chunk = sentence + " " # Start new temp chunk with the current sentence
                    temp_tokens = sentence_tokens
            if temp_chunk.strip():
                chunks.append(temp_chunk.strip())
                logger.debug(f"Saved final split sentence chunk {len(chunks)}. Tokens: {temp_tokens}")
            # Reset current chunk tracking as the oversized paragraph was handled
            current_chunk = ""
            current_chunk_tokens = 0
        else:
            # Add the paragraph to the current chunk
            # Check for chapter start for potential chunk break *before* adding the paragraph
            is_chapter_start = chapter_regex.match(paragraph) is not None

            # If this paragraph is a chapter start and we already have content in the current chunk,
            # finalize the current chunk *before* adding the new chapter paragraph.
            # This ensures chapters start in new chunks.
            if is_chapter_start and current_chunk.strip() and i > 0:
                chunks.append(current_chunk.strip())
                logger.debug(f"Saved chunk {len(chunks)} at chapter start. Tokens: {current_chunk_tokens}")
                current_chunk = paragraph + "\n\n"
                current_chunk_tokens = paragraph_tokens
            else:
                 # Add the paragraph to the current chunk
                 current_chunk += paragraph + "\n\n"
                 current_chunk_tokens += tokens_with_separator # Use tokens_with_separator for the increment

    # Add the last chunk if not empty
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
        logger.debug(f"Saved final chunk {len(chunks)}. Tokens: {current_chunk_tokens}")


    logger.info(f"Chunking complete. Created {len(chunks)} chunks.")
    return chunks


def format_with_ai(chunk, is_first_chunk=False):
    """Process a chunk of text with the OpenAI API and format it."""
    logger.debug(f"Calling API for chunk (first: {is_first_chunk}). Chunk tokens: {count_tokens(chunk)}")
    # Different instructions depending on if it's first chunk or a continuation
    context = "Você está formatando o início de um livro. " if is_first_chunk else "Você está formatando uma continuação de texto. "

    chunk_prompt = f"""
    {context}Você é um editor literário. Corrija e formate o texto a seguir como parte de um livro.

    Regras:
    - Corrija erros gramaticais, ortográficos e de concordância em português
    - Mantenha a formatação de capítulos se presente no texto
    - **NÃO** adicione títulos de capítulo se eles não estiverem explicitamente no texto
    - Se houver um marcador de quebra de página (===QUEBRA_DE_PAGINA===) dentro deste fragmento, MANTENHA-O EXATAMENTE COMO ESTÁ
    - Mantenha estilo literário fluido e bem escrito
    - **NÃO** adicione texto introdutório ou conclusivo que não faça parte do rascunho
    - O resultado deve ser APENAS o texto formatado do fragmento.
    - Mantenha fidelidade ao texto original: não omita parágrafos ou adicione conteúdo

    Texto do fragmento:
    \"\"\"{chunk}\"\"\"
    """

    # Try with exponential backoff in case of rate limiting or other temporary API errors
    max_retries = 5 # Increased retries
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": chunk_prompt}],
                temperature=TEMPERATURE,
                max_tokens=MAX_OUTPUT_TOKENS # Use the corrected constant
            )
            formatted_text = response.choices[0].message.content
            logger.debug(f"Successfully processed chunk. Output tokens: {count_tokens(formatted_text)}")
            return formatted_text

        except Exception as e:
            logger.warning(f"API error processing chunk (Attempt {attempt + 1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt + 1 # Exponential backoff with jitter: 1, 3, 5, 9 seconds
                logger.info(f"Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                logger.error(f"Failed to process chunk after {max_retries} attempts.")
                # Consider logging the failed chunk content if needed for debugging
                return chunk  # Return the original chunk if all retries fail


def apply_formatting(doc, formatted_text):
    """Apply formatting to the document based on the formatted text."""
    if not formatted_text:
        logger.warning("Received empty formatted text for a chunk. Skipping insertion.")
        return

    # Split the text by page break markers first
    parts = formatted_text.split("===QUEBRA_DE_PAGINA===")
    chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS), re.IGNORECASE)

    for part_index, part in enumerate(parts):
        part = part.strip()
        if not part:
            continue # Skip empty parts resulting from split

        # Add page breaks between parts (but not before the very first part of the entire document)
        # Check if this is not the first part overall by looking at the document's existing content
        # This logic assumes clear_content() is called initially
        if part_index > 0 or (len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip() != ''):
             doc.add_page_break()
             logger.debug("Added page break.")

        paragraphs_in_part = part.split("\n\n")
        for paragraph in paragraphs_in_part:
            paragraph = paragraph.strip()
            if not paragraph:
                continue # Skip empty paragraphs

            # Check if this is a chapter heading based on patterns
            is_chapter = chapter_regex.match(paragraph) is not None

            p = doc.add_paragraph(paragraph)

            # Apply specific formatting based on content
            if is_chapter:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Apply bold and size to all runs in the paragraph
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(14)
                logger.debug(f"Formatted paragraph as centered chapter title: '{paragraph[:50]}...'")
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Optional: Set font size for normal text
                for run in p.runs:
                    run.font.size = Pt(12) # Example normal text size
                logger.debug(f"Formatted paragraph as justified text: '{paragraph[:50]}...'")


def main():
    logger.info("Iniciando processamento do livro.")

    # === PASSO 1 – Lê o texto bruto ===
    try:
        with open(INPUT_TXT, "r", encoding="utf-8") as f:
            texto_original = f.read()
        logger.info(f"Arquivo '{INPUT_TXT}' carregado com sucesso. Tamanho: {len(texto_original)} caracteres.")
    except FileNotFoundError:
        logger.error(f"Erro: O arquivo de rascunho '{INPUT_TXT}' não foi encontrado.")
        exit(1)
    except Exception as e:
        logger.error(f"Erro ao ler o arquivo '{INPUT_TXT}': {e}")
        exit(1)

    # === PASSO 2 – Split the text into chunks ===
    logger.info(f"Fragmentando o texto em chunks de ~{MAX_CHUNK_TOKENS} tokens...")
    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS)
    logger.info(f"✅ Texto dividido em {len(text_chunks)} fragmentos.")

    # === PASSO 3 – Carrega o template ou cria um novo documento ===
    try:
        doc = Document(TEMPLATE_DOCX)
        logger.info(f"Template '{TEMPLATE_DOCX}' carregado com sucesso.")
        # Create backup of previous output file if it exists
        if os.path.exists(OUTPUT_DOCX):
            try:
                # Use shutil.copy2 to preserve metadata
                shutil.copy2(OUTPUT_DOCX, BACKUP_DOCX)
                logger.info(f"Backup criado: {BACKUP_DOCX}")
            except Exception as e:
                logger.warning(f"Não foi possível criar backup: {e}")

        # === PASSO 4 – Limpa corpo do documento (opcional) ===
        # Check if the body object exists and is not None before trying to clear
        if hasattr(doc, '_body') and doc._body:
            doc._body.clear_content()
            logger.info("Conteúdo do template limpo.")
        else:
             logger.warning("Document body not found or is None. Cannot clear content.")

    except FileNotFoundError:
        logger.warning(f"Template '{TEMPLATE_DOCX}' não encontrado. Criando um novo documento com configurações padrão.")
        doc = Document()
        # Setup basic document properties for a new document
        try:
            section = doc.sections[0]
            section.page_height = Inches(9)
            section.page_width = Inches(6)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            logger.info("Aplicadas configurações de página padrão ao novo documento.")
        except Exception as e:
             logger.warning(f"Não foi possível aplicar configurações de página padrão: {e}")


    # === PASSO 5 – Process and insert each chunk ===
    logger.info("Processando fragmentos com ChatGPT e inserindo no documento...")

    # Use tqdm for a visual progress bar
    for i, chunk in enumerate(tqdm(text_chunks, desc="Processando fragmentos", unit="chunk")):
        logger.info(f"Processando fragmento {i + 1}/{len(text_chunks)}...")

        # Process this chunk using the AI
        formatted_chunk = format_with_ai(chunk, is_first_chunk=(i == 0))

        # Apply formatting and add to the document
        apply_formatting(doc, formatted_chunk)

        # Periodic saves to prevent data loss in case of crashes
        if (i + 1) % 10 == 0 or i == len(text_chunks) - 1: # Save every 10 chunks or on the last chunk
            temp_path = f"{OUTPUT_DOCX}.temp"
            try:
                doc.save(temp_path)
                # Atomically replace the main file with the temp file
                shutil.move(temp_path, OUTPUT_DOCX)
                logger.info(f"Salvo progresso parcial ({i + 1}/{len(text_chunks)} fragmentos)")
            except Exception as e:
                logger.error(f"Erro ao salvar progresso parcial: {e}")
                # If save fails, try to keep the temp file
                if os.path.exists(temp_path):
                    logger.warning(f"Arquivo temporário de salvamento parcial '{temp_path}' mantido.")


    # === PASSO 6 – Salva o novo documento FINAL ===
    # The final save is also covered by the loop's periodic save, but explicitly save here just in case.
    # Also, ensure the final temporary file is cleaned up or moved correctly.
    final_temp_path = f"{OUTPUT_DOCX}.temp"
    if os.path.exists(final_temp_path):
         try:
             shutil.move(final_temp_path, OUTPUT_DOCX)
             logger.info(f"Arquivo temporário final movido para {OUTPUT_DOCX}.")
         except Exception as e:
             logger.error(f"Erro ao mover arquivo temporário final: {e}")


    try:
        # One final save attempt of the doc object itself
        doc.save(OUTPUT_DOCX)
        logger.info(f"📘 Processamento concluído. Livro gerado com sucesso: {OUTPUT_DOCX}")
    except Exception as e:
        logger.error(f"Erro final ao salvar documento: {e}")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.info("Processo interrompido pelo usuário.")
    except Exception as e:
        logger.exception(f"Ocorreu um erro não tratado durante a execução: {e}")
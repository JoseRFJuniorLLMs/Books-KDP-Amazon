# --- Using Google's Gemini API (gemini-1.5-pro) ---

# Standard Python Libraries
from dotenv import load_dotenv
import os
import re
import logging
from tqdm import tqdm
import time
import shutil
import traceback # Para log de erros detalhado
import glob # Para encontrar arquivos .txt
import smtplib # For email
import ssl # For email security
from email.message import EmailMessage # For constructing email

# Third-party Libraries (ensure python-docx is installed: pip install python-docx)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.shared import RGBColor

# Google Generative AI Library (ensure installed: pip install google-generativeai)
import google.generativeai as genai

# === SETUP LOGGING ===
log_dir = "logs"
if not os.path.exists(log_dir): os.makedirs(log_dir)
log_filepath = os.path.join(log_dir, "book_processor_multi_author_mem.log")
# Central log for processed files
PROCESSED_LOG_FILE = os.path.join(log_dir, "processed_books.log")

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(module)s:%(lineno)d - %(funcName)s - %(message)s',
    handlers=[ logging.FileHandler(log_filepath, encoding='utf-8'), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# === LOAD ENVIRONMENT VARIABLES ===
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
# Email configuration from .env file
EMAIL_SENDER_ADDRESS = os.getenv("EMAIL_SENDER_ADDRESS")
EMAIL_SENDER_APP_PASSWORD = os.getenv("EMAIL_SENDER_APP_PASSWORD") # Use App Password for Gmail
EMAIL_RECIPIENT_ADDRESS = "web2ajax@gmail.com" # Hardcoded recipient as requested
EMAIL_SMTP_SERVER = os.getenv("EMAIL_SMTP_SERVER", "smtp.gmail.com") # Default Gmail
EMAIL_SMTP_PORT = int(os.getenv("EMAIL_SMTP_PORT", 587)) # Default Gmail Port (TLS)

# === CONFIGURATIONS ===

# -- Base Directories --
BASE_INPUT_TXT_DIR = "txt"    # Main directory containing author subfolders
BASE_OUTPUT_DOCX_DIR = "docx"  # Main directory for DOCX output
BASE_OUTPUT_TXT_DIR = "txt"   # Main directory for final TXT outputs (Notes, Numbered)

# -- File Names --
# MODIFIED: Point to the correct location inside the docx directory
TEMPLATE_DOCX = os.path.join(BASE_OUTPUT_DOCX_DIR, "Estrutura.docx") # Template OBRIGATÓRIO

# -- Output File Basenames (will be prepended with Book Name) --
FINAL_DOCX_BASENAME = "Livro_Final_Formatado_Sem_Notas.docx"
FINAL_NUMBERED_TXT_BASENAME = "Livro_Final_Com_Notas_Numeros.txt" # Updated name
NOTES_TXT_FILE_BASENAME = "notas_rodape.txt"

# -- API and Processing Settings --
MODEL_NAME = "gemini-1.5-pro"
MAX_CHUNK_TOKENS = 1500
MAX_OUTPUT_TOKENS = 8192 # Max for gemini-1.5-pro
TEMPERATURE = 0.5 # Conservative for fidelity

# -- Styles and Patterns --
NORMAL_STYLE_NAME = "Normal" # Style name expected in TEMPLATE_DOCX
CHAPTER_PATTERNS = [
    r'^\s*Capítulo \w+', r'^\s*CAPÍTULO \w+', r'^\s*Capítulo \d+',
    r'^\s*CHAPTER \w+', r'^\s*Chapter \d+', r'^\s*LIVRO \w+', r'^\s*PARTE \w+',
    # Add other patterns if needed
]
PAGE_BREAK_MARKER = "===QUEBRA_DE_PAGINA==="
AI_FAILURE_MARKER = "*** FALHA NA IA - TEXTO ORIGINAL ABAIXO ***"
FORMATTING_ERROR_MARKER = "*** ERRO DE FORMATAÇÃO - TEXTO ORIGINAL ABAIXO ***"

# --- VALIDATE API KEY ---
if not GOOGLE_API_KEY:
    logger.error("FATAL: GOOGLE_API_KEY não encontrada nas variáveis de ambiente (.env).")
    exit(1)

# --- VALIDATE EMAIL CONFIG (Informational) ---
if not EMAIL_SENDER_ADDRESS or not EMAIL_SENDER_APP_PASSWORD:
    logger.warning("AVISO: Variáveis de ambiente para envio de e-mail (EMAIL_SENDER_ADDRESS, EMAIL_SENDER_APP_PASSWORD) não configuradas no .env. Notificação por e-mail será desativada.")
else:
    logger.info(f"Configuração de e-mail carregada. Notificações serão enviadas de '{EMAIL_SENDER_ADDRESS}' para '{EMAIL_RECIPIENT_ADDRESS}'.")

# --- SETUP GEMINI CLIENT ---
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    # Lenient safety settings - use with caution and understand the risks
    safety_settings_lenient = {
        'HATE': 'BLOCK_NONE', 'HARASSMENT': 'BLOCK_NONE',
        'SEXUAL' : 'BLOCK_NONE', 'DANGEROUS' : 'BLOCK_NONE'
    }
    generation_config = genai.GenerationConfig(
                    temperature=TEMPERATURE,
                    max_output_tokens=MAX_OUTPUT_TOKENS
                )
    gemini_model = genai.GenerativeModel(
        MODEL_NAME,
        safety_settings=safety_settings_lenient,
        generation_config=generation_config
    )
    logger.info(f"Modelo Gemini '{MODEL_NAME}' inicializado com sucesso.")
except Exception as e:
    logger.error(f"FATAL: Falha ao inicializar modelo Gemini ({MODEL_NAME}): {e}")
    logger.error(traceback.format_exc())
    exit(1)

# --- HELPER FUNCTIONS ---

def count_tokens_approx(text):
    """Estimates token count (approx. 3 chars/token)."""
    if not text: return 0
    return len(text) // 3

def create_chunks(text, max_tokens, author_name="N/A", book_name="N/A"):
    """Splits text into chunks, handling large paragraphs."""
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} Iniciando criação de chunks. Máx tokens (aprox): {max_tokens}")
    chunks = []
    current_chunk = ""
    paragraphs = text.split("\n\n") # Split by double newline
    paragraphs_stripped = [p.strip() for p in paragraphs if p.strip()]

    logger.info(f"{log_prefix} Texto dividido inicialmente em {len(paragraphs_stripped)} blocos não vazios ('\\n\\n').")

    for i, paragraph_text in enumerate(paragraphs):
        if not paragraph_text.strip(): # Skip empty blocks but preserve spacing
            if chunks and chunks[-1].strip() and not chunks[-1].endswith("\n\n"):
                chunks[-1] += "\n\n"
            continue

        paragraph_tokens = count_tokens_approx(paragraph_text)
        tokens_with_separator = paragraph_tokens + (count_tokens_approx("\n\n") if current_chunk else 0)

        # Combine if fits
        if current_chunk and (count_tokens_approx(current_chunk) + tokens_with_separator > max_tokens):
            chunks.append(current_chunk)
            logger.debug(f"{log_prefix} Chunk {len(chunks)} salvo (limite atingido).")
            current_chunk = paragraph_text # Start new chunk
        elif count_tokens_approx(current_chunk) + tokens_with_separator <= max_tokens:
            separator = "\n\n" if current_chunk else ""
            current_chunk += separator + paragraph_text
        elif paragraph_tokens > max_tokens: # Current chunk is empty, but paragraph itself is too big
             logger.warning(f"{log_prefix} Parágrafo {i+1} ({paragraph_tokens} tk) excede limite {max_tokens}. Iniciando SUBDIVISÃO.")
             # Subdivide the large paragraph (by sentence or line)
             sub_chunks_added_count = 0
             sentences = re.split(r'(?<=[.!?])\s+', paragraph_text) # Try sentences first
             if len(sentences) <= 1: sentences = paragraph_text.split('\n') # Fallback to lines

             current_sub_chunk = ""
             for sentence_num, sentence in enumerate(sentences):
                 sentence_clean = sentence.strip()
                 if not sentence_clean: continue
                 sentence_tokens = count_tokens_approx(sentence)
                 tokens_with_sub_separator = sentence_tokens + (count_tokens_approx("\n") if current_sub_chunk else 0)

                 if current_sub_chunk and (count_tokens_approx(current_sub_chunk) + tokens_with_sub_separator > max_tokens):
                     chunks.append(current_sub_chunk)
                     sub_chunks_added_count += 1
                     logger.debug(f"{log_prefix} Sub-chunk {len(chunks)} salvo (Parág. {i+1}).")
                     current_sub_chunk = sentence
                 elif sentence_tokens > max_tokens: # Sentence itself is too big
                     if current_sub_chunk: # Save previous sub-chunk first
                         chunks.append(current_sub_chunk)
                         sub_chunks_added_count += 1
                         logger.debug(f"{log_prefix} Sub-chunk {len(chunks)} salvo (antes sent. longa, Parág. {i+1}).")
                     chunks.append(sentence) # Add long sentence as its own chunk
                     sub_chunks_added_count += 1
                     logger.warning(f"{log_prefix} -> Sentença/Linha {sentence_num+1} ({sentence_tokens} tk) excede limite. Adicionada como sub-chunk (PODE FALHAR NA API).")
                     current_sub_chunk = "" # Reset
                 else: # Add to current sub-chunk
                     sub_separator = "\n" if current_sub_chunk else "" # Use single newline within subdivided paragraph
                     current_sub_chunk += sub_separator + sentence

             if current_sub_chunk: # Save the last sub-chunk
                 chunks.append(current_sub_chunk)
                 sub_chunks_added_count += 1
                 logger.debug(f"{log_prefix} Último sub-chunk {len(chunks)} salvo (Parág. {i+1}).")

             if sub_chunks_added_count == 0: # Failed to subdivide
                  logger.warning(f"{log_prefix} Parágrafo {i+1} excedeu limite, mas não subdividido. Adicionando original (PODE FALHAR NA API).")
                  chunks.append(paragraph_text)

             current_chunk = "" # Reset main chunk after handling large paragraph
        else: # Paragraph fits and starts a new chunk
            current_chunk = paragraph_text

    if current_chunk: # Add the last remaining chunk
        chunks.append(current_chunk)
        logger.debug(f"{log_prefix} Chunk final {len(chunks)} salvo.")

    # Post-processing: Merge small consecutive chunks if possible
    merged_chunks = []
    temp_chunk = ""
    for i, chunk in enumerate(chunks):
        tokens_with_separator = count_tokens_approx(chunk) + (count_tokens_approx("\n\n") if temp_chunk else 0)
        if count_tokens_approx(temp_chunk) + tokens_with_separator <= max_tokens:
            separator = "\n\n" if temp_chunk else ""
            temp_chunk += separator + chunk
        else:
            if temp_chunk: merged_chunks.append(temp_chunk)
            temp_chunk = chunk
    if temp_chunk: # Add the last merged chunk
        merged_chunks.append(temp_chunk)

    final_chunk_count = len(merged_chunks)
    logger.info(f"{log_prefix} ✅ Chunking concluído. {final_chunk_count} chunks finais (após merge).")
    return merged_chunks

def _call_gemini_api(model, prompt_text, chunk_for_log, author_name="N/A", book_name="N/A"):
    """Internal function to call Gemini API with retries and error handling."""
    log_prefix = f"[{author_name}/{book_name}]"
    max_retries = 5
    base_wait_time = 5 # Initial wait time in seconds
    log_chunk_preview = chunk_for_log[:150].replace('\n', '\\n') + ('...' if len(chunk_for_log) > 150 else '')

    for attempt in range(max_retries):
        logger.info(f"{log_prefix} Chamando API (Tentativa {attempt + 1}/{max_retries}). Chunk (início): '{log_chunk_preview}'")
        try:
            # Use generation_config defined when creating the model
            response = model.generate_content(prompt_text)

            # Detailed response logging/checking
            finish_reason = "UNKNOWN"; safety_ratings = "UNKNOWN"; block_reason = "N/A"; result_text = None

            # 1. Check for prompt blocking
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                 if hasattr(response.prompt_feedback, 'block_reason') and response.prompt_feedback.block_reason:
                      block_reason = response.prompt_feedback.block_reason.name
                      logger.error(f"{log_prefix} API BLOQUEOU O PROMPT (Tentativa {attempt + 1}). Razão: {block_reason}.")
                      return None # Immediate failure if prompt is blocked

            # 2. Check for candidates
            if not response.candidates:
                 logger.error(f"{log_prefix} API retornou SEM CANDIDATOS (Tentativa {attempt + 1}). Resposta: {response}")
                 # Continue to retry, might be a temporary issue
            else:
                 # 3. Process the first candidate
                 try:
                    candidate = response.candidates[0]
                    finish_reason = candidate.finish_reason.name if hasattr(candidate, 'finish_reason') and candidate.finish_reason else "FINISH_REASON_UNKNOWN"
                    safety_ratings = [(r.category.name, r.probability.name) for r in candidate.safety_ratings] if candidate.safety_ratings else "N/A"
                    logger.debug(f"{log_prefix} API Call OK (Tentativa {attempt + 1}). Finish: {finish_reason}. Safety: {safety_ratings}")

                    # Handle different finish reasons
                    if finish_reason == "STOP": pass # Normal completion
                    elif finish_reason == "MAX_TOKENS": logger.warning(f"{log_prefix} API TRUNCOU resposta devido a MAX_OUTPUT_TOKENS.")
                    elif finish_reason == "SAFETY": logger.warning(f"{log_prefix} API interrompeu resposta devido a SAFETY. Ratings: {safety_ratings}.")
                    elif finish_reason == "RECITATION": logger.warning(f"{log_prefix} API interrompeu resposta devido a RECITATION.")
                    elif finish_reason == "OTHER": logger.warning(f"{log_prefix} API interrompeu resposta por OUTRA RAZÃO.")
                    else: logger.warning(f"{log_prefix} API retornou com finish_reason inesperado: {finish_reason}.")

                    # 4. Extract text
                    result_text = ""
                    if hasattr(candidate, 'content') and hasattr(candidate.content, 'parts'):
                        text_parts = [part.text for part in candidate.content.parts if hasattr(part, 'text')]
                        result_text = "".join(text_parts).strip() if text_parts else ""
                    elif hasattr(response, 'text') and response.text: # Fallback
                        result_text = response.text.strip()

                    # 5. Check if text was extracted
                    if result_text:
                        # logger.debug(f"{log_prefix} Texto API recebido (100 chars): '{result_text[:100].replace('\n', '\\n')}...'")
                        # Sanity check: very short output compared to input
                        if len(result_text) < len(chunk_for_log) * 0.1 and len(chunk_for_log) > 100:
                            logger.warning(f"{log_prefix} Resposta da API parece muito curta. Input len: {len(chunk_for_log)}, Output len: {len(result_text)}.")
                        return result_text # SUCCESS!
                    else:
                         logger.warning(f"{log_prefix} Resposta API não continha texto utilizável (Tentativa {attempt+1}), mesmo com candidato. Finish Reason: {finish_reason}.")
                         # Continue to retry

                 except Exception as e_details:
                    logger.error(f"{log_prefix} Erro ao extrair detalhes/texto da resposta API (Tentativa {attempt+1}): {e_details} - Resposta Crua: {response}")
                    logger.error(traceback.format_exc())
                    # Continue to retry

            # --- Exponential backoff with jitter before retrying (only if no success or prompt block) ---
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time) # Jitter
                logger.info(f"{log_prefix} Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                 logger.error(f"{log_prefix} Falha final na API após {max_retries} tentativas para o chunk: '{log_chunk_preview}'")
                 return None # Failure after all retries

        except Exception as e: # Handle exceptions during the API call itself
            logger.warning(f"{log_prefix} Erro durante a chamada da API ({model.model_name}) (Tentativa {attempt + 1}/{max_retries}): {e}")
            logger.error(traceback.format_exc())
            # Specific handling for common transient errors
            if "RESOURCE_EXHAUSTED" in str(e) or "429" in str(e):
                 logger.warning(f"{log_prefix} Erro de cota (RESOURCE_EXHAUSTED / 429). Aumentando espera.")
                 base_wait_time = max(15, base_wait_time) # Increase base wait time for quota errors
            elif "Internal error encountered." in str(e) or "500" in str(e):
                 logger.warning(f"{log_prefix} Erro interno do servidor (500). Tentando novamente.")
                 # Maintain normal wait
            # Apply backoff for other errors too
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                logger.info(f"{log_prefix} Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                logger.error(f"{log_prefix} Falha final na API após {max_retries} tentativas (erro na chamada) para o chunk: '{log_chunk_preview}'")
                return None # Failure after all retries

    logger.error(f"{log_prefix} Loop de tentativas da API concluído sem sucesso explícito para o chunk: '{log_chunk_preview}'")
    return None

def format_with_ai_correction_only(model, chunk, author_name, book_name, is_first_chunk=False):
    """Calls Gemini API focusing ONLY on OCR/grammar correction."""
    log_prefix = f"[{author_name}/{book_name}]"
    context_start = "Você está formatando o início de um livro." if is_first_chunk else "Você está continuando a formatação de um texto de livro existente."
    ocr_errors_examples = """
        * **Troca de letras similares:** 'rn' vs 'm', 'c' vs 'e', 't' vs 'f', 'l' vs 'i', 'I' vs 'l', 'O' vs '0', 'S' vs '5', 'B' vs '8'.
        * **Hífens indevidos:** Palavras quebradas incorretamente no meio ou hífens extras.
        * **Hífens ausentes:** Palavras que deveriam ser hifenizadas (ex: "guarda-chuva") aparecem juntas ou separadas.
        * **Espaços ausentes ou extras:** Palavras coladas ("onomundo") ou espaços excessivos.
        * **Pontuação incorreta:** Pontos finais trocados por vírgulas, pontos de interrogação/exclamação mal interpretados.
        * **Acentuação:** Falta de acentos (ex: 'e' vs 'é', 'a' vs 'à'), acentos incorretos (crase onde não deve) ou caracteres estranhos no lugar de acentos.
        * **Letras duplicadas ou ausentes:** "caaasa" ou "casaa" em vez de "casa".
        * **Confusão maiúsculas/minúsculas:** Nomes próprios em minúsculas, inícios de frase em minúsculas.
        * **Caracteres especiais/ruído:** Símbolos aleatórios inseridos no texto.
        * **Quebras de linha estranhas:** Parágrafos divididos no meio sem motivo aparente. Preserve as quebras de parágrafo intencionais (duas quebras de linha).
    """
    chunk_prompt = f"""
{context_start} Você é um editor literário proficiente em português do Brasil. Sua tarefa é CORRIGIR e FORMATAR o fragmento de texto a seguir, que pertence a um livro do autor {author_name}.

**CONTEXTO IMPORTANTE:** Este texto provavelmente foi extraído via OCR de um PDF e pode conter erros de reconhecimento, digitação e gramática. O objetivo principal é obter um texto LIMPO e CORRETO em português do Brasil padrão, mantendo a estrutura e o significado originais.

**SIGA RIGOROSAMENTE ESTAS REGRAS:**

1.  **Correção Profunda:** Corrija TODOS os erros gramaticais, ortográficos, de pontuação, acentuação e concordância verbal/nominal. Use o português do Brasil como referência. FOQUE em erros comuns de OCR como os listados abaixo.
2.  **Estilo e Tom:** Mantenha o estilo literário e o tom do texto original do autor {author_name}. Seja claro, fluido e envolvente. NÃO altere o significado, a voz ou a intenção do autor.
3.  **Fidelidade Estrutural:** MANTENHA a estrutura de parágrafos original. Parágrafos são geralmente separados por UMA linha em branco (duas quebras de linha `\\n\\n`). NÃO junte parágrafos que estavam separados. NÃO divida parágrafos desnecessariamente.
4.  **Sem Adições/Remoções:** NÃO omita frases ou informações. NÃO adicione conteúdo, introduções, resumos, conclusões ou opiniões que não estavam no fragmento original. SEJA ESTRITAMENTE FIEL AO CONTEÚDO.
5.  **Marcadores de Capítulo/Quebra:** Se encontrar marcadores como 'Capítulo X', '***', '---', etc., no início de um parágrafo, MANTENHA-OS EXATAMENTE como estão, naquele parágrafo específico. NÃO adicione ou remova esses marcadores.
6.  **Quebra de Página:** Se o marcador '{PAGE_BREAK_MARKER}' aparecer, MANTENHA-O EXATAMENTE onde está, em sua própria linha, sem texto antes ou depois na mesma linha.
7.  **Erros Comuns de OCR (FOCO ESPECIAL):** Preste atenção e corrija diligentemente:
    {ocr_errors_examples}
8.  **Formato de Saída:** Retorne APENAS o texto corrigido e formatado. Use parágrafos separados por duas quebras de linha (`\\n\\n`). NÃO use NENHUMA formatação especial como Markdown (`*`, `#`, `_`), HTML, etc. Retorne TEXTO PURO. Não inclua comentários sobre o que você fez, apenas o texto resultante.

**Texto do fragmento para processar (pode conter erros):**
\"\"\"
{chunk}
\"\"\"

**Lembre-se: O resultado deve ser APENAS o texto corrigido.**
"""
    # logger.debug(f"{log_prefix} Enviando chunk para CORREÇÃO (API: {model.model_name}). Tam Aprox: {count_tokens_approx(chunk)} tk")
    return _call_gemini_api(model, chunk_prompt, chunk, author_name, book_name)


def format_with_ai_footnote_only(model, chunk, author_name, book_name):
    """Calls Gemini API focusing ONLY on identifying potential footnotes."""
    log_prefix = f"[{author_name}/{book_name}]"
    chunk_prompt = f"""
Você é um assistente de edição trabalhando no texto do autor {author_name}. Sua tarefa é analisar o fragmento de texto A SEGUIR, que JÁ FOI CORRIGIDO no passo anterior, e APENAS inserir marcadores para potenciais notas de rodapé onde estritamente necessário.

**REGRAS IMPORTANTES:**

1.  **NÃO ALTERE O TEXTO CORRIGIDO:** Não faça correções, não mude palavras, não reestruture frases. Apenas insira os marcadores.
2.  **MARCADORES DE NOTA:** Insira marcadores APENAS nos seguintes casos:
    * **Termos em Idioma Estrangeiro (não comuns):** Imediatamente APÓS uma palavra ou frase curta em latim, francês, inglês, etc., que não seja de uso corrente em português, insira: `[NOTA_IDIOMA:palavra_original][CONTEUDO_NOTA:Tradução ou breve explicação]`. Exemplo: "...uma certa *joie de vivre*[NOTA_IDIOMA:joie de vivre][CONTEUDO_NOTA:Alegria de viver (francês)]..."
    * **Citações/Referências:** APÓS uma citação direta curta ou uma referência bibliográfica no texto (ex: (Autor, Ano)), insira: `[NOTA_CITACAO:Texto citado ou referência][CONTEUDO_NOTA:Referência bibliográfica completa ou fonte, se conhecida ou inferível]`. Exemplo: "...como disse Foucault (1975)[NOTA_CITACAO:Foucault (1975)][CONTEUDO_NOTA:FOUCAULT, Michel. Vigiar e Punir. 1975.], a disciplina..."
    * **Nomes Próprios (contexto essencial):** APÓS um nome de pessoa, local ou evento histórico POUCO CONHECIDO que SEJA ESSENCIAL contextualizar brevemente para a compreensão do trecho, insira: `[NOTA_NOME:Nome Mencionado][CONTEUDO_NOTA:Breve identificação (datas, relevância)]`. Use com MODERAÇÃO. Exemplo: "...influenciado por Kropotkin[NOTA_NOME:Kropotkin][CONTEUDO_NOTA:Piotr Kropotkin (1842-1921), anarquista russo.]..."
    * **Termos Técnicos/Jargão (essencial):** APÓS um termo técnico MUITO específico de uma área, cuja definição SEJA INDISPENSÁVEL para o leitor geral entender o argumento naquele ponto, insira: `[NOTA_TERMO:Termo Técnico][CONTEUDO_NOTA:Definição concisa]`. Use com MUITA MODERAÇÃO. Exemplo: "...aplicando a análise de isotopias[NOTA_TERMO:Isotopias][CONTEUDO_NOTA:Na semiótica greimasiana, recorrência de categorias sêmicas que garante a homogeneidade de um discurso.]..."
3.  **FORMATO DOS MARCADORES:** Use EXATAMENTE `[NOTA_TIPO:Referência]` seguido IMEDIATAMENTE por `[CONTEUDO_NOTA:Explicação]`. Não adicione espaços entre eles. Não use outros formatos.
4.  **CRITÉRIO:** Seja conservador. Adicione notas apenas se a informação for realmente útil e provavelmente desconhecida para um leitor culto médio. É MELHOR ERRAR POR NÃO ADICIONAR do que por adicionar excessivamente. NÃO adicione notas para termos comuns, nomes famosos ou citações óbvias.
5.  **NÃO INVENTE CONTEÚDO:** O `[CONTEUDO_NOTA:...]` deve ser uma tradução direta, uma referência óbvia, ou uma contextualização muito breve e factual, se possível inferida do próprio texto ou conhecimento geral básico. NÃO pesquise externamente para criar notas complexas. Se não souber o conteúdo, NÃO insira a nota.
6.  **SAÍDA:** Retorne APENAS o texto original (do input) com os marcadores inseridos nos locais exatos. Mantenha a estrutura de parágrafos (`\\n\\n`). Não adicione NENHUM outro texto, comentário ou explicação.

**Texto JÁ CORRIGIDO para analisar e inserir marcadores de nota:**
\"\"\"
{chunk}
\"\"\"

**Lembre-se: NÃO altere o texto, apenas insira os marcadores `[NOTA_...][CONTEUDO_NOTA:...]` quando apropriado.**
"""
    # logger.debug(f"{log_prefix} Enviando chunk para IDENTIFICAÇÃO DE NOTAS (API: {model.model_name}). Tam Aprox: {count_tokens_approx(chunk)} tk")
    return _call_gemini_api(model, chunk_prompt, chunk, author_name, book_name)

# --- PROCESSING STEP FUNCTIONS ---

def apply_formatting_pass1(doc, formatted_chunk_text, normal_style_name, chapter_patterns, corrected_text_list, author_name, book_name):
    """
    Applies formatting to DOCX (Pass 1 - no notes) & collects text to return.
    Relies on template styles for Normal text. Applies SPECIFIC font/size
    to detected chapters, overriding template styles for chapters.
    """
    log_prefix = f"[{author_name}/{book_name}]"
    if not formatted_chunk_text or not formatted_chunk_text.strip():
        # logger.warning(f"{log_prefix} Chunk formatado vazio (Passo 1). Pulando.")
        return

    # Collect text for next step (memory)
    plain_text_for_list = formatted_chunk_text.replace(PAGE_BREAK_MARKER, "\n\n").strip()
    if plain_text_for_list:
        corrected_text_list.append(plain_text_for_list)
    else:
        if formatted_chunk_text.strip() == PAGE_BREAK_MARKER:
             logger.debug(f"{log_prefix} Chunk continha apenas marcador de página, não adicionado à lista.")
        # else: # Avoid logging empty strings from API response formatting
             # logger.warning(f"{log_prefix} Texto formatado resultou em vazio. Original: '{formatted_chunk_text[:50]}...'")
             pass

    # --- Apply formatting to DOCX ---
    normal_style = None
    try: # Fetch Normal style (existence checked once in run_correction_pass)
        if normal_style_name in doc.styles:
             normal_style = doc.styles[normal_style_name]
    except Exception as e_style:
         logger.error(f"{log_prefix} Erro ao acessar estilo '{normal_style_name}': {e_style}.")

    chapter_regex = re.compile('|'.join(chapter_patterns), re.IGNORECASE)

    # Process parts split by page breaks
    parts = formatted_chunk_text.split(PAGE_BREAK_MARKER)
    content_present_before = any(p.text.strip() for p in doc.paragraphs)

    for part_index, part in enumerate(parts):
        part_clean = part.strip()

        # Add page break before new part (except first), avoid duplicates
        if part_index > 0:
             last_para_is_page_break = False
             if doc.paragraphs:
                 last_p = doc.paragraphs[-1]
                 if not last_p.text.strip() and any(run.text and '\f' in run.text for run in last_p.runs):
                     last_para_is_page_break = True
             if not last_para_is_page_break:
                 # logger.debug(f"{log_prefix} Adicionando quebra de página DOCX (antes da parte {part_index + 1}).")
                 doc.add_page_break()
             # else: logger.debug(f"{log_prefix} Quebra de página omitida (duplicada).")
        elif content_present_before and doc.paragraphs and not doc.paragraphs[-1].text.strip() and any(run.text and '\f' in run.text for run in doc.paragraphs[-1].runs):
            # logger.debug(f"{log_prefix} Quebra de página omitida início (anterior era quebra).")
            pass

        if not part_clean: # Skip empty parts (after page break)
            if part_index > 0 : content_present_before = True
            continue

        # Process paragraphs within the part
        paragraphs_in_part = part_clean.split("\n\n")
        for paragraph_text in paragraphs_in_part:
            paragraph_text_clean = paragraph_text.strip()
            if not paragraph_text_clean: # Handle intentional empty paragraphs for spacing
                if doc.paragraphs and doc.paragraphs[-1].text.strip(): # Only add if previous wasn't empty
                     p = doc.add_paragraph()
                     if normal_style: p.style = normal_style # Style empty paragraph if possible
                continue

            # Detect special markers / chapters
            is_ai_failure_marker = paragraph_text_clean.startswith(AI_FAILURE_MARKER)
            is_formatting_error_marker = paragraph_text_clean.startswith(FORMATTING_ERROR_MARKER)
            is_chapter = not is_ai_failure_marker and not is_formatting_error_marker and chapter_regex.match(paragraph_text_clean) is not None

            # Add paragraph and run
            p = doc.add_paragraph()
            run = p.add_run(paragraph_text_clean)
            content_present_before = True

            # Apply specific formatting or styles
            try:
                if is_chapter:
                    # SPECIFIC formatting for chapters, overriding template
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.name = 'French Script MT'
                    run.font.size = Pt(48)
                    run.bold = False
                    # logger.debug(f"{log_prefix} Aplicada formatação específica de capítulo (French Script MT 48pt).")
                elif is_ai_failure_marker or is_formatting_error_marker:
                    # Specific formatting for error markers
                    if normal_style: p.style = normal_style # Base on Normal if available
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Red
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # logger.debug(f"{log_prefix} Aplicada formatação de marcador de erro.")
                else: # Normal paragraph
                    # Apply 'Normal' style ONLY if found in the template
                    if normal_style:
                        p.style = normal_style
                    # else: Rely on document defaults if Normal style missing

            except Exception as e_apply_style:
                 logger.error(f"{log_prefix} Erro ao aplicar formatação/estilo: {e_apply_style}. Texto: '{paragraph_text_clean[:50]}...'")
                 pass # Continue without specific formatting on error


def run_correction_pass(model, input_txt_path, template_docx_path, output_docx_path, author_name, book_name):
    """
    Pass 1: Corrects text using AI, generates base DOCX using template.
    Returns: Tuple (bool: success, str|None: corrected_text)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 1: Correção e Geração DOCX ---")
    logger.info(f"{log_prefix} Lendo texto original de: {input_txt_path}")
    try:
        with open(input_txt_path, "r", encoding="utf-8") as f: texto_original = f.read()
        logger.info(f"{log_prefix} Entrada '{os.path.basename(input_txt_path)}' carregada ({len(texto_original)} chars).")
    except FileNotFoundError:
        logger.error(f"{log_prefix} FATAL: Arquivo de entrada '{input_txt_path}' não encontrado."); return (False, None)
    except Exception as e:
        logger.error(f"{log_prefix} FATAL ao ler entrada '{input_txt_path}': {e}")
        logger.error(traceback.format_exc()); return (False, None)

    # Ensure output directory exists
    output_docx_dir = os.path.dirname(output_docx_path)
    os.makedirs(output_docx_dir, exist_ok=True) # Create if not exists

    # Chunk the input text
    logger.info(f"{log_prefix} Dividindo texto original em chunks...")
    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS, author_name, book_name)
    if not text_chunks:
        logger.error(f"{log_prefix} Nenhum chunk gerado. Abortando Passo 1."); return (False, None)
    logger.info(f"{log_prefix} Texto dividido em {len(text_chunks)} chunks.")

    doc = None
    normal_style_exists = False
    logger.info(f"{log_prefix} Preparando documento DOCX para: {output_docx_path} usando template OBRIGATÓRIO.")
    try:
        # Backup existing output DOCX if it exists
        if os.path.exists(output_docx_path):
             backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
             backup_docx_path = os.path.join(output_docx_dir, f"backup_{os.path.splitext(os.path.basename(output_docx_path))[0]}_{backup_timestamp}.docx")
             try:
                 shutil.copy2(output_docx_path, backup_docx_path)
                 logger.info(f"{log_prefix} Backup do DOCX anterior criado: {os.path.basename(backup_docx_path)}")
             except Exception as e_bkp:
                 logger.warning(f"{log_prefix} Falha ao criar backup de '{os.path.basename(output_docx_path)}': {e_bkp}")

        # Load MANDATORY template
        if not os.path.exists(template_docx_path):
            logger.error(f"{log_prefix} FATAL: Template OBRIGATÓRIO '{template_docx_path}' não encontrado. Abortando.")
            return (False, None)

        try:
            doc = Document(template_docx_path)
            logger.info(f"{log_prefix} Template '{template_docx_path}' carregado. Limpando conteúdo existente...")
            # Clear existing content (paragraphs, tables) from template body
            # This preserves template's styles, headers, footers, page setup etc.
            for para in doc.paragraphs: p_element = para._element; p_element.getparent().remove(p_element)
            for table in doc.tables: tbl_element = table._element; tbl_element.getparent().remove(tbl_element)
            logger.info(f"{log_prefix} Conteúdo principal do template limpo.")

            # Check for essential 'Normal' style ONCE after loading template
            if NORMAL_STYLE_NAME in doc.styles:
                normal_style_exists = True
                logger.info(f"{log_prefix} Estilo '{NORMAL_STYLE_NAME}' encontrado no template.")
            else:
                logger.warning(f"{log_prefix} AVISO: Estilo '{NORMAL_STYLE_NAME}' NÃO encontrado no template '{TEMPLATE_DOCX}'. Formatação de parágrafos normais dependerá do padrão do documento.")

        except Exception as e_load_template:
             logger.error(f"{log_prefix} FATAL: Falha ao carregar/limpar template OBRIGATÓRIO '{template_docx_path}': {e_load_template}. Abortando.")
             return (False, None)

    except Exception as e_doc:
        logger.error(f"{log_prefix} FATAL: Erro crítico ao preparar DOCX: {e_doc}")
        logger.error(traceback.format_exc()); return (False, None)

    # Process chunks through API and format DOCX
    logger.info(f"{log_prefix} Iniciando chamadas à API para CORREÇÃO de {len(text_chunks)} chunks...")
    corrected_text_list_pass1 = []
    processed_chunks_count = 0
    failed_chunks_count = 0
    # Use TQDM for progress bar
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P1: Corrigindo", unit="chunk", leave=False)

    for i, chunk in progress_bar:
        # Update progress bar description
        # progress_bar.set_description(f"{log_prefix} P1: Corrigindo Chunk {i+1}/{len(text_chunks)}")

        # Call AI for correction
        corrected_chunk = format_with_ai_correction_only(model, chunk, author_name, book_name, is_first_chunk=(i == 0))

        # Apply formatting to DOCX and collect text
        if corrected_chunk is not None: # Check for None explicitly after API call changes
            if corrected_chunk.strip(): # Process if not empty string
                 try:
                    apply_formatting_pass1(doc, corrected_chunk, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
                    processed_chunks_count += 1
                 except Exception as format_err:
                    logger.error(f"{log_prefix} Erro na apply_formatting_pass1 (Chunk {i+1}): {format_err}.")
                    logger.error(traceback.format_exc())
                    failed_chunks_count += 1
                    # Apply fallback with error marker
                    try:
                        fallback_text = f"{FORMATTING_ERROR_MARKER}\n\n{chunk}"
                        apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
                    except Exception as fallback_format_err:
                        logger.critical(f"{log_prefix} Falha CRÍTICA ao aplicar fallback de erro FORMATAÇÃO (Chunk {i+1}): {fallback_format_err}.")
            # else: Handle cases where API might return empty string correctly (already handled in apply_formatting)

        else: # API call failed (returned None)
            logger.warning(f"{log_prefix} Chunk {i+1} falhou na CORREÇÃO (API retornou None). Usando fallback com marcador.")
            failed_chunks_count += 1
            # Apply fallback with AI failure marker
            try:
                fallback_text = f"{AI_FAILURE_MARKER}\n\n{chunk}"
                apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
            except Exception as fallback_format_err:
                 logger.critical(f"{log_prefix} Falha CRÍTICA ao aplicar fallback de falha API (Chunk {i+1}): {fallback_format_err}.")

        # Save DOCX progress periodically
        processed_total = processed_chunks_count + failed_chunks_count
        if processed_total > 0 and (processed_total % 10 == 0 or (i + 1) == len(text_chunks)):
            temp_save_path = f"{output_docx_path}.{processed_total}.temp_save" # Include count in temp name
            try:
                # logger.debug(f"{log_prefix} Salvando progresso DOCX (chunk {i+1})...")
                doc.save(temp_save_path)
                # Atomic rename/move is generally safer than copy+delete
                shutil.move(temp_save_path, output_docx_path)
                logger.info(f"{log_prefix} Progresso DOCX (Passo 1) salvo ({processed_total} chunks processados).")
            except Exception as e_save:
                 logger.error(f"{log_prefix} Erro ao salvar progresso DOCX (Chunk {i+1}) para '{os.path.basename(output_docx_path)}': {e_save}")
                 # Attempt to remove potentially corrupt temp file
                 if os.path.exists(temp_save_path):
                      try: os.remove(temp_save_path)
                      except OSError: pass

    # Save final DOCX after loop
    try:
        logger.info(f"{log_prefix} Salvando DOCX final (Passo 1) em: {os.path.basename(output_docx_path)}")
        doc.save(output_docx_path)
    except Exception as e_final_save:
        logger.error(f"{log_prefix} Erro no salvamento final do DOCX (Passo 1): {e_final_save}")
        logger.error(traceback.format_exc())
        # Continue to return text, but DOCX might be incomplete/corrupt

    # Join corrected text for returning to Pass 2
    full_corrected_text = "\n\n".join(corrected_text_list_pass1)
    logger.info(f"{log_prefix} Acumulado texto corrigido para Pass 2 ({len(full_corrected_text)} chars).")

    logger.info(f"{log_prefix} --- Passo 1 concluído. Chunks OK: {processed_chunks_count}, Falhas/Fallback: {failed_chunks_count} ---")
    # Return success status (True even if some chunks failed but DOCX save attempted) and the text
    return (True, full_corrected_text)


def run_footnote_id_pass(model, corrected_text_content, author_name, book_name):
    """
    Pass 2: Identifies potential footnotes in the corrected text using AI.
    Returns: Tuple (bool: success, str|None: marked_text)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 2: Identificação de Notas ---")
    if not corrected_text_content:
        logger.error(f"{log_prefix} Texto corrigido de entrada vazio. Abortando Passo 2.")
        return (False, None)
    # logger.info(f"{log_prefix} Recebido texto corrigido para Passo 2 ({len(corrected_text_content)} chars).") # Verbose

    # Chunk the corrected text
    logger.info(f"{log_prefix} Dividindo texto corrigido em chunks para notas...")
    text_chunks = create_chunks(corrected_text_content, MAX_CHUNK_TOKENS, author_name, book_name)
    if not text_chunks:
        logger.error(f"{log_prefix} Nenhum chunk gerado do texto corrigido. Abortando Passo 2."); return (False, None)
    logger.info(f"{log_prefix} Texto corrigido dividido em {len(text_chunks)} chunks.")

    # Process chunks through API for footnote marking
    logger.info(f"{log_prefix} Iniciando chamadas à API para IDENTIFICAÇÃO DE NOTAS em {len(text_chunks)} chunks...")
    marked_text_list_pass2 = []
    processed_chunks_count = 0
    failed_chunks_count = 0
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P2: Notas", unit="chunk", leave=False)

    for i, chunk in progress_bar:
        # progress_bar.set_description(f"{log_prefix} P2: Notas Chunk {i+1}/{len(text_chunks)}")
        # Call AI for footnote marking
        marked_chunk = format_with_ai_footnote_only(model, chunk, author_name, book_name)

        if marked_chunk is not None: # API call succeeded (might have markers or not)
            marked_text_list_pass2.append(marked_chunk)
            processed_chunks_count += 1
            # Optional logging based on content
            # if "[NOTA_" in marked_chunk: logger.debug(f"{log_prefix} Chunk {i+1}: marcadores de nota encontrados.")
            # else: logger.debug(f"{log_prefix} Chunk {i+1}: NENHUM marcador de nota adicionado.")
        else: # API call failed
            logger.warning(f"{log_prefix} Chunk {i+1} falhou na IDENTIFICAÇÃO DE NOTAS (API retornou None). Usando texto original do chunk.")
            marked_text_list_pass2.append(chunk) # Use original chunk as fallback
            failed_chunks_count += 1

    # Join marked text for returning to Pass 3
    full_marked_text = "\n\n".join(marked_text_list_pass2)
    logger.info(f"{log_prefix} Acumulado texto com marcadores para Pass 3 ({len(full_marked_text)} chars).")

    logger.info(f"{log_prefix} --- Passo 2 concluído. Chunks OK: {processed_chunks_count}, Falhas/Fallback: {failed_chunks_count} ---")
    # Return success status and the marked text
    return (True, full_marked_text)


def run_final_txt_generation(marked_text_content, output_notes_path, output_numbered_txt_path, author_name, book_name):
    """
    Pass 3: Processes footnote markers [NOTA_...] to generate final numbered TXT
            and a separate notes file.
    Returns: bool: success status
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 3: Geração Final TXT (Notas e Numerado) ---")
    if marked_text_content is None: # Check for None explicitly
        logger.error(f"{log_prefix} Texto marcado de entrada é None. Abortando Passo 3.")
        return False
    # logger.info(f"{log_prefix} Recebido texto com marcadores para Passo 3 ({len(marked_text_content)} chars).") # Verbose

    # Ensure output directories exist (should be created by main already)
    os.makedirs(os.path.dirname(output_notes_path), exist_ok=True)
    os.makedirs(os.path.dirname(output_numbered_txt_path), exist_ok=True)

    footnote_counter = 1
    notes_found = []
    # Regex to find footnote marker pairs: [NOTA_TYPE:Ref][CONTEUDO_NOTA:Content]
    footnote_pattern = re.compile(
        r'(\[NOTA_(?:IDIOMA|CITACAO|NOME|TERMO):[^\]]+?\])\s*(\[CONTEUDO_NOTA:([^\]]*?)\])',
        re.IGNORECASE # Ignore case for marker tags
    )

    # Function to replace markers and collect notes
    def replace_marker_and_collect_note(match):
        nonlocal footnote_counter
        original_marker = match.group(1)
        content_marker = match.group(2)
        content = match.group(3).strip()

        if not content: # Handle empty content notes
             logger.warning(f"{log_prefix} Encontrado marcador [CONTEUDO_NOTA:] vazio após {original_marker}. Ignorando nota.")
             return "" # Remove both markers without adding number

        notes_found.append(f"{footnote_counter}. {content}") # Store numbered note
        replacement = f"[{footnote_counter}]" # Replacement text in main content
        # logger.debug(f"{log_prefix} Nota {footnote_counter}: '{content}'")
        footnote_counter += 1
        return replacement

    logger.info(f"{log_prefix} Processando marcadores e gerando arquivos finais TXT...")
    try:
        # Process the entire marked text using the replacement function
        final_numbered_text = footnote_pattern.sub(replace_marker_and_collect_note, marked_text_content)

        # --- Save the notes file ---
        logger.info(f"{log_prefix} Salvando arquivo de notas em: {os.path.basename(output_notes_path)}")
        with open(output_notes_path, "w", encoding="utf-8") as f_notes:
            f_notes.write(f"Notas de Rodapé Geradas para {author_name} - {book_name}\n")
            f_notes.write("=" * (30 + len(author_name) + len(book_name)) + "\n\n")
            if notes_found:
                f_notes.write("\n".join(notes_found))
                f_notes.write("\n") # Extra newline for clarity
                logger.info(f"{log_prefix} {len(notes_found)} notas salvas em '{os.path.basename(output_notes_path)}'.")
            else:
                f_notes.write("(Nenhuma nota de rodapé foi identificada ou extraída com sucesso)\n")
                logger.info(f"{log_prefix} Nenhuma nota de rodapé identificada/salva.")

        # --- Save the final TXT with numbered references ---
        logger.info(f"{log_prefix} Salvando TXT final com números [{footnote_counter-1}] em: {os.path.basename(output_numbered_txt_path)}")
        with open(output_numbered_txt_path, "w", encoding="utf-8") as f_numbered:
            f_numbered.write(final_numbered_text)
        logger.info(f"{log_prefix} TXT final com números salvo ({len(final_numbered_text)} chars).")

    except Exception as e_final_gen:
        logger.error(f"{log_prefix} FATAL: Erro durante a geração final dos arquivos TXT (Passo 3): {e_final_gen}")
        logger.error(traceback.format_exc())
        return False # Indicate failure

    logger.info(f"{log_prefix} --- Passo 3 concluído. ---")
    return True # Indicate success

# --- FUNCTIONS FOR MANAGING PROCESSED FILES LOG ---

def load_processed_files(filepath):
    """Reads the central log file and returns a set of processed file identifiers."""
    processed = set()
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                cleaned_line = line.strip()
                if cleaned_line: # Avoid adding empty lines
                    processed.add(cleaned_line)
        logger.info(f"Carregados {len(processed)} registros de arquivos processados de '{filepath}'.")
    except FileNotFoundError:
        logger.info(f"Arquivo de log de processados '{filepath}' não encontrado. Iniciando sem histórico.")
    except Exception as e:
        logger.error(f"Erro ao carregar log de processados '{filepath}': {e}")
    return processed

def log_processed_file(filepath, file_identifier):
    """Appends a file identifier to the central processed log file."""
    try:
        # Use append mode ('a')
        with open(filepath, 'a', encoding='utf-8') as f:
            # Basic file locking might be needed if running concurrently, but usually fine for append.
            # import fcntl
            # fcntl.flock(f, fcntl.LOCK_EX)
            f.write(f"{file_identifier}\n")
            # fcntl.flock(f, fcntl.LOCK_UN)
        # logger.debug(f"Registrado '{file_identifier}' como processado em '{filepath}'.") # Verbose
    except Exception as e:
        logger.error(f"Erro ao registrar '{file_identifier}' no log de processados '{filepath}': {e}")

# --- FUNCTION TO SEND COMPLETION EMAIL ---
def send_completion_email(sender_email, sender_password, recipient_email, smtp_server, smtp_port,
                          processed_count, skipped_count, failed_count, total_duration_seconds,
                          main_log_path, processed_log_path):
    """Sends a completion notification email."""

    if not sender_email or not sender_password:
        logger.warning("E-mail de envio ou senha não configurados no .env. Não é possível enviar notificação.")
        return

    logger.info(f"Preparando e-mail de notificação para {recipient_email}...")

    subject = "Script Processador de Livros - Conclusão"
    body = f"""
Olá,

O script de processamento de livros concluiu a execução.

Resumo da Execução:
--------------------------------------------------
- Livros Processados com Sucesso (novos): {processed_count}
- Livros Pulados (já concluídos antes): {skipped_count}
- Livros com Falha (precisam de atenção): {failed_count}
- Tempo Total de Execução: {total_duration_seconds:.2f} segundos ({total_duration_seconds/60:.2f} minutos)
--------------------------------------------------

Logs para Consulta:
- Log Detalhado da Execução: {os.path.abspath(main_log_path)}
- Log de Livros Já Concluídos: {os.path.abspath(processed_log_path)}

Atenciosamente,
Seu Script Processador de Livros
"""

    message = EmailMessage()
    message['Subject'] = subject
    message['From'] = sender_email
    message['To'] = recipient_email
    message.set_content(body)

    context = ssl.create_default_context() # For secure connection

    try:
        server = None # Initialize server variable
        logger.info(f"Conectando ao servidor SMTP: {smtp_server}:{smtp_port}...")
        if smtp_port == 465: # SSL connection
             server = smtplib.SMTP_SSL(smtp_server, smtp_port, context=context)
             server.login(sender_email, sender_password)
        else: # Assume TLS (port 587 or other)
            server = smtplib.SMTP(smtp_server, smtp_port, timeout=30) # Added timeout
            server.ehlo() # Identify client to server
            server.starttls(context=context) # Secure the connection
            server.ehlo() # Re-identify after TLS
            server.login(sender_email, sender_password)

        logger.info("Enviando e-mail...")
        server.send_message(message)
        logger.info(f"✅ E-mail de notificação enviado com sucesso para {recipient_email}.")

    except smtplib.SMTPAuthenticationError:
        logger.error("FALHA NA AUTENTICAÇÃO do e-mail. Verifique EMAIL_SENDER_ADDRESS e EMAIL_SENDER_APP_PASSWORD no .env.")
        logger.error("Lembre-se: Para Gmail com 2FA, use uma 'Senha de App'.")
    except smtplib.SMTPServerDisconnected:
         logger.error("Servidor SMTP desconectou inesperadamente. Tente novamente.")
    except smtplib.SMTPConnectError as e:
         logger.error(f"Erro ao conectar ao servidor SMTP {smtp_server}:{smtp_port}. Verifique o nome/porta e a rede. Erro: {e}")
    except smtplib.SMTPException as e:
        logger.error(f"Erro SMTP ao enviar e-mail: {e}")
        logger.error(traceback.format_exc())
    except ssl.SSLError as e:
         logger.error(f"Erro SSL/TLS ao conectar ao servidor SMTP: {e}")
         logger.error("Verifique a porta e as configurações de segurança (SSL/TLS).")
    except OSError as e:
         logger.error(f"Erro de Rede/OS (ex: Timeout, Host não encontrado) ao tentar conectar ao servidor SMTP: {e}")
         logger.error("Verifique a conexão com a internet e o endereço/porta do servidor.")
    except Exception as e:
        logger.error(f"Erro inesperado ao enviar e-mail: {e}")
        logger.error(traceback.format_exc())
    finally:
        if server:
            try:
                server.quit() # Ensure connection is closed
            except smtplib.SMTPServerDisconnected:
                 pass # Ignore if already disconnected
            except Exception as e_quit:
                 logger.warning(f"Erro ao fechar conexão SMTP: {e_quit}")


# --- MAIN EXECUTION ---
def main():
    start_time_main = time.time()
    logger.info("========================================================")
    logger.info(f"Iniciando Processador Multi-Autor/Livro (Resumo por Arquivo) - {time.strftime('%Y-%m-%d %H:%M:%S')}")
    logger.info(f"Diretório de Entrada TXT: {BASE_INPUT_TXT_DIR}")
    logger.info(f"Diretório de Saída DOCX: {BASE_OUTPUT_DOCX_DIR}")
    logger.info(f"Diretório de Saída TXT (Notas/Final): {BASE_OUTPUT_TXT_DIR}")
    logger.info(f"Template DOCX OBRIGATÓRIO: {TEMPLATE_DOCX}")
    logger.info(f"Log de Arquivos Processados: {PROCESSED_LOG_FILE}")
    logger.info("========================================================")

    # Load the set of already processed files from the central log
    processed_files_set = load_processed_files(PROCESSED_LOG_FILE)

    # --- Find Author Folders ---
    if not os.path.isdir(BASE_INPUT_TXT_DIR):
        logger.error(f"FATAL: Diretório de entrada base '{BASE_INPUT_TXT_DIR}' não encontrado! Abortando.")
        return
    try:
        author_folders = sorted([f for f in os.listdir(BASE_INPUT_TXT_DIR) if os.path.isdir(os.path.join(BASE_INPUT_TXT_DIR, f))])
    except Exception as e:
        logger.error(f"FATAL: Erro ao listar diretórios em '{BASE_INPUT_TXT_DIR}': {e}")
        return
    if not author_folders:
        logger.warning(f"Nenhuma subpasta de autor encontrada em '{BASE_INPUT_TXT_DIR}'. Saindo.")
        return
    logger.info(f"Autores encontrados ({len(author_folders)}): {', '.join(author_folders)}")

    # --- Initialize Counters for Summary ---
    total_books_processed_this_run = 0
    total_books_skipped = 0
    total_books_failed = 0

    # === MAIN LOOP: Iterate through each AUTHOR folder ===
    for author_name in author_folders:
        author_input_dir = os.path.join(BASE_INPUT_TXT_DIR, author_name)
        author_output_docx_dir = os.path.join(BASE_OUTPUT_DOCX_DIR, author_name)
        author_output_txt_dir = os.path.join(BASE_OUTPUT_TXT_DIR, author_name) # For notes and final txt

        logger.info(f"--- Verificando Autor: {author_name} em '{author_input_dir}' ---")

        # --- Find potential BOOK files (.txt) within the author folder ---
        try:
            input_txt_files = sorted(glob.glob(os.path.join(author_input_dir, "*.txt")))
            # Filter out files that look like output from this script
            input_txt_files_filtered = [
                f for f in input_txt_files if not (
                    os.path.basename(f).endswith(FINAL_NUMBERED_TXT_BASENAME) or
                    os.path.basename(f).endswith(NOTES_TXT_FILE_BASENAME) or
                    os.path.basename(f).startswith("backup_") # Exclude backups
                )
            ]
            if len(input_txt_files) != len(input_txt_files_filtered):
                 logger.debug(f"[{author_name}] Filtrados {len(input_txt_files) - len(input_txt_files_filtered)} arquivos que pareciam ser de saída.")
            input_txt_files = input_txt_files_filtered

        except Exception as e:
            logger.error(f"[{author_name}] Erro ao listar/filtrar arquivos .txt em '{author_input_dir}': {e}. Pulando autor.")
            continue # Skip to the next author

        if not input_txt_files:
            logger.warning(f"[{author_name}] Nenhum arquivo .txt de entrada (livro) encontrado/restante em '{author_input_dir}'.")
            continue

        logger.info(f"[{author_name}] Encontrados {len(input_txt_files)} arquivos .txt potenciais para processar.")

        # === INNER LOOP: Iterate through each BOOK file (.txt) ===
        for input_txt_path in input_txt_files:
            book_filename = os.path.basename(input_txt_path)
            # Create a unique identifier for the processed log file
            file_identifier = f"{author_name}/{book_filename}"
            log_prefix_book = f"[{file_identifier}]" # Specific log prefix for this book

            logger.info(f"{log_prefix_book} Verificando status...")

            # --- Check if this specific BOOK was already processed successfully ---
            if file_identifier in processed_files_set:
                logger.info(f"{log_prefix_book} Já processado anteriormente (encontrado em '{PROCESSED_LOG_FILE}'). Pulando.")
                total_books_skipped += 1
                continue # Skip to the next book file

            # --- Start Processing the Book ---
            logger.info(f"{log_prefix_book} >>> Iniciando processamento...")
            book_start_time = time.time()

            # Derive base name for output files (remove .txt extension)
            base_book_name = os.path.splitext(book_filename)[0]

            # Construct specific output file paths for this book
            output_docx_path = os.path.join(author_output_docx_dir, f"{base_book_name}_{FINAL_DOCX_BASENAME}")
            output_notes_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{NOTES_TXT_FILE_BASENAME}")
            output_numbered_txt_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{FINAL_NUMBERED_TXT_BASENAME}")

            # Ensure output directories exist for this author
            try:
                os.makedirs(author_output_docx_dir, exist_ok=True)
                os.makedirs(author_output_txt_dir, exist_ok=True)
            except Exception as e_mkdir:
                 logger.error(f"{log_prefix_book} ERRO ao criar diretórios de saída: {e_mkdir}. Pulando este livro.")
                 total_books_failed += 1
                 continue # Skip to the next book

            # --- Execute Processing Steps for this Book ---
            all_steps_successful_for_book = False # Assume failure until all steps succeed
            corrected_text_content = None
            marked_text_content = None

            try:
                # === STEP 1: Correction & DOCX Generation ===
                pass1_success, corrected_text_content = run_correction_pass(
                    gemini_model, input_txt_path, TEMPLATE_DOCX, output_docx_path,
                    author_name, base_book_name
                )
                if not pass1_success:
                    logger.error(f"{log_prefix_book} Passo 1 (Correção/DOCX) FALHOU.")
                    # No need to set all_steps_successful_for_book = False, it's already False
                else:
                    # === STEP 2: Footnote Identification ===
                    pass2_success, marked_text_content = run_footnote_id_pass(
                        gemini_model, corrected_text_content,
                        author_name, base_book_name
                    )
                    if not pass2_success:
                        logger.error(f"{log_prefix_book} Passo 2 (Identificação Notas) FALHOU.")
                    else:
                        # === STEP 3: Final TXT Generation ===
                        pass3_success = run_final_txt_generation(
                            marked_text_content, output_notes_path, output_numbered_txt_path,
                            author_name, base_book_name
                        )
                        if not pass3_success:
                            logger.error(f"{log_prefix_book} Passo 3 (Geração TXTs Finais) FALHOU.")
                        else:
                            # Only if ALL steps succeeded:
                            all_steps_successful_for_book = True

            except Exception as e_step: # Catch unexpected errors during the step calls
                 logger.error(f"{log_prefix_book} Erro inesperado durante a execução dos passos: {e_step}")
                 logger.error(traceback.format_exc())
                 all_steps_successful_for_book = False # Ensure failure

            # --- Conclude Processing for this Book ---
            book_end_time = time.time()
            book_total_time = book_end_time - book_start_time

            if all_steps_successful_for_book:
                logger.info(f"✅ {log_prefix_book} Processamento CONCLUÍDO COM SUCESSO em {book_total_time:.2f} seg.")
                # Log successful completion in the central file
                log_processed_file(PROCESSED_LOG_FILE, file_identifier)
                processed_files_set.add(file_identifier) # Update in-memory set
                total_books_processed_this_run += 1
            else:
                logger.warning(f"⚠️ {log_prefix_book} Processamento CONCLUÍDO COM FALHAS em {book_total_time:.2f} seg. O arquivo NÃO foi marcado como concluído.")
                total_books_failed += 1
                # Optional: Clean up partial output files for failed books?
                # Consider leaving them for debugging unless disk space is critical.

            logger.info(f"{log_prefix_book} --- Fim do processamento ---")
            # Optional short pause between books?
            # time.sleep(1)

        # --- End of loop for books within the current author ---
        logger.info(f"--- Concluída verificação do Autor: {author_name} ---")

    # === End of Main Loop (All Authors Processed) ===
    end_time_main = time.time()
    total_time_main = end_time_main - start_time_main

    # === Final Summary Logging ===
    logger.info("========================================================")
    logger.info("🏁 Processamento Multi-Autor/Multi-Livro Concluído!")
    logger.info(f"Tempo total geral: {total_time_main:.2f} seg ({total_time_main/60:.2f} min).")
    logger.info(f"Livros Processados com Sucesso (nesta execução): {total_books_processed_this_run}")
    logger.info(f"Livros Pulados (já processados anteriormente): {total_books_skipped}")
    logger.info(f"Livros com Falha (precisam de atenção/reprocessamento): {total_books_failed}")
    logger.info(f"Log detalhado salvo em: {os.path.abspath(log_filepath)}")
    logger.info(f"Registro de livros já concluídos: {os.path.abspath(PROCESSED_LOG_FILE)}")
    logger.info("Verifique os diretórios de saída para os resultados.")
    logger.info(f"  - DOCX: {BASE_OUTPUT_DOCX_DIR}/<NomeAutor>/<NomeLivroBase>_Livro_Final_Formatado_Sem_Notas.docx")
    logger.info(f"  - TXTs: {BASE_OUTPUT_TXT_DIR}/<NomeAutor>/<NomeLivroBase>_Livro_Final_Com_Notas_Numeros.txt")
    logger.info(f"  - TXTs: {BASE_OUTPUT_TXT_DIR}/<NomeAutor>/<NomeLivroBase>_notas_rodape.txt")
    logger.info("========================================================")

    # === Send Completion Email ===
    if EMAIL_SENDER_ADDRESS and EMAIL_SENDER_APP_PASSWORD:
        send_completion_email(
            sender_email=EMAIL_SENDER_ADDRESS,
            sender_password=EMAIL_SENDER_APP_PASSWORD,
            recipient_email=EMAIL_RECIPIENT_ADDRESS,
            smtp_server=EMAIL_SMTP_SERVER,
            smtp_port=EMAIL_SMTP_PORT,
            processed_count=total_books_processed_this_run,
            skipped_count=total_books_skipped,
            failed_count=total_books_failed,
            total_duration_seconds=total_time_main,
            main_log_path=log_filepath,
            processed_log_path=PROCESSED_LOG_FILE
        )
    else:
        logger.info("Envio de e-mail pulado (configuração ausente no .env).")


# --- Script Entry Point ---
if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("\nProcesso interrompido manualmente (Ctrl+C). E-mail de conclusão não será enviado.")
    except Exception as e_main:
        logger.critical(f"FATAL: Erro fatal inesperado durante a execução de main: {e_main}")
        logger.critical(traceback.format_exc())
        # Optionally, try to send a failure email here if configured
        # if EMAIL_SENDER_ADDRESS and EMAIL_SENDER_APP_PASSWORD:
        #     send_failure_email(...) # Would need a dedicated function
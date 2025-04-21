# --- Using Google's Gemini API (gemini-1.5-flash) ---

# from openai import OpenAI # No longer needed
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.styles.style import _ParagraphStyle # Para checagem de tipo
from docx.oxml.shared import OxmlElement # Para adicionar estilo se não existir
from docx.shared import RGBColor # Para definir cor de fonte (exemplo, não usado por padrão)

from dotenv import load_dotenv
import os
# import tiktoken # tiktoken is for OpenAI models - consider Google's token counting if needed
import re
import logging
from tqdm import tqdm
import time
import shutil

# Import the Google Generative AI library
import google.generativeai as genai

# === SETUP LOGGING ===
log_dir = "logs"
if not os.path.exists(log_dir):
    os.makedirs(log_dir)

# Ensure logs directory exists
log_filepath = os.path.join(log_dir, "book_processor_ocr.log") # Changed log filename

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(module)s - %(funcName)s - %(message)s',
    handlers=[
        logging.FileHandler(log_filepath, encoding='utf-8'), # Ensure UTF-8 for logs
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# === CARREGA VARIÁVEIS DE AMBIENTE DO .env ===
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# === CONFIGURAÇÕES ===
INPUT_TXT = "rascunho_ocr.txt" # Changed input name to reflect content
TEMPLATE_DOCX = "Estrutura.docx" # Template pode conter estilos pré-definidos
OUTPUT_DOCX = "Livro_Final_Formatado_Gemini_OCR_Corrigido_v1.docx" # Changed output name

# Model name for Gemini
MODEL_NAME = "gemini-1.5-flash" # Using the recommended model

# MAX_CHUNK_TOKENS: Max tokens for the INPUT chunk (approximate for chunking logic).
# Gemini 1.5 Flash has a 1M token context window. 10000 is very safe.
# ATENÇÃO: A função count_tokens_approx é uma APROXIMAÇÃO baseada em caracteres.
# O tamanho real em tokens pode variar. Ajuste se necessário.
MAX_CHUNK_TOKENS = 10000 # Pode precisar reduzir se prompts mais longos causarem problemas

# MAX_OUTPUT_TOKENS: Max tokens the model will GENERATE per chunk.
# Gemini 1.5 Flash supports up to 8192. 4096 is safe.
MAX_OUTPUT_TOKENS = 4096

TEMPERATURE = 0.6 # Pode ser útil reduzir um pouco a temperatura para tarefas de correção mais focadas

# Nomes dos estilos do Word a serem usados (devem existir no TEMPLATE_DOCX)
# Se não existirem, o script aplicará formatação direta como fallback.
NORMAL_STYLE_NAME = "Normal" # Ou "TextoCorpo", "Texto Normal", etc. conforme seu template
CHAPTER_STYLE_NAME = "Heading 1" # Ou "TítuloCapítulo", "Título 1", etc. conforme seu template

# Padrões Regex para identificar inícios de capítulo
CHAPTER_PATTERNS = [
    r'^\s*Capítulo \w+',
    r'^\s*CAPÍTULO \w+',
    r'^\s*Capítulo \d+',
    r'^\s*CHAPTER \w+',
    r'^\s*Chapter \d+',
    r'^\s*LIVRO \w+', # Exemplo adicional
    r'^\s*PARTE \w+', # Exemplo adicional
]

# Padrões Regex para identificar outras quebras (ex: quebra de cena)
# Estes também podem forçar o início de um novo chunk, se desejado.
OTHER_BREAK_PATTERNS = [
    r'^\s*\*\*\*\s*$', # Linha contendo apenas "***" (com espaços opcionais)
    r'^\s*---+\s*$',    # Linha contendo "---" ou mais hífens
]

# Marcador de quebra de página explícito no texto de entrada
PAGE_BREAK_MARKER = "===QUEBRA_DE_PAGINA==="

# --- Fim das Configurações ---

# --- Validação da API Key ---
if not GOOGLE_API_KEY:
    logger.error("Google API key (GOOGLE_API_KEY) não encontrada no arquivo .env ou variáveis de ambiente.")
    exit(1)

# --- Setup Gemini API Client ---
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    gemini_model = genai.GenerativeModel(MODEL_NAME)
    logger.info(f"Modelo Gemini '{MODEL_NAME}' inicializado com sucesso.")
except Exception as e:
    logger.error(f"Falha ao inicializar o cliente ou modelo Gemini: {e}")
    exit(1)

# --- Funções Auxiliares ---

def count_tokens_approx(text):
    """
    Aproxima a contagem de tokens usando contagem de caracteres.
    NOTA: Esta é uma aproximação rápida para a lógica de chunking.
    A contagem real de tokens pode diferir significativamente.
    """
    if not text:
        return 0
    # Aproximação: 1 token ≈ 4 caracteres para scripts latinos (ajuste se necessário)
    return len(text) // 4

def create_chunks(text, max_tokens):
    """
    Divide o texto em chunks respeitando parágrafos, capítulos e outras quebras.
    Retorna uma lista de chunks, cada um com texto <= max_tokens (aproximado).
    Tenta evitar dividir no meio de parágrafos ou frases, se possível.
    """
    logger.info(f"Iniciando criação de chunks. Máximo de tokens (aprox.) por chunk: {max_tokens}")
    chunks = []
    current_chunk = ""
    current_chunk_tokens = 0

    # Compila os padrões regex para eficiência
    all_break_patterns = CHAPTER_PATTERNS + OTHER_BREAK_PATTERNS
    break_regex = re.compile('|'.join(f"({p})" for p in all_break_patterns), re.IGNORECASE | re.MULTILINE)
    full_line_break_regex = re.compile(r'|'.join(all_break_patterns), re.IGNORECASE)

    # Divide por quebras de linha DUPLAS primeiro, mantendo quebras simples dentro dos parágrafos
    paragraphs = text.split("\n\n")
    logger.info(f"Texto dividido inicialmente em {len(paragraphs)} blocos (baseado em '\\n\\n').")

    processed_paragraphs = []
    # Rejunta blocos que são apenas uma linha, tratando quebras simples como parte do parágrafo anterior
    # Isso ajuda a evitar que linhas curtas (possivelmente de formatação ruim do OCR) quebrem chunks desnecessariamente.
    temp_para = ""
    for para in paragraphs:
        stripped_para = para.strip()
        if not stripped_para: # Ignora blocos completamente vazios
             if temp_para: # Se havia algo antes, fecha o parágrafo acumulado
                 processed_paragraphs.append(temp_para)
                 temp_para = ""
             continue

        # Se o parágrafo atual NÃO começa com um marcador de quebra E existe um parágrafo anterior acumulado,
        # E o parágrafo anterior não termina com pontuação forte (.!?)
        # Tenta juntar com o anterior (considerando quebras simples podem ter separado frases)
        # Esta lógica pode ser complexa e talvez seja melhor simplificar ou remover se causar problemas.
        # Por agora, vamos manter a divisão por \n\n como principal.
        # A lógica abaixo está comentada por ser potencialmente problemática.
        # should_join = False
        # if temp_para and not full_line_break_regex.match(stripped_para):
        #     last_char = temp_para.strip()[-1:]
        #     if last_char not in ['.', '!', '?', ':']:
        #         should_join = True

        # if should_join:
        #    temp_para += "\n" + para # Usa \n simples para juntar
        # else:
        #    if temp_para: # Salva o anterior antes de começar um novo
        #        processed_paragraphs.append(temp_para)
        #    temp_para = para # Começa um novo parágrafo acumulado

        # Lógica Simplificada: Trata cada bloco de \n\n como um parágrafo potencial
        if temp_para:
             processed_paragraphs.append(temp_para) # Salva o parágrafo anterior
        temp_para = para # Começa/atualiza o parágrafo atual

    if temp_para: # Salva o último parágrafo acumulado
        processed_paragraphs.append(temp_para)

    logger.info(f"Texto reprocessado em {len(processed_paragraphs)} parágrafos lógicos.")


    for i, paragraph in enumerate(processed_paragraphs):
        # Usa a contagem aproximada de tokens
        paragraph_tokens = count_tokens_approx(paragraph) # Contamos o parágrafo inteiro como veio

        # Estimativa de tokens incluindo separador (aproximadamente 1 token para "\n\n")
        # Adicionamos o separador apenas se já houver conteúdo no chunk
        tokens_with_separator = paragraph_tokens + (1 if current_chunk else 0)

        # Verifica se adicionar o parágrafo excede o limite do chunk atual
        if current_chunk_tokens > 0 and (current_chunk_tokens + tokens_with_separator > max_tokens):
            chunks.append(current_chunk.strip())
            logger.debug(f"Chunk {len(chunks)} salvo (limite atingido). Tokens (aprox.): {current_chunk_tokens}")
            current_chunk = ""
            current_chunk_tokens = 0

        # Lida com parágrafos individuais que excedem o limite
        if paragraph_tokens > max_tokens:
            logger.warning(
                f"Parágrafo {i+1} ({paragraph_tokens} tokens aprox.) excede limite de {max_tokens} tokens por chunk. "
                f"O parágrafo será adicionado como um chunk único, potencialmente excedendo o limite."
            )
            # Se havia algo no chunk atual, salva primeiro
            if current_chunk.strip():
                 chunks.append(current_chunk.strip())
                 logger.debug(f"Chunk {len(chunks)} salvo (antes do parágrafo longo). Tokens (aprox.): {current_chunk_tokens}")

            # Adiciona o parágrafo grande como seu próprio chunk
            chunks.append(paragraph.strip()) # Adiciona o parágrafo grande como está
            logger.debug(f"Chunk {len(chunks)} salvo (contendo parágrafo longo único). Tokens (aprox.): {paragraph_tokens}")
            # Reseta o chunk principal, pois o parágrafo grande foi tratado
            current_chunk = ""
            current_chunk_tokens = 0
            continue # Pula para o próximo parágrafo

        # Lógica para adicionar parágrafo normal (que não excede o limite sozinho)
        paragraph_clean_for_check = paragraph.strip() # Para checar marcadores
        is_break_marker = full_line_break_regex.match(paragraph_clean_for_check) is not None

        # Se for um marcador de quebra e já houver conteúdo no chunk atual,
        # finaliza o chunk atual ANTES de adicionar o parágrafo marcador.
        if is_break_marker and current_chunk.strip():
            chunks.append(current_chunk.strip())
            logger.debug(f"Chunk {len(chunks)} salvo (antes do marcador de quebra '{paragraph_clean_for_check[:30]}...'). Tokens (aprox.): {current_chunk_tokens}")
            current_chunk = paragraph + "\n\n" # Adiciona marcador ao novo chunk
            current_chunk_tokens = tokens_with_separator
        else:
            # Adiciona o parágrafo (ou o marcador, se for o primeiro item do chunk)
            # Garante que haja um espaço ou quebra de linha entre o conteúdo anterior e o novo parágrafo
            separator = "\n\n" if current_chunk else ""
            current_chunk += separator + paragraph
            current_chunk_tokens += tokens_with_separator

    # Adiciona o último chunk se não estiver vazio
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
        logger.debug(f"Chunk final {len(chunks)} salvo. Tokens (aprox.): {current_chunk_tokens}")

    logger.info(f"✅ Chunking concluído. Criados {len(chunks)} chunks.")
    return chunks


# ============================================================
# FUNÇÃO format_with_ai MODIFICADA (COMO MOSTRADO NA SEÇÃO 1)
# ============================================================
def format_with_ai(model, chunk, is_first_chunk=False):
    """
    Processa um chunk de texto com a API Gemini para correção e formatação,
    com foco especial na correção de erros comuns de OCR em português.
    Retorna o texto formatado ou None se ocorrer um erro irrecuperável.
    """
    context_start = "Você está formatando o início de um livro." if is_first_chunk else "Você está continuando a formatação de um texto de livro existente."

    # Prompt APRIMORADO com instruções específicas para erros de OCR
    chunk_prompt = f"""
    {context_start} Você é um editor literário proficiente em português do Brasil. Sua tarefa é corrigir e formatar o fragmento de texto a seguir como parte de um livro.

    **IMPORTANTE:** Este fragmento de texto provavelmente foi extraído usando OCR (Reconhecimento Óptico de Caracteres) de um documento PDF e pode conter erros específicos desse processo, além de possíveis erros gerais de digitação ou gramática.

    Siga RIGOROSAMENTE estas regras:

    1.  **Correção Geral:** Corrija erros gramaticais, ortográficos, de pontuação e concordância verbal/nominal. Use o português do Brasil como padrão.
    2.  **Estilo:** Mantenha um estilo literário fluido, claro e envolvente, consistente com o tom do texto original. Não altere o significado ou a intenção do autor.
    3.  **Fidelidade:** Mantenha-se o mais fiel possível ao texto original. NÃO omita parágrafos, frases ou informações importantes. NÃO adicione conteúdo, introduções, resumos ou conclusões que não estavam no fragmento original.
    4.  **Marcadores de Capítulo/Quebra:** Se houver marcadores de capítulo ou quebras de seção (como 'Capítulo X', '***', etc.) no início de um parágrafo, MANTENHA-OS EXATAMENTE como estão nesse parágrafo. NÃO adicione títulos ou marcadores que não existam no texto original.
    5.  **Quebra de Página:** Se o marcador '{PAGE_BREAK_MARKER}' aparecer no texto, MANTENHA-O EXATAMENTE nessa posição, em sua própria linha, sem alterações.

    6.  **Prioridades Adicionais de Correção (Erros Comuns de OCR em Português):** Preste atenção especial e corrija os seguintes tipos de erros frequentemente encontrados em textos OCRizados do português:
        * **Caracteres Inválidos:** Remova caracteres que claramente não pertencem ao alfabeto português ou à pontuação padrão (ex: '■', '`', '^' isolados, símbolos estranhos como '§', '¢', etc.), a menos que façam parte de uma palavra estrangeira ou notação intencional.
        * **Acentuação Deturpada:** Corrija caracteres como "a~", "a^", "a´", "a`" para "á", "â", "ã", etc., e o mesmo para outras vogais (e, i, o, u). Restaure acentos faltantes ou incorretos.
        * **Cedilha Incorreta:** Restaure 'ç' quando aparecer como 'c,', 'c;', 'c.', 'g', ou outros caracteres incorretos.
        * **Diacríticos Separados:** Junte letras e seus acentos/til quando separados (ex: "´e" para "é", "~a" para "ã", "^o" para "ô").
        * **Caracteres Específicos Deformados:** Corrija "ã", "õ", "á", "à", "â", "ê", "ô", "é", "í", "ó", "ú", "ü" (raro em PT-BR, mas pode aparecer) quando estiverem malformados.
        * **Confusão de Caracteres:** Diferencie e corrija caracteres visualmente semelhantes que o OCR confunde (ex: "rn" -> "m", "m" -> "rn", "cl" -> "d", "d" -> "cl", "vv" -> "w", "w" -> "vv", "1" -> "l", "l" -> "1", "O" -> "0", "0" -> "O"). Use o contexto da palavra.
        * **Conjugações Verbais:** Restaure terminações verbais comuns (-ção, -ões, -aram, -eram, -iam, -ava, -asse, etc.) se estiverem deturpadas.
        * **Plurais e Concordâncias:** Corrija marcadores de plural (-s, -es, -is, -eis, -ões) e garanta a concordância nominal e verbal.
        * **Artigos e Preposições:** Restaure artigos e preposições (da, do, na, no, em, para, com, por, etc.) se estiverem como símbolos ou fragmentados.
        * **Hifenização:** Corrija ou adicione hífens corretamente em palavras compostas, pronomes oblíquos (ex: "disse-lhe"), e conforme a norma culta. Cuidado com hífens que resultam de quebras de linha do OCR.
        * **Aspas e Diálogos:** Padronize o uso de aspas (geralmente "" ou «») e travessões (—) para diálogos, corrigindo usos incorretos ou caracteres substitutos.
        * **Pontuação:** Corrija o uso de vírgulas, pontos, ponto-e-vírgula, dois-pontos, reticências, parênteses, etc., especialmente se parecerem erros de OCR (ex: '.' em vez de ',', ';' onde não cabe).
        * **Estrangeirismos vs. Erros:** Preserve palavras estrangeiras que fazem sentido no contexto, mas corrija sequências que parecem ser erros de OCR.
        * **Abreviaturas:** Mantenha ou expanda abreviaturas comuns (Sr., Dr., Prof., etc.) de forma consistente, se o contexto permitir.
        * **Pronomes de Tratamento:** Corrija formas como "V. Exa.", "Sr.", "D.", etc., se estiverem malformadas.
        * **Numerais:** Corrija números escritos por extenso ou algarismos se estiverem deturpados pelo OCR.
        * **Grafia Antiga:** Se identificar claramente grafia antiga (ex: "pharmacia"), modernize-a para a ortografia atual do português do Brasil (ex: "farmácia"), a menos que o estilo do livro sugira manter a forma arcaica.
        * **Contrações:** Restaure contrações (dele, neste, àquela, etc.) se estiverem separadas ou malformadas (ex: "a aquela").
        * **Prefixos e Sufixos:** Corrija prefixos (pré-, pós-, sub-) e sufixos (-mente, -ável, -ismo) se estiverem incorretos ou separados.
        * **Siglas e Acrônimos:** Mantenha siglas e acrônimos, corrigindo se parecerem erros de OCR.
        * **Expressões Idiomáticas:** Preserve e corrija expressões idiomáticas que possam ter sido quebradas pelo OCR.

    7.  **Formato de Saída:** O resultado deve ser APENAS o texto formatado e corrigido do fragmento. Use parágrafos separados por duas quebras de linha (`\\n\\n`). NÃO use formatação Markdown (como `*`, `#`, `_`) ou qualquer outra marcação. Retorne texto puro.

    Texto do fragmento para processar (pode conter erros de OCR):
    \"\"\"
    {chunk}
    \"\"\"
    """

    # Contagem de tokens de entrada (opcional, mas bom para depuração)
    # try:
    #     input_token_count = model.count_tokens(chunk_prompt).total_tokens
    #     logger.debug(f"Chamando API para chunk (Primeiro: {is_first_chunk}). Tokens de entrada estimados para o prompt: {input_token_count}")
    # except Exception as e:
    #     logger.warning(f"Não foi possível contar tokens de entrada para o prompt: {e}")

    logger.debug(f"Enviando chunk (Primeiro: {is_first_chunk}) para a API Gemini com instruções de OCR...")

    max_retries = 5
    for attempt in range(max_retries):
        try:
            # Usa o método generate_content para Gemini
            response = model.generate_content(
                chunk_prompt,
                generation_config=genai.GenerationConfig(
                    temperature=TEMPERATURE,
                    max_output_tokens=MAX_OUTPUT_TOKENS
                ),
                # safety_settings=[...] # Ajuste se necessário
            )

            # Verifica se há conteúdo bloqueado
            if not response.candidates:
                block_reason = "Não especificado"
                # Tenta obter detalhes do bloqueio (pode variar com a API)
                if hasattr(response, 'prompt_feedback') and response.prompt_feedback and hasattr(response.prompt_feedback, 'block_reason'):
                     block_reason = response.prompt_feedback.block_reason.name # Ou .name, dependendo da estrutura
                logger.error(f"API bloqueou o prompt no chunk (Tentativa {attempt + 1}/{max_retries}). Razão: {block_reason}")
                logger.error(f"Conteúdo do chunk problemático (primeiros 500 chars): {chunk[:500]}")
                return None # Falha irrecuperável se bloqueado

            # Acessa o conteúdo de texto da resposta
            # Verifica se 'text' existe no candidato
            if hasattr(response.candidates[0].content.parts[0], 'text'):
                formatted_text = response.text.strip()
            else:
                logger.warning(f"Resposta da API para chunk {attempt+1} não continha 'text'. Resposta: {response.candidates[0].content.parts[0]}")
                formatted_text = "" # Ou trata como erro

            # Contagem de tokens de saída (opcional)
            # try:
            #     output_token_count = model.count_tokens(formatted_text).total_tokens
            #     logger.debug(f"Chunk processado com sucesso. Tokens de saída: {output_token_count}")
            # except Exception as e:
            #     logger.warning(f"Não foi possível contar tokens de saída: {e}")
            #     logger.debug(f"Chunk processado com sucesso.")

            logger.debug(f"Chunk processado com sucesso pela API.")
            return formatted_text

        except Exception as e:
            logger.warning(f"Erro na API ao processar chunk (Tentativa {attempt + 1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                wait_time = (2 ** attempt) + (os.urandom(1)[0] / 255.0) # Exponential backoff com jitter
                logger.info(f"Tentando novamente em {wait_time:.2f} segundos...")
                time.sleep(wait_time)
            else:
                logger.error(f"Falha ao processar chunk após {max_retries} tentativas.")
                logger.error(f"Conteúdo do chunk com falha (primeiros 500 chars): {chunk[:500]}...")
                return None # Retorna None para indicar falha definitiva


def apply_formatting(doc, formatted_text, normal_style_name, chapter_style_name):
    """
    Aplica formatação ao documento Word usando estilos (com fallback para formatação direta).
    """
    if not formatted_text:
        logger.warning("Texto formatado vazio recebido para um chunk. Pulando inserção.")
        return

    # Tenta obter os estilos do documento
    normal_style = None
    chapter_style = None
    try:
        normal_style = doc.styles[normal_style_name]
        if not isinstance(normal_style, _ParagraphStyle):
             logger.warning(f"Estilo '{normal_style_name}' encontrado, mas não é um estilo de parágrafo. Usando fallback.")
             normal_style = None # Força fallback
    except KeyError:
        logger.warning(f"Estilo '{normal_style_name}' não encontrado no documento. Usando fallback de formatação direta para texto normal.")

    try:
        chapter_style = doc.styles[chapter_style_name]
        if not isinstance(chapter_style, _ParagraphStyle):
             logger.warning(f"Estilo '{chapter_style_name}' encontrado, mas não é um estilo de parágrafo. Usando fallback.")
             chapter_style = None # Força fallback
    except KeyError:
        logger.warning(f"Estilo '{chapter_style_name}' não encontrado no documento. Usando fallback de formatação direta para títulos de capítulo.")


    # Compila regex de capítulo aqui também para identificar títulos
    chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS), re.IGNORECASE)

    # Divide o texto pelo marcador de quebra de página primeiro
    parts = formatted_text.split(PAGE_BREAK_MARKER)

    # Flag para controlar se já adicionamos algum conteúdo (para lógica de quebra de página)
    # Verifica se o documento já tem parágrafos com texto.
    content_added_in_this_run = any(p.text.strip() for p in doc.paragraphs)

    for part_index, part in enumerate(parts):
        part_clean = part.strip()
        if not part_clean:
            # Se for uma parte vazia e não for a primeira parte, significa que era um
            # marcador de quebra de página explícito. Adiciona a quebra.
            if part_index > 0 or content_added_in_this_run:
                doc.add_page_break()
                logger.debug("Quebra de página explícita (de marcador) adicionada.")
            continue # Pula para a próxima parte

        # Adiciona quebra de página ANTES da nova parte com conteúdo,
        # EXCETO se for a primeiríssima parte a ser adicionada no documento.
        if part_index > 0 or content_added_in_this_run:
             doc.add_page_break()
             logger.debug("Quebra de página (antes da parte de texto) adicionada.")

        paragraphs_in_part = part_clean.split("\n\n")
        for paragraph_text in paragraphs_in_part:
            paragraph_text_clean = paragraph_text.strip()
            if not paragraph_text_clean:
                continue # Pula parágrafos vazios resultantes do split \n\n

            # Verifica se é um título de capítulo
            is_chapter = chapter_regex.match(paragraph_text_clean) is not None

            p = doc.add_paragraph()
            # Adiciona o texto como uma única 'run' inicialmente
            run = p.add_run(paragraph_text_clean)
            content_added_in_this_run = True # Marca que adicionamos conteúdo

            # Aplica estilo ou formatação direta
            if is_chapter:
                if chapter_style:
                    p.style = chapter_style
                    logger.debug(f"Aplicado estilo '{chapter_style.name}' ao título: '{paragraph_text_clean[:50]}...'")
                else:
                    # Fallback: Formatação direta para título
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.size = Pt(14)
                    run.bold = True
                    # Poderia adicionar mais formatação fallback aqui (espaçamento, etc.)
                    logger.debug(f"Aplicada formatação direta (fallback) ao título: '{paragraph_text_clean[:50]}...'")
            else:
                if normal_style:
                    p.style = normal_style
                    # Garante justificação se o estilo não o fizer (ou se quiser forçar)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    logger.debug(f"Aplicado estilo '{normal_style.name}' ao texto: '{paragraph_text_clean[:50]}...'")
                else:
                    # Fallback: Formatação direta para texto normal
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run.font.size = Pt(12) # Tamanho padrão fallback
                    run.bold = False      # Garante que não está em negrito
                    # run.font.name = 'Times New Roman' # Exemplo de fonte fallback
                    logger.debug(f"Aplicada formatação direta (fallback) ao texto: '{paragraph_text_clean[:50]}...'")


def main():
    """Função principal para orquestrar o processo."""
    logger.info("========================================================")
    logger.info(f"Iniciando processamento do livro (com foco em OCR) com Gemini ({MODEL_NAME})")
    logger.info("========================================================")

    start_time = time.time()

    # Define backup filename here, after OUTPUT_DOCX is set
    backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
    BACKUP_DOCX = f"backup_{os.path.splitext(OUTPUT_DOCX)[0]}_{backup_timestamp}.docx"

    # === PASSO 1 – Lê o texto bruto ===
    try:
        with open(INPUT_TXT, "r", encoding="utf-8") as f:
            texto_original = f.read()
        logger.info(f"Arquivo de entrada '{INPUT_TXT}' carregado ({len(texto_original)} caracteres).")
    except FileNotFoundError:
        logger.error(f"Erro Fatal: Arquivo de entrada '{INPUT_TXT}' não encontrado.")
        return
    except Exception as e:
        logger.error(f"Erro Fatal ao ler o arquivo '{INPUT_TXT}': {e}")
        return

    # === PASSO 2 – Divide o texto em chunks ===
    logger.info(f"Dividindo o texto em chunks (máx. {MAX_CHUNK_TOKENS} tokens aprox.)...")
    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS)
    if not text_chunks:
        logger.error("Nenhum chunk de texto foi gerado. Verifique o arquivo de entrada.")
        return
    logger.info(f"Texto dividido em {len(text_chunks)} chunks.")

    # === PASSO 3 – Carrega o template ou cria novo documento ===
    doc = None
    try:
        # Cria backup do arquivo de saída anterior, se existir, ANTES de carregar/limpar
        if os.path.exists(OUTPUT_DOCX):
            try:
                shutil.copy2(OUTPUT_DOCX, BACKUP_DOCX) # copy2 preserva metadados
                logger.info(f"Backup do arquivo anterior criado: {BACKUP_DOCX}")
            except Exception as e:
                logger.warning(f"Não foi possível criar backup de '{OUTPUT_DOCX}': {e}")

        doc = Document(TEMPLATE_DOCX)
        logger.info(f"Template '{TEMPLATE_DOCX}' carregado.")

        # === PASSO 4 – Limpa o corpo do documento (opcional, mas recomendado com template) ===
        if hasattr(doc, '_body') and doc._body is not None:
            # Iterar e remover parágrafos e tabelas explicitamente é mais seguro
            # do que clear_content() que pode ser menos robusto em algumas versões/casos.
            for para in reversed(doc.paragraphs): # Iterar ao contrário para remoção segura
                p_element = para._element
                p_element.getparent().remove(p_element)
            for table in reversed(doc.tables):
                 t_element = table._element
                 t_element.getparent().remove(t_element)
            logger.info("Conteúdo principal do template limpo (parágrafos e tabelas).")
        else:
            logger.warning("Não foi possível acessar/limpar o corpo do documento do template de forma robusta.")

    except FileNotFoundError:
        logger.warning(f"Template '{TEMPLATE_DOCX}' não encontrado.")
        logger.info("Criando um novo documento Word com configurações padrão.")
        doc = Document()
        # Define configurações básicas de página para um novo documento
        try:
            section = doc.sections[0]
            # Exemplo A5 (aproximado)
            section.page_height = Inches(8.27) # Altura A5
            section.page_width = Inches(5.83)  # Largura A5
            # Margens comuns para livros (ajuste conforme necessário)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.6)
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            logger.info("Configurações de página (A5 aprox.) e margens aplicadas ao novo documento.")
            # Nota: Os estilos NORMAL_STYLE_NAME e CHAPTER_STYLE_NAME podem não existir.
        except Exception as e:
             logger.warning(f"Não foi possível aplicar configurações de página ao novo documento: {e}")
    except Exception as e:
        logger.error(f"Erro ao carregar ou processar o template '{TEMPLATE_DOCX}': {e}")
        return

    # === PASSO 5 – Processa e insere cada chunk ===
    logger.info(f"Iniciando processamento de {len(text_chunks)} chunks com a API Gemini...")
    processed_chunks_count = 0
    failed_chunks_count = 0

    # Barra de progresso
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc="Processando Chunks", unit="chunk")

    for i, chunk in progress_bar:
        chunk_start_time = time.time()
        progress_bar.set_description(f"Processando Chunk {i+1}/{len(text_chunks)}")

        # Processa o chunk com a IA (usando a função atualizada)
        formatted_chunk = format_with_ai(gemini_model, chunk, is_first_chunk=(i == 0))

        if formatted_chunk is not None:
            # Aplica formatação e adiciona ao documento
            apply_formatting(doc, formatted_chunk, NORMAL_STYLE_NAME, CHAPTER_STYLE_NAME)
            processed_chunks_count += 1
            logger.debug(f"Chunk {i+1} processado e adicionado em {time.time() - chunk_start_time:.2f}s.")
        else:
            failed_chunks_count += 1
            logger.error(f"Chunk {i + 1} falhou no processamento pela API e foi pulado.")
            # OPCIONAL: Adicionar o chunk original ou um placeholder em caso de falha
            # placeholder_text = f"[ERRO PROCESSANDO CHUNK {i+1} - CONTEÚDO ORIGINAL ABAIXO]\n\n{chunk}"
            # apply_formatting(doc, placeholder_text, NORMAL_STYLE_NAME, CHAPTER_STYLE_NAME)
            # logger.warning(f"Chunk {i+1} original (ou placeholder) inserido no documento devido à falha.")


        # === PASSO 5.1 – Salva progresso periodicamente ===
        # Salva a cada 5 chunks ou no último chunk (mais frequente pode ser útil)
        if (i + 1) % 5 == 0 or (i + 1) == len(text_chunks):
            temp_save_path = f"{OUTPUT_DOCX}.temp"
            try:
                doc.save(temp_save_path)
                # Move atomicamente para o arquivo final (sobrescreve)
                shutil.move(temp_save_path, OUTPUT_DOCX)
                logger.info(f"Progresso salvo ({i + 1}/{len(text_chunks)} chunks processados). Arquivo: {OUTPUT_DOCX}")
            except Exception as e:
                logger.error(f"Erro ao salvar progresso parcial em '{OUTPUT_DOCX}': {e}")
                if os.path.exists(temp_save_path):
                    logger.warning(f"Arquivo temporário de salvamento parcial '{temp_save_path}' mantido.")

    # === PASSO 6 – Conclusão e Salvamento Final ===
    # O salvamento periódico já cuidou do último passo, mas verificamos o .temp
    final_temp_path = f"{OUTPUT_DOCX}.temp"
    if os.path.exists(final_temp_path):
        try:
            # Tentativa final de mover se o último save falhou na movimentação
            shutil.move(final_temp_path, OUTPUT_DOCX)
            logger.info(f"Arquivo temporário final movido para '{OUTPUT_DOCX}'.")
        except Exception as e:
            logger.error(f"Erro ao mover arquivo temporário final para '{OUTPUT_DOCX}': {e}. O arquivo pode estar em '{final_temp_path}'.")

    end_time = time.time()
    total_time = end_time - start_time

    logger.info("========================================================")
    logger.info("✅ Processamento Concluído!")
    logger.info(f"Tempo total: {total_time:.2f} segundos ({total_time/60:.2f} minutos).")
    logger.info(f"Chunks processados com sucesso: {processed_chunks_count}")
    logger.info(f"Chunks com falha (pulados ou com placeholder): {failed_chunks_count}")
    logger.info(f"Livro final gerado: {OUTPUT_DOCX}")
    if os.path.exists(BACKUP_DOCX):
        logger.info(f"Backup do arquivo anterior (se existia): {BACKUP_DOCX}")
    logger.info(f"Log detalhado disponível em: {log_filepath}")
    logger.info("========================================================")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("Processo interrompido pelo usuário (Ctrl+C). O último progresso salvo pode estar disponível.")
    except Exception as e:
        # Captura qualquer outra exceção não tratada no main()
        logger.exception(f"Erro fatal inesperado durante a execução: {e}")
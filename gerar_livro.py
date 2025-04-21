import openai
from docx import Document
from dotenv import load_dotenv
import os

# === CARREGA VARIÁVEIS DE AMBIENTE DO .env ===
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")

# === CONFIGURAÇÕES ===
INPUT_TXT = "rascunho.txt"
TEMPLATE_DOCX = "Estrutura.docx"
OUTPUT_DOCX = "Livro_Final_Formatado.docx"

# === SETUP OPENAI (Updated for openai>=1.0.0) ===
# Create a client object
client = openai.OpenAI(api_key=API_KEY)

# === PASSO 1 – Lê o texto bruto ===
try:
    with open(INPUT_TXT, "r", encoding="utf-8") as f:
        texto_original = f.read()
except FileNotFoundError:
    print(f"Erro: O arquivo de rascunho '{INPUT_TXT}' não foi encontrado.")
    exit()


# === PASSO 2 – Prompt para formatar como livro ===
prompt = f"""
Você é um editor literário. Corrija e formate o texto a seguir como um livro pronto para publicação.

Regras:
- Corrija erros gramaticais, ortográficos e de concordância
- Utilize títulos centralizados para capítulos
- Use parágrafos justificados com espaçamento apropriado
- Separe capítulos com quebras de página (escreva como: ===QUEBRA_DE_PAGINA===)
- Mantenha estilo literário fluido e bem escrito

Texto:
\"\"\"{texto_original}\"\"\"
"""

print("🔁 Enviando para ChatGPT...")
try:
    # Use the client object for chat completions (Updated syntax)
    response = client.chat.completions.create(
        model="gpt-4-turbo", # You might use "gpt-4-turbo-preview" or other available models
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=4096,
)

    # Access the content from the response object (Updated syntax)
    texto_formatado = response.choices[0].message.content
    print("✅ Texto formatado recebido.")

except Exception as e:
    print(f"Ocorreu um erro ao chamar a API da OpenAI: {e}")
    exit()


# === PASSO 3 – Carrega o template ===
try:
    doc = Document(TEMPLATE_DOCX)
except FileNotFoundError:
    print(f"Erro: O arquivo de template '{TEMPLATE_DOCX}' não foi encontrado.")
    exit()


# === PASSO 4 – Limpa corpo do documento (opcional) ===
# Check if the body is not None before clearing
if doc._body:
    doc._body.clear_content()
else:
     # If body is None, just add sections if needed or handle as appropriate
     # For simple documents, _body will likely not be None
    pass # Or add a section if needed


# === PASSO 5 – Insere texto formatado no corpo ===
# Process text line by line or paragraph by paragraph based on your preferred splitting
# Splitting by "\n\n" assumes paragraphs are separated by two newlines
paragraphs = texto_formatado.split("\n\n")

for i, par_block in enumerate(paragraphs):
    par_block = par_block.strip()
    if not par_block:
        continue # Skip empty blocks

    # Check for the page break marker within the block
    # It's safer to add the page break before the content of the new chapter
    if "===QUEBRA_DE_PAGINA===" in par_block:
        # Add a page break. If this is the very first block, don't add a break before content.
        # We can add the page break *before* the content of the next chapter starts.
        # Let's assume the marker is typically *at the beginning* of a block indicating the next content is a new page.
        # Or, it might be a block *by itself*.
        # A robust way is to split by the marker and process parts.
        parts = par_block.split("===QUEBRA_DE_PAGINA===")
        for part in parts:
             part = part.strip()
             if part: # Add any text that was around the marker
                 p = doc.add_paragraph(part)
                 # You might need to apply specific styles here for titles etc.
                 # This basic example just adds paragraphs
             # If the marker was found, add a page break *after* any content from the previous block/part
             doc.add_page_break()

        # If the block *was* just the marker, the loop will add a page break.
        # If it had text around the marker, it will add the text, then the break.
        # We need to be careful not to add double page breaks if the marker block is followed by another page break block.
        # A better approach might be to process the *entire* text_formatado and split by the marker *first*.

# --- Alternative and potentially better way to process with page breaks ---
print("📝 Inserindo texto formatado no documento...")
full_text_parts = texto_formatado.split("===QUEBRA_DE_PAGINA===")

# Add content before the first possible page break
if full_text_parts:
    # Process the first part (content before the first marker)
    paragraphs_first_part = full_text_parts[0].strip().split("\n\n")
    for par in paragraphs_first_part:
        if par.strip():
            p = doc.add_paragraph(par.strip())
            # You'll need logic here to identify and center titles
            # Example: If the paragraph starts with "Capítulo ", center it
            if par.strip().startswith("Capítulo ") or par.strip().startswith("CHAPTER "):
                 p.alignment = 1 # 1 for center alignment


    # Process the remaining parts, each starting with a page break
    for i in range(1, len(full_text_parts)):
        doc.add_page_break() # Add page break before the new chapter's content
        chapter_content = full_text_parts[i].strip()
        if chapter_content:
            paragraphs_chapter = chapter_content.split("\n\n")
            for par in paragraphs_chapter:
                 if par.strip():
                     p = doc.add_paragraph(par.strip())
                     # Apply styling, e.g., centering for titles
                     if par.strip().startswith("Capítulo ") or par.strip().startswith("CHAPTER "):
                          p.alignment = 1 # 1 for center alignment


# === PASSO 6 – Salva o novo documento ===
doc.save(OUTPUT_DOCX)
print(f"📘 Livro gerado com sucesso: {OUTPUT_DOCX}")
# --- Using Google's Gemini API (gemini-1.5-pro) ---

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.styles.style import _ParagraphStyle # Para checagem de tipo
from docx.shared import RGBColor

from dotenv import load_dotenv
import os
import re
import logging
from tqdm import tqdm
import time
import shutil

# Import the Google Generative AI library
import google.generativeai as genai

# === SETUP LOGGING ===
log_dir = "logs"
if not os.path.exists(log_dir): os.makedirs(log_dir)
log_filepath = os.path.join(log_dir, "book_processor_ocr_pro.log")
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(module)s - %(funcName)s - %(message)s',
    handlers=[ logging.FileHandler(log_filepath, encoding='utf-8'), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# === CARREGA VARIÁVEIS DE AMBIENTE ===
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# === CONFIGURAÇÕES ===
INPUT_TXT = "rascunho.txt"
TEMPLATE_DOCX = "Estrutura.docx"
OUTPUT_DOCX = "Livro_Final_Formatado_Gemini_OCR_Corrigido_Pro_v1.docx"
MODEL_NAME = "gemini-1.5-pro"
# !!!!! IMPORTANTE: AJUSTE CONFORME NECESSÁRIO APÓS TESTAR A NOVA LÓGICA !!!!!
# Pode tentar aumentar um pouco (ex: 4000) se a subdivisão funcionar bem,
# mas mantenha um valor seguro abaixo de MAX_OUTPUT_TOKENS.
MAX_CHUNK_TOKENS = 1000 # Mantenha baixo inicialmente para forçar subdivisão
# !!!!! ------------------------------------------------------------- !!!!!
MAX_OUTPUT_TOKENS = 8192 # Limite da API para a resposta
TEMPERATURE = 0.6
NORMAL_STYLE_NAME = "Normal"
CHAPTER_STYLE_NAME = "Heading 1" # Ou o nome real do seu estilo de capítulo
CHAPTER_PATTERNS = [
    r'^\s*Capítulo \w+', r'^\s*CAPÍTULO \w+', r'^\s*Capítulo \d+',
    r'^\s*CHAPTER \w+', r'^\s*Chapter \d+', r'^\s*LIVRO \w+', r'^\s*PARTE \w+',
]
OTHER_BREAK_PATTERNS = [r'^\s*\*\*\*\s*$', r'^\s*---+\s*$']
PAGE_BREAK_MARKER = "===QUEBRA_DE_PAGINA==="
AI_FAILURE_MARKER = "*** FALHA NA IA - TEXTO ORIGINAL ABAIXO ***"
FORMATTING_ERROR_MARKER = "*** ERRO DE FORMATAÇÃO - TEXTO ORIGINAL ABAIXO ***"

# --- Validação API Key ---
if not GOOGLE_API_KEY: logger.error("GOOGLE_API_KEY não encontrada."); exit(1)

# --- Setup Gemini Client ---
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    gemini_model = genai.GenerativeModel(MODEL_NAME)
    logger.info(f"Modelo Gemini '{MODEL_NAME}' inicializado.")
except Exception as e: logger.error(f"Falha ao inicializar modelo Gemini ({MODEL_NAME}): {e}"); exit(1)

# --- Funções Auxiliares ---

def count_tokens_approx(text):
    """Estima a contagem de tokens (aproximadamente 4 caracteres por token)."""
    if not text: return 0
    # Uma estimativa muito básica. Para maior precisão, use a API count_tokens do Gemini.
    return len(text) // 4

# --- FUNÇÃO create_chunks ATUALIZADA ---
def create_chunks(text, max_tokens):
    """Divide o texto em chunks, subdividindo parágrafos grandes."""
    logger.info(f"Iniciando criação de chunks. Máx tokens (aprox): {max_tokens}")
    chunks = []
    current_chunk = ""
    current_chunk_tokens = 0

    # Inicialmente, divide por blocos maiores (parágrafos separados por linha dupla)
    # Remove espaços em branco extras de cada bloco e ignora blocos vazios
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    logger.info(f"Texto dividido inicialmente em {len(paragraphs)} blocos não vazios ('\\n\\n').")

    for i, paragraph_text in enumerate(paragraphs):
        paragraph_tokens = count_tokens_approx(paragraph_text)
        # Adiciona tokens para a separação \n\n que será reintroduzida ao juntar
        tokens_with_separator = paragraph_tokens + (count_tokens_approx("\n\n") if current_chunk else 0)

        # --- LÓGICA DE COMBINAÇÃO ---
        # Se o chunk atual + o novo parágrafo (com separador) exceder o limite
        if current_chunk_tokens > 0 and (current_chunk_tokens + tokens_with_separator > max_tokens):
            # Salva o chunk atual antes de iniciar um novo
            chunks.append(current_chunk)
            logger.debug(f"Chunk {len(chunks)} salvo (limite atingido ao tentar adicionar Parágrafo {i+1}). Tokens: {current_chunk_tokens}")
            current_chunk = paragraph_text # Começa novo chunk com o parágrafo atual
            current_chunk_tokens = paragraph_tokens
        # Se o chunk atual + o novo parágrafo couberem
        elif current_chunk_tokens + tokens_with_separator <= max_tokens:
             # Adiciona o parágrafo ao chunk atual (com separador se não for o primeiro)
            separator = "\n\n" if current_chunk else ""
            current_chunk += separator + paragraph_text
            current_chunk_tokens += tokens_with_separator # Atualiza contagem

        # --- LÓGICA DE SUBDIVISÃO ---
        # Se o parágrafo ATUAL SOZINHO já excede o limite (mesmo começando um chunk novo)
        # Esta condição é verificada DEPOIS da tentativa de adição,
        # garantindo que parágrafos grandes sejam tratados mesmo se forem os primeiros
        # de um chunk.
        if paragraph_tokens > max_tokens:
            logger.warning(f"Parágrafo {i+1} ({paragraph_tokens} tk) excede limite {max_tokens}. Iniciando SUBDIVISÃO.")

            # Se havia algo no 'current_chunk' ANTES deste parágrafo grande, salva.
            # Isso pode acontecer se o parágrafo anterior coube, mas este não.
            # Nota: A lógica anterior já pode ter iniciado o current_chunk com este parágrafo.
            # Precisamos verificar se o current_chunk contém APENAS este parágrafo grande
            # ou se continha algo antes. Se continha algo antes, o current_chunk
            # (sem este parágrafo) precisa ser salvo.

            # Se current_chunk NÃO é IGUAL ao parágrafo problemático, significa
            # que havia conteúdo anterior que precisa ser salvo.
            if current_chunk != paragraph_text and current_chunk.strip():
                 # Remove o parágrafo grande que foi adicionado na etapa anterior
                 # (se foi adicionado) para salvar o que veio antes.
                 if current_chunk.endswith("\n\n" + paragraph_text):
                     chunk_to_save = current_chunk[:-len("\n\n" + paragraph_text)]
                     chunks.append(chunk_to_save)
                     logger.debug(f"Chunk {len(chunks)} salvo (antes do parág. longo subdividido). Tokens: {count_tokens_approx(chunk_to_save)}")
                 elif current_chunk == paragraph_text:
                      # Se o current_chunk É o parágrafo grande, não há nada antes para salvar.
                      pass
                 else:
                      # Caso inesperado, logar. Pode indicar erro na lógica.
                      logger.warning(f"Lógica de salvamento pré-subdivisão encontrou estado inesperado. Current Chunk: '{current_chunk[:50]}...', Parágrafo: '{paragraph_text[:50]}...'")


            # --- Início da Subdivisão do Parágrafo Grande ---
            sub_chunks_added_count = 0
            # Tenta dividir por linhas '\n' dentro do parágrafo grande.
            # Filtra linhas vazias que podem existir.
            lines = [line for line in paragraph_text.split('\n') if line.strip()]
            current_sub_chunk = ""
            current_sub_chunk_tokens = 0

            for line_num, line in enumerate(lines):
                line_tokens = count_tokens_approx(line)
                # Tokens com separador '\n' (exceto para a primeira linha do sub-chunk)
                tokens_with_line_separator = line_tokens + (count_tokens_approx("\n") if current_sub_chunk else 0)

                # Se adicionar a linha estourar o limite do sub-chunk atual
                if current_sub_chunk_tokens > 0 and (current_sub_chunk_tokens + tokens_with_line_separator > max_tokens):
                    chunks.append(current_sub_chunk) # Salva o sub-chunk completo
                    sub_chunks_added_count += 1
                    logger.debug(f"Sub-chunk {len(chunks)} salvo (parág. longo {i+1}). Tokens: {current_sub_chunk_tokens}")
                    current_sub_chunk = line # Começa novo sub-chunk com a linha atual
                    current_sub_chunk_tokens = line_tokens
                # Se a linha SOZINHA estoura o limite (caso extremo)
                elif line_tokens > max_tokens:
                     # Salva o sub-chunk anterior se houver
                    if current_sub_chunk:
                        chunks.append(current_sub_chunk)
                        sub_chunks_added_count += 1
                        logger.debug(f"Sub-chunk {len(chunks)} salvo (antes linha longa, parág. {i+1}). Tokens: {current_sub_chunk_tokens}")
                    # Adiciona a linha longa como um chunk próprio
                    chunks.append(line)
                    sub_chunks_added_count += 1
                    logger.warning(f"  -> Linha {line_num+1} dentro do parág. {i+1} ({line_tokens} tk) excede limite {max_tokens}. Adicionando como sub-chunk único.")
                    current_sub_chunk = "" # Reseta para próximo sub-chunk
                    current_sub_chunk_tokens = 0
                # Se a linha cabe no sub-chunk atual
                else:
                    line_separator = "\n" if current_sub_chunk else ""
                    current_sub_chunk += line_separator + line
                    current_sub_chunk_tokens = count_tokens_approx(current_sub_chunk) # Recalcula tokens

            # Salva o último sub-chunk restante da divisão por linhas
            if current_sub_chunk:
                chunks.append(current_sub_chunk)
                sub_chunks_added_count += 1
                logger.debug(f"Último sub-chunk {len(chunks)} salvo (parág. longo {i+1}). Tokens: {current_sub_chunk_tokens}")

            if sub_chunks_added_count == 0:
                 logger.warning(f"Parágrafo {i+1} excedeu limite, mas nenhuma subdivisão por linha foi feita (talvez uma única linha longa?). Adicionando original como chunk.")
                 chunks.append(paragraph_text) # Fallback: adiciona o parágrafo original se a subdivisão falhou

            # --- Fim da Subdivisão ---
            current_chunk = "" # Reseta o chunk principal após processar o parágrafo subdividido
            current_chunk_tokens = 0
            # continue # Não precisa de continue aqui, o loop for principal continuará

    # Salva o último chunk restante que pode não ter atingido o limite
    if current_chunk:
        chunks.append(current_chunk)
        logger.debug(f"Chunk final {len(chunks)} salvo. Tokens: {current_chunk_tokens}")

    logger.info(f"✅ Chunking concluído. {len(chunks)} chunks.")
    return chunks
# --- FIM DA FUNÇÃO create_chunks ATUALIZADA ---


def format_with_ai(model, chunk, is_first_chunk=False):
    # (Função format_with_ai permanece igual à versão anterior que corrigiu o acesso ao texto)
    context_start = "Você está formatando o início de um livro." if is_first_chunk else "Você está continuando a formatação de um texto de livro existente."
    chunk_prompt = f"""
    {context_start} Você é um editor literário proficiente em português do Brasil. Sua tarefa é corrigir e formatar o fragmento de texto a seguir como parte de um livro.

    **IMPORTANTE:** Este fragmento de texto provavelmente foi extraído usando OCR (Reconhecimento Óptico de Caracteres) de um documento PDF e pode conter erros específicos desse processo, além de possíveis erros gerais de digitação ou gramática.

    Siga RIGOROSAMENTE estas regras:

    1.  **Correção Geral:** Corrija erros gramaticais, ortográficos, de pontuação e concordância verbal/nominal. Use o português do Brasil como padrão.
    2.  **Estilo:** Mantenha um estilo literário fluido, claro e envolvente, consistente com o tom do texto original. Não altere o significado ou a intenção do autor.
    3.  **Fidelidade:** Mantenha-se o mais fiel possível ao texto original. NÃO omita parágrafos, frases ou informações importantes. NÃO adicione conteúdo, introduções, resumos ou conclusões que não estavam no fragmento original.
    4.  **Marcadores de Capítulo/Quebra:** Se houver marcadores de capítulo ou quebras de seção (como 'Capítulo X', '***', etc.) no início de um parágrafo, MANTENHA-OS EXATAMENTE como estão nesse parágrafo. NÃO adicione títulos ou marcadores que não existam no texto original.
    5.  **Quebra de Página:** Se o marcador '{PAGE_BREAK_MARKER}' aparecer no texto, MANTENHA-O EXATAMENTE nessa posição, em sua própria linha, sem alterações.

    6.  **Prioridades Adicionais de Correção (Erros Comuns de OCR em Português):** Preste atenção especial e corrija os seguintes tipos de erros frequentemente encontrados em textos OCRizados do português:
        * **Caracteres Inválidos:** Remova caracteres que claramente não pertencem ao alfabeto português ou à pontuação padrão (ex: '■', '`', '^' isolados, símbolos estranhos como '§', '¢', etc.), a menos que façam parte de uma palavra estrangeira ou notação intencional.
        * **Acentuação Deturpada:** Corrija caracteres como "a~", "a^", "a´", "a`" para "á", "â", "ã", etc., e o mesmo para outras vogais (e, i, o, u). Restaure acentos faltantes ou incorretos.
        * **Cedilha Incorreta:** Restaure 'ç' quando aparecer como 'c,', 'c;', 'c.', 'g', ou outros caracteres incorretos.
        * **Diacríticos Separados:** Junte letras e seus acentos/til quando separados (ex: "´e" para "é", "~a" para "ã", "^o" para "ô").
        * **Caracteres Específicos Deformados:** Corrija "ã", "õ", "á", "à", "â", "ê", "ô", "é", "í", "ó", "ú", "ü" (raro em PT-BR, mas pode aparecer) quando estiverem malformados.
        * **Confusão de Caracteres:** Diferencie e corrija caracteres visualmente semelhantes que o OCR confunde (ex: "rn" -> "m", "m" -> "rn", "cl" -> "d", "d" -> "cl", "vv" -> "w", "w" -> "vv", "1" -> "l", "l" -> "1", "O" -> "0", "0" -> "O"). Use o contexto da palavra.
        * **Conjugações Verbais:** Restaure terminações verbais comuns (-ção, -ões, -aram, -eram, -iam, -ava, -asse, etc.) se estiverem deturpadas.
        * **Plurais e Concordâncias:** Corrija marcadores de plural (-s, -es, -is, -eis, -ões) e garanta a concordância nominal e verbal.
        * **Artigos e Preposições:** Restaure artigos e preposições (da, do, na, no, em, para, com, por, etc.) se estiverem como símbolos ou fragmentados.
        * **Hifenização:** Corrija ou adicione hífens corretamente em palavras compostas, pronomes oblíquos (ex: "disse-lhe"), e conforme a norma culta. Cuidado com hífens que resultam de quebras de linha do OCR.
        * **Aspas e Diálogos:** Padronize o uso de aspas (geralmente "" ou «») e travessões (—) para diálogos, corrigindo usos incorretos ou caracteres substitutos.
        * **Pontuação:** Corrija o uso de vírgulas, pontos, ponto-e-vírgula, dois-pontos, reticências, parênteses, etc., especialmente se parecerem erros de OCR (ex: '.' em vez de ',', ';' onde não cabe).
        * **Estrangeirismos vs. Erros:** Preserve palavras estrangeiras que fazem sentido no contexto, mas corrija sequências que parecem ser erros de OCR.
        * **Abreviaturas:** Mantenha ou expanda abreviaturas comuns (Sr., Dr., Prof., etc.) de forma consistente, se o contexto permitir.
        * **Pronomes de Tratamento:** Corrija formas como "V. Exa.", "Sr.", "D.", etc., se estiverem malformadas.
        * **Numerais:** Corrija números escritos por extenso ou algarismos se estiverem deturpados pelo OCR.
        * **Grafia Antiga:** Se identificar claramente grafia antiga (ex: "pharmacia"), modernize-a para a ortografia atual do português do Brasil (ex: "farmácia"), a menos que o estilo do livro sugira manter a forma arcaica.
        * **Contrações:** Restaure contrações (dele, neste, àquela, etc.) se estiverem separadas ou malformadas (ex: "a aquela").
        * **Prefixos e Sufixos:** Corrija prefixos (pré-, pós-, sub-) e sufixos (-mente, -ável, -ismo) se estiverem incorretos ou separados.
        * **Siglas e Acrônimos:** Mantenha siglas e acrônimos, corrigindo se parecerem erros de OCR.
        * **Expressões Idiomáticas:** Preserve e corrija expressões idiomáticas que possam ter sido quebradas pelo OCR.

    7.  **Formato de Saída:** O resultado deve ser APENAS o texto formatado e corrigido do fragmento. Use parágrafos separados por duas quebras de linha (`\\n\\n`). NÃO use formatação Markdown (como `*`, `#`, `_`) ou qualquer outra marcação. Retorne texto puro.

    Texto do fragmento para processar (pode conter erros de OCR):
    \"\"\"
    {chunk}
    \"\"\"
    """
    logger.debug(f"Enviando chunk (Primeiro: {is_first_chunk}) para API ({model.model_name}). Tam Aprox: {count_tokens_approx(chunk)} tk") # Log tamanho
    max_retries = 5
    for attempt in range(max_retries):
        try:
            response = model.generate_content(
                chunk_prompt,
                generation_config=genai.GenerationConfig(
                    temperature=TEMPERATURE,
                    max_output_tokens=MAX_OUTPUT_TOKENS # Garante que o limite de SAÍDA está definido
                ),
                # Adicione safety_settings se necessário para evitar bloqueios,
                # mas pode permitir conteúdo indesejado. EX:
                # safety_settings={
                #     'HATE': 'BLOCK_NONE',
                #     'HARASSMENT': 'BLOCK_NONE',
                #     'SEXUAL' : 'BLOCK_NONE',
                #     'DANGEROUS' : 'BLOCK_NONE'
                # }
            )
            finish_reason = "UNKNOWN"; safety_ratings = "UNKNOWN"; block_reason = "N/A"; formatted_text = ""

            # Verifica se houve bloqueio antes de tentar acessar 'candidates'
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback and hasattr(response.prompt_feedback, 'block_reason'):
                  block_reason = response.prompt_feedback.block_reason.name
                  logger.error(f"API bloqueou prompt (Tentativa {attempt + 1}/{max_retries}). Razão: {block_reason}. Chunk: {chunk[:200]}...")
                  # Não retorna None imediatamente, continua para o próximo retry
            # Verifica se a resposta foi efetivamente vazia ou não tem candidatos
            elif not response.candidates:
                 logger.error(f"API retornou sem candidatos (Tentativa {attempt + 1}/{max_retries}). Resposta: {response}. Chunk: {chunk[:200]}...")
                 # Não retorna None imediatamente, continua para o próximo retry
            else:
                # Se chegou aqui, há candidatos, tenta processar
                try:
                    candidate = response.candidates[0] # Assume o primeiro candidato é o melhor
                    finish_reason = candidate.finish_reason.name if hasattr(candidate.finish_reason, 'name') else "FINISH_REASON_UNKNOWN"
                    safety_ratings = [(r.category.name, r.probability.name) for r in candidate.safety_ratings] if candidate.safety_ratings else "N/A"
                    logger.debug(f"Chunk processado (Tentativa {attempt + 1}). Finish: {finish_reason}. Safety: {safety_ratings}")

                    if finish_reason == "MAX_TOKENS": logger.warning(f"API TRUNCOU resposta (MAX_OUTPUT_TOKENS: {MAX_OUTPUT_TOKENS}). Final pode faltar. Chunk Input Aprox: {count_tokens_approx(chunk)} tk")
                    if finish_reason == "SAFETY": logger.warning(f"API interrompeu resposta (SAFETY). Conteúdo pode estar incompleto.")
                    if finish_reason == "RECITATION": logger.warning(f"API interrompeu resposta (RECITATION).") # Novo finish_reason
                    if finish_reason == "OTHER": logger.warning(f"API interrompeu resposta (OTHER REASON).")

                    # Tenta obter o texto via response.text primeiro (mais comum e direto)
                    if hasattr(response, 'text') and response.text:
                        formatted_text = response.text.strip()
                        logger.debug(f"Texto API (via response.text, 100 chars): '{formatted_text[:100]}...'")
                        return formatted_text # SUCESSO

                    # Fallback: Tenta juntar as partes do conteúdo do candidato (menos comum)
                    elif hasattr(candidate, 'content') and hasattr(candidate.content, 'parts'):
                        logger.debug("response.text não encontrado ou vazio. Tentando juntar partes do candidato.")
                        text_parts = [part.text for part in candidate.content.parts if hasattr(part, 'text')]
                        if text_parts:
                            formatted_text = "".join(text_parts).strip()
                            logger.debug(f"Texto API (via parts, 100 chars): '{formatted_text[:100]}...'")
                            return formatted_text # SUCESSO via parts
                        else:
                            logger.warning(f"Resposta API sem 'text' em response ou nas partes (Tentativa {attempt+1}). Content: {candidate.content}")
                    else:
                         logger.warning(f"Resposta API sem 'text' e sem 'content.parts' utilizáveis (Tentativa {attempt+1}). Candidate: {candidate}")


                except AttributeError as ae:
                    logger.error(f"Erro de Atributo ao acessar resposta API (Tentativa {attempt+1}): {ae} - Resposta: {response}")
                except IndexError:
                     logger.error(f"Erro de Índice: Sem candidatos na resposta (Tentativa {attempt+1}). Resposta: {response}")
                except Exception as e_details:
                    logger.error(f"Erro Genérico ao extrair detalhes/texto API (Tentativa {attempt+1}): {e_details} - Resposta: {response}")

            # Se chegou aqui (erro, bloqueio, sem texto), espera e tenta novamente
            if attempt < max_retries - 1:
                wait_time = (2 ** attempt) + (os.urandom(1)[0] / 255.0) # Backoff exponencial com jitter
                logger.info(f"Tentando novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                 logger.error(f"Falha ao processar chunk após {max_retries} tentativas (vazio/bloqueado/erro extração).")
                 return None # Retorna None após todas as tentativas falharem

        except Exception as e:
            logger.warning(f"Erro chamada API ({model.model_name}) (Tentativa {attempt + 1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                wait_time = (2 ** attempt) + (os.urandom(1)[0] / 255.0)
                logger.info(f"Tentando novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                logger.error(f"Falha ao processar chunk após {max_retries} tentativas (erro chamada API). Chunk: {chunk[:200]}...")
                return None # Retorna None após todas as tentativas falharem

    logger.error(f"Loop de tentativas concluído sem sucesso explícito ou erro capturado.") # Segurança
    return None


def apply_formatting(doc, formatted_text, normal_style_name, chapter_style_name):
    """Aplica formatação ao documento Word usando estilos (com fallback)."""
    if not formatted_text or not formatted_text.strip():
        logger.warning("Texto formatado vazio ou apenas espaços recebido. Pulando inserção.")
        return

    # Cache dos estilos para evitar buscas repetidas
    normal_style = None
    chapter_style = None
    try:
        style_candidate = doc.styles[normal_style_name]
        if isinstance(style_candidate, _ParagraphStyle): normal_style = style_candidate
        else: logger.warning(f"'{normal_style_name}' existe mas NÃO é estilo de parágrafo. Usando fallback.")
    except KeyError: logger.warning(f"Estilo '{normal_style_name}' NÃO encontrado. Usando fallback.")

    try:
        style_candidate = doc.styles[chapter_style_name]
        if isinstance(style_candidate, _ParagraphStyle): chapter_style = style_candidate
        else: logger.warning(f"'{chapter_style_name}' existe mas NÃO é estilo de parágrafo. Usando fallback.")
    except KeyError: logger.warning(f"Estilo '{chapter_style_name}' NÃO encontrado. Usando fallback.")

    chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS), re.IGNORECASE)
    # Divide o texto processado pelo marcador de quebra de página
    parts = formatted_text.split(PAGE_BREAK_MARKER)
    # Verifica se já existe algum conteúdo no documento antes de adicionar quebras
    content_present_before = any(p.text.strip() for p in doc.paragraphs)

    for part_index, part in enumerate(parts):
        part_clean = part.strip()

        # Adiciona quebra de página ANTES de cada parte (exceto a primeira se o doc estiver vazio)
        # Garante que não adicione quebra dupla se a última ação foi adicionar uma.
        if part_index > 0 or content_present_before:
             # Verifica se o último parágrafo não é já uma quebra de página "invisível"
             last_para_is_page_break = False
             if doc.paragraphs:
                 last_p = doc.paragraphs[-1]
                 # Um parágrafo vazio com um run contendo '\f' é como o add_page_break() funciona
                 if not last_p.text.strip() and any(run.text == '\f' for run in last_p.runs):
                     last_para_is_page_break = True

             if not last_para_is_page_break:
                  doc.add_page_break()
                  logger.debug(f"Quebra de página adicionada antes da parte {part_index + 1}.")
                  content_present_before = True # Agora temos conteúdo (a quebra)

        if not part_clean: continue # Pula partes vazias

        # Divide a parte atual em parágrafos (separados por \n\n na saída da IA)
        paragraphs_in_part = part_clean.split("\n\n")
        for paragraph_text in paragraphs_in_part:
            paragraph_text_clean = paragraph_text.strip()
            if not paragraph_text_clean: continue # Pula parágrafos vazios

            is_chapter = chapter_regex.match(paragraph_text_clean) is not None
            is_ai_failure_marker = paragraph_text_clean.startswith(AI_FAILURE_MARKER)
            is_formatting_error_marker = paragraph_text_clean.startswith(FORMATTING_ERROR_MARKER)

            # Adiciona o parágrafo ao documento
            p = doc.add_paragraph()
            run = p.add_run(paragraph_text_clean)
            content_present_before = True # Marcamos que adicionamos conteúdo

            # Aplica Estilos ou formatação fallback
            if is_chapter and not is_ai_failure_marker and not is_formatting_error_marker:
                if chapter_style:
                    p.style = chapter_style
                    logger.debug(f"Aplicado estilo '{chapter_style.name}' ao capítulo: '{paragraph_text_clean[:50]}...'")
                else: # Fallback formatação capítulo
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.size = Pt(14)
                    run.bold = True
                    logger.debug(f"Aplicada formatação fallback de capítulo a: '{paragraph_text_clean[:50]}...'")
            elif is_ai_failure_marker or is_formatting_error_marker:
                 # Formatação especial para marcadores de erro
                 if normal_style: p.style = normal_style # Usa base normal se disponível
                 run.font.italic = True
                 run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Vermelho
                 p.alignment = WD_ALIGN_PARAGRAPH.LEFT # Alinha à esquerda para destaque
                 logger.debug(f"Aplicada formatação de ERRO a: '{paragraph_text_clean[:50]}...'")
            else:
                # Parágrafo normal
                if normal_style:
                    p.style = normal_style
                    # Garante justificado se o estilo Normal não for (opcional)
                    # p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # logger.debug(f"Aplicado estilo '{normal_style.name}' a: '{paragraph_text_clean[:50]}...'")
                else: # Fallback formatação normal
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run.font.size = Pt(12)
                    run.bold = False
                    logger.debug(f"Aplicada formatação fallback normal a: '{paragraph_text_clean[:50]}...'")


def main():
    logger.info("========================================================")
    logger.info(f"Iniciando processamento (OCR) com Gemini ({MODEL_NAME})")
    logger.info("========================================================")
    start_time = time.time()
    backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
    # Garante que o nome do backup não contenha caracteres inválidos (ex: ':')
    backup_timestamp_safe = backup_timestamp.replace(":", "-")
    base_output_name = os.path.splitext(OUTPUT_DOCX)[0]
    BACKUP_DOCX = f"backup_{base_output_name}_{backup_timestamp_safe}.docx"


    # === PASSO 1 – Lê o texto bruto ===
    try:
        with open(INPUT_TXT, "r", encoding="utf-8") as f: texto_original = f.read()
        logger.info(f"Entrada '{INPUT_TXT}' carregada ({len(texto_original)} chars).")
    except FileNotFoundError: logger.error(f"Fatal: Entrada '{INPUT_TXT}' não encontrada."); return
    except Exception as e: logger.error(f"Fatal ao ler '{INPUT_TXT}': {e}"); return

    # === PASSO 2 – Divide o texto em chunks (usando a nova função) ===
    logger.info(f"Dividindo texto (máx. {MAX_CHUNK_TOKENS} tk aprox.)...")
    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS)
    if not text_chunks: logger.error("Nenhum chunk gerado. Verifique o texto de entrada e a lógica de chunking."); return
    logger.info(f"Texto dividido em {len(text_chunks)} chunks.")

    # === PASSO 3 – Carrega o template ou cria novo documento ===
    doc = None
    template_used = False
    try:
        # --- Backup ---
        if os.path.exists(OUTPUT_DOCX):
            try:
                shutil.copy2(OUTPUT_DOCX, BACKUP_DOCX)
                logger.info(f"Backup do arquivo anterior criado: {BACKUP_DOCX}")
            except Exception as e:
                logger.warning(f"Falha ao criar backup de '{OUTPUT_DOCX}': {e}. O arquivo existente será sobrescrito se o processo falhar.")

        # --- Carrega Template ---
        doc = Document(TEMPLATE_DOCX)
        logger.info(f"Template '{TEMPLATE_DOCX}' carregado.")
        template_used = True

        # === PASSO 4 – Limpa o corpo do documento (APENAS SE USOU TEMPLATE) ===
        try:
            # Acessa o corpo do documento de forma segura
            body_element = doc._body._body
            # Remove todos os parágrafos e tabelas existentes no corpo
            for child in reversed(body_element):
                body_element.remove(child)
            logger.info("Conteúdo principal do template limpo (parágrafos/tabelas).")
        except Exception as clean_err:
             logger.error(f"Erro durante limpeza do template: {clean_err}")
             logger.warning("Continuando apesar do erro na limpeza. Conteúdo do template pode permanecer.")

    except FileNotFoundError:
        logger.warning(f"Template '{TEMPLATE_DOCX}' não encontrado. Criando novo documento A5.")
        doc = Document()
        # Configurações de página e margens para A5 (aproximado)
        try:
            section = doc.sections[0]
            section.page_height = Inches(8.27) # A5 Altura
            section.page_width = Inches(5.83)  # A5 Largura
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.6)
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            logger.info("Novo documento criado com configurações de página (A5) e margens.")

            # Cria estilos básicos se não existirem (Normal e Heading 1 base)
            styles = doc.styles
            if NORMAL_STYLE_NAME not in styles:
                 style = styles.add_style(NORMAL_STYLE_NAME, 1) # WD_STYLE_TYPE.PARAGRAPH = 1
                 style.font.name = 'Times New Roman'
                 style.font.size = Pt(12)
                 style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                 style.paragraph_format.space_after = Pt(6)
                 logger.info(f"Estilo '{NORMAL_STYLE_NAME}' criado.")
            else:
                 logger.info(f"Estilo '{NORMAL_STYLE_NAME}' já existe no novo documento (padrão).")

            # Tenta criar estilo de capítulo baseado no Heading 1 se existir
            if CHAPTER_STYLE_NAME not in styles:
                 try:
                     base_style = styles['Heading 1']
                     style = styles.add_style(CHAPTER_STYLE_NAME, 1) # WD_STYLE_TYPE.PARAGRAPH = 1
                     style.base_style = base_style # Herda do Heading 1
                     style.font.name = 'Times New Roman' # Pode sobrescrever se necessário
                     style.font.size = Pt(14)
                     style.font.bold = True
                     style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                     style.paragraph_format.space_before = Pt(12)
                     style.paragraph_format.space_after = Pt(12)
                     logger.info(f"Estilo '{CHAPTER_STYLE_NAME}' criado (baseado em Heading 1).")
                 except KeyError:
                     # Se nem 'Heading 1' existir, cria um estilo de capítulo básico
                     style = styles.add_style(CHAPTER_STYLE_NAME, 1)
                     style.font.name = 'Times New Roman'
                     style.font.size = Pt(14); style.font.bold = True
                     style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                     style.paragraph_format.space_before = Pt(12); style.paragraph_format.space_after = Pt(12)
                     logger.warning(f"Estilo base 'Heading 1' não encontrado. Estilo '{CHAPTER_STYLE_NAME}' criado sem herança.")
            else:
                 logger.info(f"Estilo '{CHAPTER_STYLE_NAME}' já existe no novo documento.")

        except Exception as e:
            logger.warning(f"Falha ao aplicar configs/estilos ao novo doc: {e}")

    except Exception as e:
        logger.error(f"Erro crítico ao carregar template ou criar novo documento: {e}"); return


    # === PASSO 5 – Processa e insere cada chunk ===
    logger.info(f"Iniciando processamento de {len(text_chunks)} chunks com API ({MODEL_NAME})...")
    processed_chunks_count = 0
    failed_chunks_count = 0
    total_chunks_processed_or_failed = 0 # Contador para salvar progresso
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc="Processando Chunks", unit="chunk")

    for i, chunk in progress_bar:
        chunk_start_time = time.time()
        progress_bar.set_description(f"Processando Chunk {i+1}/{len(text_chunks)}")
        formatted_chunk = None
        try:
            # Passa False para is_first_chunk se já adicionamos conteúdo (mesmo que seja fallback)
            formatted_chunk = format_with_ai(gemini_model, chunk, is_first_chunk=(total_chunks_processed_or_failed == 0))
        except Exception as api_err:
            logger.error(f"Erro INESPERADO na chamada format_with_ai para chunk {i+1}: {api_err}")
            # Considera como falha da API para fins de fallback
            formatted_chunk = None

        # ---- LÓGICA DE FALLBACK ----
        if formatted_chunk and formatted_chunk.strip():
            try:
                # Tenta aplicar formatação ao texto da IA
                apply_formatting(doc, formatted_chunk, NORMAL_STYLE_NAME, CHAPTER_STYLE_NAME)
                processed_chunks_count += 1
                total_chunks_processed_or_failed += 1
            except Exception as format_err:
                # Se falhar AQUI (raro, mas possível), loga o erro e tenta o fallback com texto original
                logger.error(f"Erro RARO na função apply_formatting para chunk {i+1} (texto da IA): {format_err}. Usando fallback com texto original.")
                failed_chunks_count += 1
                total_chunks_processed_or_failed += 1
                try:
                    # Tenta aplicar formatação ao texto ORIGINAL + marcador de erro de formatação
                    apply_formatting(doc, f"{FORMATTING_ERROR_MARKER}\n\n{chunk}", NORMAL_STYLE_NAME, CHAPTER_STYLE_NAME)
                    logger.warning(f"Chunk {i+1} adicionado como original devido a erro na formatação pós-IA.")
                except Exception as fallback_format_err:
                    # Se o fallback falhar TAMBÉM, loga erro crítico
                    logger.critical(f"Falha CRÍTICA ao aplicar fallback para chunk {i+1} (erro formatação pós-IA): {fallback_format_err}. CONTEÚDO PERDIDO.")
                    # Não incrementa total_chunks_processed_or_failed aqui, pois foi perdido
        else:
            # Se a API falhou (retornou None) ou retornou vazio, usa o fallback diretamente
            if formatted_chunk is None:
                logger.warning(f"Chunk {i+1} falhou na API após retentativas. Usando fallback com texto original.")
            else: # formatted_chunk era "" ou só espaços
                logger.warning(f"Chunk {i+1} retornou vazio da API. Usando fallback com texto original.")

            failed_chunks_count += 1
            total_chunks_processed_or_failed += 1
            try:
                # Tenta aplicar formatação ao texto ORIGINAL + marcador de falha da IA
                apply_formatting(doc, f"{AI_FAILURE_MARKER}\n\n{chunk}", NORMAL_STYLE_NAME, CHAPTER_STYLE_NAME)
            except Exception as fallback_format_err:
                 # Se o fallback falhar, loga erro crítico
                logger.critical(f"Falha CRÍTICA ao aplicar fallback para chunk {i+1} (falha API): {fallback_format_err}. CONTEÚDO PERDIDO.")
                # Decrementa o contador, pois o chunk foi perdido
                total_chunks_processed_or_failed -= 1
        # ---- FIM FALLBACK ----

        chunk_end_time = time.time()
        logger.debug(f"Chunk {i+1} processado em {chunk_end_time - chunk_start_time:.2f} seg.")

        # === PASSO 5.1 – Salva progresso periodicamente ===
        # Salva a cada 5 chunks processados (OK ou Falha com fallback) OU no último chunk
        if total_chunks_processed_or_failed > 0 and \
           (total_chunks_processed_or_failed % 5 == 0 or (i + 1) == len(text_chunks)):
            temp_save_path = f"{OUTPUT_DOCX}.temp_save" # Nome temporário diferente
            try:
                doc.save(temp_save_path)
                # Usa replace para atomicidade (reduz chance de corromper em falha)
                # No Windows, os.replace pode falhar se o destino existir, então removemos primeiro.
                if os.path.exists(OUTPUT_DOCX):
                    os.remove(OUTPUT_DOCX)
                shutil.move(temp_save_path, OUTPUT_DOCX)
                logger.info(f"Progresso salvo ({total_chunks_processed_or_failed} chunks no doc). Arquivo: {OUTPUT_DOCX}")
            except Exception as e:
                logger.error(f"Erro ao salvar progresso parcial (chunk {i+1}): {e}")
                # Tenta remover o .temp_save se ele existir para evitar confusão
                if os.path.exists(temp_save_path):
                    try: os.remove(temp_save_path)
                    except Exception: pass
                logger.warning(f"O arquivo '{OUTPUT_DOCX}' pode não conter o progresso do último lote salvo.")


    # === PASSO 6 – Conclusão e Salvamento Final ===
    # Garante um salvamento final após o loop, mesmo que o último chunk não tenha disparado o save periódico
    final_temp_path = f"{OUTPUT_DOCX}.final_temp"
    try:
        doc.save(final_temp_path)
        if os.path.exists(OUTPUT_DOCX):
             os.remove(OUTPUT_DOCX)
        shutil.move(final_temp_path, OUTPUT_DOCX)
        logger.info(f"Salvamento final concluído: {OUTPUT_DOCX}")
    except Exception as e:
        logger.error(f"Erro no salvamento final: {e}")
        # Informa o usuário sobre o arquivo temporário se ele existir
        if os.path.exists(final_temp_path):
             logger.warning(f"ATENÇÃO: O salvamento final falhou, mas o arquivo temporário '{final_temp_path}' PODE conter a versão completa.")
        else:
             logger.warning(f"ATENÇÃO: O salvamento final falhou e o arquivo temporário não foi encontrado. '{OUTPUT_DOCX}' pode estar incompleto.")

    end_time = time.time(); total_time = end_time - start_time
    logger.info("========================================================")
    logger.info("✅ Processamento Concluído!")
    logger.info(f"Modelo: {MODEL_NAME}")
    logger.info(f"Tempo total: {total_time:.2f} seg ({total_time/60:.2f} min).")
    logger.info(f"Chunks enviados para API: {len(text_chunks)}")
    logger.info(f"Chunks processados OK pela IA (sem fallback): {processed_chunks_count}")
    logger.info(f"Chunks com falha na IA ou formatação (usado fallback): {failed_chunks_count}")
    # Verifica quantos chunks realmente estão no documento
    final_chunks_in_doc = processed_chunks_count + failed_chunks_count
    logger.info(f"Total de chunks (originais ou processados) incluídos no documento: {final_chunks_in_doc}")
    if final_chunks_in_doc < len(text_chunks):
        logger.error(f"ATENÇÃO: {len(text_chunks) - final_chunks_in_doc} chunks podem ter sido PERDIDOS devido a erros críticos no fallback ou salvamento. Verifique logs CRITICAL.")
    logger.info(f"Livro final salvo em: {OUTPUT_DOCX}")
    if os.path.exists(BACKUP_DOCX): logger.info(f"Backup do arquivo anterior (se existia): {BACKUP_DOCX}")
    logger.info(f"Log detalhado salvo em: {log_filepath}")
    logger.info("========================================================")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("Processo interrompido manualmente (Ctrl+C). O último salvamento pode estar incompleto.")
    except Exception as e:
        # Loga a exceção completa com traceback no arquivo
        logger.exception(f"Erro fatal inesperado durante a execução: {e}")
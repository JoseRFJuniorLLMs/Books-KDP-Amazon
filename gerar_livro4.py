from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from dotenv import load_dotenv
import os
import tiktoken
import re
import logging
from tqdm import tqdm

# === SETUP LOGGING ===
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("book_processor.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# === CARREGA VARIÁVEIS DE AMBIENTE DO .env ===
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    logger.error("API key not found. Please set OPENAI_API_KEY in your .env file.")
    exit(1)

# === CONFIGURAÇÕES ===
INPUT_TXT = "rascunho.txt"
TEMPLATE_DOCX = "Estrutura.docx"
OUTPUT_DOCX = "Livro_Final_Formatado.docx"
BACKUP_DOCX = f"backup_{os.path.splitext(OUTPUT_DOCX)[0]}_{os.path.getmtime(OUTPUT_DOCX) if os.path.exists(OUTPUT_DOCX) else 'initial'}.docx"

# Configuration constants
MODEL_NAME = "gpt-4-turbo"  # or "gpt-4o"
MAX_CHUNK_TOKENS = 8000  # Reduced to be safer
MAX_OUTPUT_TOKENS = 10000
TEMPERATURE = 0.7
CHAPTER_PATTERNS = [
    r'^Capítulo \w+',
    r'^CAPÍTULO \w+',
    r'^Capítulo \d+',
    r'^CHAPTER \w+',
    r'^Chapter \d+'
]

# Setup token counter
try:
    tokenizer = tiktoken.encoding_for_model(MODEL_NAME)
except Exception as e:
    logger.warning(f"Could not load tokenizer for {MODEL_NAME}: {e}. Using approximate character count instead.")
    tokenizer = None

# === SETUP OPENAI ===
client = OpenAI(api_key=API_KEY)


def count_tokens(text):
    """Count tokens in text, either with tiktoken or approximately with characters."""
    if tokenizer:
        return len(tokenizer.encode(text))
    else:
        # Fallback: rough approximation (1 token ≈ 4 chars for Latin scripts)
        return len(text) // 4


def create_chunks(text, max_tokens):
    """
    Split text into chunks respecting paragraph and chapter boundaries.
    Returns a list of chunks, each containing no more than max_tokens.
    """
    chunks = []
    current_chunk = ""
    current_chunk_tokens = 0

    # Compile the chapter patterns for efficiency
    chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS), re.IGNORECASE)

    paragraphs = text.split("\n\n")

    for i, paragraph in enumerate(paragraphs):
        paragraph = paragraph.strip()
        if not paragraph:
            continue

        paragraph_tokens = count_tokens(paragraph)

        # If adding this paragraph would exceed the limit, save chunk and start new one
        if current_chunk_tokens + paragraph_tokens > max_tokens:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
                current_chunk_tokens = 0

        # Special case: if a single paragraph is larger than max_tokens
        if paragraph_tokens > max_tokens:
            logger.warning(
                f"Paragraph exceeds token limit ({paragraph_tokens} > {max_tokens}). Splitting at sentence boundaries.")
            # Split the paragraph at sentence boundaries
            sentences = re.split(r'(?<=[.!?])\s+', paragraph)
            temp_chunk = ""
            temp_tokens = 0

            for sentence in sentences:
                sentence_tokens = count_tokens(sentence)
                if temp_tokens + sentence_tokens <= max_tokens:
                    temp_chunk += sentence + " "
                    temp_tokens += sentence_tokens
                else:
                    if temp_chunk:
                        chunks.append(temp_chunk.strip())
                    temp_chunk = sentence + " "
                    temp_tokens = sentence_tokens

            if temp_chunk:
                chunks.append(temp_chunk.strip())
        else:
            # Check if this paragraph starts a new chapter
            is_chapter_start = chapter_regex.match(paragraph) is not None

            # If this is a chapter start and not the first chunk, start a new chunk
            if is_chapter_start and current_chunk and i > 0:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
                current_chunk_tokens = paragraph_tokens
            else:
                current_chunk += paragraph + "\n\n"
                current_chunk_tokens += paragraph_tokens

    # Add the last chunk if not empty
    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks


def format_with_ai(chunk, is_first_chunk=False):
    """Process a chunk of text with the OpenAI API and format it."""
    # Different instructions depending on if it's first chunk or a continuation
    context = "Você está formatando o início de um livro. " if is_first_chunk else "Você está formatando uma continuação de texto. "

    chunk_prompt = f"""
    {context}Você é um editor literário. Corrija e formate o texto a seguir como parte de um livro.

    Regras:
    - Corrija erros gramaticais, ortográficos e de concordância
    - Mantenha a formatação de capítulos se presente no texto
    - **NÃO** adicione títulos de capítulo se eles não estiverem explicitamente no texto
    - Se houver um marcador de quebra de página (===QUEBRA_DE_PAGINA===) dentro deste fragmento, mantenha-o
    - Mantenha estilo literário fluido e bem escrito
    - **NÃO** adicione texto introdutório ou conclusivo que não faça parte do rascunho
    - O resultado deve ser APENAS o texto formatado
    - Mantenha fidelidade ao texto original: não omita parágrafos ou adicione conteúdo

    Texto do fragmento:
    \"\"\"{chunk}\"\"\"
    """

    # Try with exponential backoff in case of rate limiting
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": chunk_prompt}],
                temperature=TEMPERATURE,
                max_tokens=MAX_OUTPUT_TOKENS
            )
            return response.choices[0].message.content
        except Exception as e:
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt  # Exponential backoff: 1, 2, 4 seconds
                logger.warning(f"API error: {e}. Retrying in {wait_time} seconds...")
                import time
                time.sleep(wait_time)
            else:
                logger.error(f"Failed to process chunk after {max_retries} attempts: {e}")
                return chunk  # Return the original chunk if all retries fail


def apply_formatting(doc, formatted_text):
    """Apply formatting to the document based on the formatted text."""
    # Split the text by page break markers first
    parts = formatted_text.split("===QUEBRA_DE_PAGINA===")
    chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS))

    for part_index, part in enumerate(parts):
        # Add page breaks between parts (but not before the first part)
        if part_index > 0:
            doc.add_page_break()

        paragraphs = part.split("\n\n")
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            # Check if this is a chapter heading
            is_chapter = chapter_regex.match(paragraph) is not None

            p = doc.add_paragraph(paragraph)

            # Apply formatting to chapters
            if is_chapter:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(14)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def main():
    # === PASSO 1 – Lê o texto bruto ===
    try:
        with open(INPUT_TXT, "r", encoding="utf-8") as f:
            texto_original = f.read()
        logger.info(f"Arquivo '{INPUT_TXT}' carregado com sucesso.")
    except FileNotFoundError:
        logger.error(f"Erro: O arquivo de rascunho '{INPUT_TXT}' não foi encontrado.")
        exit(1)
    except Exception as e:
        logger.error(f"Erro ao ler o arquivo '{INPUT_TXT}': {e}")
        exit(1)

    # === PASSO 2 – Split the text into chunks ===
    logger.info("Fragmentando o texto...")
    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS)
    logger.info(f"✅ Texto dividido em {len(text_chunks)} fragmentos.")

    # === PASSO 3 – Carrega o template ou cria um novo documento ===
    try:
        doc = Document(TEMPLATE_DOCX)
        logger.info(f"Template '{TEMPLATE_DOCX}' carregado com sucesso.")
        # Create backup of template if needed
        if os.path.exists(OUTPUT_DOCX):
            try:
                import shutil
                shutil.copy2(OUTPUT_DOCX, BACKUP_DOCX)
                logger.info(f"Backup criado: {BACKUP_DOCX}")
            except Exception as e:
                logger.warning(f"Não foi possível criar backup: {e}")

        # === PASSO 4 – Limpa corpo do documento (opcional) ===
        if hasattr(doc, '_body') and doc._body:
            doc._body.clear_content()
            logger.info("Conteúdo do template limpo.")
    except FileNotFoundError:
        logger.warning(f"Template '{TEMPLATE_DOCX}' não encontrado. Criando um novo documento.")
        doc = Document()
        # Setup default document properties
        section = doc.sections[0]
        section.page_height = Inches(9)
        section.page_width = Inches(6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # === PASSO 5 – Process and insert each chunk ===
    logger.info("Processando fragmentos com ChatGPT e inserindo no documento...")

    for i, chunk in enumerate(tqdm(text_chunks, desc="Processando fragmentos", unit="chunk")):
        logger.info(f"Processando fragmento {i + 1}/{len(text_chunks)}...")

        # Process this chunk
        formatted_chunk = format_with_ai(chunk, is_first_chunk=(i == 0))

        # Apply formatting to document
        apply_formatting(doc, formatted_chunk)

        # Periodic saves to prevent data loss
        if (i + 1) % 5 == 0 or i == len(text_chunks) - 1:
            temp_path = f"{OUTPUT_DOCX}.temp"
            try:
                doc.save(temp_path)
                os.replace(temp_path, OUTPUT_DOCX)
                logger.info(f"Salvo progresso parcial ({i + 1}/{len(text_chunks)} fragmentos)")
            except Exception as e:
                logger.error(f"Erro ao salvar progresso parcial: {e}")

    # === PASSO 6 – Salva o novo documento ===
    try:
        doc.save(OUTPUT_DOCX)
        logger.info(f"📘 Livro gerado com sucesso: {OUTPUT_DOCX}")
    except Exception as e:
        logger.error(f"Erro ao salvar documento final: {e}")
        # Try to recover from temp file if exists
        if os.path.exists(f"{OUTPUT_DOCX}.temp"):
            try:
                os.replace(f"{OUTPUT_DOCX}.temp", OUTPUT_DOCX)
                logger.info(f"Recuperado último salvamento parcial: {OUTPUT_DOCX}")
            except:
                logger.error("Não foi possível recuperar último salvamento parcial.")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.info("Processo interrompido pelo usuário.")
    except Exception as e:
        logger.exception(f"Erro não tratado: {e}")
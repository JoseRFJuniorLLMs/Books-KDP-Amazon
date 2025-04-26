# -*- coding: utf-8 -*-
# --- Using Google's Gemini API (gemini-1.5-pro) ---

# Standard Python Libraries
import sys # << Importado para sys.executable >>
from dotenv import load_dotenv
import os
import re
import logging
from tqdm import tqdm
import time
import shutil
import traceback # Para log de erros detalhado
import glob # Para encontrar arquivos .txt
import smtplib # For email
import ssl # For email security
from email.message import EmailMessage # For constructing email
import subprocess # << IMPORTADO PARA CHAMAR O SCRIPT TRADUTOR >>

# Third-party Libraries (ensure installed: pip install python-docx google-generativeai python-dotenv tqdm)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.shared import RGBColor
import google.generativeai as genai # Google Generative AI Library

# === SETUP LOGGING ===
log_dir = "logs"
if not os.path.exists(log_dir): os.makedirs(log_dir)
log_filepath = os.path.join(log_dir, "book_processor_multi_author_mem.log")
# Log para arquivos com CORREÇÃO concluída
PROCESSED_LOG_FILE = os.path.join(log_dir, "processed_books.log")
# Log para arquivos com TRADUÇÃO concluída (NOVO)
TRANSLATED_LOG_FILE = os.path.join(log_dir, "translated_books.log")

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(module)s:%(lineno)d - %(funcName)s - %(message)s',
    handlers=[ logging.FileHandler(log_filepath, encoding='utf-8'), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# === CARREGA VARIÁVEIS DE AMBIENTE ===
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
# Email configuration from .env file (CARREGADO AQUI)
EMAIL_SENDER_ADDRESS = os.getenv("EMAIL_SENDER_ADDRESS")
EMAIL_SENDER_APP_PASSWORD = os.getenv("EMAIL_SENDER_APP_PASSWORD") # Use App Password for Gmail
EMAIL_RECIPIENT_ADDRESS = os.getenv("EMAIL_RECIPIENT_ADDRESS", "web2ajax@gmail.com") # Destinatário padrão
EMAIL_SMTP_SERVER = os.getenv("EMAIL_SMTP_SERVER", "smtp.gmail.com") # Default Gmail
EMAIL_SMTP_PORT = int(os.getenv("EMAIL_SMTP_PORT", 587)) # Default Gmail Port (TLS)


# === CONFIGURAÇÕES ===

# -- Diretórios Base --
BASE_INPUT_TXT_DIR = "txt"
BASE_OUTPUT_DOCX_DIR = "docx"
BASE_OUTPUT_TXT_DIR = "txt"

# -- Nomes de Arquivos Base --
# !! AJUSTE SE NECESSÁRIO !!
TEMPLATE_DOCX = "Estrutura.docx" # Template OBRIGATÓRIO

# -- Nomes Base dos Arquivos de Saída (Correção) --
FINAL_DOCX_BASENAME = "Livro_Final_Formatado_Sem_Notas.docx"
FINAL_NUMBERED_TXT_BASENAME = "Livro_Final_Com_Notas_Numeros.txt"
NOTES_TXT_FILE_BASENAME = "notas_rodape.txt"

# -- Nomes Base dos Arquivos de Saída (Tradução - NOVO) --
# O nome exato será gerado dentro do loop principal
TRANSLATED_DOCX_SUFFIX = "-A0.docx"

# -- Configurações da API e Processamento (Correção) --
MODEL_NAME = "gemini-1.5-pro" # Modelo para correção e notas
MAX_CHUNK_TOKENS = 1500 # Aprox. limite de tokens por chunk para API
MAX_OUTPUT_TOKENS = 8192 # Limite de saída do Gemini
TEMPERATURE = 0.5 # Temperatura para a correção

# -- Configurações para o Script Tradutor (NOVO) --
# !! AJUSTE ESTE CAMINHO para onde você salvou o script_tradutor_hibrido.py !!
PATH_TO_TRANSLATOR_SCRIPT = "script_tradutor_hibrido.py"
NUM_WORDS_TO_TRANSLATE = 100 # Quantidade de palavras para o script tradutor processar

# -- Estilos e Padrões (Correção) --
NORMAL_STYLE_NAME = "Normal" # Estilo esperado no template
CHAPTER_PATTERNS = [
    r'^\s*Capítulo \w+', r'^\s*CAPÍTULO \w+', r'^\s*Capítulo \d+',
    r'^\s*CHAPTER \w+', r'^\s*Chapter \d+', r'^\s*LIVRO \w+', r'^\s*PARTE \w+',
    # Adicionar outros padrões se necessário
]
PAGE_BREAK_MARKER = "===QUEBRA_DE_PAGINA===" # Marcador para quebra de página manual no texto
AI_FAILURE_MARKER = "*** FALHA NA IA - TEXTO ORIGINAL ABAIXO ***" # Marcador de falha da API
FORMATTING_ERROR_MARKER = "*** ERRO DE FORMATAÇÃO - TEXTO ORIGINAL ABAIXO ***" # Marcador de erro de formatação

# --- Validação API Key ---
if not GOOGLE_API_KEY:
    logger.error("FATAL: GOOGLE_API_KEY não encontrada nas variáveis de ambiente (.env).")
    exit(1)

# --- Validação Config Email (Informativo) ---
email_configured = bool(EMAIL_SENDER_ADDRESS and EMAIL_SENDER_APP_PASSWORD and EMAIL_RECIPIENT_ADDRESS)
if not email_configured:
    logger.warning("AVISO: Configurações de e-mail (EMAIL_SENDER_ADDRESS, EMAIL_SENDER_APP_PASSWORD, EMAIL_RECIPIENT_ADDRESS) incompletas no .env. Notificação por e-mail será desativada.")
else:
    logger.info(f"Configuração de e-mail carregada. Notificações serão enviadas de '{EMAIL_SENDER_ADDRESS}' para '{EMAIL_RECIPIENT_ADDRESS}'.")

# --- Setup Gemini Client ---
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    # Configurações de segurança mais permissivas (use com cautela)
    safety_settings_lenient = {
        'HATE': 'BLOCK_NONE', 'HARASSMENT': 'BLOCK_NONE',
        'SEXUAL' : 'BLOCK_NONE', 'DANGEROUS' : 'BLOCK_NONE'
    }
    generation_config = genai.GenerationConfig(
                    temperature=TEMPERATURE,
                    max_output_tokens=MAX_OUTPUT_TOKENS
                )
    gemini_model = genai.GenerativeModel(
        MODEL_NAME,
        safety_settings=safety_settings_lenient,
        generation_config=generation_config
    )
    logger.info(f"Modelo Gemini '{MODEL_NAME}' inicializado com sucesso.")
except Exception as e:
    logger.error(f"FATAL: Falha ao inicializar modelo Gemini ({MODEL_NAME}): {e}")
    logger.error(traceback.format_exc())
    exit(1)

# --- Funções Auxiliares ---

def count_tokens_approx(text):
    """Estima a contagem de tokens (aproximadamente 3 chars/token)."""
    if not text: return 0
    return len(text) // 3

def create_chunks(text, max_tokens, author_name="N/A", book_name="N/A"):
    """Divide o texto em chunks, subdividindo parágrafos grandes."""
    log_prefix = f"[{author_name}/{book_name}]"
    # logger.info(f"{log_prefix} Iniciando criação de chunks. Máx tokens (aprox): {max_tokens}") # Verbose
    chunks = []
    current_chunk = ""
    current_chunk_tokens = 0
    # Divide por parágrafos (duas quebras de linha)
    paragraphs = text.split("\n\n")

    for i, paragraph_text in enumerate(paragraphs):
        # Ignora blocos vazios mas tenta manter o espaçamento
        if not paragraph_text.strip():
            if chunks and chunks[-1].strip() and not chunks[-1].endswith("\n\n"):
                chunks[-1] += "\n\n"
            continue

        paragraph_tokens = count_tokens_approx(paragraph_text)
        tokens_with_separator = paragraph_tokens + (count_tokens_approx("\n\n") if current_chunk else 0)

        # Se o parágrafo cabe no chunk atual, adiciona
        if current_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if current_chunk else ""
            current_chunk += separator + paragraph_text
            current_chunk_tokens = count_tokens_approx(current_chunk)
        # Se não cabe, salva o chunk atual e inicia um novo com o parágrafo
        else:
            if current_chunk: # Salva o chunk anterior se ele não estiver vazio
                chunks.append(current_chunk)
                # logger.debug(f"{log_prefix} Chunk {len(chunks)} salvo (limite atingido). Tokens: {current_chunk_tokens}.") # Verbose
            # O parágrafo atual começa um novo chunk
            current_chunk = paragraph_text
            current_chunk_tokens = paragraph_tokens

            # --- Subdivisão se o *próprio* parágrafo atual for muito grande ---
            if paragraph_tokens > max_tokens:
                logger.warning(f"{log_prefix} Parágrafo {i+1} ({paragraph_tokens} tk) excede limite {max_tokens}. Iniciando SUBDIVISÃO.")
                # Remove o parágrafo grande do current_chunk (que acabamos de adicionar)
                current_chunk = ""
                current_chunk_tokens = 0

                sub_chunks_added_count = 0
                # Tenta dividir por frases, senão por linhas
                sentences = re.split(r'(?<=[.!?])\s+', paragraph_text)
                if len(sentences) <= 1 :
                    sentences = paragraph_text.split('\n') # Fallback para linhas

                current_sub_chunk = ""
                current_sub_chunk_tokens = 0
                for sentence_num, sentence in enumerate(sentences):
                    sentence_clean = sentence.strip()
                    if not sentence_clean: continue

                    sentence_tokens = count_tokens_approx(sentence)
                    tokens_with_sub_separator = sentence_tokens + (count_tokens_approx("\n") if current_sub_chunk else 0)

                    # Se a sentença cabe no sub-chunk atual
                    if current_sub_chunk_tokens + tokens_with_sub_separator <= max_tokens:
                        sub_separator = "\n" if current_sub_chunk else "" # Usa \n dentro do parágrafo subdividido
                        current_sub_chunk += sub_separator + sentence
                        current_sub_chunk_tokens = count_tokens_approx(current_sub_chunk)
                    # Se não cabe, salva o sub-chunk e inicia um novo
                    else:
                        if current_sub_chunk: # Salva o sub-chunk anterior
                            chunks.append(current_sub_chunk)
                            sub_chunks_added_count += 1
                            # logger.debug(f"{log_prefix} Sub-chunk {len(chunks)} salvo (limite sub). Tokens: {current_sub_chunk_tokens}.") # Verbose

                        # Verifica se a *própria* sentença é grande demais
                        if sentence_tokens > max_tokens:
                            chunks.append(sentence) # Adiciona a sentença longa como um chunk próprio
                            sub_chunks_added_count += 1
                            logger.warning(f"{log_prefix}  -> Sentença/Linha {sentence_num+1} ({sentence_tokens} tk) excede limite. Adicionada como chunk individual (PODE FALHAR NA API).")
                            current_sub_chunk = "" # Reseta, pois ela foi adicionada separadamente
                            current_sub_chunk_tokens = 0
                        else:
                            # A sentença não é grande demais, ela inicia o novo sub-chunk
                            current_sub_chunk = sentence
                            current_sub_chunk_tokens = sentence_tokens

                # Salva o último sub-chunk restante
                if current_sub_chunk:
                    chunks.append(current_sub_chunk)
                    sub_chunks_added_count += 1
                    # logger.debug(f"{log_prefix} Último sub-chunk {len(chunks)} salvo (Parág. {i+1}). Tokens: {current_sub_chunk_tokens}.") # Verbose

                # Se não conseguiu subdividir (caso raro), adiciona o parágrafo original
                if sub_chunks_added_count == 0:
                     logger.warning(f"{log_prefix} Parágrafo {i+1} excedeu limite, mas não foi subdividido. Adicionando original como chunk (PODE FALHAR NA API).")
                     chunks.append(paragraph_text)

                # Reseta o chunk principal após lidar com o parágrafo grande
                current_chunk = ""
                current_chunk_tokens = 0
            # --- Fim da Subdivisão ---

    # Adiciona o último chunk que sobrou
    if current_chunk:
        chunks.append(current_chunk)
        # logger.debug(f"{log_prefix} Chunk final {len(chunks)} salvo. Tokens: {current_chunk_tokens}.") # Verbose

    # --- Pós-processamento: Junta chunks pequenos consecutivos ---
    # logger.debug(f"{log_prefix} Iniciando merge de chunks pequenos...") # Verbose
    merged_chunks = []
    temp_chunk = ""
    temp_chunk_tokens = 0
    for i, chunk in enumerate(chunks):
        chunk_tokens = count_tokens_approx(chunk)
        tokens_with_separator = chunk_tokens + (count_tokens_approx("\n\n") if temp_chunk else 0)

        # Se o chunk atual cabe junto com o temporário
        if temp_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if temp_chunk else ""
            temp_chunk += separator + chunk
            temp_chunk_tokens = count_tokens_approx(temp_chunk)
        # Se não cabe, salva o temporário e inicia um novo
        else:
            if temp_chunk: merged_chunks.append(temp_chunk)
            temp_chunk = chunk
            temp_chunk_tokens = chunk_tokens

    # Salva o último chunk temporário
    if temp_chunk:
        merged_chunks.append(temp_chunk)

    final_chunk_count = len(merged_chunks)
    # Comentado log de merge pois pode ser verboso
    # if final_chunk_count < len(chunks):
    #      logger.info(f"{log_prefix} Merge concluído. De {len(chunks)} para {final_chunk_count} chunks.")

    # logger.info(f"{log_prefix} ✅ Chunking concluído. {final_chunk_count} chunks finais.")
    return merged_chunks

# <<< MODIFICADO _call_gemini_api para retornar latência e tokens >>>
def _call_gemini_api(model, prompt_text, chunk_for_log, author_name="N/A", book_name="N/A"):
    """
    Função interna para chamar a API Gemini com retries, tratamento de erro
    e retorno de estatísticas (latência, tokens).
    Retorna: (result_text|None, latency_secs, prompt_tokens, output_tokens, total_tokens)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    max_retries = 5
    base_wait_time = 5
    log_chunk_preview = chunk_for_log[:100].replace('\n', '\\n') + ('...' if len(chunk_for_log) > 100 else '')
    default_return = (None, 0, 0, 0, 0) # Retorno padrão em caso de falha completa

    for attempt in range(max_retries):
        # logger.info(f"{log_prefix} Chamando API (Tentativa {attempt + 1}/{max_retries}).") # Verbose
        start_time = time.time() # Inicia cronômetro
        response = None
        latency = 0
        prompt_tokens = 0
        output_tokens = 0
        total_tokens = 0
        result_text = None

        try:
            response = model.generate_content(prompt_text)
            latency = time.time() - start_time # Calcula latência

            # Tenta extrair metadados de uso (tokens)
            try:
                if hasattr(response, 'usage_metadata'):
                    usage = response.usage_metadata
                    prompt_tokens = usage.prompt_token_count if hasattr(usage, 'prompt_token_count') else 0
                    output_tokens = usage.candidates_token_count if hasattr(usage, 'candidates_token_count') else 0
                    # Calcula total se não existir (pode ser mais preciso somar)
                    total_tokens = usage.total_token_count if hasattr(usage, 'total_token_count') else (prompt_tokens + output_tokens)
                    logger.debug(f"{log_prefix} API OK ({latency:.2f}s). Tokens: P{prompt_tokens} + O{output_tokens} = T{total_tokens}")
                else: logger.debug(f"{log_prefix} API OK ({latency:.2f}s). Metadados de uso (tokens) não encontrados na resposta.")
            except AttributeError as e_usage:
                logger.warning(f"{log_prefix} Atributo esperado não encontrado em usage_metadata: {e_usage}")
            except Exception as e_usage_other:
                 logger.error(f"{log_prefix} Erro inesperado ao processar usage_metadata: {e_usage_other}")

            # Verifica bloqueio de prompt
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback and \
               hasattr(response.prompt_feedback, 'block_reason') and response.prompt_feedback.block_reason:
                block_reason = response.prompt_feedback.block_reason.name
                logger.error(f"{log_prefix} API BLOQUEOU O PROMPT (Tentativa {attempt + 1}). Razão: {block_reason}. Latência: {latency:.2f}s.")
                return default_return # Falha

            # Verifica candidatos
            if not response.candidates:
                 logger.error(f"{log_prefix} API retornou SEM CANDIDATOS (Tentativa {attempt + 1}). Latência: {latency:.2f}s. Resposta: {response}")
            else:
                 try:
                    candidate = response.candidates[0]
                    finish_reason = candidate.finish_reason.name if hasattr(candidate, 'finish_reason') and candidate.finish_reason else "FINISH_REASON_UNKNOWN"

                    # Extrai texto
                    if hasattr(candidate, 'content') and hasattr(candidate.content, 'parts'):
                        text_parts = [part.text for part in candidate.content.parts if hasattr(part, 'text')]
                        if text_parts: result_text = "".join(text_parts).strip()
                    elif hasattr(response, 'text') and response.text: result_text = response.text.strip()

                    # Checa se o término foi normal e se extraiu texto
                    if finish_reason == "STOP" and result_text is not None: # Checa se result_text não é None
                        return (result_text, latency, prompt_tokens, output_tokens, total_tokens) # SUCESSO
                    else:
                         # Loga motivos de término não ideais ou falha na extração
                         if finish_reason != "STOP":
                             logger.warning(f"{log_prefix} API terminou não normalmente (Tentativa {attempt+1}). Finish: {finish_reason}. Latência: {latency:.2f}s.")
                             if finish_reason == "SAFETY" and hasattr(candidate, 'safety_ratings') and candidate.safety_ratings:
                                 logger.warning(f" -> Safety Ratings: {[(r.category.name, r.probability.name) for r in candidate.safety_ratings]}")
                         if result_text is None: # Verifica None explicitamente
                             logger.warning(f"{log_prefix} Resposta API não continha texto utilizável (Tentativa {attempt+1}). Latência: {latency:.2f}s. Finish: {finish_reason}.")

                 except Exception as e_details:
                    logger.error(f"{log_prefix} Erro ao extrair detalhes/texto da resposta API (Tentativa {attempt+1}): {e_details}. Latência: {latency:.2f}s. Resposta: {response}")
                    logger.error(traceback.format_exc())

            # Espera Exponencial se falhou em obter texto nesta tentativa
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                logger.info(f"{log_prefix} Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                 logger.error(f"{log_prefix} Falha final na API após {max_retries} tentativas para o chunk: '{log_chunk_preview}'")
                 return default_return

        except Exception as e: # Erro na chamada da API
            latency = time.time() - start_time
            logger.warning(f"{log_prefix} Erro durante chamada API ({model.model_name}) (Tentativa {attempt + 1}/{max_retries}): {e}. Latência parcial: {latency:.2f}s")
            logger.debug(traceback.format_exc()) # Log completo no debug
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                if "RESOURCE_EXHAUSTED" in str(e) or "429" in str(e):
                    logger.warning(f"{log_prefix} Erro de cota. Aumentando espera.")
                    base_wait_time = max(15, base_wait_time)
                    wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                logger.info(f"{log_prefix} Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                logger.error(f"{log_prefix} Falha final API (erro na chamada) para chunk: '{log_chunk_preview}'")
                return default_return

    logger.error(f"{log_prefix} Loop de tentativas da API concluído sem sucesso explícito.")
    return default_return

# <<< MODIFICADO format_with_ai_correction_only para retornar stats >>>
def format_with_ai_correction_only(model, chunk, author_name, book_name, is_first_chunk=False):
    """
    Chama a API Gemini focando APENAS na correção de OCR/gramática.
    Retorna: (corrected_chunk|None, latency, p_tokens, o_tokens, t_tokens)
    """
    context_start = "Você está formatando o início de um livro." if is_first_chunk else "Você está continuando a formatação de um texto de livro existente."
    ocr_errors_examples = """
        * **Troca de letras similares:** 'rn' vs 'm', 'c' vs 'e', 't' vs 'f', 'l' vs 'i', 'I' vs 'l', 'O' vs '0', 'S' vs '5', 'B' vs '8'.
        * **Hífens indevidos:** Palavras quebradas incorretamente no meio ou hífens extras.
        * **Hífens ausentes:** Palavras que deveriam ser hifenizadas (ex: "guarda-chuva") aparecem juntas ou separadas.
        * **Espaços ausentes ou extras:** Palavras coladas ("onomundo") ou espaços excessivos.
        * **Pontuação incorreta:** Pontos finais trocados por vírgulas, pontos de interrogação/exclamação mal interpretados.
        * **Acentuação:** Falta de acentos (ex: 'e' vs 'é', 'a' vs 'à'), acentos incorretos (crase onde não deve) ou caracteres estranhos no lugar de acentos.
        * **Letras duplicadas ou ausentes:** "caaasa" ou "casaa" em vez de "casa".
        * **Confusão maiúsculas/minúsculas:** Nomes próprios em minúsculas, inícios de frase em minúsculas.
        * **Caracteres especiais/ruído:** Símbolos aleatórios '%', '#', '@' inseridos no texto.
        * **Quebras de linha estranhas:** Parágrafos divididos no meio sem motivo aparente. Preserve as quebras de parágrafo intencionais (duas quebras de linha).
    """
    chunk_prompt = f"""
{context_start} Você é um editor literário proficiente em português do Brasil. Sua tarefa é CORRIGIR e FORMATAR o fragmento de texto a seguir, que pertence a um livro do autor {author_name}.
**CONTEXTO IMPORTANTE:** Este texto provavelmente foi extraído via OCR de um PDF e pode conter erros de reconhecimento, digitação e gramática. O objetivo principal é obter um texto LIMPO e CORRETO em português do Brasil padrão, mantendo a estrutura e o significado originais.
**SIGA RIGOROSAMENTE ESTAS REGRAS:**
1.  **Correção Profunda:** Corrija TODOS os erros gramaticais, ortográficos, de pontuação, acentuação e concordância verbal/nominal. Use o português do Brasil como referência. FOQUE em erros comuns de OCR como os listados abaixo.
2.  **Estilo e Tom:** Mantenha o estilo literário e o tom do texto original do autor {author_name}. Seja claro, fluido e envolvente. NÃO altere o significado, a voz ou a intenção do autor.
3.  **Fidelidade Estrutural:** MANTENHA a estrutura de parágrafos original. Parágrafos são geralmente separados por UMA linha em branco (duas quebras de linha `\\n\\n`). NÃO junte parágrafos que estavam separados. NÃO divida parágrafos desnecessariamente.
4.  **Sem Adições/Remoções:** NÃO omita frases ou informações. NÃO adicione conteúdo, introduções, resumos, conclusões ou opiniões que não estavam no fragmento original. SEJA ESTRITAMENTE FIEL AO CONTEÚDO.
5.  **Marcadores de Capítulo/Quebra:** Se encontrar marcadores como 'Capítulo X', '***', '---', etc., no início de um parágrafo, MANTENHA-OS EXATAMENTE como estão, naquele parágrafo específico. NÃO adicione ou remova esses marcadores.
6.  **Quebra de Página:** Se o marcador '{PAGE_BREAK_MARKER}' aparecer, MANTENHA-O EXATAMENTE onde está, em sua própria linha, sem texto antes ou depois na mesma linha.
7.  **Erros Comuns de OCR (FOCO ESPECIAL):** Preste atenção e corrija diligentemente:
    {ocr_errors_examples}
8.  **Formato de Saída:** Retorne APENAS o texto corrigido e formatado. Use parágrafos separados por duas quebras de linha (`\\n\\n`). NÃO use NENHUMA formatação especial como Markdown (`*`, `#`, `_`), HTML, etc. Retorne TEXTO PURO. Não inclua comentários sobre o que você fez, apenas o texto resultante.
**Texto do fragmento para processar (pode conter erros):**
\"\"\"
{chunk}
\"\"\"
**Lembre-se: O resultado deve ser APENAS o texto corrigido.**
"""
    return _call_gemini_api(model, chunk_prompt, chunk, author_name, book_name)

# <<< MODIFICADO format_with_ai_footnote_only para retornar stats >>>
def format_with_ai_footnote_only(model, chunk, author_name, book_name):
    """
    Chama a API Gemini focando APENAS na identificação de notas.
    Retorna: (marked_chunk|None, latency, p_tokens, o_tokens, t_tokens)
    """
    chunk_prompt = f"""
Você é um assistente de edição trabalhando no texto do autor {author_name}. Sua tarefa é analisar o fragmento de texto A SEGUIR, que JÁ FOI CORRIGIDO no passo anterior, e APENAS inserir marcadores para potenciais notas de rodapé onde estritamente necessário.
**REGRAS IMPORTANTES:**
1.  **NÃO ALTERE O TEXTO CORRIGIDO:** Não faça correções, não mude palavras, não reestruture frases. Apenas insira os marcadores.
2.  **MARCADORES DE NOTA:** Insira marcadores APENAS nos seguintes casos:
    * **Termos em Idioma Estrangeiro (não comuns):** Imediatamente APÓS uma palavra ou frase curta em latim, francês, inglês, etc., que não seja de uso corrente em português, insira: `[NOTA_IDIOMA:palavra_original][CONTEUDO_NOTA:Tradução ou breve explicação]`. Exemplo: "...uma certa *joie de vivre*[NOTA_IDIOMA:joie de vivre][CONTEUDO_NOTA:Alegria de viver (francês)]..."
    * **Citações/Referências:** APÓS uma citação direta curta ou uma referência bibliográfica no texto (ex: (Autor, Ano)), insira: `[NOTA_CITACAO:Texto citado ou referência][CONTEUDO_NOTA:Referência bibliográfica completa ou fonte, se conhecida ou inferível]`. Exemplo: "...como disse Foucault (1975)[NOTA_CITACAO:Foucault (1975)][CONTEUDO_NOTA:FOUCAULT, Michel. Vigiar e Punir. 1975.], a disciplina..."
    * **Nomes Próprios (contexto essencial):** APÓS um nome de pessoa, local ou evento histórico POUCO CONHECIDO que SEJA ESSENCIAL contextualizar brevemente para a compreensão do trecho, insira: `[NOTA_NOME:Nome Mencionado][CONTEUDO_NOTA:Breve identificação (datas, relevância)]`. Use com MODERAÇÃO. Exemplo: "...influenciado por Kropotkin[NOTA_NOME:Kropotkin][CONTEUDO_NOTA:Piotr Kropotkin (1842-1921), anarquista russo.]..."
    * **Termos Técnicos/Jargão (essencial):** APÓS um termo técnico MUITO específico de uma área, cuja definição SEJA INDISPENSÁVEL para o leitor geral entender o argumento naquele ponto, insira: `[NOTA_TERMO:Termo Técnico][CONTEUDO_NOTA:Definição concisa]`. Use com MUITA MODERAÇÃO. Exemplo: "...aplicando a análise de isotopias[NOTA_TERMO:Isotopias][CONTEUDO_NOTA:Na semiótica greimasiana, recorrência de categorias sêmicas que garante a homogeneidade de um discurso.]..."
3.  **FORMATO DOS MARCADORES:** Use EXATAMENTE `[NOTA_TIPO:Referência]` seguido IMEDIATAMENTE por `[CONTEUDO_NOTA:Explicação]`. Não adicione espaços entre eles. Não use outros formatos.
4.  **CRITÉRIO:** Seja conservador. Adicione notas apenas se a informação for realmente útil e provavelmente desconhecida para um leitor culto médio. É MELHOR ERRAR POR NÃO ADICIONAR do que por adicionar excessivamente. NÃO adicione notas para termos comuns, nomes famosos ou citações óbvias.
5.  **NÃO INVENTE CONTEÚDO:** O `[CONTEUDO_NOTA:...]` deve ser uma tradução direta, uma referência óbvia, ou uma contextualização muito breve e factual, se possível inferida do próprio texto ou conhecimento geral básico. NÃO pesquise externamente para criar notas complexas. Se não souber o conteúdo, NÃO insira a nota.
6.  **SAÍDA:** Retorne APENAS o texto original (do input) com os marcadores inseridos nos locais exatos. Mantenha a estrutura de parágrafos (`\\n\\n`). Não adicione NENHUM outro texto, comentário ou explicação.
**Texto JÁ CORRIGIDO para analisar e inserir marcadores de nota:**
\"\"\"
{chunk}
\"\"\"
**Lembre-se: NÃO altere o texto, apenas insira os marcadores `[NOTA_...][CONTEUDO_NOTA:...]` quando apropriado.**
"""
    return _call_gemini_api(model, chunk_prompt, chunk, author_name, book_name)

# --- FUNÇÕES DE PROCESSAMENTO DOS PASSOS ---

def apply_formatting_pass1(doc, formatted_chunk_text, normal_style_name, chapter_patterns, corrected_text_list, author_name, book_name):
    """
    Aplica formatação ao DOCX (Passo 1 - sem notas) e coleta texto para retornar.
    (Função inalterada)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    if not formatted_chunk_text or not formatted_chunk_text.strip(): return
    plain_text_for_list = formatted_chunk_text.replace(PAGE_BREAK_MARKER, "\n\n").strip()
    if plain_text_for_list: corrected_text_list.append(plain_text_for_list)
    normal_style = None
    try:
        if normal_style_name in doc.styles: normal_style = doc.styles[normal_style_name]
    except Exception as e_style: logger.error(f"{log_prefix} Erro ao acessar estilo '{normal_style_name}': {e_style}.")
    chapter_regex = re.compile('|'.join(chapter_patterns), re.IGNORECASE)
    parts = formatted_chunk_text.split(PAGE_BREAK_MARKER)
    content_present_before = any(p.text.strip() for p in doc.paragraphs)
    for part_index, part in enumerate(parts):
        part_clean = part.strip()
        if part_index > 0:
             last_para_is_page_break = False
             if doc.paragraphs:
                 last_p = doc.paragraphs[-1]
                 if not last_p.text.strip() and any(run.text and '\f' in run.text for run in last_p.runs): last_para_is_page_break = True
             if not last_para_is_page_break: doc.add_page_break()
        if not part_clean:
            if part_index > 0 : content_present_before = True
            continue
        paragraphs_in_part = part_clean.split("\n\n")
        for paragraph_text in paragraphs_in_part:
            paragraph_text_clean = paragraph_text.strip()
            if not paragraph_text_clean:
                if doc.paragraphs and doc.paragraphs[-1].text.strip():
                     p = doc.add_paragraph()
                     if normal_style: p.style = normal_style
                continue
            is_ai_failure_marker = paragraph_text_clean.startswith(AI_FAILURE_MARKER)
            is_formatting_error_marker = paragraph_text_clean.startswith(FORMATTING_ERROR_MARKER)
            is_chapter = not is_ai_failure_marker and not is_formatting_error_marker and chapter_regex.match(paragraph_text_clean) is not None
            p = doc.add_paragraph()
            run = p.add_run(paragraph_text_clean)
            content_present_before = True
            try:
                if is_chapter:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.name = 'French Script MT'; run.font.size = Pt(48); run.bold = False
                elif is_ai_failure_marker or is_formatting_error_marker:
                    if normal_style: p.style = normal_style
                    run.font.italic = True; run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif normal_style: p.style = normal_style
            except Exception as e_apply_style: logger.error(f"{log_prefix} Erro ao aplicar estilo/formatação: {e_apply_style}. Texto: '{paragraph_text_clean[:50]}...'")

# <<< MODIFICADO run_correction_pass para acumular e retornar stats >>>
def run_correction_pass(model, input_txt_path, template_docx_path, output_docx_path, author_name, book_name):
    """
    Executa o Passo 1: Corrige texto e gera DOCX base.
    Retorna: (success_bool, corrected_text_str|None, total_latency, total_prompt_tokens, total_output_tokens, total_tokens)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 1: Correção e Geração DOCX ---")
    default_return = (False, None, 0, 0, 0, 0) # Retorno padrão em caso de falha
    try:
        with open(input_txt_path, "r", encoding="utf-8") as f: texto_original = f.read()
        logger.info(f"{log_prefix} Entrada '{os.path.basename(input_txt_path)}' carregada ({len(texto_original)} chars).")
    except Exception as e: logger.error(f"{log_prefix} FATAL ao ler entrada '{input_txt_path}': {e}"); return default_return

    output_docx_dir = os.path.dirname(output_docx_path); os.makedirs(output_docx_dir, exist_ok=True)
    logger.info(f"{log_prefix} Dividindo texto original em chunks...")
    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS, author_name, book_name)
    if not text_chunks: logger.error(f"{log_prefix} Nenhum chunk gerado. Abortando Passo 1."); return default_return

    doc = None; logger.info(f"{log_prefix} Preparando DOCX: {os.path.basename(output_docx_path)}")
    try:
        if os.path.exists(output_docx_path): # Backup logic
             backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
             backup_docx_path = os.path.join(output_docx_dir, f"backup_{os.path.splitext(os.path.basename(output_docx_path))[0]}_{backup_timestamp}.docx")
             try: shutil.copy2(output_docx_path, backup_docx_path); logger.info(f"{log_prefix} Backup DOCX criado: {os.path.basename(backup_docx_path)}")
             except Exception as e_bkp: logger.warning(f"{log_prefix} Falha ao criar backup: {e_bkp}")
        if not os.path.exists(template_docx_path): # Check template
            logger.error(f"{log_prefix} FATAL: Template '{template_docx_path}' não encontrado."); return default_return
        try: # Load and clear template
            doc = Document(template_docx_path); logger.info(f"{log_prefix} Template '{os.path.basename(template_docx_path)}' carregado.")
            for para in doc.paragraphs: p_element = para._element; p_element.getparent().remove(p_element)
            for table in doc.tables: tbl_element = table._element; tbl_element.getparent().remove(tbl_element)
            logger.info(f"{log_prefix} Conteúdo principal do template limpo.")
            if NORMAL_STYLE_NAME in doc.styles: logger.info(f"{log_prefix} Estilo '{NORMAL_STYLE_NAME}' encontrado.")
            else: logger.warning(f"{log_prefix} AVISO: Estilo '{NORMAL_STYLE_NAME}' NÃO encontrado.")
        except Exception as e_load_template: logger.error(f"{log_prefix} FATAL: Falha ao carregar/limpar template: {e_load_template}."); return default_return
    except Exception as e_doc: logger.error(f"{log_prefix} FATAL: Erro crítico ao preparar DOCX: {e_doc}"); return default_return

    logger.info(f"{log_prefix} Iniciando chamadas à API para CORREÇÃO de {len(text_chunks)} chunks...")
    corrected_text_list_pass1 = []
    processed_chunks_count = 0; failed_chunks_count = 0
    total_latency_pass1 = 0; total_prompt_tokens_pass1 = 0; total_output_tokens_pass1 = 0; total_tokens_pass1 = 0

    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P1: Corrigindo", unit="chunk", leave=False)
    for i, chunk in progress_bar:
        corrected_chunk, latency, p_tokens, o_tokens, t_tokens = format_with_ai_correction_only(
            model, chunk, author_name, book_name, is_first_chunk=(i == 0)
        )
        total_latency_pass1 += latency; total_prompt_tokens_pass1 += p_tokens; total_output_tokens_pass1 += o_tokens; total_tokens_pass1 += t_tokens
        if corrected_chunk is not None:
             try: apply_formatting_pass1(doc, corrected_chunk, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name); processed_chunks_count += 1
             except Exception as format_err:
                logger.error(f"{log_prefix} Erro apply_formatting_pass1 (Chunk {i+1}): {format_err}."); failed_chunks_count += 1
                try: fallback_text = f"{FORMATTING_ERROR_MARKER}\n\n{chunk}"; apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
                except Exception as ff_err: logger.critical(f"{log_prefix} Falha CRÍTICA fallback FORMATAÇÃO: {ff_err}.")
        else:
            logger.warning(f"{log_prefix} Chunk {i+1} falhou na CORREÇÃO (API). Usando fallback.")
            failed_chunks_count += 1
            try: fallback_text = f"{AI_FAILURE_MARKER}\n\n{chunk}"; apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_PATTERNS, corrected_text_list_pass1, author_name, book_name)
            except Exception as ff_err: logger.critical(f"{log_prefix} Falha CRÍTICA fallback API: {ff_err}.")

        processed_total = processed_chunks_count + failed_chunks_count
        if processed_total > 0 and (processed_total % 10 == 0 or (i + 1) == len(text_chunks)):
            temp_save_path = f"{output_docx_path}.{processed_total}.temp_save"
            try: doc.save(temp_save_path); shutil.move(temp_save_path, output_docx_path); logger.info(f"{log_prefix} Progresso DOCX salvo ({processed_total} chunks).")
            except Exception as e_save: logger.error(f"{log_prefix} Erro salvar progresso DOCX: {e_save}")

    try: logger.info(f"{log_prefix} Salvando DOCX final (Passo 1): {os.path.basename(output_docx_path)}"); doc.save(output_docx_path)
    except Exception as e_final_save: logger.error(f"{log_prefix} Erro salvamento final DOCX (Passo 1): {e_final_save}")

    full_corrected_text = "\n\n".join(corrected_text_list_pass1)
    #logger.info(f"{log_prefix} Acumulado texto corrigido para Pass 2 ({len(full_corrected_text)} chars).") # Menos verboso
    logger.info(f"{log_prefix} --- Passo 1 concluído. Chunks OK: {processed_chunks_count}, Falhas: {failed_chunks_count} ---")
    logger.info(f"{log_prefix} Stats API P1: Lat:{total_latency_pass1:.2f}s, Toks:{total_tokens_pass1}(P:{total_prompt_tokens_pass1},O:{total_output_tokens_pass1})")
    # Retorna sucesso (mesmo com falhas parciais, pois o processo tentou), texto e estatísticas
    success_overall = failed_chunks_count == 0 # Considera sucesso geral apenas se NENHUM chunk falhou? Ou se >0 ok? Vamos considerar True se rodou.
    return (True, full_corrected_text, total_latency_pass1, total_prompt_tokens_pass1, total_output_tokens_pass1, total_tokens_pass1)

# <<< MODIFICADO run_footnote_id_pass para acumular e retornar stats >>>
def run_footnote_id_pass(model, corrected_text_content, author_name, book_name):
    """
    Executa o Passo 2: Identifica notas no texto já corrigido.
    Retorna: (success_bool, marked_text_str|None, total_latency, total_prompt_tokens, total_output_tokens, total_tokens)
    """
    log_prefix = f"[{author_name}/{book_name}]"
    logger.info(f"{log_prefix} --- Iniciando Passo 2: Identificação de Notas ---")
    default_return = (False, None, 0, 0, 0, 0)
    if corrected_text_content is None: logger.error(f"{log_prefix} Input None. Abortando P2."); return default_return

    logger.info(f"{log_prefix} Dividindo texto corrigido em chunks...")
    text_chunks = create_chunks(corrected_text_content, MAX_CHUNK_TOKENS, author_name, book_name)
    if not text_chunks: logger.error(f"{log_prefix} Nenhum chunk gerado. Abortando P2."); return default_return

    logger.info(f"{log_prefix} Iniciando API ID Notas em {len(text_chunks)} chunks...")
    marked_text_list_pass2 = []
    processed_chunks_count = 0; failed_chunks_count = 0
    total_latency_pass2 = 0; total_prompt_tokens_pass2 = 0; total_output_tokens_pass2 = 0; total_tokens_pass2 = 0

    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P2: Notas", unit="chunk", leave=False)
    for i, chunk in progress_bar:
        marked_chunk, latency, p_tokens, o_tokens, t_tokens = format_with_ai_footnote_only( model, chunk, author_name, book_name )
        total_latency_pass2 += latency; total_prompt_tokens_pass2 += p_tokens; total_output_tokens_pass2 += o_tokens; total_tokens_pass2 += t_tokens
        if marked_chunk is not None: marked_text_list_pass2.append(marked_chunk); processed_chunks_count += 1
        else: logger.warning(f"{log_prefix} Chunk {i+1} falhou ID Notas (API). Usando original."); marked_text_list_pass2.append(chunk); failed_chunks_count += 1

    full_marked_text = "\n\n".join(marked_text_list_pass2)
    #logger.info(f"{log_prefix} Acumulado texto com marcadores para Pass 3 ({len(full_marked_text)} chars).") # Menos verboso
    logger.info(f"{log_prefix} --- Passo 2 concluído. Chunks OK: {processed_chunks_count}, Falhas: {failed_chunks_count} ---")
    logger.info(f"{log_prefix} Stats API P2: Lat:{total_latency_pass2:.2f}s, Toks:{total_tokens_pass2}(P:{total_prompt_tokens_pass2},O:{total_output_tokens_pass2})")
    # Retorna sucesso (mesmo com falhas parciais), texto e estatísticas acumuladas
    success_overall = failed_chunks_count == 0
    return (True, full_marked_text, total_latency_pass2, total_prompt_tokens_pass2, total_output_tokens_pass2, total_tokens_pass2)

# run_final_txt_generation permanece inalterado
def run_final_txt_generation(marked_text_content, output_notes_path, output_numbered_txt_path, author_name, book_name):
    """
    Executa o Passo 3: Processa marcadores [NOTA_...] para gerar TXTs. (Inalterado)
    Retorna: bool -> success_status
    """
    log_prefix = f"[{author_name}/{book_name}]"
    #logger.info(f"{log_prefix} --- Iniciando Passo 3: Geração TXTs Finais ---") # Menos verboso
    if marked_text_content is None: logger.error(f"{log_prefix} Input None. Abortando P3."); return False
    os.makedirs(os.path.dirname(output_notes_path), exist_ok=True)
    os.makedirs(os.path.dirname(output_numbered_txt_path), exist_ok=True)
    footnote_counter = 1; notes_found = []
    footnote_pattern = re.compile(r'(\[NOTA_(?:IDIOMA|CITACAO|NOME|TERMO):[^\]]+?\])\s*(\[CONTEUDO_NOTA:([^\]]*?)\])', re.IGNORECASE)
    def replace_marker_and_collect_note(match):
        nonlocal footnote_counter; content = match.group(3).strip()
        if not content: logger.warning(f"{log_prefix} CONTEUDO_NOTA vazio: {match.group(1)}"); return ""
        notes_found.append(f"{footnote_counter}. {content}"); replacement = f"[{footnote_counter}]"; footnote_counter += 1; return replacement
    #logger.info(f"{log_prefix} Processando marcadores e gerando arquivos finais TXT...")
    try:
        final_numbered_text = footnote_pattern.sub(replace_marker_and_collect_note, marked_text_content)
        #logger.info(f"{log_prefix} Salvando notas: {os.path.basename(output_notes_path)}")
        with open(output_notes_path, "w", encoding="utf-8") as f_notes:
            f_notes.write(f"Notas para {author_name} - {book_name}\n" + "=" * 30 + "\n\n")
            if notes_found: f_notes.write("\n".join(notes_found) + "\n"); #logger.info(f"{log_prefix} {len(notes_found)} notas salvas.")
            else: f_notes.write("(Nenhuma nota identificada)\n"); #logger.info(f"{log_prefix} Nenhuma nota salva.")
        #logger.info(f"{log_prefix} Salvando TXT numerado [{footnote_counter-1}]: {os.path.basename(output_numbered_txt_path)}")
        with open(output_numbered_txt_path, "w", encoding="utf-8") as f_numbered: f_numbered.write(final_numbered_text)
        #logger.info(f"{log_prefix} TXT numerado salvo ({len(final_numbered_text)} chars).")
    except Exception as e_final_gen: logger.error(f"{log_prefix} Erro Geração TXT (Passo 3): {e_final_gen}"); return False
    logger.info(f"{log_prefix} --- Passo 3 concluído. ---")
    return True

# --- Funções para Gerenciar Logs de Processados (Corrigidas)---

def load_processed_files(filepath):
    """Lê o arquivo de log de CORREÇÃO e retorna um set."""
    processed = set()
    try: # << CORRIGIDO: try: em nova linha >>
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                cleaned = line.strip()
                if cleaned: processed.add(cleaned)
        logger.info(f"Carregados {len(processed)} registros de CORREÇÕES concluídas de '{filepath}'.")
    except FileNotFoundError: logger.info(f"Log de correções '{filepath}' não encontrado.")
    except Exception as e: logger.error(f"Erro ao carregar log de correções '{filepath}': {e}")
    return processed

def log_processed_file(filepath, file_identifier):
    """Adiciona um identificador ao log de CORREÇÃO."""
    try:
        with open(filepath, 'a', encoding='utf-8') as f: f.write(f"{file_identifier}\n")
    except Exception as e: logger.error(f"Erro ao registrar '{file_identifier}' no log de correções '{filepath}': {e}")

def load_translated_files(filepath):
    """Lê o arquivo de log de TRADUÇÃO e retorna um set."""
    processed = set()
    try: # << CORRIGIDO: try: em nova linha >>
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                cleaned = line.strip()
                if cleaned: processed.add(cleaned)
        logger.info(f"Carregados {len(processed)} registros de TRADUÇÕES concluídas de '{filepath}'.")
    except FileNotFoundError: logger.info(f"Log de traduções '{filepath}' não encontrado.")
    except Exception as e: logger.error(f"Erro ao carregar log de traduções '{filepath}': {e}")
    return processed

def log_translated_file(filepath, file_identifier):
    """Adiciona um identificador ao log de TRADUÇÃO."""
    try:
        with open(filepath, 'a', encoding='utf-8') as f: f.write(f"{file_identifier}\n")
        logger.debug(f"Registrado '{file_identifier}' como TRADUZIDO com sucesso em '{filepath}'.")
    except Exception as e: logger.error(f"Erro ao registrar '{file_identifier}' no log de traduções '{filepath}': {e}")


# --- FUNÇÃO DE ENVIO DE E-MAIL (MODIFICADA para incluir mais stats) ---
def send_completion_email(sender_email, sender_password, recipient_email, smtp_server, smtp_port,
                          processed_correction, skipped_correction, failed_correction, # Contadores Correção
                          processed_translation, skipped_translation, failed_translation, # Contadores Tradução
                          total_duration_seconds,
                          main_log_path, processed_log_path, translated_log_path,
                          # << NOVOS Parâmetros de Stats >>
                          total_correction_latency_secs, total_correction_tokens,
                          total_footnote_latency_secs, total_footnote_tokens,
                          avg_correction_time_secs, total_correction_time_secs,
                          avg_translation_time_secs, total_translation_time_secs,
                          # << NOVAS Listas de Livros >>
                          processed_correction_books, skipped_correction_books, failed_correction_books,
                          processed_translation_books, skipped_translation_books, failed_translation_books
                          ):
    """Envia um e-mail de notificação de conclusão com resumo detalhado e estatísticas."""

    global email_configured # Acessa a flag global definida no início
    if not email_configured: # Checa a flag em vez de re-ler variáveis
        logger.warning("Envio de e-mail desativado (configuração inicial incompleta).")
        return

    logger.info(f"Preparando e-mail de notificação para {recipient_email}...")

    subject = "Script Processador Livros (Correção+Tradução) - Concluído"
    body = f"""
Olá,

O script de processamento de livros (Correção + Tradução Híbrida) concluiu a execução.

Resumo Geral:
--------------------------------------------------
- Tempo Total de Execução: {total_duration_seconds:.2f} seg ({total_duration_seconds/60:.2f} min)

Resumo Etapa de Correção:
--------------------------------------------------
- Livros Corrigidos OK (nesta execução): {processed_correction}
- Livros Pulados (correção já feita): {skipped_correction}
- Livros com Falha na Correção: {failed_correction}
"""
    if processed_correction > 0:
        body += f"- Tempo Total Correção (livros OK): {total_correction_time_secs:.2f} seg ({total_correction_time_secs/60:.2f} min)\n"
        body += f"- Tempo Médio por Correção: {avg_correction_time_secs:.2f} seg\n"
        # Adiciona stats da API para Correção (Passo 1)
        body += f"- Tempo Total API (Passo 1 Correção): {total_correction_latency_secs:.2f} seg ({total_correction_latency_secs/60:.2f} min)\n"
        body += f"- Tokens Totais API (Passo 1 Correção): {total_correction_tokens}\n"
        # Adiciona stats da API para Notas (Passo 2)
        body += f"- Tempo Total API (Passo 2 Notas): {total_footnote_latency_secs:.2f} seg ({total_footnote_latency_secs/60:.2f} min)\n"
        body += f"- Tokens Totais API (Passo 2 Notas): {total_footnote_tokens}\n"


    body += f"""
Resumo Etapa de Tradução Híbrida:
--------------------------------------------------
- Livros Traduzidos OK (nesta execução): {processed_translation}
- Livros Pulados (tradução já feita): {skipped_translation}
- Livros com Falha na Tradução: {failed_translation}
"""
    if processed_translation > 0:
        body += f"- Tempo Total Tradução (livros OK): {total_translation_time_secs:.2f} seg ({total_translation_time_secs/60:.2f} min)\n"
        body += f"- Tempo Médio por Tradução: {avg_translation_time_secs:.2f} seg\n"
        # Stats API do tradutor não são coletados aqui

    # --- Listas de Livros ---
    # Adiciona as listas ao corpo do e-mail, se não estiverem vazias
    body += "\n--------------------------------------------------\n"
    body += "Detalhes por Livro:\n"
    body += "--------------------------------------------------\n"

    if processed_correction_books: body += f"\nLivros Corrigidos OK ({len(processed_correction_books)}):\n - " + "\n - ".join(processed_correction_books) + "\n"
    if skipped_correction_books: body += f"\nLivros Pulados na Correção ({len(skipped_correction_books)}):\n - " + "\n - ".join(skipped_correction_books) + "\n"
    if failed_correction_books: body += f"\nLivros com Falha na Correção ({len(failed_correction_books)}):\n - " + "\n - ".join(failed_correction_books) + "\n"

    if processed_translation_books: body += f"\nLivros Traduzidos OK ({len(processed_translation_books)}):\n - " + "\n - ".join(processed_translation_books) + "\n"
    if skipped_translation_books: body += f"\nLivros Pulados na Tradução ({len(skipped_translation_books)}):\n - " + "\n - ".join(skipped_translation_books) + "\n"
    if failed_translation_books: body += f"\nLivros com Falha na Tradução ({len(failed_translation_books)}):\n - " + "\n - ".join(failed_translation_books) + "\n"

    # Aviso sobre tamanho potencial
    total_listed = len(processed_correction_books) + len(skipped_correction_books) + len(failed_correction_books) + \
                   len(processed_translation_books) + len(skipped_translation_books) + len(failed_translation_books)
    if total_listed > 50: # Limiar arbitrário
        body += "\n(Nota: Listas de livros podem estar longas. Consulte os logs para detalhes completos.)\n"

    body += f"""
--------------------------------------------------

Logs para Consulta:
- Log Detalhado da Execução: {os.path.abspath(main_log_path)}
- Log de Correções Concluídas: {os.path.abspath(processed_log_path)}
- Log de Traduções Concluídas: {os.path.abspath(translated_log_path)}

Atenciosamente,
Seu Script Processador de Livros
"""

    message = EmailMessage()
    message['Subject'] = subject
    message['From'] = sender_email
    message['To'] = recipient_email
    message.set_content(body)

    context = ssl.create_default_context()

    try:
        server = None
        logger.info(f"Conectando ao servidor SMTP: {smtp_server}:{smtp_port}...")
        if smtp_port == 465:
             server = smtplib.SMTP_SSL(smtp_server, smtp_port, context=context, timeout=30)
             server.login(sender_email, sender_password)
        else:
            server = smtplib.SMTP(smtp_server, smtp_port, timeout=30)
            server.ehlo(); server.starttls(context=context); server.ehlo()
            server.login(sender_email, sender_password)

        logger.info("Enviando e-mail de resumo final...")
        server.send_message(message)
        logger.info(f"✅ E-mail de resumo final enviado com sucesso para {recipient_email}.")

    except smtplib.SMTPAuthenticationError:
        logger.error("FALHA NA AUTENTICAÇÃO do e-mail. Verifique .env e Senha de App.")
    except Exception as e:
        logger.error(f"ERRO ao enviar e-mail de resumo final: {e}")
        logger.debug(traceback.format_exc())
    finally:
        if server:
            try: server.quit()
            except Exception: pass


# --- FUNÇÃO PRINCIPAL (main - CORRIGIDA PARA BUSCA RECURSIVA E PATHS) ---
def main():
    start_time_main = time.time()
    logger.info("========================================================")
    logger.info(f"Iniciando Processador (Correção + Tradução) - {time.strftime('%Y-%m-%d %H:%M:%S')}")
    logger.info(f"Diretório Entrada: {BASE_INPUT_TXT_DIR}, Saída DOCX: {BASE_OUTPUT_DOCX_DIR}, Saída TXT: {BASE_OUTPUT_TXT_DIR}")
    logger.info(f"Template: {TEMPLATE_DOCX}, Script Tradutor: {PATH_TO_TRANSLATOR_SCRIPT}")
    logger.info(f"Log Correção: {PROCESSED_LOG_FILE}, Log Tradução: {TRANSLATED_LOG_FILE}")
    logger.info("========================================================")

    processed_files_set = load_processed_files(PROCESSED_LOG_FILE)
    translated_files_set = load_translated_files(TRANSLATED_LOG_FILE)

    if not os.path.isdir(BASE_INPUT_TXT_DIR): logger.error(f"FATAL: Diretório '{BASE_INPUT_TXT_DIR}' não encontrado!"); return
    try: author_folders = sorted([f for f in os.listdir(BASE_INPUT_TXT_DIR) if os.path.isdir(os.path.join(BASE_INPUT_TXT_DIR, f))])
    except Exception as e: logger.error(f"FATAL: Erro ao listar autores: {e}"); return
    if not author_folders: logger.warning(f"Nenhuma pasta de autor encontrada."); return
    logger.info(f"Autores encontrados ({len(author_folders)}): {', '.join(author_folders)}")

    # --- Inicializa Contadores e Acumuladores de Stats ---
    total_books_processed_correction = 0; total_books_skipped_correction = 0; total_books_failed_correction = 0
    total_translation_processed = 0; total_translation_skipped = 0; total_translation_failed = 0
    grand_total_correction_latency = 0; grand_total_correction_prompt_tokens = 0; grand_total_correction_output_tokens = 0; grand_total_correction_total_tokens = 0
    grand_total_footnote_latency = 0; grand_total_footnote_prompt_tokens = 0; grand_total_footnote_output_tokens = 0; grand_total_footnote_total_tokens = 0
    correction_times = []
    translation_times = []
    processed_correction_list = []; skipped_correction_list = []; failed_correction_list = []
    processed_translation_list = []; skipped_translation_list = []; failed_translation_list = []

    # === LOOP PRINCIPAL: AUTOR ===
    for author_name in author_folders:
        author_input_dir = os.path.join(BASE_INPUT_TXT_DIR, author_name)
        # O diretório de saída base do autor ainda é útil para logs gerais do autor
        # author_output_docx_dir_base = os.path.join(BASE_OUTPUT_DOCX_DIR, author_name) # Removido pois não usado diretamente
        # author_output_txt_dir_base = os.path.join(BASE_OUTPUT_TXT_DIR, author_name) # Removido pois não usado diretamente
        logger.info(f"--- Verificando Autor: {author_name} em '{author_input_dir}' ---")

        # <<< MODIFICAÇÃO: Busca Recursiva por TXT >>>
        try:
            # Procura por *.txt em author_input_dir E TODAS as subpastas
            search_pattern = os.path.join(author_input_dir, '**', '*.txt')
            # IMPORTANTE: glob.glob retorna caminhos completos
            input_txt_files_found = sorted(glob.glob(search_pattern, recursive=True))

            # Filtra arquivos que parecem ser de saída ou backup
            input_txt_files = [
                f for f in input_txt_files_found if not (
                    os.path.basename(f).endswith(FINAL_NUMBERED_TXT_BASENAME) or
                    os.path.basename(f).endswith(NOTES_TXT_FILE_BASENAME) or
                    os.path.basename(f).startswith("backup_")
                    # Poderia adicionar mais filtros se necessário
                )
            ]
            logger.info(f"[{author_name}] Encontrados {len(input_txt_files_found)} arquivos .txt (antes de filtrar).")
        except Exception as e:
            logger.error(f"[{author_name}] Erro ao buscar arquivos .txt recursivamente: {e}")
            continue # Pula para o próximo autor

        if not input_txt_files:
            logger.warning(f"[{author_name}] Nenhum arquivo .txt de entrada válido encontrado (incluindo subpastas).")
            continue

        logger.info(f"[{author_name}] Processando {len(input_txt_files)} arquivos .txt válidos.")
        # <<< FIM MODIFICAÇÃO BUSCA >>>

        # === LOOP INTERNO: LIVRO ===
        for input_txt_path in input_txt_files: # input_txt_path AGORA é o caminho completo

            # <<< MODIFICAÇÃO: Tratamento de Caminho Relativo e Saída >>>
            try:
                # Cria um identificador único baseado no caminho relativo a BASE_INPUT_TXT_DIR
                # Ex: 'Nietzsche/Zaratustra/Livro.txt'
                relative_path = os.path.relpath(input_txt_path, BASE_INPUT_TXT_DIR)
                file_identifier = relative_path.replace('\\', '/') # Garante / como separador

                # Extrai partes do caminho relativo para construir saídas
                path_parts = file_identifier.split('/')
                author_name_from_path = path_parts[0] # Ex: 'Nietzsche'
                book_subpath_parts = path_parts[1:-1] # Ex: ['Zaratustra'] ou [] se não houver subpasta
                book_filename = path_parts[-1]         # Ex: 'Livro.txt'
                base_book_name = os.path.splitext(book_filename)[0] # Ex: 'Livro'

                # Constrói diretórios de saída ESPELHANDO a estrutura de entrada
                # Ex: os.path.join('Zaratustra') -> 'Zaratustra'
                # Ex: os.path.join() -> '' (para arquivos direto na pasta do autor)
                book_subdir_rel = os.path.join(*book_subpath_parts)

                # Diretórios de saída completos para este livro específico
                author_output_docx_book_dir = os.path.join(BASE_OUTPUT_DOCX_DIR, author_name_from_path, book_subdir_rel)
                author_output_txt_book_dir = os.path.join(BASE_OUTPUT_TXT_DIR, author_name_from_path, book_subdir_rel)

                # Cria os diretórios de saída específicos do livro/subpasta, se não existirem
                os.makedirs(author_output_docx_book_dir, exist_ok=True)
                os.makedirs(author_output_txt_book_dir, exist_ok=True)

                # Constrói nomes dos arquivos de saída DENTRO das subpastas corretas
                output_docx_path = os.path.join(author_output_docx_book_dir, f"{base_book_name}_{FINAL_DOCX_BASENAME}")
                output_notes_path = os.path.join(author_output_txt_book_dir, f"{base_book_name}_{NOTES_TXT_FILE_BASENAME}")
                output_numbered_txt_path = os.path.join(author_output_txt_book_dir, f"{base_book_name}_{FINAL_NUMBERED_TXT_BASENAME}")
                translated_docx_path = os.path.join(author_output_docx_book_dir, f"{base_book_name}_{FINAL_DOCX_BASENAME.replace('.docx', TRANSLATED_DOCX_SUFFIX)}")

            except Exception as e_path:
                logger.error(f"Erro ao processar caminhos para '{input_txt_path}': {e}")
                logger.error(traceback.format_exc())
                continue # Pula para o próximo livro em caso de erro no path
            # <<< FIM MODIFICAÇÃO PATHS >>>

            log_prefix_book = f"[{file_identifier}]" # Usa identificador mais completo
            logger.info(f"--------------------------------------------------------")
            logger.info(f"{log_prefix_book} Processando Livro...")

            correction_successful_this_run = False
            # output_docx_path já definido acima

            # --- Verifica Status da CORREÇÃO ---
            #logger.info(f"{log_prefix_book} Verificando status (Correção)...") # Menos verboso
            if file_identifier in processed_files_set:
                logger.info(f"{log_prefix_book} CORREÇÃO já feita. Pulando etapa.")
                total_books_skipped_correction += 1
                skipped_correction_list.append(file_identifier)
                correction_successful_this_run = True
                # output_docx_path já foi definido
            else:
                # --- Executa a CORREÇÃO ---
                logger.info(f"{log_prefix_book} Iniciando processamento (Correção)...")
                book_start_time = time.time()
                # Paths já definidos

                all_correction_steps_successful = True
                book_corr_latency=0; book_corr_p_tokens=0; book_corr_o_tokens=0; book_corr_t_tokens=0
                book_note_latency=0; book_note_p_tokens=0; book_note_o_tokens=0; book_note_t_tokens=0

                try:
                    # PASSO 1: CORREÇÃO DOCX
                    pass1_success, corrected_text_content, lat1, p1, o1, t1 = run_correction_pass(
                        gemini_model, input_txt_path, TEMPLATE_DOCX, output_docx_path,
                        author_name_from_path, base_book_name # Usa nomes extraídos do path
                    )
                    book_corr_latency+=lat1; book_corr_p_tokens+=p1; book_corr_o_tokens+=o1; book_corr_t_tokens+=t1
                    if not pass1_success or corrected_text_content is None: all_correction_steps_successful = False
                    else:
                        # PASSO 2: IDENTIFICAÇÃO DE NOTAS
                        pass2_success, marked_text_content, lat2, p2, o2, t2 = run_footnote_id_pass(
                            gemini_model, corrected_text_content,
                            author_name_from_path, base_book_name # Usa nomes extraídos
                        )
                        book_note_latency+=lat2; book_note_p_tokens+=p2; book_note_o_tokens+=o2; book_note_t_tokens+=t2
                        if not pass2_success or marked_text_content is None: all_correction_steps_successful = False
                        else:
                            # PASSO 3: GERAÇÃO FINAL TXT
                            pass3_success = run_final_txt_generation(
                                marked_text_content, output_notes_path, output_numbered_txt_path,
                                author_name_from_path, base_book_name # Usa nomes extraídos
                            )
                            if not pass3_success: all_correction_steps_successful = False
                except Exception as e_corr_steps:
                     logger.error(f"{log_prefix_book} Erro inesperado CORREÇÃO: {e_corr_steps}"); logger.error(traceback.format_exc()); all_correction_steps_successful = False

                book_end_time = time.time()
                book_total_time = book_end_time - book_start_time

                if all_correction_steps_successful:
                    logger.info(f"✅ {log_prefix_book} Etapa de CORREÇÃO SUCESSO em {book_total_time:.2f} seg.")
                    log_processed_file(PROCESSED_LOG_FILE, file_identifier); processed_files_set.add(file_identifier)
                    total_books_processed_correction += 1; correction_successful_this_run = True
                    processed_correction_list.append(file_identifier)
                    correction_times.append(book_total_time)
                    # Acumula stats GERAIS
                    grand_total_correction_latency += book_corr_latency; grand_total_correction_prompt_tokens += book_corr_p_tokens; grand_total_correction_output_tokens += book_corr_o_tokens; grand_total_correction_total_tokens += book_corr_t_tokens
                    grand_total_footnote_latency += book_note_latency; grand_total_footnote_prompt_tokens += book_note_p_tokens; grand_total_footnote_output_tokens += book_note_o_tokens; grand_total_footnote_total_tokens += book_note_t_tokens
                else:
                    logger.warning(f"⚠️ {log_prefix_book} Etapa de CORREÇÃO FALHAS em {book_total_time:.2f} seg.")
                    total_books_failed_correction += 1; correction_successful_this_run = False
                    failed_correction_list.append(file_identifier)

            # --- ETAPA DE TRADUÇÃO ---
            if correction_successful_this_run:
                #logger.info(f"{log_prefix_book} Verificando status (Tradução)...") # Menos verboso
                # Usa output_docx_path e translated_docx_path definidos acima
                if output_docx_path is None or not os.path.exists(output_docx_path):
                     logger.warning(f"{log_prefix_book} Input DOCX tradução não encontrado. Pulando.");
                     total_translation_failed += 1
                     failed_translation_list.append(f"{file_identifier} (Input DOCX ausente)")
                elif file_identifier in translated_files_set:
                    logger.info(f"{log_prefix_book} TRADUÇÃO já feita. Pulando.");
                    total_translation_skipped += 1
                    skipped_translation_list.append(file_identifier)
                else:
                    logger.info(f"{log_prefix_book} >>> Iniciando TRADUÇÃO HÍBRIDA...")
                    translation_start_time = time.time()
                    if not os.path.exists(PATH_TO_TRANSLATOR_SCRIPT):
                         logger.error(f"{log_prefix_book} ERRO CRÍTICO: Script tradutor '{PATH_TO_TRANSLATOR_SCRIPT}' não encontrado.");
                         total_translation_failed += 1
                         failed_translation_list.append(f"{file_identifier} (Script tradutor não encontrado)")
                    else:
                        try:
                            command = [ sys.executable, PATH_TO_TRANSLATOR_SCRIPT, '--input', output_docx_path, '--output', translated_docx_path, '--words', str(NUM_WORDS_TO_TRANSLATE) ]
                            logger.info(f"{log_prefix_book} Executando: {' '.join(command)}")
                            result = subprocess.run(command, capture_output=True, text=True, encoding='utf-8', check=False)
                            translation_end_time = time.time()
                            translation_total_time = translation_end_time - translation_start_time
                            if result.returncode == 0:
                                logger.info(f"✅ {log_prefix_book} TRADUÇÃO HÍBRIDA SUCESSO em {translation_total_time:.2f} seg.")
                                log_translated_file(TRANSLATED_LOG_FILE, file_identifier); translated_files_set.add(file_identifier)
                                total_translation_processed += 1
                                processed_translation_list.append(file_identifier)
                                translation_times.append(translation_total_time)
                                if result.stdout: logger.debug(f"{log_prefix_book} Saída tradutor:\n{result.stdout}")
                            else:
                                logger.error(f"❌ {log_prefix_book} TRADUÇÃO HÍBRIDA FALHOU (código: {result.returncode}) em {translation_total_time:.2f} seg.")
                                if result.stderr: logger.error(f"{log_prefix_book} Erro tradutor:\n{result.stderr}")
                                else: logger.error(f"{log_prefix_book} Tradutor não reportou erro específico.")
                                total_translation_failed += 1
                                failed_translation_list.append(f"{file_identifier} (Erro script: {result.returncode})")
                        except Exception as e_translate_sub:
                             logger.error(f"{log_prefix_book} Erro CRÍTICO subprocesso tradução: {e_translate_sub}"); logger.error(traceback.format_exc());
                             total_translation_failed += 1
                             failed_translation_list.append(f"{file_identifier} (Exceção subprocesso)")
            # --- Fim Tradução ---
            logger.info(f"{log_prefix_book} --- Fim processamento livro ---")
        # Fim loop livros
        logger.info(f"--- Concluída verificação Autor: {author_name} ---")
        # logger.info(f"--------------------------------------------------------") # Menos verboso
    # --- Fim loop autores ---

    end_time_main = time.time()
    total_time_main = end_time_main - start_time_main

    # --- Cálculos Finais de Tempo e Stats ---
    total_corr_time_ok = sum(correction_times)
    avg_corr_time_ok = total_corr_time_ok / len(correction_times) if correction_times else 0
    total_trans_time_ok = sum(translation_times)
    avg_trans_time_ok = total_trans_time_ok / len(translation_times) if translation_times else 0

    # --- Resumo Final Logging ---
    logger.info("===================== RESUMO FINAL =====================")
    logger.info(f"Tempo total geral: {total_time_main:.2f} seg ({total_time_main/60:.2f} min).")
    # Correção
    logger.info("--- Resumo Etapa de Correção ---")
    logger.info(f"Livros Corrigidos OK: {total_books_processed_correction}")
    logger.info(f"Livros Pulados (já corrigidos): {total_books_skipped_correction}")
    logger.info(f"Livros com Falha na Correção: {total_books_failed_correction}")
    if correction_times:
        logger.info(f"Tempo Total Correção (livros OK): {total_corr_time_ok:.2f} seg ({total_corr_time_ok/60:.2f} min)")
        logger.info(f"Tempo Médio por Correção: {avg_corr_time_ok:.2f} seg")
    # Stats API Correção (Passo 1 + Passo 2)
    logger.info(f"API Correção (Passo 1) - Latência Total: {grand_total_correction_latency:.2f}s / Tokens Totais: {grand_total_correction_total_tokens} (P: {grand_total_correction_prompt_tokens}, O: {grand_total_correction_output_tokens})")
    logger.info(f"API Notas (Passo 2)    - Latência Total: {grand_total_footnote_latency:.2f}s / Tokens Totais: {grand_total_footnote_total_tokens} (P: {grand_total_footnote_prompt_tokens}, O: {grand_total_footnote_output_tokens})")
    # Tradução
    logger.info("--- Resumo Etapa de Tradução ---")
    logger.info(f"Livros Traduzidos OK: {total_translation_processed}")
    logger.info(f"Livros Pulados (já traduzidos): {total_translation_skipped}")
    logger.info(f"Livros com Falha na Tradução: {total_translation_failed}")
    if translation_times:
        logger.info(f"Tempo Total Tradução (livros OK): {total_trans_time_ok:.2f} seg ({total_trans_time_ok/60:.2f} min)")
        logger.info(f"Tempo Médio por Tradução: {avg_trans_time_ok:.2f} seg")
    # Logs e Arquivos
    logger.info("--- Logs ---")
    logger.info(f"Log detalhado: {os.path.abspath(log_filepath)}")
    logger.info(f"Log de correções concluídas: {os.path.abspath(PROCESSED_LOG_FILE)}")
    logger.info(f"Log de traduções concluídas: {os.path.abspath(TRANSLATED_LOG_FILE)}")
    logger.info("--- Arquivos Gerados (Estrutura Exemplo) ---") # Texto exemplo atualizado
    logger.info(f"  - DOCX Corrigido: {BASE_OUTPUT_DOCX_DIR}/<Autor>/<SubpastaLivro>/<Livro>_{FINAL_DOCX_BASENAME}")
    logger.info(f"  - TXT Numerado:   {BASE_OUTPUT_TXT_DIR}/<Autor>/<SubpastaLivro>/<Livro>_{FINAL_NUMBERED_TXT_BASENAME}")
    logger.info(f"  - TXT Notas:      {BASE_OUTPUT_TXT_DIR}/<Autor>/<SubpastaLivro>/<Livro>_{NOTES_TXT_FILE_BASENAME}")
    logger.info(f"  - DOCX Traduzido: {BASE_OUTPUT_DOCX_DIR}/<Autor>/<SubpastaLivro>/<Livro>_{FINAL_DOCX_BASENAME.replace('.docx', TRANSLATED_DOCX_SUFFIX)}")
    logger.info("========================================================")

    # === Envio de E-mail FINAL (com mais stats e listas de livros) ===
    if email_configured: # Usa a flag definida no início
        send_completion_email(
            sender_email=EMAIL_SENDER_ADDRESS, sender_password=EMAIL_SENDER_APP_PASSWORD,
            recipient_email=EMAIL_RECIPIENT_ADDRESS, smtp_server=EMAIL_SMTP_SERVER, smtp_port=EMAIL_SMTP_PORT,
            # Contadores
            processed_correction=total_books_processed_correction, skipped_correction=total_books_skipped_correction, failed_correction=total_books_failed_correction,
            processed_translation=total_translation_processed, skipped_translation=total_translation_skipped, failed_translation=total_translation_failed,
            # Tempos Gerais e Logs
            total_duration_seconds=total_time_main,
            main_log_path=log_filepath, processed_log_path=PROCESSED_LOG_FILE, translated_log_path=TRANSLATED_LOG_FILE,
            # Tempos Específicos
            avg_correction_time_secs=avg_corr_time_ok, total_correction_time_secs=total_corr_time_ok,
            avg_translation_time_secs=avg_trans_time_ok, total_translation_time_secs=total_trans_time_ok,
            # Stats API
            total_correction_latency_secs=grand_total_correction_latency, total_correction_tokens=grand_total_correction_total_tokens,
            total_footnote_latency_secs=grand_total_footnote_latency, total_footnote_tokens=grand_total_footnote_total_tokens,
            # << Listas de Livros >>
            processed_correction_books=processed_correction_list, skipped_correction_books=skipped_correction_list, failed_correction_books=failed_correction_list,
            processed_translation_books=processed_translation_list, skipped_translation_books=skipped_translation_list, failed_translation_books=failed_translation_list
        )
    else:
        logger.info("Envio de e-mail de resumo final pulado (configuração ausente ou incompleta no .env).")

# --- Ponto de Entrada ---
if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("\nProcesso interrompido manualmente (Ctrl+C).")
    except Exception as e_main:
        logger.critical(f"FATAL: Erro inesperado na execução de main(): {e_main}")
        logger.critical(traceback.format_exc())

# --- Ponto de Entrada ---
if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("\nProcesso interrompido manualmente (Ctrl+C).")
    except Exception as e_main:
        logger.critical(f"FATAL: Erro inesperado na execução de main(): {e_main}")
        logger.critical(traceback.format_exc())
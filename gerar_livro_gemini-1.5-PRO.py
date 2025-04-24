# --- Using Google's Gemini API (gemini-1.5-pro) ---

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.shared import RGBColor

from dotenv import load_dotenv
import os
import re
import logging
from tqdm import tqdm
import time
import shutil
import traceback # Para log de erros detalhado
import glob # Para encontrar arquivos .txt

# Import the Google Generative AI library
import google.generativeai as genai

# === SETUP LOGGING ===
log_dir = "logs"
if not os.path.exists(log_dir): os.makedirs(log_dir)
# Usar um nome de log diferente para este fluxo
log_filepath = os.path.join(log_dir, "book_processor_multi_author.log")
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(module)s:%(lineno)d - %(funcName)s - %(message)s', # Adicionado lineno
    handlers=[ logging.FileHandler(log_filepath, encoding='utf-8'), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# === CARREGA VARIÁVEIS DE AMBIENTE ===
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# === CONFIGURAÇÕES ===

# -- Diretórios Base --
BASE_INPUT_TXT_DIR = "txt"    # Diretório principal contendo subpastas de autores
BASE_OUTPUT_DOCX_DIR = "docx"  # Diretório principal para saída DOCX
BASE_OUTPUT_TXT_DIR = "txt"   # Diretório principal para saída TXT (sim, o mesmo do input)

# -- Nomes de Arquivos Base (serão colocados nas subpastas dos autores) --
TEMPLATE_DOCX = "Estrutura.docx" # Template geral, pode estar na raiz do script

# -- Nomes Base dos Arquivos de Saída --
FINAL_DOCX_BASENAME = "Livro_Final_Formatado_Sem_Notas.docx"
FINAL_NUMBERED_TXT_BASENAME = "Livro_Final_Com_Numeros.txt"
NOTES_TXT_FILE_BASENAME = "notas_rodape.txt"
CORRECTED_TXT_BASENAME = "Livro_Corrigido_Passo1.txt"
MARKED_TXT_BASENAME = "Livro_Marcado_Passo2.txt"
PROCESSED_MARKER_FILENAME = ".processed_ok" # Marcador para evitar reprocessamento

# -- Configurações da API e Processamento --
MODEL_NAME = "gemini-1.5-pro"
MAX_CHUNK_TOKENS = 1500 # Aumentado um pouco, ajuste conforme necessidade e testes
MAX_OUTPUT_TOKENS = 8192 # Máximo suportado pelo gemini-1.5-pro
TEMPERATURE = 0.5 # Um pouco mais conservador para manter fidelidade

# -- Estilos e Padrões --
NORMAL_STYLE_NAME = "Normal"
CHAPTER_STYLE_NAME = "Heading 1" # Ou o nome real do seu estilo de capítulo
CHAPTER_PATTERNS = [
    r'^\s*Capítulo \w+', r'^\s*CAPÍTULO \w+', r'^\s*Capítulo \d+',
    r'^\s*CHAPTER \w+', r'^\s*Chapter \d+', r'^\s*LIVRO \w+', r'^\s*PARTE \w+',
    # Adicionar outros padrões se necessário
]
OTHER_BREAK_PATTERNS = [r'^\s*\*\*\*\s*$', r'^\s*---+\s*$'] # Não usado diretamente, mas mantido
PAGE_BREAK_MARKER = "===QUEBRA_DE_PAGINA==="
AI_FAILURE_MARKER = "*** FALHA NA IA - TEXTO ORIGINAL ABAIXO ***"
FORMATTING_ERROR_MARKER = "*** ERRO DE FORMATAÇÃO - TEXTO ORIGINAL ABAIXO ***"

# --- Validação API Key ---
if not GOOGLE_API_KEY:
    logger.error("GOOGLE_API_KEY não encontrada nas variáveis de ambiente.")
    exit(1)

# --- Setup Gemini Client ---
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    safety_settings_lenient = {
        'HATE': 'BLOCK_NONE', 'HARASSMENT': 'BLOCK_NONE',
        'SEXUAL' : 'BLOCK_NONE', 'DANGEROUS' : 'BLOCK_NONE'
    }
    # Configurações de geração
    generation_config = genai.GenerationConfig(
                    temperature=TEMPERATURE,
                    max_output_tokens=MAX_OUTPUT_TOKENS
                )
    gemini_model = genai.GenerativeModel(
        MODEL_NAME,
        safety_settings=safety_settings_lenient,
        generation_config=generation_config
    )
    logger.info(f"Modelo Gemini '{MODEL_NAME}' inicializado.")
except Exception as e:
    logger.error(f"Falha ao inicializar modelo Gemini ({MODEL_NAME}): {e}")
    logger.error(traceback.format_exc())
    exit(1)

# --- Funções Auxiliares (Reutilizadas) ---

def count_tokens_approx(text):
    """Estima a contagem de tokens (aproximadamente 4 caracteres por token)."""
    if not text: return 0
    # Uma estimativa um pouco mais refinada, considerando quebras de linha
    return len(text) // 3 # Ajustar se necessário, 3 tende a ser mais seguro que 4

def create_chunks(text, max_tokens, author_name="N/A"):
    """Divide o texto em chunks, subdividindo parágrafos grandes."""
    logger.info(f"[{author_name}] Iniciando criação de chunks. Máx tokens (aprox): {max_tokens}")
    chunks = []
    current_chunk = ""
    current_chunk_tokens = 0
    # Divide por parágrafos (duas quebras de linha), mantendo quebras internas
    paragraphs = text.split("\n\n")
    paragraphs_stripped = [p.strip() for p in paragraphs if p.strip()] # Lista para log

    logger.info(f"[{author_name}] Texto dividido inicialmente em {len(paragraphs_stripped)} blocos não vazios ('\\n\\n').")

    # Usaremos os parágrafos originais (com espaços iniciais/finais dentro deles)
    # para preservar a estrutura original o máximo possível antes da IA.
    # A IA deve lidar com espaços extras no início/fim de parágrafos se necessário.
    for i, paragraph_text in enumerate(paragraphs):
        # Ignora blocos completamente vazios resultantes de múltiplas quebras de linha
        if not paragraph_text.strip():
            # Se o bloco anterior não era vazio, adicionamos uma quebra dupla para manter o espaço
            if chunks and chunks[-1].strip():
                 # Evita adicionar múltiplas quebras se já houver
                 if not chunks[-1].endswith("\n\n"):
                      chunks[-1] += "\n\n"
            continue

        paragraph_tokens = count_tokens_approx(paragraph_text)
        # Adiciona tokens para o separador '\n\n' se o chunk atual não estiver vazio
        tokens_with_separator = paragraph_tokens + (count_tokens_approx("\n\n") if current_chunk else 0)

        # Lógica de Combinação
        # Se o chunk atual + separador + novo parágrafo exceder o limite E o chunk atual já tem algo
        if current_chunk and (current_chunk_tokens + tokens_with_separator > max_tokens):
            chunks.append(current_chunk)
            logger.debug(f"[{author_name}] Chunk {len(chunks)} salvo (limite atingido). Tokens: {current_chunk_tokens}. Conteúdo: '{current_chunk[:50]}...'")
            current_chunk = paragraph_text
            current_chunk_tokens = paragraph_tokens
        # Se couber ou se o chunk atual estiver vazio
        elif current_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if current_chunk else ""
            current_chunk += separator + paragraph_text
            current_chunk_tokens = count_tokens_approx(current_chunk) # Recalcula total

        # Lógica de Subdivisão (se o parágrafo SOZINHO já excede o limite)
        # Isso acontece se o parágrafo for adicionado a um chunk vazio e ainda assim for muito grande
        if not chunks or (chunks and chunks[-1] != current_chunk) : # Evita re-processar o mesmo parágrafo já subdividido
            if paragraph_tokens > max_tokens:
                logger.warning(f"[{author_name}] Parágrafo {i+1} ({paragraph_tokens} tk) excede limite {max_tokens}. Iniciando SUBDIVISÃO.")

                # Se havia algo no current_chunk ANTES deste parágrafo grande, salva esse pedaço primeiro
                if current_chunk != paragraph_text and current_chunk.strip():
                    # Tenta remover o parágrafo grande que acabamos de adicionar
                    prefix_to_save = ""
                    if current_chunk.endswith("\n\n" + paragraph_text):
                        prefix_to_save = current_chunk[:-len("\n\n" + paragraph_text)]
                    elif current_chunk == paragraph_text : # Não deveria acontecer aqui, mas por segurança
                         pass
                    else:
                         logger.warning(f"[{author_name}] Lógica de prefixo de subdivisão inesperada. Chunk atual pode conter parte do parágrafo grande.")
                         # Neste caso, salvamos o que temos e começamos a subdivisão do parágrafo grande do zero
                         prefix_to_save = current_chunk

                    if prefix_to_save.strip():
                         chunks.append(prefix_to_save)
                         logger.debug(f"[{author_name}] Chunk {len(chunks)} salvo (prefixo antes da subdivisão). Tokens: {count_tokens_approx(prefix_to_save)}. Conteúdo: '{prefix_to_save[:50]}...'")


                # Subdivide o parágrafo grande por linhas ou frases (tentar por frases primeiro '.')
                sub_chunks_added_count = 0
                sentences = re.split(r'(?<=[.!?])\s+', paragraph_text) # Tenta dividir por frases
                if len(sentences) <= 1 : # Se não dividiu bem por frases, tenta por linhas
                    sentences = paragraph_text.split('\n')

                current_sub_chunk = ""
                current_sub_chunk_tokens = 0

                for sentence_num, sentence in enumerate(sentences):
                    sentence_clean = sentence.strip()
                    if not sentence_clean: continue

                    sentence_tokens = count_tokens_approx(sentence) # Usar a frase original para token count
                    # Adiciona token para o espaço/quebra de linha entre sub-chunks
                    tokens_with_sub_separator = sentence_tokens + (count_tokens_approx("\n") if current_sub_chunk else 0)

                    # Se o sub-chunk atual + separador + nova sentença exceder o limite E o sub-chunk não está vazio
                    if current_sub_chunk and (current_sub_chunk_tokens + tokens_with_sub_separator > max_tokens):
                        chunks.append(current_sub_chunk)
                        sub_chunks_added_count += 1
                        logger.debug(f"[{author_name}] Sub-chunk {len(chunks)} salvo (Parág. {i+1}). Tokens: {current_sub_chunk_tokens}. Conteúdo: '{current_sub_chunk[:50]}...'")
                        current_sub_chunk = sentence # Inicia novo sub-chunk
                        current_sub_chunk_tokens = sentence_tokens
                    # Se a própria sentença excede o limite (caso extremo)
                    elif sentence_tokens > max_tokens:
                        # Salva o que tinha antes
                        if current_sub_chunk:
                            chunks.append(current_sub_chunk)
                            sub_chunks_added_count += 1
                            logger.debug(f"[{author_name}] Sub-chunk {len(chunks)} salvo (antes sentença longa, Parág. {i+1}). Tokens: {current_sub_chunk_tokens}. Conteúdo: '{current_sub_chunk[:50]}...'")
                        # Adiciona a sentença longa como um chunk próprio (pode falhar na API)
                        chunks.append(sentence)
                        sub_chunks_added_count += 1
                        logger.warning(f"[{author_name}]  -> Sentença/Linha {sentence_num+1} ({sentence_tokens} tk) excede limite. Adicionada como sub-chunk individual (PODE FALHAR NA API).")
                        current_sub_chunk = "" # Reseta o sub-chunk
                        current_sub_chunk_tokens = 0
                    # Se couber ou se o sub-chunk atual estiver vazio
                    else:
                        sub_separator = "\n" if current_sub_chunk else "" # Usar quebra de linha simples na subdivisão
                        current_sub_chunk += sub_separator + sentence
                        current_sub_chunk_tokens = count_tokens_approx(current_sub_chunk) # Recalcula

                # Salva o último sub-chunk se houver algo nele
                if current_sub_chunk:
                    chunks.append(current_sub_chunk)
                    sub_chunks_added_count += 1
                    logger.debug(f"[{author_name}] Último sub-chunk {len(chunks)} salvo (Parág. {i+1}). Tokens: {current_sub_chunk_tokens}. Conteúdo: '{current_sub_chunk[:50]}...'")

                # Se, apesar de exceder o limite, não conseguiu subdividir (parágrafo gigante sem quebras?)
                if sub_chunks_added_count == 0:
                     logger.warning(f"[{author_name}] Parágrafo {i+1} excedeu limite, mas não foi subdividido (sem frases/linhas?). Adicionando original como chunk (PODE FALHAR NA API).")
                     chunks.append(paragraph_text) # Adiciona o original mesmo assim

                # Após subdividir, reseta o chunk principal para o próximo parágrafo
                current_chunk = ""
                current_chunk_tokens = 0

            # Se o parágrafo não excede o limite E não foi combinado (ou seja, é o início de um chunk)
            elif not current_chunk:
                 current_chunk = paragraph_text
                 current_chunk_tokens = paragraph_tokens


    # Adiciona o último chunk restante, se houver
    if current_chunk:
        chunks.append(current_chunk)
        logger.debug(f"[{author_name}] Chunk final {len(chunks)} salvo. Tokens: {current_chunk_tokens}. Conteúdo: '{current_chunk[:50]}...'")

    # Pós-processamento: Junta chunks pequenos consecutivos se possível
    merged_chunks = []
    temp_chunk = ""
    temp_chunk_tokens = 0
    for i, chunk in enumerate(chunks):
        chunk_tokens = count_tokens_approx(chunk)
        tokens_with_separator = chunk_tokens + (count_tokens_approx("\n\n") if temp_chunk else 0)

        if temp_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if temp_chunk else ""
            temp_chunk += separator + chunk
            temp_chunk_tokens = count_tokens_approx(temp_chunk)
        else:
            merged_chunks.append(temp_chunk)
            logger.debug(f"[{author_name}] Merged chunk {len(merged_chunks)} salvo. Tokens: {temp_chunk_tokens}.")
            temp_chunk = chunk
            temp_chunk_tokens = chunk_tokens

    if temp_chunk: # Adiciona o último temp_chunk
        merged_chunks.append(temp_chunk)
        logger.debug(f"[{author_name}] Merged chunk final {len(merged_chunks)} salvo. Tokens: {temp_chunk_tokens}.")

    logger.info(f"[{author_name}] ✅ Chunking concluído. {len(merged_chunks)} chunks finais (após merge).")
    return merged_chunks


# --- FUNÇÕES DA API GEMINI (Genérica e Específicas por Passo) ---

def _call_gemini_api(model, prompt_text, chunk_for_log, author_name="N/A"):
    """Função interna para chamar a API Gemini com retries."""
    max_retries = 5
    base_wait_time = 5 # Segundos iniciais de espera

    # Reduzir o tamanho do chunk logado para evitar logs muito grandes
    log_chunk_preview = chunk_for_log[:150].replace('\n', '\\n') + ('...' if len(chunk_for_log) > 150 else '')


    for attempt in range(max_retries):
        logger.info(f"[{author_name}] Chamando API (Tentativa {attempt + 1}/{max_retries}). Chunk (início): '{log_chunk_preview}'")
        try:
            # Usa as generation_config definidas no modelo
            response = model.generate_content(prompt_text)

            # Log detalhado da resposta
            finish_reason = "UNKNOWN"; safety_ratings = "UNKNOWN"; block_reason = "N/A"; result_text = None

            # 1. Checar bloqueio no prompt_feedback
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                 # Verificar se block_reason existe e tem valor
                 if hasattr(response.prompt_feedback, 'block_reason') and response.prompt_feedback.block_reason:
                      block_reason = response.prompt_feedback.block_reason.name
                      logger.error(f"[{author_name}] API BLOQUEOU O PROMPT (Tentativa {attempt + 1}). Razão: {block_reason}. Chunk: '{log_chunk_preview}'")
                      # Se bloqueado, não adianta tentar de novo com o mesmo prompt
                      return None # Retorna falha imediatamente

            # 2. Checar se há candidatos
            if not response.candidates:
                 logger.error(f"[{author_name}] API retornou SEM CANDIDATOS (Tentativa {attempt + 1}). Resposta: {response}. Chunk: '{log_chunk_preview}'")
                 # Continuar para retry, pode ser erro temporário

            else:
                 # 3. Processar o primeiro candidato (geralmente o único)
                 try:
                    candidate = response.candidates[0]
                    finish_reason = candidate.finish_reason.name if hasattr(candidate, 'finish_reason') and candidate.finish_reason else "FINISH_REASON_UNKNOWN"
                    safety_ratings = [(r.category.name, r.probability.name) for r in candidate.safety_ratings] if candidate.safety_ratings else "N/A"

                    logger.debug(f"[{author_name}] API Call OK (Tentativa {attempt + 1}). Finish: {finish_reason}. Safety: {safety_ratings}")

                    if finish_reason == "STOP": # Significa que terminou normalmente
                         pass # OK
                    elif finish_reason == "MAX_TOKENS":
                         logger.warning(f"[{author_name}] API TRUNCOU resposta devido a MAX_OUTPUT_TOKENS ({model.generation_config.max_output_tokens}). Pode haver perda de conteúdo no final deste chunk.")
                    elif finish_reason == "SAFETY":
                         logger.warning(f"[{author_name}] API interrompeu resposta devido a SAFETY settings no CANDIDATO. Safety Ratings: {safety_ratings}. Conteúdo pode estar incompleto ou ausente.")
                         # Considerar retornar None aqui se o bloqueio de segurança for crítico? Por enquanto, tentaremos obter o texto parcial.
                    elif finish_reason == "RECITATION":
                         logger.warning(f"[{author_name}] API interrompeu resposta devido a RECITATION.")
                    elif finish_reason == "OTHER":
                         logger.warning(f"[{author_name}] API interrompeu resposta por OUTRA RAZÃO não especificada.")
                    else: # Inclui UNKNOWN e outros não listados
                         logger.warning(f"[{author_name}] API retornou com finish_reason inesperado: {finish_reason}.")

                    # 4. Tentar extrair o texto da resposta
                    result_text = ""
                    if hasattr(candidate, 'content') and hasattr(candidate.content, 'parts'):
                        text_parts = [part.text for part in candidate.content.parts if hasattr(part, 'text')]
                        if text_parts:
                            result_text = "".join(text_parts).strip()
                        else: # Se 'parts' existe mas está vazia ou sem 'text'
                             logger.warning(f"[{author_name}] Resposta API tem 'parts' mas não foi possível extrair texto (Tentativa {attempt+1}). Parts: {candidate.content.parts}")
                    # Fallback para response.text (menos comum em modelos mais novos, mas seguro verificar)
                    elif hasattr(response, 'text') and response.text:
                        result_text = response.text.strip()
                        logger.debug(f"[{author_name}] Texto extraído via response.text (fallback).")

                    # 5. Verificar se obtivemos algum texto
                    if result_text:
                        logger.debug(f"[{author_name}] Texto API recebido (100 chars): '{result_text[:100].replace('\n', '\\n')}...'")
                        # Pequena verificação de sanidade: se o resultado for extremamente curto comparado ao input, pode ser um erro
                        if len(result_text) < len(chunk_for_log) * 0.1 and len(chunk_for_log) > 100: # Se for menos de 10% e input > 100 chars
                            logger.warning(f"[{author_name}] Resposta da API parece muito curta comparada ao input. Input len: {len(chunk_for_log)}, Output len: {len(result_text)}. Verifique o resultado.")
                        return result_text # SUCESSO!

                    else: # Não conseguiu extrair texto, mesmo com candidato
                         logger.warning(f"[{author_name}] Resposta API não continha texto utilizável (Tentativa {attempt+1}), embora tenha retornado candidato. Finish Reason: {finish_reason}. Candidate: {candidate}")
                         # Continuar para retry

                 except Exception as e_details:
                    logger.error(f"[{author_name}] Erro ao extrair detalhes/texto da resposta API (Tentativa {attempt+1}): {e_details} - Resposta Crua: {response}")
                    logger.error(traceback.format_exc())
                    # Continuar para retry

            # --- Fim do processamento da resposta ---

            # Espera exponencial com jitter antes de tentar novamente (apenas se não houve sucesso ou bloqueio)
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time) # Jitter para evitar thundering herd
                logger.info(f"[{author_name}] Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                 logger.error(f"[{author_name}] Falha final na API após {max_retries} tentativas para o chunk: '{log_chunk_preview}'")
                 return None # Falha após todas as tentativas

        except Exception as e:
            logger.warning(f"[{author_name}] Erro durante a chamada da API ({model.model_name}) (Tentativa {attempt + 1}/{max_retries}): {e}")
            logger.error(traceback.format_exc()) # Log completo do erro
            if "RESOURCE_EXHAUSTED" in str(e) or "429" in str(e):
                 logger.warning(f"[{author_name}] Erro de cota (RESOURCE_EXHAUSTED / 429). Aumentando espera.")
                 base_wait_time = 15 # Aumenta espera base para erros de cota
            elif "Internal error encountered." in str(e) or "500" in str(e):
                 logger.warning(f"[{author_name}] Erro interno do servidor (500). Tentando novamente.")
                 # Mantém espera normal
            # Outros erros podem ser tentados novamente também

            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                logger.info(f"[{author_name}] Tentando API novamente em {wait_time:.2f} seg...")
                time.sleep(wait_time)
            else:
                logger.error(f"[{author_name}] Falha final na API após {max_retries} tentativas (erro na chamada) para o chunk: '{log_chunk_preview}'")
                return None # Falha após todas as tentativas

    logger.error(f"[{author_name}] Loop de tentativas da API concluído sem sucesso explícito para o chunk: '{log_chunk_preview}'")
    return None


def format_with_ai_correction_only(model, chunk, author_name, is_first_chunk=False):
    """Chama a API Gemini focando APENAS na correção de OCR/gramática."""
    context_start = "Você está formatando o início de um livro." if is_first_chunk else "Você está continuando a formatação de um texto de livro existente."

    # Lista de erros de OCR comuns em Português (Exemplo - Expanda conforme necessário)
    ocr_errors_examples = """
        * **Troca de letras similares:** 'rn' vs 'm', 'c' vs 'e', 't' vs 'f', 'l' vs 'i', 'I' vs 'l', 'O' vs '0', 'S' vs '5', 'B' vs '8'.
        * **Hífens indevidos:** Palavras quebradas incorretamente no meio ou hífens extras.
        * **Hífens ausentes:** Palavras que deveriam ser hifenizadas (ex: "guarda-chuva") aparecem juntas ou separadas.
        * **Espaços ausentes ou extras:** Palavras coladas ("onomundo") ou espaços excessivos.
        * **Pontuação incorreta:** Pontos finais trocados por vírgulas, pontos de interrogação/exclamação mal interpretados.
        * **Acentuação:** Falta de acentos (ex: 'e' vs 'é', 'a' vs 'à'), acentos incorretos (crase onde não deve) ou caracteres estranhos no lugar de acentos.
        * **Letras duplicadas ou ausentes:** "caaasa" ou "casaa" em vez de "casa".
        * **Confusão maiúsculas/minúsculas:** Nomes próprios em minúsculas, inícios de frase em minúsculas.
        * **Caracteres especiais/ruído:** Símbolos aleatórios inseridos no texto.
        * **Quebras de linha estranhas:** Parágrafos divididos no meio sem motivo aparente. Preserve as quebras de parágrafo intencionais (duas quebras de linha).
    """

    chunk_prompt = f"""
{context_start} Você é um editor literário proficiente em português do Brasil. Sua tarefa é CORRIGIR e FORMATAR o fragmento de texto a seguir, que pertence a um livro do autor {author_name}.

**CONTEXTO IMPORTANTE:** Este texto provavelmente foi extraído via OCR de um PDF e pode conter erros de reconhecimento, digitação e gramática. O objetivo principal é obter um texto LIMPO e CORRETO em português do Brasil padrão, mantendo a estrutura e o significado originais.

**SIGA RIGOROSAMENTE ESTAS REGRAS:**

1.  **Correção Profunda:** Corrija TODOS os erros gramaticais, ortográficos, de pontuação, acentuação e concordância verbal/nominal. Use o português do Brasil como referência. FOQUE em erros comuns de OCR como os listados abaixo.
2.  **Estilo e Tom:** Mantenha o estilo literário e o tom do texto original do autor {author_name}. Seja claro, fluido e envolvente. NÃO altere o significado, a voz ou a intenção do autor.
3.  **Fidelidade Estrutural:** MANTENHA a estrutura de parágrafos original. Parágrafos são geralmente separados por UMA linha em branco (duas quebras de linha `\\n\\n`). NÃO junte parágrafos que estavam separados. NÃO divida parágrafos desnecessariamente.
4.  **Sem Adições/Remoções:** NÃO omita frases ou informações. NÃO adicione conteúdo, introduções, resumos, conclusões ou opiniões que não estavam no fragmento original. SEJA ESTRITAMENTE FIEL AO CONTEÚDO.
5.  **Marcadores de Capítulo/Quebra:** Se encontrar marcadores como 'Capítulo X', '***', '---', etc., no início de um parágrafo, MANTENHA-OS EXATAMENTE como estão, naquele parágrafo específico. NÃO adicione ou remova esses marcadores.
6.  **Quebra de Página:** Se o marcador '{PAGE_BREAK_MARKER}' aparecer, MANTENHA-O EXATAMENTE onde está, em sua própria linha, sem texto antes ou depois na mesma linha.
7.  **Erros Comuns de OCR (FOCO ESPECIAL):** Preste atenção e corrija diligentemente:
    {ocr_errors_examples}
8.  **Formato de Saída:** Retorne APENAS o texto corrigido e formatado. Use parágrafos separados por duas quebras de linha (`\\n\\n`). NÃO use NENHUMA formatação especial como Markdown (`*`, `#`, `_`), HTML, etc. Retorne TEXTO PURO. Não inclua comentários sobre o que você fez, apenas o texto resultante.

**Texto do fragmento para processar (pode conter erros):**
\"\"\"
{chunk}
\"\"\"

**Lembre-se: O resultado deve ser APENAS o texto corrigido.**
"""
    logger.debug(f"[{author_name}] Enviando chunk para CORREÇÃO (API: {model.model_name}). Tam Aprox: {count_tokens_approx(chunk)} tk")
    return _call_gemini_api(model, chunk_prompt, chunk, author_name)


def format_with_ai_footnote_only(model, chunk, author_name):
    """Chama a API Gemini focando APENAS na identificação de notas."""
    # Prompt focado apenas na regra de notas, assumindo texto já corrigido
    chunk_prompt = f"""
Você é um assistente de edição trabalhando no texto do autor {author_name}. Sua tarefa é analisar o fragmento de texto A SEGUIR, que JÁ FOI CORRIGIDO no passo anterior, e APENAS inserir marcadores para potenciais notas de rodapé onde estritamente necessário.

**REGRAS IMPORTANTES:**

1.  **NÃO ALTERE O TEXTO CORRIGIDO:** Não faça correções, não mude palavras, não reestruture frases. Apenas insira os marcadores.
2.  **MARCADORES DE NOTA:** Insira marcadores APENAS nos seguintes casos:
    * **Termos em Idioma Estrangeiro (não comuns):** Imediatamente APÓS uma palavra ou frase curta em latim, francês, inglês, etc., que não seja de uso corrente em português, insira: `[NOTA_IDIOMA:palavra_original][CONTEUDO_NOTA:Tradução ou breve explicação]`. Exemplo: "...uma certa *joie de vivre*[NOTA_IDIOMA:joie de vivre][CONTEUDO_NOTA:Alegria de viver (francês)]..."
    * **Citações/Referências:** APÓS uma citação direta curta ou uma referência bibliográfica no texto (ex: (Autor, Ano)), insira: `[NOTA_CITACAO:Texto citado ou referência][CONTEUDO_NOTA:Referência bibliográfica completa ou fonte, se conhecida ou inferível]`. Exemplo: "...como disse Foucault (1975)[NOTA_CITACAO:Foucault (1975)][CONTEUDO_NOTA:FOUCAULT, Michel. Vigiar e Punir. 1975.], a disciplina..."
    * **Nomes Próprios (contexto essencial):** APÓS um nome de pessoa, local ou evento histórico POUCO CONHECIDO que SEJA ESSENCIAL contextualizar brevemente para a compreensão do trecho, insira: `[NOTA_NOME:Nome Mencionado][CONTEUDO_NOTA:Breve identificação (datas, relevância)]`. Use com MODERAÇÃO. Exemplo: "...influenciado por Kropotkin[NOTA_NOME:Kropotkin][CONTEUDO_NOTA:Piotr Kropotkin (1842-1921), anarquista russo.]..."
    * **Termos Técnicos/Jargão (essencial):** APÓS um termo técnico MUITO específico de uma área, cuja definição SEJA INDISPENSÁVEL para o leitor geral entender o argumento naquele ponto, insira: `[NOTA_TERMO:Termo Técnico][CONTEUDO_NOTA:Definição concisa]`. Use com MUITA MODERAÇÃO. Exemplo: "...aplicando a análise de isotopias[NOTA_TERMO:Isotopias][CONTEUDO_NOTA:Na semiótica greimasiana, recorrência de categorias sêmicas que garante a homogeneidade de um discurso.]..."
3.  **FORMATO DOS MARCADORES:** Use EXATAMENTE `[NOTA_TIPO:Referência]` seguido IMEDIATAMENTE por `[CONTEUDO_NOTA:Explicação]`. Não adicione espaços entre eles. Não use outros formatos.
4.  **CRITÉRIO:** Seja conservador. Adicione notas apenas se a informação for realmente útil e provavelmente desconhecida para um leitor culto médio. É MELHOR ERRAR POR NÃO ADICIONAR do que por adicionar excessivamente. NÃO adicione notas para termos comuns, nomes famosos ou citações óbvias.
5.  **NÃO INVENTE CONTEÚDO:** O `[CONTEUDO_NOTA:...]` deve ser uma tradução direta, uma referência óbvia, ou uma contextualização muito breve e factual, se possível inferida do próprio texto ou conhecimento geral básico. NÃO pesquise externamente para criar notas complexas. Se não souber o conteúdo, NÃO insira a nota.
6.  **SAÍDA:** Retorne APENAS o texto original (do input) com os marcadores inseridos nos locais exatos. Mantenha a estrutura de parágrafos (`\\n\\n`). Não adicione NENHUM outro texto, comentário ou explicação.

**Texto JÁ CORRIGIDO para analisar e inserir marcadores de nota:**
\"\"\"
{chunk}
\"\"\"

**Lembre-se: NÃO altere o texto, apenas insira os marcadores `[NOTA_...][CONTEUDO_NOTA:...]` quando apropriado.**
"""
    logger.debug(f"[{author_name}] Enviando chunk para IDENTIFICAÇÃO DE NOTAS (API: {model.model_name}). Tam Aprox: {count_tokens_approx(chunk)} tk")
    return _call_gemini_api(model, chunk_prompt, chunk, author_name)


# --- FUNÇÕES DE PROCESSAMENTO DOS PASSOS ---

def apply_formatting_pass1(doc, formatted_chunk_text, normal_style_name, chapter_style_name, corrected_text_list, author_name):
    """Aplica formatação ao DOCX (Passo 1 - sem notas) e coleta texto para o TXT corrigido."""
    if not formatted_chunk_text or not formatted_chunk_text.strip():
        logger.warning(f"[{author_name}] Chunk formatado vazio ou apenas espaços recebido (Passo 1). Pulando.")
        return

    # Texto para adicionar à lista do arquivo TXT (Passo 1 - corrigido)
    # Substitui marcador por quebra dupla no TXT e remove espaços extras no início/fim
    plain_text_for_list = formatted_chunk_text.replace(PAGE_BREAK_MARKER, "\n\n").strip()
    if plain_text_for_list:
        corrected_text_list.append(plain_text_for_list)
    else:
        # Se o texto original era só o marcador de página, não adiciona nada à lista
        if formatted_chunk_text.strip() == PAGE_BREAK_MARKER:
             logger.debug(f"[{author_name}] Chunk continha apenas marcador de página, não adicionado à lista de texto corrigido.")
        else:
             logger.warning(f"[{author_name}] Texto formatado resultou em vazio após strip/replace. Original: '{formatted_chunk_text[:50]}...'")


    # Aplica formatação ao DOCX
    normal_style = None
    chapter_style = None
    try:
        # Tenta buscar os estilos (pode falhar se o template não os tiver)
        if normal_style_name in doc.styles:
             normal_style = doc.styles[normal_style_name]
        else:
             logger.warning(f"[{author_name}] Estilo '{normal_style_name}' não encontrado no template. Usando formatação padrão.")
        if chapter_style_name in doc.styles:
             chapter_style = doc.styles[chapter_style_name]
        else:
            logger.warning(f"[{author_name}] Estilo '{chapter_style_name}' não encontrado no template. Usando formatação manual para capítulos.")
    except Exception as e_style:
         logger.error(f"[{author_name}] Erro ao acessar estilos do documento: {e_style}. Usando formatação padrão/manual.")


    chapter_regex = re.compile('|'.join(CHAPTER_PATTERNS), re.IGNORECASE)
    # Divide o chunk processado pelo marcador de quebra de página
    parts = formatted_chunk_text.split(PAGE_BREAK_MARKER)
    # Verifica se já existe algum conteúdo no documento antes deste chunk
    # Considera parágrafos que não sejam completamente vazios
    content_present_before = any(p.text.strip() for p in doc.paragraphs)

    for part_index, part in enumerate(parts):
        part_clean = part.strip() # Remove espaços no início/fim da parte

        # Adiciona quebra de página ANTES da nova parte, EXCETO para a primeira parte
        # E somente se já havia conteúdo antes ou se não for a primeira parte do chunk.
        # E se o último parágrafo não for já uma quebra de página.
        if part_index > 0: # Sempre adiciona quebra antes da parte 2, 3, etc.
             last_para_is_page_break = False
             if doc.paragraphs:
                 last_p = doc.paragraphs[-1]
                 # Verifica se o último parágrafo está vazio E contém um run com form feed ('\f')
                 if not last_p.text.strip() and any(run.text and '\f' in run.text for run in last_p.runs):
                     last_para_is_page_break = True
             if not last_para_is_page_break:
                 logger.debug(f"[{author_name}] Adicionando quebra de página ao DOCX (antes da parte {part_index + 1} do chunk).")
                 doc.add_page_break()
             else:
                 logger.debug(f"[{author_name}] Quebra de página omitida (último parágrafo já era page break).")
        elif content_present_before and not doc.paragraphs[-1].text.strip() and any(run.text and '\f' in run.text for run in doc.paragraphs[-1].runs):
             # Caso especial: se já havia conteúdo e o ÚLTIMO parágrafo era uma quebra, não adiciona outra.
             logger.debug(f"[{author_name}] Quebra de página omitida no início do chunk (último parágrafo do doc já era page break).")


        # Se a parte (após strip) estiver vazia, pula para a próxima
        if not part_clean:
            # Mesmo se a parte estiver vazia, ela pode ter sido precedida por uma quebra de página,
            # então marcamos que agora há conteúdo (a quebra) para a lógica da próxima parte.
            if part_index > 0 : content_present_before = True
            continue

        # Divide a parte em parágrafos (baseado em \n\n)
        paragraphs_in_part = part_clean.split("\n\n")
        for paragraph_text in paragraphs_in_part:
            paragraph_text_clean = paragraph_text.strip() # Remove espaços de cada parágrafo
            if not paragraph_text_clean:
                # Adiciona um parágrafo vazio para manter o espaçamento se o original era "\n\n"
                # Mas evita adicionar múltiplos parágrafos vazios consecutivos
                if doc.paragraphs and doc.paragraphs[-1].text.strip(): # Só adiciona se o último NÃO era vazio
                     p = doc.add_paragraph()
                     if normal_style: p.style = normal_style # Aplica estilo normal ao parágrafo vazio
                continue

            # Verifica marcadores especiais (Falha da IA, Erro de Formatação)
            is_ai_failure_marker = paragraph_text_clean.startswith(AI_FAILURE_MARKER)
            is_formatting_error_marker = paragraph_text_clean.startswith(FORMATTING_ERROR_MARKER)
            # Verifica se é um capítulo (DEPOIS de verificar falhas)
            is_chapter = not is_ai_failure_marker and not is_formatting_error_marker and chapter_regex.match(paragraph_text_clean) is not None

            # Adiciona o parágrafo ao documento
            p = doc.add_paragraph()
            # Adiciona o texto como um único run neste passo (sem processamento de notas)
            run = p.add_run(paragraph_text_clean)
            content_present_before = True # Marca que agora temos conteúdo

            # Aplica estilos e formatação específica
            try:
                if is_chapter:
                    if chapter_style:
                        p.style = chapter_style
                        logger.debug(f"[{author_name}] Aplicado estilo de capítulo '{chapter_style.name}'.")
                    else: # Fallback para formatação manual
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(14)
                        run.bold = True
                        logger.debug(f"[{author_name}] Aplicada formatação manual de capítulo.")
                elif is_ai_failure_marker or is_formatting_error_marker:
                    if normal_style: p.style = normal_style
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Vermelho
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT # Alinha à esquerda para destaque
                    logger.debug(f"[{author_name}] Aplicada formatação de marcador de erro.")
                else: # Parágrafo normal
                    if normal_style:
                        p.style = normal_style
                        # logger.debug(f"[{author_name}] Aplicado estilo normal '{normal_style.name}'.") # Log muito verboso
                    else: # Fallback para formatação manual
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        run.font.size = Pt(12)
                        run.bold = False
                        # logger.debug(f"[{author_name}] Aplicada formatação manual normal.")
            except Exception as e_apply_style:
                 logger.error(f"[{author_name}] Erro ao aplicar estilo/formatação ao parágrafo: {e_apply_style}. Texto: '{paragraph_text_clean[:50]}...'")
                 # Tenta garantir uma formatação mínima
                 p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def run_correction_pass(model, input_txt_path, template_docx_path, output_docx_path, output_corrected_txt_path, author_name):
    """Executa o Passo 1: Corrige texto e gera DOCX base e TXT corrigido."""
    logger.info(f"[{author_name}] --- Iniciando Passo 1: Correção ---")
    logger.info(f"[{author_name}] Lendo texto original de: {input_txt_path}")
    try:
        with open(input_txt_path, "r", encoding="utf-8") as f: texto_original = f.read()
        logger.info(f"[{author_name}] Entrada '{os.path.basename(input_txt_path)}' carregada ({len(texto_original)} chars).")
    except FileNotFoundError:
        logger.error(f"[{author_name}] Fatal: Arquivo de entrada '{input_txt_path}' não encontrado."); return False
    except Exception as e:
        logger.error(f"[{author_name}] Fatal ao ler entrada '{input_txt_path}': {e}")
        logger.error(traceback.format_exc()); return False

    # Garante que o diretório de saída para o DOCX existe
    output_docx_dir = os.path.dirname(output_docx_path)
    if not os.path.exists(output_docx_dir):
        logger.info(f"[{author_name}] Criando diretório de saída DOCX: {output_docx_dir}")
        os.makedirs(output_docx_dir, exist_ok=True)

    # Garante que o diretório de saída para o TXT existe (pode ser o mesmo do input)
    output_txt_dir = os.path.dirname(output_corrected_txt_path)
    if not os.path.exists(output_txt_dir):
        logger.info(f"[{author_name}] Criando diretório de saída TXT: {output_txt_dir}")
        os.makedirs(output_txt_dir, exist_ok=True)


    logger.info(f"[{author_name}] Dividindo texto original em chunks...")
    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS, author_name)
    if not text_chunks:
        logger.error(f"[{author_name}] Nenhum chunk gerado a partir do texto original. Abortando Passo 1."); return False
    logger.info(f"[{author_name}] Texto dividido em {len(text_chunks)} chunks.")

    doc = None
    logger.info(f"[{author_name}] Preparando documento DOCX para: {output_docx_path}")
    try:
        # Backup do DOCX final se já existir (mesma lógica de antes)
        if os.path.exists(output_docx_path):
            backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
            backup_docx_path = os.path.join(output_docx_dir, f"backup_{os.path.splitext(os.path.basename(output_docx_path))[0]}_{backup_timestamp}.docx")
            try:
                shutil.copy2(output_docx_path, backup_docx_path)
                logger.info(f"[{author_name}] Backup do DOCX anterior criado: {backup_docx_path}")
            except Exception as e_bkp:
                logger.warning(f"[{author_name}] Falha ao criar backup de '{output_docx_path}': {e_bkp}")

        # Carrega template ou cria novo
        if os.path.exists(template_docx_path):
            try:
                doc = Document(template_docx_path)
                logger.info(f"[{author_name}] Template '{template_docx_path}' carregado.")
                # Limpa corpo do template (essencial se o template tiver conteúdo)
                try:
                    for para in doc.paragraphs:
                         p_element = para._element
                         p_element.getparent().remove(p_element)
                    # Tentar remover tabelas também, se houver
                    for table in doc.tables:
                         tbl_element = table._element
                         tbl_element.getparent().remove(tbl_element)
                    logger.info(f"[{author_name}] Conteúdo principal do template limpo (parágrafos e tabelas).")
                except Exception as clean_err:
                    logger.warning(f"[{author_name}] Erro durante limpeza do template (pode ser inofensivo se já estava vazio): {clean_err}")
            except Exception as e_load_template:
                 logger.warning(f"[{author_name}] Falha ao carregar template '{template_docx_path}': {e_load_template}. Criando novo documento.")
                 doc = Document()
        else:
            logger.warning(f"[{author_name}] Template '{template_docx_path}' não encontrado. Criando novo documento.")
            doc = Document()

        # Aplica configurações de página e estilos padrão se for um novo documento
        # ou se precisar garantir (pode ser redundante se o template já tiver, mas seguro)
        if not doc.styles: # Checa se o doc tem estilos (indicador de novo doc ou falha no template)
            logger.info(f"[{author_name}] Aplicando configurações de página A5 e estilos básicos.")
            try:
                section = doc.sections[0]
                section.page_height = Inches(8.27)
                section.page_width = Inches(5.83)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.6)
                section.top_margin = Inches(0.7)
                section.bottom_margin = Inches(0.7)

                styles = doc.styles
                if NORMAL_STYLE_NAME not in styles:
                    normal_style = styles.add_style(NORMAL_STYLE_NAME, 1) # WD_STYLE_TYPE.PARAGRAPH = 1
                    normal_style.font.name = 'Times New Roman'
                    normal_style.font.size = Pt(12)
                    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    normal_style.paragraph_format.space_after = Pt(0) # Sem espaço extra após parágrafo normal
                    normal_style.paragraph_format.line_spacing = 1.15 # Espaçamento entre linhas

                if CHAPTER_STYLE_NAME not in styles:
                    chapter_style = styles.add_style(CHAPTER_STYLE_NAME, 1)
                    chapter_style.base_style = styles['Heading 1'] if 'Heading 1' in styles else None # Tenta basear no Heading 1
                    chapter_style.font.name = 'Times New Roman'
                    chapter_style.font.size = Pt(14)
                    chapter_style.font.bold = True
                    chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    chapter_style.paragraph_format.space_before = Pt(12) # Espaço antes do título do capítulo
                    chapter_style.paragraph_format.space_after = Pt(12) # Espaço depois do título do capítulo

                logger.info(f"[{author_name}] Novo documento criado/configurado com A5 e estilos básicos.")
            except Exception as e_cfg:
                logger.warning(f"[{author_name}] Falha ao aplicar configs/estilos ao novo doc: {e_cfg}")

    except Exception as e_doc:
        logger.error(f"[{author_name}] Erro crítico ao carregar/criar documento DOCX: {e_doc}")
        logger.error(traceback.format_exc()); return False

    # Processa chunks para correção via API
    logger.info(f"[{author_name}] Iniciando chamadas à API para CORREÇÃO de {len(text_chunks)} chunks...")
    corrected_text_list_pass1 = []
    processed_chunks_count = 0
    failed_chunks_count = 0
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"[{author_name}] Passo 1: Corrigindo", unit="chunk")

    for i, chunk in progress_bar:
        progress_bar.set_description(f"[{author_name}] P1: Corrigindo Chunk {i+1}/{len(text_chunks)}")
        corrected_chunk = format_with_ai_correction_only(model, chunk, author_name, is_first_chunk=(i == 0))

        if corrected_chunk and corrected_chunk.strip():
            try:
                # Passa author_name para a função de formatação
                apply_formatting_pass1(doc, corrected_chunk, NORMAL_STYLE_NAME, CHAPTER_STYLE_NAME, corrected_text_list_pass1, author_name)
                processed_chunks_count += 1
            except Exception as format_err:
                logger.error(f"[{author_name}] Erro na apply_formatting_pass1 (Chunk {i+1}): {format_err}. Usando fallback com marcador de erro.")
                logger.error(traceback.format_exc())
                failed_chunks_count += 1
                try:
                    fallback_text = f"{FORMATTING_ERROR_MARKER}\n\n{chunk}" # Adiciona marcador de erro de formatação
                    apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_STYLE_NAME, corrected_text_list_pass1, author_name)
                except Exception as fallback_format_err:
                    logger.critical(f"[{author_name}] Falha CRÍTICA ao aplicar fallback de erro de FORMATAÇÃO (Chunk {i+1}): {fallback_format_err}. CONTEÚDO PODE ESTAR INCONSISTENTE.")
        else:
            logger.warning(f"[{author_name}] Chunk {i+1} falhou na CORREÇÃO (API retornou vazio ou erro). Usando fallback com texto original e marcador de falha.")
            failed_chunks_count += 1
            try:
                fallback_text = f"{AI_FAILURE_MARKER}\n\n{chunk}" # Adiciona marcador de falha da IA
                apply_formatting_pass1(doc, fallback_text, NORMAL_STYLE_NAME, CHAPTER_STYLE_NAME, corrected_text_list_pass1, author_name)
            except Exception as fallback_format_err:
                 logger.critical(f"[{author_name}] Falha CRÍTICA ao aplicar fallback de falha da API (Chunk {i+1}): {fallback_format_err}. CONTEÚDO PODE ESTAR INCONSISTENTE.")

        # Salva progresso do DOCX periodicamente para evitar perda total
        # A cada 10 chunks ou no último chunk
        if (processed_chunks_count + failed_chunks_count) > 0 and \
           (((processed_chunks_count + failed_chunks_count) % 10 == 0) or ((i + 1) == len(text_chunks))):
            temp_save_path = f"{output_docx_path}.temp_save"
            try:
                logger.debug(f"[{author_name}] Salvando progresso DOCX (chunk {i+1})...")
                doc.save(temp_save_path)
                # Tenta mover atomicamente (melhor que copiar e deletar)
                shutil.move(temp_save_path, output_docx_path)
                logger.info(f"[{author_name}] Progresso DOCX (Passo 1) salvo ({processed_chunks_count + failed_chunks_count} chunks processados).")
            except Exception as e_save:
                 logger.error(f"[{author_name}] Erro ao salvar progresso DOCX (Chunk {i+1}) para '{output_docx_path}': {e_save}")
                 # Não interrompe, mas o arquivo pode estar desatualizado

    # Salva versão final do DOCX após o loop
    try:
        logger.info(f"[{author_name}] Salvando DOCX final (Passo 1) em: {output_docx_path}")
        doc.save(output_docx_path)
    except Exception as e_final_save:
        logger.error(f"[{author_name}] Erro no salvamento final do DOCX (Passo 1): {e_final_save}")
        logger.error(traceback.format_exc())
        # Não retorna False aqui, pois o TXT corrigido ainda pode ser útil para o próximo passo

    # Junta e salva o TXT corrigido (resultado do Passo 1)
    try:
        logger.info(f"[{author_name}] Salvando TXT corrigido (Passo 1) em: {output_corrected_txt_path}")
        full_corrected_text = "\n\n".join(corrected_text_list_pass1)
        with open(output_corrected_txt_path, "w", encoding="utf-8") as f_txt:
            f_txt.write(full_corrected_text)
        logger.info(f"[{author_name}] TXT corrigido salvo ({len(full_corrected_text)} chars).")
    except Exception as e_txt_save:
        logger.error(f"[{author_name}] Erro ao salvar TXT corrigido (Passo 1) para '{output_corrected_txt_path}': {e_txt_save}")
        logger.error(traceback.format_exc())
        return False # Falha aqui impede o próximo passo, então retorna False

    logger.info(f"[{author_name}] --- Passo 1 concluído. Chunks OK: {processed_chunks_count}, Falhas/Fallback: {failed_chunks_count} ---")
    # Retorna True apenas se não houve falha crítica ao salvar o TXT corrigido
    return True


def run_footnote_id_pass(model, input_corrected_txt_path, output_marked_txt_path, author_name):
    """Executa o Passo 2: Identifica notas no texto já corrigido e gera TXT marcado."""
    logger.info(f"[{author_name}] --- Iniciando Passo 2: Identificação de Notas ---")
    logger.info(f"[{author_name}] Lendo texto corrigido (saída P1) de: {input_corrected_txt_path}")
    try:
        # Verifica se o arquivo de entrada (corrigido) existe
        if not os.path.exists(input_corrected_txt_path):
             logger.error(f"[{author_name}] Fatal: Arquivo de texto corrigido '{input_corrected_txt_path}' não encontrado. Abortando Passo 2.")
             return False
        with open(input_corrected_txt_path, "r", encoding="utf-8") as f: texto_corrigido = f.read()
        logger.info(f"[{author_name}] Texto corrigido '{os.path.basename(input_corrected_txt_path)}' carregado ({len(texto_corrigido)} chars).")
    except Exception as e:
        logger.error(f"[{author_name}] Fatal ao ler texto corrigido '{input_corrected_txt_path}': {e}")
        logger.error(traceback.format_exc()); return False

    # Garante que o diretório de saída para o TXT marcado existe
    output_txt_dir = os.path.dirname(output_marked_txt_path)
    if not os.path.exists(output_txt_dir):
        logger.info(f"[{author_name}] Criando diretório de saída TXT (para P2): {output_txt_dir}")
        os.makedirs(output_txt_dir, exist_ok=True)

    logger.info(f"[{author_name}] Dividindo texto corrigido em chunks para identificação de notas...")
    text_chunks = create_chunks(texto_corrigido, MAX_CHUNK_TOKENS, author_name) # Usa a mesma lógica de chunking
    if not text_chunks:
        logger.error(f"[{author_name}] Nenhum chunk gerado a partir do texto corrigido. Abortando Passo 2."); return False
    logger.info(f"[{author_name}] Texto corrigido dividido em {len(text_chunks)} chunks.")

    # Processa chunks para identificação de notas via API
    logger.info(f"[{author_name}] Iniciando chamadas à API para IDENTIFICAÇÃO DE NOTAS em {len(text_chunks)} chunks...")
    marked_text_list_pass2 = []
    processed_chunks_count = 0
    failed_chunks_count = 0
    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"[{author_name}] Passo 2: Notas", unit="chunk")

    for i, chunk in progress_bar:
        progress_bar.set_description(f"[{author_name}] P2: Notas Chunk {i+1}/{len(text_chunks)}")
        # Passa author_name para a função da API
        marked_chunk = format_with_ai_footnote_only(model, chunk, author_name)

        if marked_chunk: # API retornou algo (pode ser o chunk original se não achou notas, ou com marcadores)
            marked_text_list_pass2.append(marked_chunk)
            processed_chunks_count += 1
            # Verifica se marcadores foram realmente adicionados (opcional, para log)
            if "[NOTA_" in marked_chunk:
                logger.debug(f"[{author_name}] Chunk {i+1} processado, marcadores de nota encontrados/adicionados.")
            else:
                 logger.debug(f"[{author_name}] Chunk {i+1} processado, NENHUM marcador de nota adicionado pela IA.")
        else:
            logger.warning(f"[{author_name}] Chunk {i+1} falhou na IDENTIFICAÇÃO DE NOTAS (API retornou vazio ou erro). Usando texto original do chunk (sem marcadores) como fallback.")
            marked_text_list_pass2.append(chunk) # Adiciona o chunk original sem marcadores
            failed_chunks_count += 1

    # Junta e salva o TXT com marcadores (resultado do Passo 2)
    try:
        logger.info(f"[{author_name}] Salvando TXT com marcadores de nota (Passo 2) em: {output_marked_txt_path}")
        full_marked_text = "\n\n".join(marked_text_list_pass2)
        with open(output_marked_txt_path, "w", encoding="utf-8") as f_mark:
            f_mark.write(full_marked_text)
        logger.info(f"[{author_name}] TXT com marcadores salvo ({len(full_marked_text)} chars).")
    except Exception as e_mark_save:
        logger.error(f"[{author_name}] Erro ao salvar TXT com marcadores (Passo 2) para '{output_marked_txt_path}': {e_mark_save}")
        logger.error(traceback.format_exc())
        return False # Falha aqui impede o próximo passo

    logger.info(f"[{author_name}] --- Passo 2 concluído. Chunks OK: {processed_chunks_count}, Falhas/Fallback: {failed_chunks_count} ---")
    # Retorna True apenas se não houve falha crítica ao salvar o TXT marcado
    return True


def run_final_txt_generation(input_marked_txt_path, output_notes_path, output_numbered_txt_path, author_name):
    """Executa o Passo 3: Processa marcadores [NOTA_...] para gerar TXT final numerado [N] e arquivo de notas."""
    logger.info(f"[{author_name}] --- Iniciando Passo 3: Geração Final TXT (Notas e Numerado) ---")
    logger.info(f"[{author_name}] Lendo texto com marcadores (saída P2) de: {input_marked_txt_path}")
    try:
        # Verifica se o arquivo de entrada (marcado) existe
        if not os.path.exists(input_marked_txt_path):
            logger.error(f"[{author_name}] Fatal: Arquivo de texto marcado '{input_marked_txt_path}' não encontrado. Abortando Passo 3.")
            return False
        with open(input_marked_txt_path, "r", encoding="utf-8") as f: marked_text = f.read()
        logger.info(f"[{author_name}] Texto com marcadores '{os.path.basename(input_marked_txt_path)}' carregado ({len(marked_text)} chars).")
    except Exception as e:
        logger.error(f"[{author_name}] Fatal ao ler texto com marcadores '{input_marked_txt_path}': {e}")
        logger.error(traceback.format_exc()); return False

    # Garante que os diretórios de saída existem (embora P1 e P2 já devam ter criado)
    output_notes_dir = os.path.dirname(output_notes_path)
    output_numbered_dir = os.path.dirname(output_numbered_txt_path)
    if not os.path.exists(output_notes_dir): os.makedirs(output_notes_dir, exist_ok=True)
    # Não precisa criar output_numbered_dir separadamente se for o mesmo que output_notes_dir

    footnote_counter = 1
    notes_found = []
    # Regex para encontrar os pares de marcadores [NOTA_...] e [CONTEUDO_NOTA:...]
    # Captura o tipo (opcional), a referência e o conteúdo
    # A referência [NOTA_...] é capturada no grupo 1, o conteúdo em si no grupo 3
    footnote_pattern = re.compile(
        r'(\[NOTA_(?:IDIOMA|CITACAO|NOME|TERMO):[^\]]+?\])\s*(\[CONTEUDO_NOTA:([^\]]*?)\])', # Tornar os conteúdos não-gananciosos
        re.IGNORECASE # Ignorar case para [NOTA_...] e [CONTEUDO_...]
    )

    # Função interna para substituir o marcador e coletar a nota
    def replace_marker_and_collect_note(match):
        nonlocal footnote_counter # Permite modificar o contador externo
        # Grupo 1: Marcador completo [NOTA_...] - Não usamos no texto final
        # Grupo 3: Conteúdo da nota (texto dentro de [CONTEUDO_NOTA:...])
        original_marker = match.group(1)
        content_marker = match.group(2)
        content = match.group(3).strip()

        if not content:
             logger.warning(f"[{author_name}] Encontrado marcador [CONTEUDO_NOTA:] vazio após {original_marker}. Ignorando esta nota.")
             # Retorna string vazia para remover ambos os marcadores sem adicionar número
             return ""

        # Adiciona a nota formatada à lista
        notes_found.append(f"{footnote_counter}. {content}")
        # O texto que substituirá os dois marcadores no texto principal
        replacement = f"[{footnote_counter}]"
        logger.debug(f"[{author_name}] Nota {footnote_counter} encontrada. Marcador: '{original_marker}{content_marker}', Conteúdo: '{content}', Substituição: '{replacement}'")
        footnote_counter += 1
        return replacement

    logger.info(f"[{author_name}] Processando marcadores e gerando arquivos finais...")
    try:
        # Usa re.sub com a função para processar o texto inteiro
        final_numbered_text = footnote_pattern.sub(replace_marker_and_collect_note, marked_text)

        # --- Salva o arquivo de notas ---
        logger.info(f"[{author_name}] Salvando arquivo de notas em: {output_notes_path}")
        with open(output_notes_path, "w", encoding="utf-8") as f_notes:
            f_notes.write(f"Notas de Rodapé Geradas para {author_name}\n")
            f_notes.write("=" * (25 + len(author_name)) + "\n\n")
            if notes_found:
                f_notes.write("\n".join(notes_found))
                f_notes.write("\n") # Linha extra no final por clareza
                logger.info(f"[{author_name}] {len(notes_found)} notas salvas.")
            else:
                f_notes.write("(Nenhuma nota de rodapé foi identificada ou extraída com sucesso)\n")
                logger.info(f"[{author_name}] Nenhuma nota de rodapé identificada/salva.")

        # --- Salva o TXT final com números [N] ---
        logger.info(f"[{author_name}] Salvando TXT final com números [{footnote_counter-1}] em: {output_numbered_txt_path}")
        with open(output_numbered_txt_path, "w", encoding="utf-8") as f_numbered:
            f_numbered.write(final_numbered_text)
        logger.info(f"[{author_name}] TXT final com números salvo ({len(final_numbered_text)} chars).")

    except Exception as e_final_gen:
        logger.error(f"[{author_name}] Erro durante a geração final dos arquivos TXT (Passo 3): {e_final_gen}")
        logger.error(traceback.format_exc())
        return False # Falha crítica

    logger.info(f"[{author_name}] --- Passo 3 concluído. ---")
    return True


# --- FUNÇÃO PRINCIPAL (main) ---
def main():
    start_time_main = time.time()
    logger.info("========================================================")
    logger.info(f"Iniciando Processador Multi-Autor - {time.strftime('%Y-%m-%d %H:%M:%S')}")
    logger.info(f"Diretório de Entrada TXT: {BASE_INPUT_TXT_DIR}")
    logger.info(f"Diretório de Saída DOCX: {BASE_OUTPUT_DOCX_DIR}")
    logger.info(f"Diretório de Saída TXT: {BASE_OUTPUT_TXT_DIR}")
    logger.info(f"Template DOCX: {TEMPLATE_DOCX}")
    logger.info("========================================================")

    # Verifica se o diretório base de entrada existe
    if not os.path.isdir(BASE_INPUT_TXT_DIR):
        logger.error(f"Diretório de entrada base '{BASE_INPUT_TXT_DIR}' não encontrado! Abortando.")
        return

    # Encontra todas as subpastas (autores) no diretório de entrada
    try:
        author_folders = [f for f in os.listdir(BASE_INPUT_TXT_DIR) if os.path.isdir(os.path.join(BASE_INPUT_TXT_DIR, f))]
    except Exception as e:
        logger.error(f"Erro ao listar diretórios em '{BASE_INPUT_TXT_DIR}': {e}")
        return

    if not author_folders:
        logger.warning(f"Nenhuma subpasta de autor encontrada em '{BASE_INPUT_TXT_DIR}'. Saindo.")
        return

    logger.info(f"Autores encontrados: {len(author_folders)} -> {', '.join(author_folders)}")

    authors_processed_count = 0
    authors_skipped_count = 0
    authors_failed_count = 0

    # Loop principal para processar cada autor
    for author_name in author_folders:
        author_start_time = time.time()
        logger.info(f"--- Processando Autor: {author_name} ---")

        # --- Define os caminhos específicos para este autor ---
        author_input_dir = os.path.join(BASE_INPUT_TXT_DIR, author_name)
        author_output_docx_dir = os.path.join(BASE_OUTPUT_DOCX_DIR, author_name)
        author_output_txt_dir = os.path.join(BASE_OUTPUT_TXT_DIR, author_name) # TXTs vão para a pasta do autor em 'txt'

        # --- Verifica se já foi processado (marcador no diretório de saída DOCX) ---
        processed_marker_path = os.path.join(author_output_docx_dir, PROCESSED_MARKER_FILENAME)
        if os.path.exists(processed_marker_path):
            logger.info(f"[{author_name}] Já processado anteriormente (marcador encontrado: '{processed_marker_path}'). Pulando.")
            authors_skipped_count += 1
            continue # Pula para o próximo autor

        # --- Encontra o arquivo .txt de entrada dentro da pasta do autor ---
        # Usamos glob para encontrar qualquer arquivo .txt na pasta do autor
        input_txt_files = glob.glob(os.path.join(author_input_dir, "*.txt"))

        # Filtra arquivos intermediários ou de notas que possam estar lá
        input_txt_files = [f for f in input_txt_files if not os.path.basename(f).startswith("Livro_Corrigido") \
                                                    and not os.path.basename(f).startswith("Livro_Marcado") \
                                                    and not os.path.basename(f).startswith("Livro_Final") \
                                                    and not os.path.basename(f).startswith("notas_rodape")]

        if not input_txt_files:
            logger.warning(f"[{author_name}] Nenhum arquivo .txt de entrada encontrado em '{author_input_dir}'. Pulando autor.")
            authors_skipped_count += 1
            continue
        elif len(input_txt_files) > 1:
            logger.warning(f"[{author_name}] Múltiplos arquivos .txt encontrados em '{author_input_dir}': {input_txt_files}. Usando o primeiro: '{input_txt_files[0]}'.")
            # Poderia adicionar lógica para escolher baseado em nome ou tamanho se necessário
        input_txt_path = input_txt_files[0]
        logger.info(f"[{author_name}] Arquivo de entrada selecionado: {input_txt_path}")

        # Deriva o nome base do livro a partir do nome do arquivo de entrada (removendo extensão e autor se presente)
        base_book_name = os.path.splitext(os.path.basename(input_txt_path))[0]
        # Tenta remover " - author_name" do final, se existir
        if base_book_name.lower().endswith(f" - {author_name.lower()}"):
             base_book_name = base_book_name[:-len(f" - {author_name}")] # Remove sufixo
        base_book_name = base_book_name.strip() # Remove espaços extras

        # --- Constrói os nomes dos arquivos de saída usando o nome base do livro ---
        # Usa o nome base do livro + sufixo padrão para clareza
        output_docx_path = os.path.join(author_output_docx_dir, f"{base_book_name}_{FINAL_DOCX_BASENAME}")
        output_corrected_txt_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{CORRECTED_TXT_BASENAME}")
        output_marked_txt_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{MARKED_TXT_BASENAME}")
        output_notes_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{NOTES_TXT_FILE_BASENAME}")
        output_numbered_txt_path = os.path.join(author_output_txt_dir, f"{base_book_name}_{FINAL_NUMBERED_TXT_BASENAME}")

        # --- Cria os diretórios de saída específicos do autor se não existirem ---
        try:
            if not os.path.exists(author_output_docx_dir):
                logger.info(f"[{author_name}] Criando diretório de saída DOCX: {author_output_docx_dir}")
                os.makedirs(author_output_docx_dir)
            if not os.path.exists(author_output_txt_dir):
                 # Verifica se é o mesmo diretório do input para não logar duas vezes
                if author_output_txt_dir != author_input_dir:
                    logger.info(f"[{author_name}] Criando diretório de saída TXT: {author_output_txt_dir}")
                    os.makedirs(author_output_txt_dir)
        except Exception as e_mkdir:
             logger.error(f"[{author_name}] Erro ao criar diretórios de saída: {e_mkdir}. Pulando autor.")
             authors_failed_count += 1
             continue

        # --- Executa a Sequência de Passos para o Autor ---
        all_steps_successful_for_author = True

        # === PASSO 1: CORREÇÃO ===
        pass1_success = run_correction_pass(gemini_model, input_txt_path, TEMPLATE_DOCX, output_docx_path, output_corrected_txt_path, author_name)
        if not pass1_success:
            logger.error(f"[{author_name}] Passo 1 (Correção) FALHOU. Abortando processamento para este autor.")
            all_steps_successful_for_author = False
        else:
            # === PASSO 2: IDENTIFICAÇÃO DE NOTAS ===
            pass2_success = run_footnote_id_pass(gemini_model, output_corrected_txt_path, output_marked_txt_path, author_name)
            if not pass2_success:
                logger.error(f"[{author_name}] Passo 2 (Identificação de Notas) FALHOU. Abortando passos restantes para este autor.")
                all_steps_successful_for_author = False
            else:
                # === PASSO 3: GERAÇÃO FINAL TXT (Notas e Numerado) ===
                pass3_success = run_final_txt_generation(output_marked_txt_path, output_notes_path, output_numbered_txt_path, author_name)
                if not pass3_success:
                    logger.error(f"[{author_name}] Passo 3 (Geração Final TXT) FALHOU.")
                    all_steps_successful_for_author = False

        # --- Conclusão do Processamento do Autor ---
        author_end_time = time.time()
        author_total_time = author_end_time - author_start_time
        if all_steps_successful_for_author:
            logger.info(f"✅ [{author_name}] Processamento concluído com SUCESSO em {author_total_time:.2f} seg.")
            # Cria o marcador de sucesso no diretório DOCX de saída do autor
            try:
                with open(processed_marker_path, 'w') as f_marker:
                    f_marker.write(time.strftime('%Y-%m-%d %H:%M:%S'))
                logger.info(f"[{author_name}] Marcador de sucesso criado: '{processed_marker_path}'")
            except Exception as e_marker:
                logger.warning(f"[{author_name}] Falha ao criar marcador de sucesso '{processed_marker_path}': {e_marker}")
            authors_processed_count += 1
        else:
            logger.warning(f"⚠️ [{author_name}] Processamento concluído com FALHAS em {author_total_time:.2f} seg.")
            authors_failed_count += 1
            # Opcional: Remover arquivos intermediários/finais incompletos em caso de falha?
            # try:
            #     if os.path.exists(output_docx_path): os.remove(output_docx_path)
            #     if os.path.exists(output_corrected_txt_path): os.remove(output_corrected_txt_path)
            #     # ... remover outros ...
            #     logger.info(f"[{author_name}] Arquivos de saída parciais removidos devido à falha.")
            # except Exception as e_clean:
            #     logger.warning(f"[{author_name}] Erro ao tentar remover arquivos parciais: {e_clean}")

        logger.info(f"--- Fim do processamento para: {author_name} ---")

    # --- Fim do Loop Principal ---
    end_time_main = time.time()
    total_time_main = end_time_main - start_time_main
    logger.info("========================================================")
    logger.info("🏁 Processamento Multi-Autor Concluído!")
    logger.info(f"Tempo total geral: {total_time_main:.2f} seg ({total_time_main/60:.2f} min).")
    logger.info(f"Autores Processados com Sucesso: {authors_processed_count}")
    logger.info(f"Autores Pulados (já processados ou sem input): {authors_skipped_count}")
    logger.info(f"Autores com Falha em algum passo: {authors_failed_count}")
    logger.info(f"Log detalhado salvo em: {log_filepath}")
    logger.info("Verifique os diretórios de saída para os resultados:")
    logger.info(f"  - DOCX: {BASE_OUTPUT_DOCX_DIR}/<nome_do_autor>/")
    logger.info(f"  - TXTs: {BASE_OUTPUT_TXT_DIR}/<nome_do_autor>/")
    logger.info("========================================================")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("\nProcesso interrompido manualmente (Ctrl+C). Arquivos podem estar incompletos.")
    except Exception as e_main:
        logger.critical(f"Erro fatal inesperado durante a execução de main: {e_main}")
        logger.critical(traceback.format_exc())
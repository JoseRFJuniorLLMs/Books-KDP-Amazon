# -*- coding: utf-8 -*-
"""
PIPELINE GEMINI 2.5 - Processamento Massivo com Vertex AI
=========================================================
Usa Gemini 2.5 Flash para TODO processamento:
- Tradução
- Correção OCR
- Notas de rodapé

Uso:
    python pipeline_gemini.py --run
    python pipeline_gemini.py --run --limit 10
"""

import sys
import asyncio
import aiohttp
import hashlib
import sqlite3
import argparse
import json
import time
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
import re

from tqdm import tqdm

# Adiciona raiz do projeto ao path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

# Configuracao centralizada
from config.settings import (
    VERTEX_PROJECT, VERTEX_REGION, BUCKET_NAME, PROJECT_ID
)

# =============================================================================
# CONFIGURAÇÕES
# =============================================================================

# Endpoint Vertex AI
VERTEX_ENDPOINT = "https://aiplatform.googleapis.com/v1/publishers/google/models"
MODEL = "gemini-2.0-flash"

# GCS
GCS_BUCKET = BUCKET_NAME or "testmassivo-kdp-2024"

# Diretórios
BASE_DIR = Path(__file__).parent
TXT_DIR = BASE_DIR / "txt"
DOCX_DIR = BASE_DIR / "docx"
DATA_DIR = BASE_DIR / "data"
TRANSLATED_DIR = BASE_DIR / "translated-portugues"

for d in [TXT_DIR, DOCX_DIR, DATA_DIR, TRANSLATED_DIR]:
    d.mkdir(exist_ok=True)

# Paralelismo MASSIVO - Gemini aguenta!
MAX_CONCURRENT = 50   # Requisições simultâneas
CHUNK_SIZE = 6000     # Caracteres por chunk
REQUESTS_PER_MINUTE = 1000  # Limite alto para Vertex AI

# =============================================================================
# VERTEX AI CLIENT
# =============================================================================

class VertexAIClient:
    """Cliente Vertex AI com aiohttp para máxima velocidade."""

    def __init__(self, api_key: str = VERTEX_API_KEY):
        self.api_key = api_key
        self.base_url = f"{VERTEX_ENDPOINT}/{MODEL}:generateContent"
        self.session = None
        self.semaphore = asyncio.Semaphore(MAX_CONCURRENT)

    async def _get_session(self):
        if self.session is None:
            timeout = aiohttp.ClientTimeout(total=120)
            self.session = aiohttp.ClientSession(timeout=timeout)
        return self.session

    async def generate(self, prompt: str, max_retries: int = 3) -> str:
        """Gera resposta com retry e rate limiting."""
        async with self.semaphore:
            session = await self._get_session()

            url = f"{self.base_url}?key={self.api_key}"
            payload = {
                "contents": [{
                    "role": "user",
                    "parts": [{"text": prompt}]
                }],
                "generationConfig": {
                    "temperature": 0.3,
                    "maxOutputTokens": 8192,
                }
            }

            for attempt in range(max_retries):
                try:
                    async with session.post(url, json=payload) as resp:
                        if resp.status == 200:
                            data = await resp.json()
                            if "candidates" in data and data["candidates"]:
                                parts = data["candidates"][0].get("content", {}).get("parts", [])
                                if parts:
                                    return parts[0].get("text", "").strip()
                            return ""
                        elif resp.status == 429:
                            wait = 2 ** attempt
                            print(f"Rate limit, aguardando {wait}s...")
                            await asyncio.sleep(wait)
                        else:
                            error = await resp.text()
                            print(f"Erro {resp.status}: {error[:200]}")
                            if attempt < max_retries - 1:
                                await asyncio.sleep(1)
                except Exception as e:
                    if attempt < max_retries - 1:
                        await asyncio.sleep(1)
                    else:
                        print(f"Erro: {e}")
            return ""

    async def close(self):
        if self.session:
            await self.session.close()

# Cliente global
client = None

def get_client():
    global client
    if client is None:
        client = VertexAIClient()
    return client

# =============================================================================
# DATABASE / CACHE
# =============================================================================

DB_PATH = DATA_DIR / "gemini_cache.db"

def init_db():
    conn = sqlite3.connect(DB_PATH)
    conn.executescript('''
        CREATE TABLE IF NOT EXISTS traducoes (
            hash TEXT PRIMARY KEY,
            texto TEXT
        );
        CREATE TABLE IF NOT EXISTS correcoes (
            hash TEXT PRIMARY KEY,
            texto TEXT
        );
        CREATE TABLE IF NOT EXISTS notas (
            termo TEXT PRIMARY KEY,
            explicacao TEXT
        );
    ''')
    conn.close()

def get_cache(table: str, key: str) -> Optional[str]:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.execute(f'SELECT texto FROM {table} WHERE hash = ?', (key,))
    row = cur.fetchone()
    conn.close()
    return row[0] if row else None

def set_cache(table: str, key: str, value: str):
    conn = sqlite3.connect(DB_PATH)
    conn.execute(f'INSERT OR REPLACE INTO {table} (hash, texto) VALUES (?, ?)', (key, value))
    conn.commit()
    conn.close()

# =============================================================================
# DETECÇÃO DE IDIOMA
# =============================================================================

LANG_WORDS = {
    'pt': ['que', 'não', 'para', 'uma', 'com', 'por', 'mais', 'você', 'isso', 'está', 'seu', 'sua'],
    'es': ['que', 'los', 'las', 'una', 'para', 'con', 'más', 'pero', 'está', 'tiene', 'sus'],
    'en': ['the', 'and', 'that', 'have', 'for', 'with', 'this', 'but', 'from', 'they', 'was'],
    'fr': ['les', 'des', 'une', 'que', 'pour', 'dans', 'qui', 'sur', 'avec', 'cette', 'sont'],
    'de': ['der', 'die', 'und', 'den', 'das', 'von', 'ist', 'mit', 'nicht', 'sich', 'auf'],
    'ru': ['что', 'как', 'это', 'его', 'она', 'они', 'все', 'для', 'был', 'быть', 'при'],
}

def detectar_idioma(texto: str) -> str:
    texto = texto[:5000].lower()
    scores = {lang: sum(1 for w in words if f' {w} ' in f' {texto} ')
              for lang, words in LANG_WORDS.items()}
    return max(scores, key=scores.get)

# =============================================================================
# PROCESSAMENTO COM GEMINI
# =============================================================================

async def corrigir_ocr(texto: str) -> str:
    """Corrige erros de OCR."""
    text_hash = hashlib.sha256(texto.encode()).hexdigest()[:32]

    cached = get_cache('correcoes', text_hash)
    if cached:
        return cached

    ai = get_client()

    prompt = f"""Corrija erros de OCR neste texto. Mantenha estrutura original.
Retorne APENAS o texto corrigido.

{texto}"""

    resultado = await ai.generate(prompt)
    if resultado:
        set_cache('correcoes', text_hash, resultado)
    return resultado or texto

async def traduzir_chunk(texto: str, idioma: str) -> str:
    """Traduz um chunk."""
    if idioma == 'pt':
        return texto

    text_hash = hashlib.sha256(texto.encode()).hexdigest()[:32]

    cached = get_cache('traducoes', text_hash)
    if cached:
        return cached

    ai = get_client()

    lang_map = {'en': 'inglês', 'es': 'espanhol', 'fr': 'francês', 'de': 'alemão', 'ru': 'russo'}

    prompt = f"""Traduza de {lang_map.get(idioma, idioma)} para português brasileiro.
Mantenha estilo original. Retorne APENAS a tradução.

{texto}"""

    resultado = await ai.generate(prompt)
    if resultado:
        set_cache('traducoes', text_hash, resultado)
    return resultado or texto

async def gerar_notas(texto: str) -> List[Tuple[str, str]]:
    """Gera notas de rodapé."""
    ai = get_client()

    prompt = f"""Identifique termos que precisam de notas de rodapé neste texto:
- Termos estrangeiros (latim, francês, alemão)
- Conceitos filosóficos/técnicos
- Nomes históricos importantes

Para cada termo, explicação BREVE (max 25 palavras).
Formato: termo|explicação
Máximo 20 termos.

{texto[:6000]}"""

    resultado = await ai.generate(prompt)

    notas = []
    if resultado:
        for linha in resultado.strip().split('\n'):
            if '|' in linha:
                partes = linha.split('|', 1)
                if len(partes) == 2:
                    termo = partes[0].strip().strip('- ')
                    explicacao = partes[1].strip()
                    if termo and explicacao and len(termo) < 100:
                        notas.append((termo, explicacao))

    return notas[:20]

# =============================================================================
# DOCX
# =============================================================================

def criar_docx(texto: str, titulo: str, autor: str, notas: List[Tuple[str, str]],
               output_path: Path) -> bool:
    """Cria DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Título
        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(24)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Autor
        p = doc.add_paragraph()
        p.add_run(autor)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Conteúdo
        for para in texto.split('\n\n'):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Notas
        if notas:
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run("NOTAS")
            run.bold = True

            for i, (termo, explicacao) in enumerate(notas, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{i}. {termo}: ")
                run.bold = True
                run.font.size = Pt(9)
                run = p.add_run(explicacao)
                run.font.size = Pt(9)

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Erro DOCX: {e}")
        return False

# =============================================================================
# PIPELINE
# =============================================================================

@dataclass
class Livro:
    titulo: str
    autor: str
    idioma: str
    txt_path: Path
    texto: str = ""

def dividir_chunks(texto: str, tamanho: int = CHUNK_SIZE) -> List[str]:
    chunks = []
    current = ""
    for para in texto.split('\n\n'):
        if len(current) + len(para) < tamanho:
            current += '\n\n' + para if current else para
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)
    return chunks

async def processar_livro(livro: Livro) -> Dict:
    """Processa um livro completo."""
    resultado = {'titulo': livro.titulo, 'status': 'erro', 'notas': 0}

    try:
        texto = livro.texto

        # 1. Correção OCR (paralelo)
        chunks = dividir_chunks(texto)
        tasks = [corrigir_ocr(c) for c in chunks]
        resultados_ocr = await asyncio.gather(*tasks)
        texto = '\n\n'.join(resultados_ocr)

        # 2. Tradução (paralelo)
        if livro.idioma != 'pt':
            chunks = dividir_chunks(texto)
            tasks = [traduzir_chunk(c, livro.idioma) for c in chunks]
            resultados_trad = await asyncio.gather(*tasks)
            texto = '\n\n'.join(resultados_trad)

        # 3. Notas
        notas = await gerar_notas(texto)
        resultado['notas'] = len(notas)

        # 4. DOCX
        docx_path = DOCX_DIR / f"{livro.titulo}.docx"
        if criar_docx(texto, livro.titulo, livro.autor, notas, docx_path):
            resultado['docx'] = str(docx_path)
            resultado['status'] = 'sucesso'

        # Salva TXT traduzido
        txt_out = TRANSLATED_DIR / f"{livro.titulo}_pt.txt"
        with open(txt_out, 'w', encoding='utf-8') as f:
            f.write(texto)

    except Exception as e:
        resultado['error'] = str(e)

    return resultado

async def processar_batch(livros: List[Livro], max_concurrent: int = 10) -> List[Dict]:
    """Processa livros em paralelo."""
    semaphore = asyncio.Semaphore(max_concurrent)

    async def process(livro):
        async with semaphore:
            return await processar_livro(livro)

    tasks = [process(l) for l in livros]
    resultados = []

    pbar = tqdm(total=len(tasks), desc="Processando")
    for coro in asyncio.as_completed(tasks):
        resultado = await coro
        resultados.append(resultado)
        status = "✓" if resultado['status'] == 'sucesso' else "✗"
        pbar.set_postfix_str(f"{status} {resultado['titulo'][:30]}")
        pbar.update(1)
    pbar.close()

    return resultados

def escanear_livros() -> List[Livro]:
    """Escaneia txt/ e retorna livros."""
    livros = []

    for folder in TXT_DIR.iterdir():
        if not folder.is_dir():
            continue

        txt_files = list(folder.glob("*.txt"))
        if not txt_files:
            continue

        main_file = max(txt_files, key=lambda f: f.stat().st_size)

        try:
            with open(main_file, 'r', encoding='utf-8', errors='ignore') as f:
                texto = f.read()

            if len(texto) < 500:
                continue

            livros.append(Livro(
                titulo=folder.name,
                autor="Autor",
                idioma=detectar_idioma(texto),
                txt_path=main_file,
                texto=texto
            ))
        except:
            pass

    return livros

# =============================================================================
# MAIN
# =============================================================================

async def main_async(args):
    init_db()

    print("\n" + "="*60)
    print("PIPELINE GEMINI 2.5 - PROCESSAMENTO MASSIVO")
    print("="*60)
    print(f"Modelo: {MODEL}")
    print(f"Parallel: {args.concurrent}")
    print("="*60)

    # Escaneia
    print("\nEscaneando livros...")
    livros = escanear_livros()

    if args.limit > 0:
        livros = livros[:args.limit]

    print(f"Livros: {len(livros)}")

    # Por idioma
    por_idioma = {}
    for l in livros:
        por_idioma[l.idioma] = por_idioma.get(l.idioma, 0) + 1
    print("Por idioma:", dict(sorted(por_idioma.items(), key=lambda x: -x[1])))

    # Processa
    print("\nProcessando...")
    start = time.time()

    resultados = await processar_batch(livros, args.concurrent)

    elapsed = time.time() - start

    # Fecha cliente
    await get_client().close()

    # Resumo
    sucesso = sum(1 for r in resultados if r['status'] == 'sucesso')
    total_notas = sum(r.get('notas', 0) for r in resultados)

    print("\n" + "="*60)
    print("RESUMO")
    print("="*60)
    print(f"Total: {len(resultados)}")
    print(f"Sucesso: {sucesso}")
    print(f"Erros: {len(resultados) - sucesso}")
    print(f"Notas geradas: {total_notas}")
    print(f"Tempo: {elapsed:.1f}s ({elapsed/60:.1f} min)")
    print(f"Velocidade: {len(resultados)/elapsed*60:.1f} livros/min")
    print(f"DOCX: {DOCX_DIR}")
    print("="*60)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--run', action='store_true')
    parser.add_argument('--limit', type=int, default=0)
    parser.add_argument('--concurrent', type=int, default=10)
    args = parser.parse_args()

    if args.run:
        asyncio.run(main_async(args))
    else:
        parser.print_help()

if __name__ == "__main__":
    main()

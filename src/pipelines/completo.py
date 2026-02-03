# -*- coding: utf-8 -*-
"""
PIPELINE COMPLETO - Do TXT bruto ao DOCX final para Amazon KDP
==============================================================
Etapas:
1. Correção de OCR (erros de digitalização)
2. Tradução para Português (se necessário)
3. Geração de Notas de Rodapé
4. Criação do DOCX formatado

Uso:
    python pipeline_completo.py --input livro.txt --output livro.docx
    python pipeline_completo.py --batch --input-dir txt/ --output-dir docx/
"""

import sys
import os
import re
import zipfile
import argparse
import hashlib
import sqlite3
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Tuple, Dict
from lxml import etree
from copy import deepcopy
import requests
from tqdm import tqdm

# Configurações
OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:14b")
FOOTNOTE_FONT_SIZE = 9
MAX_FOOTNOTES = 50
CHUNK_SIZE = 2000  # caracteres por chunk

# Diretórios padrão
BASE_DIR = Path(__file__).parent
INPUT_DIR = BASE_DIR / "txt"
OUTPUT_DIR = BASE_DIR / "docx"
CACHE_DIR = BASE_DIR / "data"
TEMPLATE_PATH = BASE_DIR / "Estrutura.docx"

# Cria diretórios
OUTPUT_DIR.mkdir(exist_ok=True)
CACHE_DIR.mkdir(exist_ok=True)

# Namespaces DOCX
NAMESPACES = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# ============================================================================
# CACHE
# ============================================================================

class PipelineCache:
    """Cache SQLite para evitar reprocessamento."""

    def __init__(self, db_path: Path = CACHE_DIR / "pipeline_cache.db"):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS corrections (
                    hash TEXT PRIMARY KEY,
                    corrected TEXT,
                    model TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('''
                CREATE TABLE IF NOT EXISTS translations (
                    hash TEXT PRIMARY KEY,
                    translated TEXT,
                    source_lang TEXT,
                    model TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('''
                CREATE TABLE IF NOT EXISTS footnotes (
                    term TEXT PRIMARY KEY,
                    explanation TEXT,
                    model TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

    def _hash(self, text: str) -> str:
        return hashlib.sha256(text.encode()).hexdigest()[:32]

    def get_correction(self, text: str) -> Optional[str]:
        with sqlite3.connect(self.db_path) as conn:
            cur = conn.execute('SELECT corrected FROM corrections WHERE hash = ?', (self._hash(text),))
            row = cur.fetchone()
            return row[0] if row else None

    def set_correction(self, original: str, corrected: str, model: str):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('INSERT OR REPLACE INTO corrections (hash, corrected, model) VALUES (?, ?, ?)',
                        (self._hash(original), corrected, model))

    def get_translation(self, text: str) -> Optional[str]:
        with sqlite3.connect(self.db_path) as conn:
            cur = conn.execute('SELECT translated FROM translations WHERE hash = ?', (self._hash(text),))
            row = cur.fetchone()
            return row[0] if row else None

    def set_translation(self, original: str, translated: str, source_lang: str, model: str):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('INSERT OR REPLACE INTO translations (hash, translated, source_lang, model) VALUES (?, ?, ?, ?)',
                        (self._hash(original), translated, source_lang, model))

    def get_footnote(self, term: str) -> Optional[str]:
        with sqlite3.connect(self.db_path) as conn:
            cur = conn.execute('SELECT explanation FROM footnotes WHERE term = ?', (term.lower(),))
            row = cur.fetchone()
            return row[0] if row else None

    def set_footnote(self, term: str, explanation: str, model: str):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('INSERT OR REPLACE INTO footnotes (term, explanation, model) VALUES (?, ?, ?)',
                        (term.lower(), explanation, model))

cache = PipelineCache()

# ============================================================================
# DETECÇÃO DE IDIOMA
# ============================================================================

def detect_language(text: str) -> str:
    """Detecta idioma do texto (simplificado)."""
    text_lower = text[:5000].lower()

    # Palavras características
    pt_words = ['que', 'não', 'para', 'uma', 'com', 'por', 'mais', 'como', 'seu', 'sua', 'ele', 'ela', 'são', 'está', 'isso', 'este', 'essa', 'muito', 'também', 'pode', 'fazer', 'tempo', 'ainda', 'bem', 'onde', 'depois', 'sobre', 'mesmo', 'aos', 'já', 'porque', 'entre', 'quando', 'essa']
    en_words = ['the', 'and', 'that', 'have', 'for', 'not', 'with', 'you', 'this', 'but', 'his', 'from', 'they', 'she', 'which', 'their', 'would', 'there', 'been', 'have', 'many', 'some', 'them', 'these', 'than', 'first', 'other', 'into', 'could', 'made']
    es_words = ['que', 'los', 'las', 'una', 'para', 'con', 'por', 'más', 'como', 'pero', 'sus', 'está', 'son', 'tiene', 'este', 'todo', 'esta', 'entre', 'cuando', 'muy', 'sin', 'sobre', 'también', 'fue', 'había', 'hasta', 'desde', 'puede']
    fr_words = ['les', 'des', 'une', 'que', 'pour', 'dans', 'qui', 'sur', 'avec', 'sont', 'par', 'mais', 'cette', 'vous', 'nous', 'elle', 'tout', 'bien', 'fait', 'comme', 'était', 'être', 'avoir', 'leur', 'même', 'aussi', 'autres']
    de_words = ['der', 'die', 'und', 'den', 'das', 'von', 'ist', 'des', 'auf', 'für', 'mit', 'dem', 'nicht', 'sich', 'eine', 'ein', 'als', 'auch', 'nach', 'wird', 'bei', 'oder', 'nur', 'werden', 'sind', 'war', 'haben']
    ru_words = ['что', 'как', 'это', 'его', 'она', 'они', 'все', 'для', 'так', 'был', 'при', 'уже', 'вот', 'меня', 'тебя', 'свой', 'есть', 'быть', 'были']

    scores = {
        'pt': sum(1 for w in pt_words if f' {w} ' in f' {text_lower} '),
        'en': sum(1 for w in en_words if f' {w} ' in f' {text_lower} '),
        'es': sum(1 for w in es_words if f' {w} ' in f' {text_lower} '),
        'fr': sum(1 for w in fr_words if f' {w} ' in f' {text_lower} '),
        'de': sum(1 for w in de_words if f' {w} ' in f' {text_lower} '),
        'ru': sum(1 for w in ru_words if f' {w} ' in f' {text_lower} '),
    }

    return max(scores, key=scores.get)

# ============================================================================
# CHUNKING
# ============================================================================

def create_chunks(text: str, max_size: int = CHUNK_SIZE) -> List[str]:
    """Divide texto em chunks respeitando parágrafos."""
    chunks = []
    current = ""

    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue

        if len(current) + len(para) + 2 <= max_size:
            current += ("\n\n" if current else "") + para
        else:
            if current:
                chunks.append(current)

            if len(para) > max_size:
                # Divide parágrafos grandes por sentenças
                sentences = re.split(r'(?<=[.!?])\s+', para)
                sub = ""
                for sent in sentences:
                    if len(sub) + len(sent) + 1 <= max_size:
                        sub += (" " if sub else "") + sent
                    else:
                        if sub:
                            chunks.append(sub)
                        sub = sent
                current = sub
            else:
                current = para

    if current:
        chunks.append(current)

    return chunks

# ============================================================================
# OLLAMA API
# ============================================================================

def call_ollama(prompt: str, model: str = OLLAMA_MODEL, max_tokens: int = 2000) -> Optional[str]:
    """Chama a API do Ollama."""
    try:
        r = requests.post(
            OLLAMA_URL,
            json={
                'model': model,
                'prompt': prompt,
                'stream': False,
                'options': {
                    'temperature': 0.3,
                    'num_predict': max_tokens
                }
            },
            timeout=120
        )
        if r.status_code == 200:
            return r.json().get('response', '').strip()
    except Exception as e:
        print(f"Erro Ollama: {e}")
    return None

# ============================================================================
# ETAPA 1: CORREÇÃO DE OCR
# ============================================================================

def correct_ocr_chunk(chunk: str, model: str = OLLAMA_MODEL) -> str:
    """Corrige erros de OCR em um chunk."""
    # Verifica cache
    cached = cache.get_correction(chunk)
    if cached:
        return cached

    prompt = f"""Corrija os erros de OCR e ortografia neste texto.

ERROS COMUNS DE OCR PARA CORRIGIR:
- 'rn' que deveria ser 'm' (ex: 'corno' → 'como')
- 'cl' que deveria ser 'd' (ex: 'cle' → 'de')
- 'l' que deveria ser 'i' (ex: 'llvre' → 'livre')
- Espaços errados (ex: 'pa lavra' → 'palavra')
- Letras trocadas ou faltando
- Pontuação incorreta
- Acentuação errada ou faltando

REGRAS:
1. Corrija TODOS os erros ortográficos e de OCR
2. Mantenha a estrutura de parágrafos
3. NÃO altere o significado ou estilo
4. NÃO adicione ou remova conteúdo
5. Retorne APENAS o texto corrigido

TEXTO:
{chunk}

TEXTO CORRIGIDO:"""

    result = call_ollama(prompt, model)

    if result:
        cache.set_correction(chunk, result, model)
        return result
    return chunk

def correct_ocr(text: str, model: str = OLLAMA_MODEL) -> str:
    """Corrige erros de OCR no texto completo."""
    chunks = create_chunks(text)
    corrected = []

    for chunk in tqdm(chunks, desc="Corrigindo OCR", leave=False):
        corrected.append(correct_ocr_chunk(chunk, model))

    return "\n\n".join(corrected)

# ============================================================================
# ETAPA 2: TRADUÇÃO
# ============================================================================

def translate_chunk(chunk: str, source_lang: str, model: str = OLLAMA_MODEL) -> str:
    """Traduz um chunk para português."""
    # Verifica cache
    cached = cache.get_translation(chunk)
    if cached:
        return cached

    lang_names = {
        'en': 'inglês',
        'es': 'espanhol',
        'fr': 'francês',
        'de': 'alemão',
        'ru': 'russo',
        'it': 'italiano',
    }

    lang_name = lang_names.get(source_lang, source_lang)

    prompt = f"""Traduza este texto de {lang_name} para português brasileiro.

REGRAS:
1. Traduza de forma natural e fluida
2. Mantenha o estilo e tom do original
3. Preserve nomes próprios
4. Mantenha a estrutura de parágrafos
5. Retorne APENAS a tradução

TEXTO EM {lang_name.upper()}:
{chunk}

TRADUÇÃO EM PORTUGUÊS:"""

    result = call_ollama(prompt, model, max_tokens=3000)

    if result:
        cache.set_translation(chunk, result, source_lang, model)
        return result
    return chunk

def translate_text(text: str, source_lang: str, model: str = OLLAMA_MODEL) -> str:
    """Traduz texto completo para português."""
    if source_lang == 'pt':
        return text

    chunks = create_chunks(text)
    translated = []

    for chunk in tqdm(chunks, desc=f"Traduzindo ({source_lang}→pt)", leave=False):
        translated.append(translate_chunk(chunk, source_lang, model))

    return "\n\n".join(translated)

# ============================================================================
# ETAPA 3: NOTAS DE RODAPÉ
# ============================================================================

# Termos para detectar
TECHNICAL_TERMS = {
    'dialética', 'metafísica', 'ontologia', 'epistemologia', 'fenomenologia',
    'existencialismo', 'niilismo', 'empirismo', 'racionalismo', 'idealismo',
    'materialismo', 'positivismo', 'pragmatismo', 'hedonismo', 'estoicismo',
    'determinismo', 'transcendental', 'inconsciente', 'subconsciente',
    'ego', 'superego', 'id', 'libido', 'arquétipo', 'karma', 'dharma',
    'nirvana', 'chakra', 'kundalini', 'tantra', 'mantra', 'zen', 'tao',
    'entropia', 'paradigma', 'axioma', 'algoritmo',
}

FOREIGN_PATTERNS = [
    r'\b(?:a priori|a posteriori|ad hoc|alter ego|carpe diem|cogito ergo sum|'
    r'de facto|status quo|tabula rasa|mea culpa|per se|ipso facto)\b',
    r'\b(?:déjà vu|coup d\'état|raison d\'être|vis-à-vis|laissez-faire)\b',
    r'\b(?:zeitgeist|weltanschauung|schadenfreude|übermensch|gestalt)\b',
]

def detect_footnote_terms(text: str) -> Dict[str, str]:
    """Detecta termos que precisam de notas."""
    terms = {}
    text_lower = text.lower()

    for term in TECHNICAL_TERMS:
        if term in text_lower and len(terms) < MAX_FOOTNOTES:
            terms[term] = 'technical'

    for pattern in FOREIGN_PATTERNS:
        if len(terms) >= MAX_FOOTNOTES:
            break
        for match in re.finditer(pattern, text, re.IGNORECASE):
            term = match.group(0).lower()
            if term not in terms:
                terms[term] = 'foreign'

    return terms

def generate_footnote(term: str, term_type: str, model: str = OLLAMA_MODEL) -> str:
    """Gera explicação para nota de rodapé."""
    # Verifica cache
    cached = cache.get_footnote(term)
    if cached:
        return cached

    type_inst = {
        'technical': 'Defina este conceito de forma breve e acessível.',
        'foreign': 'Traduza e explique a origem deste termo estrangeiro.',
    }

    prompt = f"""Nota de rodapé para: {term}

{type_inst.get(term_type, 'Explique brevemente.')}

Regras:
- Máximo 20 palavras
- Seja factual
- Se não souber, escreva: "Termo do contexto literário"

Nota:"""

    result = call_ollama(prompt, model, max_tokens=80)

    if result:
        cache.set_footnote(term, result, model)
        return result
    return "Termo do contexto literário."

def generate_footnotes(text: str, model: str = OLLAMA_MODEL) -> List[Tuple[str, str]]:
    """Gera todas as notas de rodapé."""
    terms = detect_footnote_terms(text)
    footnotes = []

    for term, term_type in tqdm(terms.items(), desc="Gerando notas", leave=False):
        explanation = generate_footnote(term, term_type, model)
        footnotes.append((term, explanation))

    return footnotes

# ============================================================================
# ETAPA 4: CRIAÇÃO DO DOCX
# ============================================================================

def create_footnotes_xml(footnotes: List[Tuple[str, str]], font_size: int = FOOTNOTE_FONT_SIZE) -> bytes:
    """Cria XML das notas de rodapé."""
    nsmap = {'w': NAMESPACES['w']}
    root = etree.Element('{%s}footnotes' % NAMESPACES['w'], nsmap=nsmap)

    # Separators
    for fn_id, fn_type in [('-1', 'separator'), ('0', 'continuationSeparator')]:
        fn = etree.SubElement(root, '{%s}footnote' % NAMESPACES['w'])
        fn.set('{%s}type' % NAMESPACES['w'], fn_type)
        fn.set('{%s}id' % NAMESPACES['w'], fn_id)
        p = etree.SubElement(fn, '{%s}p' % NAMESPACES['w'])
        r = etree.SubElement(p, '{%s}r' % NAMESPACES['w'])
        tag = 'separator' if fn_type == 'separator' else 'continuationSeparator'
        etree.SubElement(r, '{%s}%s' % (NAMESPACES['w'], tag))

    # Notas
    for note_id, (term, text) in enumerate(footnotes, start=1):
        fn = etree.SubElement(root, '{%s}footnote' % NAMESPACES['w'])
        fn.set('{%s}id' % NAMESPACES['w'], str(note_id))

        p = etree.SubElement(fn, '{%s}p' % NAMESPACES['w'])

        # Style
        pPr = etree.SubElement(p, '{%s}pPr' % NAMESPACES['w'])
        pStyle = etree.SubElement(pPr, '{%s}pStyle' % NAMESPACES['w'])
        pStyle.set('{%s}val' % NAMESPACES['w'], 'FootnoteText')

        # Número
        r_ref = etree.SubElement(p, '{%s}r' % NAMESPACES['w'])
        rPr_ref = etree.SubElement(r_ref, '{%s}rPr' % NAMESPACES['w'])
        rStyle = etree.SubElement(rPr_ref, '{%s}rStyle' % NAMESPACES['w'])
        rStyle.set('{%s}val' % NAMESPACES['w'], 'FootnoteReference')
        etree.SubElement(r_ref, '{%s}footnoteRef' % NAMESPACES['w'])

        # Espaço
        r_space = etree.SubElement(p, '{%s}r' % NAMESPACES['w'])
        t_space = etree.SubElement(r_space, '{%s}t' % NAMESPACES['w'])
        t_space.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_space.text = ' '

        # Termo em negrito
        r_term = etree.SubElement(p, '{%s}r' % NAMESPACES['w'])
        rPr_term = etree.SubElement(r_term, '{%s}rPr' % NAMESPACES['w'])
        etree.SubElement(rPr_term, '{%s}b' % NAMESPACES['w'])
        sz = etree.SubElement(rPr_term, '{%s}sz' % NAMESPACES['w'])
        sz.set('{%s}val' % NAMESPACES['w'], str(font_size * 2))
        t_term = etree.SubElement(r_term, '{%s}t' % NAMESPACES['w'])
        t_term.text = f"{term}: "

        # Texto
        r_text = etree.SubElement(p, '{%s}r' % NAMESPACES['w'])
        rPr_text = etree.SubElement(r_text, '{%s}rPr' % NAMESPACES['w'])
        sz2 = etree.SubElement(rPr_text, '{%s}sz' % NAMESPACES['w'])
        sz2.set('{%s}val' % NAMESPACES['w'], str(font_size * 2))
        t_text = etree.SubElement(r_text, '{%s}t' % NAMESPACES['w'])
        t_text.text = text

    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

def insert_footnote_refs(doc_xml: bytes, footnotes: List[Tuple[str, str]]) -> Tuple[bytes, int]:
    """Insere referências de notas no documento."""
    tree = etree.fromstring(doc_xml)
    inserted = set()

    for note_id, (term, _) in enumerate(footnotes, start=1):
        if term.lower() in inserted:
            continue

        for t_elem in tree.iter('{%s}t' % NAMESPACES['w']):
            if t_elem.text and term.lower() in t_elem.text.lower():
                text = t_elem.text
                idx = text.lower().find(term.lower())
                if idx >= 0:
                    before = text[:idx + len(term)]
                    after = text[idx + len(term):]
                    t_elem.text = before

                    r_parent = t_elem.getparent()
                    p_parent = r_parent.getparent()

                    r_ref = etree.Element('{%s}r' % NAMESPACES['w'])
                    rPr = etree.SubElement(r_ref, '{%s}rPr' % NAMESPACES['w'])
                    vertAlign = etree.SubElement(rPr, '{%s}vertAlign' % NAMESPACES['w'])
                    vertAlign.set('{%s}val' % NAMESPACES['w'], 'superscript')
                    fn_ref = etree.SubElement(r_ref, '{%s}footnoteReference' % NAMESPACES['w'])
                    fn_ref.set('{%s}id' % NAMESPACES['w'], str(note_id))

                    idx_in_parent = list(p_parent).index(r_parent)
                    p_parent.insert(idx_in_parent + 1, r_ref)

                    if after:
                        r_after = etree.Element('{%s}r' % NAMESPACES['w'])
                        orig_rPr = r_parent.find('{%s}rPr' % NAMESPACES['w'])
                        if orig_rPr is not None:
                            r_after.append(deepcopy(orig_rPr))
                        t_after = etree.SubElement(r_after, '{%s}t' % NAMESPACES['w'])
                        t_after.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        t_after.text = after
                        p_parent.insert(idx_in_parent + 2, r_after)

                    inserted.add(term.lower())
                    break

    return etree.tostring(tree), len(inserted)

def create_docx(text: str, footnotes: List[Tuple[str, str]], template_path: Path, output_path: Path,
                book_title: str = "Livro", author_name: str = "Autor", original_title: str = "") -> bool:
    """Cria DOCX final com notas de rodapé.

    Estrutura:
    - Páginas 1-3: Template (capa, rosto, ficha técnica) - PRESERVADO
    - Página 4+: Conteúdo do livro - INSERIDO
    - Última página: Informações da editora - PRESERVADO
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import tempfile

    try:
        # Cria documento base
        doc = Document(template_path)

        # Identifica seções do template
        # Início: parágrafos 0-30 (capa, rosto, créditos)
        # Fim: parágrafos com "Also", "Capa:", "ISBN", "EDITORA" (últimas ~15 linhas)

        # Encontra onde começa a página final (informações da editora)
        final_page_start = None
        for i, p in enumerate(doc.paragraphs):
            if 'Also ' in p.text or 'Capa:' in p.text:
                final_page_start = i
                break

        # Guarda parágrafos finais (editora)
        final_paragraphs = []
        if final_page_start:
            for p in doc.paragraphs[final_page_start:]:
                final_paragraphs.append(p.text)

        # Encontra onde termina o cabeçalho (após créditos iniciais)
        header_end = 0
        for i, p in enumerate(doc.paragraphs[:35]):
            text = p.text.strip()

            # Substitui placeholders
            if 'Nome do' in text:
                p.clear()
            elif text == 'Livro':
                p.clear()
                run = p.add_run(book_title)
                run.bold = True
                run.font.size = Pt(24)
            elif 'Nome do Altor' in text or 'Nome do Autor' in text:
                p.clear()
                p.add_run(author_name)
            elif 'Also ' in text and original_title:
                p.clear()
                p.add_run(f"Also {original_title}")

            # Marca fim do cabeçalho
            if 'web2ajax' in text or 'Jose R F Junior' in text:
                header_end = i + 5  # Alguns parágrafos vazios depois

        # Remove parágrafos do meio (entre cabeçalho e rodapé)
        # Mantém só cabeçalho + rodapé
        if final_page_start and header_end:
            # Remove do header_end até final_page_start
            paragraphs_to_remove = []
            for i in range(header_end, final_page_start):
                if i < len(doc.paragraphs):
                    paragraphs_to_remove.append(doc.paragraphs[i])

            for p in paragraphs_to_remove:
                p._element.getparent().remove(p._element)

        # Encontra posição para inserir conteúdo (após cabeçalho)
        insert_after = min(header_end, len(doc.paragraphs) - 1)

        # Adiciona quebra de página
        if insert_after < len(doc.paragraphs):
            run = doc.paragraphs[insert_after].add_run()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run._r.append(br)

        # Insere conteúdo do livro
        insert_element = doc.paragraphs[insert_after]._element

        for para_text in text.split('\n\n'):
            para_text = para_text.strip()
            if para_text:
                # Cria novo parágrafo
                new_p = doc.add_paragraph()
                new_p.text = para_text
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Move para posição correta (antes da página final)
                insert_element.addnext(new_p._element)
                insert_element = new_p._element

        # Salva temporário
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        # Reabre e adiciona footnotes
        with zipfile.ZipFile(tmp_path, 'r') as zin:
            files = {name: zin.read(name) for name in zin.namelist()}

        os.unlink(tmp_path)

        if footnotes:
            files['word/footnotes.xml'] = create_footnotes_xml(footnotes)
            files['word/document.xml'], _ = insert_footnote_refs(files['word/document.xml'], footnotes)

            # Content_Types
            ct = files['[Content_Types].xml'].decode('utf-8')
            if 'footnotes.xml' not in ct:
                pos = ct.find('</Types>')
                ct = ct[:pos] + '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>' + ct[pos:]
                files['[Content_Types].xml'] = ct.encode('utf-8')

            # Relationships
            rels_path = 'word/_rels/document.xml.rels'
            if rels_path in files:
                rels = files[rels_path].decode('utf-8')
                if 'footnotes.xml' not in rels:
                    rids = re.findall(r'Id="rId(\d+)"', rels)
                    next_rid = max(int(r) for r in rids) + 1 if rids else 1
                    pos = rels.find('</Relationships>')
                    rels = rels[:pos] + f'<Relationship Id="rId{next_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>' + rels[pos:]
                    files[rels_path] = rels.encode('utf-8')

        # Salva final
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name, content in files.items():
                zout.writestr(name, content)

        return True

    except Exception as e:
        print(f"Erro ao criar DOCX: {e}")
        return False

# ============================================================================
# PIPELINE PRINCIPAL
# ============================================================================

def process_book(input_path: Path, output_path: Path, template_path: Path,
                 skip_correction: bool = False, skip_translation: bool = False,
                 skip_footnotes: bool = False, model: str = OLLAMA_MODEL) -> Dict:
    """Processa um livro completo pelo pipeline."""

    result = {
        'input': str(input_path),
        'output': str(output_path),
        'status': 'success',
        'language': None,
        'chars_original': 0,
        'chars_final': 0,
        'footnotes': 0,
        'error': None
    }

    try:
        # Lê arquivo
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()

        result['chars_original'] = len(text)

        if len(text) < 500:
            result['status'] = 'skipped'
            result['error'] = 'Texto muito curto'
            return result

        # Detecta idioma
        lang = detect_language(text)
        result['language'] = lang

        # Etapa 1: Correção OCR
        if not skip_correction:
            print(f"  [1/4] Corrigindo OCR...")
            text = correct_ocr(text, model)

        # Etapa 2: Tradução
        if not skip_translation and lang != 'pt':
            print(f"  [2/4] Traduzindo ({lang} → pt)...")
            text = translate_text(text, lang, model)
        else:
            print(f"  [2/4] Tradução: pulada (já em português)")

        # Etapa 3: Notas de rodapé
        footnotes = []
        if not skip_footnotes:
            print(f"  [3/4] Gerando notas de rodapé...")
            footnotes = generate_footnotes(text, model)
            result['footnotes'] = len(footnotes)

        # Etapa 4: Criar DOCX
        print(f"  [4/4] Criando DOCX...")
        success = create_docx(text, footnotes, template_path, output_path)

        if not success:
            result['status'] = 'error'
            result['error'] = 'Falha ao criar DOCX'

        result['chars_final'] = len(text)

    except Exception as e:
        result['status'] = 'error'
        result['error'] = str(e)

    return result

# ============================================================================
# MAIN
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Pipeline Completo: TXT → Correção → Tradução → Notas → DOCX"
    )
    parser.add_argument('--input', '-i', type=Path, help='Arquivo TXT de entrada')
    parser.add_argument('--output', '-o', type=Path, help='Arquivo DOCX de saída')
    parser.add_argument('--template', '-t', type=Path, default=TEMPLATE_PATH, help='Template DOCX')
    parser.add_argument('--model', '-m', default=OLLAMA_MODEL, help='Modelo Ollama')

    parser.add_argument('--batch', action='store_true', help='Modo batch (processa pasta)')
    parser.add_argument('--input-dir', type=Path, default=INPUT_DIR, help='Pasta de entrada (batch)')
    parser.add_argument('--output-dir', type=Path, default=OUTPUT_DIR, help='Pasta de saída (batch)')

    parser.add_argument('--skip-correction', action='store_true', help='Pular correção OCR')
    parser.add_argument('--skip-translation', action='store_true', help='Pular tradução')
    parser.add_argument('--skip-footnotes', action='store_true', help='Pular notas de rodapé')

    parser.add_argument('--limit', '-l', type=int, default=0, help='Limitar arquivos (batch)')

    args = parser.parse_args()

    print("="*60)
    print("PIPELINE COMPLETO - TXT para DOCX Amazon KDP")
    print("="*60)
    print(f"Modelo: {args.model}")
    print(f"Template: {args.template}")
    print(f"Correção OCR: {'Não' if args.skip_correction else 'Sim'}")
    print(f"Tradução: {'Não' if args.skip_translation else 'Sim'}")
    print(f"Notas de rodapé: {'Não' if args.skip_footnotes else 'Sim'}")
    print()

    # Modo arquivo único
    if args.input:
        if not args.input.exists():
            print(f"Arquivo não encontrado: {args.input}")
            return 1

        output = args.output or args.input.with_suffix('.docx')

        print(f"Processando: {args.input.name}")
        result = process_book(
            args.input, output, args.template,
            args.skip_correction, args.skip_translation, args.skip_footnotes,
            args.model
        )

        print()
        print(f"Status: {result['status']}")
        print(f"Idioma: {result['language']}")
        print(f"Caracteres: {result['chars_original']:,} → {result['chars_final']:,}")
        print(f"Notas de rodapé: {result['footnotes']}")
        if result['error']:
            print(f"Erro: {result['error']}")
        print(f"Saída: {output}")

        return 0 if result['status'] == 'success' else 1

    # Modo batch
    if args.batch:
        files = list(args.input_dir.rglob('*.txt'))
        if args.limit > 0:
            files = files[:args.limit]

        print(f"Arquivos: {len(files)}")
        print()

        success = 0
        errors = 0

        for f in files:
            book_name = f.stem.replace('_pt', '')
            output = args.output_dir / f"{book_name}.docx"

            if output.exists():
                print(f"[SKIP] {f.name} (já existe)")
                continue

            print(f"[PROC] {f.name}")
            result = process_book(
                f, output, args.template,
                args.skip_correction, args.skip_translation, args.skip_footnotes,
                args.model
            )

            if result['status'] == 'success':
                success += 1
                print(f"       → OK ({result['footnotes']} notas)")
            else:
                errors += 1
                print(f"       → ERRO: {result['error']}")

        print()
        print("="*60)
        print(f"Concluído: {success} sucesso, {errors} erros")
        print(f"Saída: {args.output_dir}")

        return 0 if errors == 0 else 1

    parser.print_help()
    return 0


if __name__ == "__main__":
    sys.exit(main())

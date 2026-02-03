# -*- coding: utf-8 -*-
"""
PIPELINE KDP COMPLETO - Com Validação e Template
================================================
1. Valida se TXT está completo
2. Traduz/Corrige com Gemini
3. Gera notas de rodapé REAIS (no pé da página)
4. Aplica template Estrutura.docx

Uso:
    python pipeline_kdp.py --validar              # Apenas valida TXTs
    python pipeline_kdp.py --processar --limit 5  # Processa livros válidos
"""

import os
import sys
import asyncio
import aiohttp
import hashlib
import sqlite3
import argparse
import zipfile
import re
import json
import time
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from lxml import etree
from copy import deepcopy

# =============================================================================
# CONFIGURAÇÕES
# =============================================================================

VERTEX_API_KEY = os.getenv("VERTEX_API_KEY", "AQ.Ab8RN6I5q1OHyDoU18jD39-8LwCZwPlXHd9SyGsPtLzmzjF_hw")
VERTEX_ENDPOINT = "https://aiplatform.googleapis.com/v1/publishers/google/models"
MODEL = "gemini-2.0-flash"

BASE_DIR = Path(__file__).parent
TXT_DIR = BASE_DIR / "txt"
DOCX_DIR = BASE_DIR / "docx"
DATA_DIR = BASE_DIR / "data"
TEMPLATE_PATH = BASE_DIR / "Estrutura.docx"

for d in [DOCX_DIR, DATA_DIR]:
    d.mkdir(exist_ok=True)

# DOCX Namespaces
NAMESPACES = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# Critérios de validação
MIN_CHARS = 10000          # Mínimo 10k caracteres
MIN_PARAGRAPHS = 20        # Mínimo 20 parágrafos
MIN_WORDS = 2000           # Mínimo 2000 palavras

# =============================================================================
# VALIDAÇÃO DE LIVROS
# =============================================================================

@dataclass
class ValidacaoResult:
    titulo: str
    valido: bool
    chars: int
    palavras: int
    paragrafos: int
    idioma: str
    tem_inicio: bool
    tem_fim: bool
    motivo: str = ""
    txt_path: str = ""

def detectar_idioma(texto: str) -> str:
    """Detecta idioma do texto."""
    texto = texto[:5000].lower()
    langs = {
        'pt': ['que', 'não', 'para', 'uma', 'com', 'por', 'mais', 'você', 'isso', 'está'],
        'es': ['que', 'los', 'las', 'una', 'para', 'con', 'más', 'pero', 'está'],
        'en': ['the', 'and', 'that', 'have', 'for', 'with', 'this', 'but', 'from'],
        'fr': ['les', 'des', 'une', 'que', 'pour', 'dans', 'qui', 'sur'],
        'de': ['der', 'die', 'und', 'den', 'das', 'von', 'ist', 'mit'],
        'ru': ['что', 'как', 'это', 'его', 'она', 'они', 'все', 'для'],
    }
    scores = {lang: sum(1 for w in words if f' {w} ' in f' {texto} ')
              for lang, words in langs.items()}
    return max(scores, key=scores.get)

def validar_txt(txt_path: Path) -> ValidacaoResult:
    """Valida se um TXT está completo."""
    titulo = txt_path.parent.name

    try:
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            texto = f.read()
    except Exception as e:
        return ValidacaoResult(
            titulo=titulo, valido=False, chars=0, palavras=0, paragrafos=0,
            idioma='?', tem_inicio=False, tem_fim=False,
            motivo=f"Erro ao ler: {e}", txt_path=str(txt_path)
        )

    # Métricas básicas
    chars = len(texto)
    palavras = len(texto.split())
    paragrafos = len([p for p in texto.split('\n\n') if p.strip()])
    idioma = detectar_idioma(texto)

    # Verifica início (tem título, prefácio, introdução, capítulo 1, etc)
    inicio_patterns = [
        r'(?i)(prefácio|preface|prólogo|prologue|introdução|introduction)',
        r'(?i)(capítulo|chapter|parte|part)\s*(1|i|um|one|primeiro|first)',
        r'(?i)^[IVX]+[\.\s]',  # Números romanos
    ]
    tem_inicio = any(re.search(p, texto[:5000]) for p in inicio_patterns)

    # Verifica fim (tem epílogo, conclusão, fim, etc)
    fim_patterns = [
        r'(?i)(epílogo|epilogue|conclusão|conclusion|posfácio|afterword)',
        r'(?i)(fim|end|the end|fine|finis)[\s\.\n]*$',
        r'(?i)(notas|notes|bibliografia|bibliography|referências|references)',
    ]
    tem_fim = any(re.search(p, texto[-5000:]) for p in fim_patterns)

    # Validação
    motivos = []

    if chars < MIN_CHARS:
        motivos.append(f"Muito curto ({chars} chars, min {MIN_CHARS})")

    if palavras < MIN_WORDS:
        motivos.append(f"Poucas palavras ({palavras}, min {MIN_WORDS})")

    if paragrafos < MIN_PARAGRAPHS:
        motivos.append(f"Poucos parágrafos ({paragrafos}, min {MIN_PARAGRAPHS})")

    # Verifica se parece truncado (termina no meio de frase)
    ultimo_para = texto.strip().split('\n\n')[-1].strip() if texto.strip() else ""
    if ultimo_para and not ultimo_para[-1] in '.!?"\'»':
        motivos.append("Parece truncado (não termina com pontuação)")

    valido = len(motivos) == 0

    return ValidacaoResult(
        titulo=titulo,
        valido=valido,
        chars=chars,
        palavras=palavras,
        paragrafos=paragrafos,
        idioma=idioma,
        tem_inicio=tem_inicio,
        tem_fim=tem_fim,
        motivo="; ".join(motivos) if motivos else "OK",
        txt_path=str(txt_path)
    )

def validar_todos_txts(limite: int = 0) -> List[ValidacaoResult]:
    """Valida todos os TXTs."""
    resultados = []

    folders = list(TXT_DIR.iterdir())
    if limite > 0:
        folders = folders[:limite]

    for folder in folders:
        if not folder.is_dir():
            continue

        txt_files = list(folder.glob("*.txt"))
        if not txt_files:
            continue

        # Usa o maior arquivo
        main_file = max(txt_files, key=lambda f: f.stat().st_size)
        resultado = validar_txt(main_file)
        resultados.append(resultado)

    return resultados

# =============================================================================
# GEMINI API
# =============================================================================

async def call_gemini(session: aiohttp.ClientSession, prompt: str,
                      semaphore: asyncio.Semaphore) -> str:
    """Chama Gemini com retry."""
    async with semaphore:
        url = f"{VERTEX_ENDPOINT}/{MODEL}:generateContent?key={VERTEX_API_KEY}"
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.3, "maxOutputTokens": 8192}
        }

        for attempt in range(5):
            try:
                async with session.post(url, json=payload) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        if "candidates" in data:
                            parts = data["candidates"][0].get("content", {}).get("parts", [])
                            if parts:
                                return parts[0].get("text", "").strip()
                        return ""
                    elif resp.status == 429:
                        await asyncio.sleep(2 ** attempt)
                    else:
                        await asyncio.sleep(1)
            except:
                await asyncio.sleep(1)
        return ""

# =============================================================================
# NOTAS DE RODAPÉ REAIS (XML)
# =============================================================================

def criar_footnotes_xml(notas: List[Tuple[str, str]], font_size: int = 9) -> bytes:
    """Cria XML das notas de rodapé para DOCX."""
    nsmap = {'w': NAMESPACES['w']}
    root = etree.Element('{%s}footnotes' % NAMESPACES['w'], nsmap=nsmap)

    # Separator (id=-1)
    sep = etree.SubElement(root, '{%s}footnote' % NAMESPACES['w'])
    sep.set('{%s}type' % NAMESPACES['w'], 'separator')
    sep.set('{%s}id' % NAMESPACES['w'], '-1')
    p1 = etree.SubElement(sep, '{%s}p' % NAMESPACES['w'])
    r1 = etree.SubElement(p1, '{%s}r' % NAMESPACES['w'])
    etree.SubElement(r1, '{%s}separator' % NAMESPACES['w'])

    # Continuation separator (id=0)
    cont = etree.SubElement(root, '{%s}footnote' % NAMESPACES['w'])
    cont.set('{%s}type' % NAMESPACES['w'], 'continuationSeparator')
    cont.set('{%s}id' % NAMESPACES['w'], '0')
    p2 = etree.SubElement(cont, '{%s}p' % NAMESPACES['w'])
    r2 = etree.SubElement(p2, '{%s}r' % NAMESPACES['w'])
    etree.SubElement(r2, '{%s}continuationSeparator' % NAMESPACES['w'])

    # Notas
    for note_id, (termo, texto) in enumerate(notas, start=1):
        fn = etree.SubElement(root, '{%s}footnote' % NAMESPACES['w'])
        fn.set('{%s}id' % NAMESPACES['w'], str(note_id))

        p = etree.SubElement(fn, '{%s}p' % NAMESPACES['w'])

        # Style
        pPr = etree.SubElement(p, '{%s}pPr' % NAMESPACES['w'])
        pStyle = etree.SubElement(pPr, '{%s}pStyle' % NAMESPACES['w'])
        pStyle.set('{%s}val' % NAMESPACES['w'], 'FootnoteText')

        # Número da nota
        r_ref = etree.SubElement(p, '{%s}r' % NAMESPACES['w'])
        rPr_ref = etree.SubElement(r_ref, '{%s}rPr' % NAMESPACES['w'])
        rStyle = etree.SubElement(rPr_ref, '{%s}rStyle' % NAMESPACES['w'])
        rStyle.set('{%s}val' % NAMESPACES['w'], 'FootnoteReference')
        etree.SubElement(r_ref, '{%s}footnoteRef' % NAMESPACES['w'])

        # Espaço
        r_space = etree.SubElement(p, '{%s}r' % NAMESPACES['w'])
        t_space = etree.SubElement(r_space, '{%s}t' % NAMESPACES['w'])
        t_space.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_space.text = ' '

        # Termo em negrito
        r_term = etree.SubElement(p, '{%s}r' % NAMESPACES['w'])
        rPr_term = etree.SubElement(r_term, '{%s}rPr' % NAMESPACES['w'])
        etree.SubElement(rPr_term, '{%s}b' % NAMESPACES['w'])
        sz = etree.SubElement(rPr_term, '{%s}sz' % NAMESPACES['w'])
        sz.set('{%s}val' % NAMESPACES['w'], str(font_size * 2))
        t_term = etree.SubElement(r_term, '{%s}t' % NAMESPACES['w'])
        t_term.text = f"{termo}: "

        # Texto da nota
        r_text = etree.SubElement(p, '{%s}r' % NAMESPACES['w'])
        rPr_text = etree.SubElement(r_text, '{%s}rPr' % NAMESPACES['w'])
        sz2 = etree.SubElement(rPr_text, '{%s}sz' % NAMESPACES['w'])
        sz2.set('{%s}val' % NAMESPACES['w'], str(font_size * 2))
        t_text = etree.SubElement(r_text, '{%s}t' % NAMESPACES['w'])
        t_text.text = texto

    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

def inserir_refs_notas(doc_xml: bytes, notas: List[Tuple[str, str]]) -> Tuple[bytes, set]:
    """Insere referências de notas no documento."""
    tree = etree.fromstring(doc_xml)
    inseridas = set()

    for note_id, (termo, _) in enumerate(notas, start=1):
        if termo.lower() in inseridas:
            continue

        for t_elem in tree.iter('{%s}t' % NAMESPACES['w']):
            if t_elem.text and termo.lower() in t_elem.text.lower():
                text = t_elem.text
                idx = text.lower().find(termo.lower())
                if idx >= 0:
                    before = text[:idx + len(termo)]
                    after = text[idx + len(termo):]
                    t_elem.text = before

                    r_parent = t_elem.getparent()
                    p_parent = r_parent.getparent()

                    # Cria referência de nota
                    r_ref = etree.Element('{%s}r' % NAMESPACES['w'])
                    rPr = etree.SubElement(r_ref, '{%s}rPr' % NAMESPACES['w'])
                    vertAlign = etree.SubElement(rPr, '{%s}vertAlign' % NAMESPACES['w'])
                    vertAlign.set('{%s}val' % NAMESPACES['w'], 'superscript')
                    fn_ref = etree.SubElement(r_ref, '{%s}footnoteReference' % NAMESPACES['w'])
                    fn_ref.set('{%s}id' % NAMESPACES['w'], str(note_id))

                    idx_in_parent = list(p_parent).index(r_parent)
                    p_parent.insert(idx_in_parent + 1, r_ref)

                    if after:
                        r_after = etree.Element('{%s}r' % NAMESPACES['w'])
                        orig_rPr = r_parent.find('{%s}rPr' % NAMESPACES['w'])
                        if orig_rPr is not None:
                            r_after.append(deepcopy(orig_rPr))
                        t_after = etree.SubElement(r_after, '{%s}t' % NAMESPACES['w'])
                        t_after.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        t_after.text = after
                        p_parent.insert(idx_in_parent + 2, r_after)

                    inseridas.add(termo.lower())
                    break

    return etree.tostring(tree), inseridas

# =============================================================================
# CRIAÇÃO DOCX COM TEMPLATE
# =============================================================================

def criar_docx_com_template(texto: str, titulo: str, autor: str,
                            notas: List[Tuple[str, str]],
                            output_path: Path) -> bool:
    """Cria DOCX usando template e com notas de rodapé reais."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        # Carrega template
        if TEMPLATE_PATH.exists():
            doc = Document(TEMPLATE_PATH)
        else:
            doc = Document()

        # Limpa conteúdo do template (mantém primeiras 3 páginas se houver)
        # Para simplificar, vamos adicionar após o conteúdo existente

        # Adiciona quebra de página
        doc.add_page_break()

        # Título do livro
        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(24)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Autor
        p = doc.add_paragraph()
        p.add_run(autor)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Conteúdo
        for para in texto.split('\n\n'):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Define fonte
                for run in p.runs:
                    run.font.size = Pt(11)

        # Salva temporário
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        # Adiciona notas de rodapé via XML
        if notas:
            with zipfile.ZipFile(tmp_path, 'r') as zin:
                files = {name: zin.read(name) for name in zin.namelist()}

            os.unlink(tmp_path)

            # Cria footnotes.xml
            files['word/footnotes.xml'] = criar_footnotes_xml(notas)

            # Insere referências no documento
            files['word/document.xml'], inseridas = inserir_refs_notas(
                files['word/document.xml'], notas
            )

            # Atualiza Content_Types
            ct = files['[Content_Types].xml'].decode('utf-8')
            if 'footnotes.xml' not in ct:
                pos = ct.find('</Types>')
                override = '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                ct = ct[:pos] + override + ct[pos:]
                files['[Content_Types].xml'] = ct.encode('utf-8')

            # Atualiza relationships
            rels_path = 'word/_rels/document.xml.rels'
            if rels_path in files:
                rels = files[rels_path].decode('utf-8')
                if 'footnotes.xml' not in rels:
                    rids = re.findall(r'Id="rId(\d+)"', rels)
                    next_rid = max(int(r) for r in rids) + 1 if rids else 1
                    pos = rels.find('</Relationships>')
                    rel = f'<Relationship Id="rId{next_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                    rels = rels[:pos] + rel + rels[pos:]
                    files[rels_path] = rels.encode('utf-8')

            # Salva final
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for name, content in files.items():
                    zout.writestr(name, content)
        else:
            # Sem notas, só renomeia
            os.rename(tmp_path, output_path)

        return True

    except Exception as e:
        print(f"Erro DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False

# =============================================================================
# PROCESSAMENTO
# =============================================================================

def dividir_chunks(texto: str, tamanho: int = 5000) -> List[str]:
    """Divide texto em chunks."""
    chunks = []
    current = ""
    for para in texto.split('\n\n'):
        if len(current) + len(para) < tamanho:
            current += '\n\n' + para if current else para
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)
    return chunks if chunks else [texto]

async def processar_livro(validacao: ValidacaoResult) -> Dict:
    """Processa um livro validado."""
    resultado = {
        'titulo': validacao.titulo,
        'status': 'erro',
        'notas': 0
    }

    print(f"\n{'='*50}")
    print(f"Processando: {validacao.titulo}")
    print(f"Idioma: {validacao.idioma}")
    print('='*50)

    # Lê texto
    with open(validacao.txt_path, 'r', encoding='utf-8', errors='ignore') as f:
        texto = f.read()

    timeout = aiohttp.ClientTimeout(total=120)
    semaphore = asyncio.Semaphore(5)

    async with aiohttp.ClientSession(timeout=timeout) as session:
        try:
            # 1. Correção OCR
            print("  [1/4] Corrigindo OCR...")
            chunks = dividir_chunks(texto)
            tasks = []
            for chunk in chunks:
                prompt = f"Corrija erros de OCR neste texto. Retorne APENAS o texto corrigido.\n\n{chunk}"
                tasks.append(call_gemini(session, prompt, semaphore))

            resultados_ocr = await asyncio.gather(*tasks)
            texto = '\n\n'.join([r for r in resultados_ocr if r])

            # 2. Tradução (se necessário)
            if validacao.idioma != 'pt':
                print(f"  [2/4] Traduzindo ({validacao.idioma} → pt)...")
                lang_map = {'en': 'inglês', 'es': 'espanhol', 'fr': 'francês', 'de': 'alemão', 'ru': 'russo'}
                chunks = dividir_chunks(texto)
                tasks = []
                for chunk in chunks:
                    prompt = f"Traduza de {lang_map.get(validacao.idioma, validacao.idioma)} para português brasileiro. Retorne APENAS a tradução.\n\n{chunk}"
                    tasks.append(call_gemini(session, prompt, semaphore))

                resultados_trad = await asyncio.gather(*tasks)
                texto = '\n\n'.join([r for r in resultados_trad if r])
            else:
                print("  [2/4] Tradução: pulada (já em PT)")

            # 3. Notas de rodapé
            print("  [3/4] Gerando notas de rodapé...")
            prompt = f"""Analise este texto e identifique termos que precisam de notas de rodapé:
- Termos estrangeiros (latim, francês, alemão)
- Conceitos filosóficos/técnicos importantes
- Nomes de pessoas históricas relevantes
- Referências a obras

Para cada termo, forneça uma explicação FACTUAL e BREVE (máximo 25 palavras).
Formato: termo|explicação
Máximo 20 termos mais importantes.

TEXTO:
{texto[:8000]}"""

            notas_raw = await call_gemini(session, prompt, semaphore)

            notas = []
            if notas_raw:
                for linha in notas_raw.split('\n'):
                    if '|' in linha:
                        partes = linha.split('|', 1)
                        if len(partes) == 2:
                            termo = partes[0].strip().strip('- *')
                            explicacao = partes[1].strip()
                            if termo and explicacao and len(termo) < 80:
                                notas.append((termo, explicacao))

            notas = notas[:20]
            resultado['notas'] = len(notas)

            # 4. Criar DOCX com template e notas reais
            print(f"  [4/4] Criando DOCX com template ({len(notas)} notas)...")
            docx_path = DOCX_DIR / f"{validacao.titulo}.docx"

            if criar_docx_com_template(texto, validacao.titulo, "Autor", notas, docx_path):
                resultado['status'] = 'sucesso'
                resultado['docx'] = str(docx_path)

        except Exception as e:
            resultado['error'] = str(e)
            print(f"ERRO: {e}")

    return resultado

# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description="Pipeline KDP Completo")
    parser.add_argument('--validar', action='store_true', help='Valida TXTs')
    parser.add_argument('--processar', action='store_true', help='Processa livros válidos')
    parser.add_argument('--limit', type=int, default=0, help='Limite de livros')
    args = parser.parse_args()

    if args.validar:
        print("="*60)
        print("VALIDAÇÃO DE LIVROS")
        print("="*60)

        resultados = validar_todos_txts(args.limit)

        validos = [r for r in resultados if r.valido]
        invalidos = [r for r in resultados if not r.valido]

        print(f"\nTotal: {len(resultados)}")
        print(f"Válidos: {len(validos)}")
        print(f"Inválidos: {len(invalidos)}")

        # Por idioma (válidos)
        por_idioma = {}
        for r in validos:
            por_idioma[r.idioma] = por_idioma.get(r.idioma, 0) + 1
        print(f"\nVálidos por idioma: {por_idioma}")

        # Mostra inválidos
        if invalidos:
            print(f"\n{'='*60}")
            print("LIVROS INVÁLIDOS:")
            print('='*60)
            for r in invalidos[:20]:
                print(f"  {r.titulo[:40]}: {r.motivo}")

        # Salva resultado
        output = DATA_DIR / "validacao.json"
        with open(output, 'w', encoding='utf-8') as f:
            json.dump([{
                'titulo': r.titulo,
                'valido': r.valido,
                'chars': r.chars,
                'palavras': r.palavras,
                'idioma': r.idioma,
                'motivo': r.motivo,
                'txt_path': r.txt_path
            } for r in resultados], f, ensure_ascii=False, indent=2)
        print(f"\nSalvo em: {output}")

    elif args.processar:
        print("="*60)
        print("PROCESSAMENTO KDP")
        print("="*60)

        # Carrega validação
        validacao_path = DATA_DIR / "validacao.json"
        if not validacao_path.exists():
            print("Execute --validar primeiro!")
            return

        with open(validacao_path, 'r') as f:
            dados = json.load(f)

        # Filtra válidos
        validos = [d for d in dados if d['valido']]

        if args.limit > 0:
            validos = validos[:args.limit]

        print(f"Livros a processar: {len(validos)}")

        # Processa
        for i, d in enumerate(validos):
            validacao = ValidacaoResult(
                titulo=d['titulo'],
                valido=True,
                chars=d['chars'],
                palavras=d['palavras'],
                paragrafos=0,
                idioma=d['idioma'],
                tem_inicio=True,
                tem_fim=True,
                txt_path=d['txt_path']
            )

            print(f"\n[{i+1}/{len(validos)}]")
            resultado = asyncio.run(processar_livro(validacao))

            status = "OK" if resultado['status'] == 'sucesso' else "ERRO"
            print(f"  → {status} ({resultado.get('notas', 0)} notas)")

        print("\n" + "="*60)
        print(f"CONCLUÍDO! DOCX em: {DOCX_DIR}")
        print("="*60)

    else:
        parser.print_help()
        print("\nFluxo:")
        print("  1. python pipeline_kdp.py --validar")
        print("  2. python pipeline_kdp.py --processar --limit 5")

if __name__ == "__main__":
    main()

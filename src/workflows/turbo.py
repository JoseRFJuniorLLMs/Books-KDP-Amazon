# -*- coding: utf-8 -*-
"""
PIPELINE TURBO - Múltiplos processos paralelos
==============================================
Divide os livros em N workers e processa em paralelo.

Uso:
    python pipeline_turbo.py --workers 10 --run
    python pipeline_turbo.py --workers 20 --run --limit 100
"""

import sys
import json
import multiprocessing as mp
import asyncio
import aiohttp
import hashlib
import sqlite3
import argparse
import time
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass

# Adiciona raiz do projeto ao path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

# Configuracao centralizada
from config.settings import VERTEX_PROJECT, VERTEX_REGION

# =============================================================================
# CONFIGURAÇÕES
# =============================================================================

VERTEX_ENDPOINT = "https://aiplatform.googleapis.com/v1/publishers/google/models"
MODEL = "gemini-2.0-flash"

BASE_DIR = Path(__file__).parent
TXT_DIR = BASE_DIR / "txt"
DOCX_DIR = BASE_DIR / "docx"
DATA_DIR = BASE_DIR / "data"
TRANSLATED_DIR = BASE_DIR / "translated-portugues"

for d in [DOCX_DIR, DATA_DIR, TRANSLATED_DIR]:
    d.mkdir(exist_ok=True)

# =============================================================================
# FUNÇÕES DO WORKER
# =============================================================================

LANG_WORDS = {
    'pt': ['que', 'não', 'para', 'uma', 'com', 'por', 'mais', 'você'],
    'es': ['que', 'los', 'las', 'una', 'para', 'con', 'más', 'pero'],
    'en': ['the', 'and', 'that', 'have', 'for', 'with', 'this', 'but'],
    'fr': ['les', 'des', 'une', 'que', 'pour', 'dans', 'qui', 'sur'],
    'de': ['der', 'die', 'und', 'den', 'das', 'von', 'ist', 'mit'],
    'ru': ['что', 'как', 'это', 'его', 'она', 'они', 'все', 'для'],
}

def detectar_idioma(texto: str) -> str:
    texto = texto[:5000].lower()
    scores = {lang: sum(1 for w in words if f' {w} ' in f' {texto} ')
              for lang, words in LANG_WORDS.items()}
    return max(scores, key=scores.get)

def dividir_chunks(texto: str, tamanho: int = 5000) -> List[str]:
    chunks = []
    current = ""
    for para in texto.split('\n\n'):
        if len(current) + len(para) < tamanho:
            current += '\n\n' + para if current else para
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)
    return chunks if chunks else [texto]

async def call_gemini(session: aiohttp.ClientSession, prompt: str, semaphore: asyncio.Semaphore) -> str:
    """Chama Gemini com retry."""
    async with semaphore:
        url = f"{VERTEX_ENDPOINT}/{MODEL}:generateContent?key={VERTEX_API_KEY}"
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.3, "maxOutputTokens": 8192}
        }

        for attempt in range(5):
            try:
                async with session.post(url, json=payload) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        if "candidates" in data:
                            parts = data["candidates"][0].get("content", {}).get("parts", [])
                            if parts:
                                return parts[0].get("text", "").strip()
                        return ""
                    elif resp.status == 429:
                        await asyncio.sleep(2 ** attempt)
                    else:
                        await asyncio.sleep(1)
            except Exception as e:
                await asyncio.sleep(1)
        return ""

async def processar_livro_async(titulo: str, texto: str, idioma: str, worker_id: int) -> Dict:
    """Processa um livro de forma assíncrona."""
    resultado = {'titulo': titulo, 'status': 'erro', 'notas': 0}

    timeout = aiohttp.ClientTimeout(total=120)
    semaphore = asyncio.Semaphore(5)  # 5 requests paralelos por livro

    async with aiohttp.ClientSession(timeout=timeout) as session:
        try:
            # 1. Correção OCR
            chunks = dividir_chunks(texto)
            tasks = []
            for chunk in chunks:
                prompt = f"Corrija erros de OCR. Retorne APENAS o texto corrigido.\n\n{chunk}"
                tasks.append(call_gemini(session, prompt, semaphore))

            resultados_ocr = await asyncio.gather(*tasks)
            texto = '\n\n'.join([r for r in resultados_ocr if r])

            # 2. Tradução
            if idioma != 'pt' and texto:
                lang_map = {'en': 'inglês', 'es': 'espanhol', 'fr': 'francês', 'de': 'alemão', 'ru': 'russo'}
                chunks = dividir_chunks(texto)
                tasks = []
                for chunk in chunks:
                    prompt = f"Traduza de {lang_map.get(idioma, idioma)} para português. Retorne APENAS a tradução.\n\n{chunk}"
                    tasks.append(call_gemini(session, prompt, semaphore))

                resultados_trad = await asyncio.gather(*tasks)
                texto = '\n\n'.join([r for r in resultados_trad if r])

            # 3. Notas
            if texto:
                prompt = f"""Identifique termos para notas de rodapé (termos estrangeiros, conceitos, nomes históricos).
Formato: termo|explicação breve (max 25 palavras)
Máximo 15 termos.

{texto[:5000]}"""
                notas_raw = await call_gemini(session, prompt, semaphore)

                notas = []
                for linha in notas_raw.split('\n'):
                    if '|' in linha:
                        partes = linha.split('|', 1)
                        if len(partes) == 2:
                            termo = partes[0].strip().strip('- ')
                            explicacao = partes[1].strip()
                            if termo and explicacao:
                                notas.append((termo, explicacao))

                resultado['notas'] = len(notas[:15])

                # 4. DOCX
                criar_docx(texto, titulo, "Autor", notas[:15], DOCX_DIR / f"{titulo}.docx")

                # Salva TXT
                with open(TRANSLATED_DIR / f"{titulo}_pt.txt", 'w', encoding='utf-8') as f:
                    f.write(texto)

                resultado['status'] = 'sucesso'

        except Exception as e:
            resultado['error'] = str(e)

    return resultado

def criar_docx(texto: str, titulo: str, autor: str, notas: List[Tuple[str, str]], output_path: Path):
    """Cria DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(24)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.add_run(autor)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        for para in texto.split('\n\n'):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if notas:
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run("NOTAS")
            run.bold = True

            for i, (termo, explicacao) in enumerate(notas, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{i}. {termo}: ")
                run.bold = True
                run.font.size = Pt(9)
                run = p.add_run(explicacao)
                run.font.size = Pt(9)

        doc.save(output_path)
    except Exception as e:
        print(f"[W{os.getpid()}] Erro DOCX: {e}")

def worker_process(livros: List[Dict], worker_id: int, result_queue: mp.Queue):
    """Processo worker que processa lista de livros."""
    print(f"[Worker {worker_id}] Iniciando com {len(livros)} livros")

    resultados = []
    for i, livro in enumerate(livros):
        titulo = livro['titulo']
        texto = livro['texto']
        idioma = livro['idioma']

        print(f"[W{worker_id}] ({i+1}/{len(livros)}) {titulo[:40]}...")

        resultado = asyncio.run(processar_livro_async(titulo, texto, idioma, worker_id))
        resultados.append(resultado)

        status = "OK" if resultado['status'] == 'sucesso' else "ERRO"
        print(f"[W{worker_id}] {status} - {titulo[:40]}")

    result_queue.put((worker_id, resultados))
    print(f"[Worker {worker_id}] Finalizado!")

# =============================================================================
# MAIN
# =============================================================================

def escanear_livros(limite: int = 0) -> List[Dict]:
    """Escaneia livros."""
    livros = []

    folders = list(TXT_DIR.iterdir())
    if limite > 0:
        folders = folders[:limite]

    for folder in folders:
        if not folder.is_dir():
            continue

        txt_files = list(folder.glob("*.txt"))
        if not txt_files:
            continue

        main_file = max(txt_files, key=lambda f: f.stat().st_size)

        try:
            with open(main_file, 'r', encoding='utf-8', errors='ignore') as f:
                texto = f.read()

            if len(texto) < 500:
                continue

            livros.append({
                'titulo': folder.name,
                'idioma': detectar_idioma(texto),
                'texto': texto
            })
        except Exception:
            pass

    return livros

def main():
    parser = argparse.ArgumentParser(description="Pipeline Turbo - Múltiplos Workers")
    parser.add_argument('--run', action='store_true', help='Executa processamento')
    parser.add_argument('--workers', type=int, default=5, help='Número de workers')
    parser.add_argument('--limit', type=int, default=0, help='Limite de livros')
    args = parser.parse_args()

    if not args.run:
        parser.print_help()
        return

    print("="*60)
    print("PIPELINE TURBO - PROCESSAMENTO PARALELO")
    print("="*60)
    print(f"Workers: {args.workers}")
    print(f"Modelo: {MODEL}")
    print("="*60)

    # Escaneia
    print("\nEscaneando livros...")
    livros = escanear_livros(args.limit)
    print(f"Total: {len(livros)} livros")

    # Por idioma
    por_idioma = {}
    for l in livros:
        por_idioma[l['idioma']] = por_idioma.get(l['idioma'], 0) + 1
    print(f"Por idioma: {por_idioma}")

    # Divide entre workers
    chunks = [[] for _ in range(args.workers)]
    for i, livro in enumerate(livros):
        chunks[i % args.workers].append(livro)

    print(f"\nDistribuição: {[len(c) for c in chunks]}")

    # Inicia workers
    start = time.time()
    result_queue = mp.Queue()
    processes = []

    for i, chunk in enumerate(chunks):
        if chunk:
            p = mp.Process(target=worker_process, args=(chunk, i, result_queue))
            processes.append(p)
            p.start()

    # Aguarda resultados
    resultados = []
    for _ in processes:
        worker_id, worker_results = result_queue.get()
        resultados.extend(worker_results)

    for p in processes:
        p.join()

    elapsed = time.time() - start

    # Resumo
    sucesso = sum(1 for r in resultados if r['status'] == 'sucesso')
    total_notas = sum(r.get('notas', 0) for r in resultados)

    print("\n" + "="*60)
    print("RESUMO")
    print("="*60)
    print(f"Total: {len(resultados)}")
    print(f"Sucesso: {sucesso}")
    print(f"Erros: {len(resultados) - sucesso}")
    print(f"Notas: {total_notas}")
    print(f"Tempo: {elapsed:.1f}s ({elapsed/60:.1f} min)")
    print(f"Velocidade: {len(resultados)/elapsed*60:.1f} livros/min")
    print(f"DOCX: {DOCX_DIR}")
    print("="*60)

if __name__ == "__main__":
    main()

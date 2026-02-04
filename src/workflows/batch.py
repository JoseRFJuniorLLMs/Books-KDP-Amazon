# -*- coding: utf-8 -*-
"""
PIPELINE VERTEX AI BATCH - Processamento Ultra-Rápido
=====================================================
Usa Batch Prediction do Vertex AI - sem rate limit!

Fluxo:
1. Prepara requests em JSONL
2. Upload para GCS
3. Submete batch job
4. Aguarda conclusão
5. Baixa resultados
6. Gera DOCX

Uso:
    python pipeline_batch.py --prepare     # Prepara arquivos
    python pipeline_batch.py --submit      # Submete batch
    python pipeline_batch.py --status      # Verifica status
    python pipeline_batch.py --process     # Processa resultados
"""

import os
import sys
import json
import hashlib
import argparse
import time
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple
from dataclasses import dataclass, asdict
import subprocess

# =============================================================================
# CONFIGURAÇÕES
# =============================================================================

PROJECT_ID = "aurorav2-484411"
LOCATION = "us-central1"
GCS_BUCKET = "testmassivo-kdp-2024"
MODEL = "gemini-2.0-flash-001"

# Diretórios
BASE_DIR = Path(__file__).parent
TXT_DIR = BASE_DIR / "txt"
DOCX_DIR = BASE_DIR / "docx"
DATA_DIR = BASE_DIR / "data"
BATCH_DIR = DATA_DIR / "batch"
TRANSLATED_DIR = BASE_DIR / "translated-portugues"

for d in [DOCX_DIR, DATA_DIR, BATCH_DIR, TRANSLATED_DIR]:
    d.mkdir(exist_ok=True)

# =============================================================================
# DETECÇÃO DE IDIOMA
# =============================================================================

LANG_WORDS = {
    'pt': ['que', 'não', 'para', 'uma', 'com', 'por', 'mais', 'você', 'isso'],
    'es': ['que', 'los', 'las', 'una', 'para', 'con', 'más', 'pero', 'está'],
    'en': ['the', 'and', 'that', 'have', 'for', 'with', 'this', 'but', 'from'],
    'fr': ['les', 'des', 'une', 'que', 'pour', 'dans', 'qui', 'sur', 'avec'],
    'de': ['der', 'die', 'und', 'den', 'das', 'von', 'ist', 'mit', 'nicht'],
    'ru': ['что', 'как', 'это', 'его', 'она', 'они', 'все', 'для', 'был'],
}

def detectar_idioma(texto: str) -> str:
    texto = texto[:5000].lower()
    scores = {lang: sum(1 for w in words if f' {w} ' in f' {texto} ')
              for lang, words in LANG_WORDS.items()}
    return max(scores, key=scores.get)

# =============================================================================
# PREPARAÇÃO DOS REQUESTS
# =============================================================================

@dataclass
class LivroRequest:
    id: str
    titulo: str
    idioma: str
    tipo: str  # 'traducao', 'ocr', 'notas'
    texto: str
    chunk_idx: int = 0

def dividir_chunks(texto: str, tamanho: int = 5000) -> List[str]:
    """Divide texto em chunks."""
    chunks = []
    current = ""
    for para in texto.split('\n\n'):
        if len(current) + len(para) < tamanho:
            current += '\n\n' + para if current else para
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)
    return chunks if chunks else [texto]

def criar_prompt(req: LivroRequest) -> str:
    """Cria prompt baseado no tipo de request."""
    if req.tipo == 'ocr':
        return f"""Corrija erros de OCR neste texto. Mantenha estrutura.
Retorne APENAS o texto corrigido, sem explicações.

{req.texto}"""

    elif req.tipo == 'traducao':
        lang_map = {'en': 'inglês', 'es': 'espanhol', 'fr': 'francês', 'de': 'alemão', 'ru': 'russo'}
        return f"""Traduza de {lang_map.get(req.idioma, req.idioma)} para português brasileiro.
Mantenha o estilo. Retorne APENAS a tradução.

{req.texto}"""

    elif req.tipo == 'notas':
        return f"""Identifique termos que precisam de notas de rodapé:
- Termos estrangeiros (latim, francês, etc)
- Conceitos filosóficos/técnicos
- Nomes históricos

Explicação BREVE (max 25 palavras cada).
Formato: termo|explicação
Máximo 15 termos.

{req.texto[:5000]}"""

    return req.texto

def preparar_batch_requests(limite: int = 0) -> Tuple[List[Dict], Dict]:
    """Prepara todos os requests para batch."""
    print("Escaneando livros...")

    livros_info = {}
    requests = []

    folders = list(TXT_DIR.iterdir())
    if limite > 0:
        folders = folders[:limite]

    for folder in folders:
        if not folder.is_dir():
            continue

        txt_files = list(folder.glob("*.txt"))
        if not txt_files:
            continue

        main_file = max(txt_files, key=lambda f: f.stat().st_size)

        try:
            with open(main_file, 'r', encoding='utf-8', errors='ignore') as f:
                texto = f.read()

            if len(texto) < 500:
                continue

            titulo = folder.name
            idioma = detectar_idioma(texto)
            livro_id = hashlib.md5(titulo.encode()).hexdigest()[:8]

            livros_info[livro_id] = {
                'titulo': titulo,
                'idioma': idioma,
                'txt_path': str(main_file),
                'chunks_ocr': [],
                'chunks_trad': [],
            }

            chunks = dividir_chunks(texto)

            # Requests de OCR
            for i, chunk in enumerate(chunks):
                req_id = f"{livro_id}_ocr_{i}"
                livros_info[livro_id]['chunks_ocr'].append(req_id)

                requests.append({
                    "request": {
                        "contents": [{"role": "user", "parts": [{"text": criar_prompt(
                            LivroRequest(req_id, titulo, idioma, 'ocr', chunk, i)
                        )}]}],
                        "generationConfig": {"temperature": 0.2, "maxOutputTokens": 8192}
                    },
                    "metadata": {"id": req_id, "livro_id": livro_id, "tipo": "ocr", "chunk": i}
                })

            # Requests de tradução (se não for PT)
            if idioma != 'pt':
                for i, chunk in enumerate(chunks):
                    req_id = f"{livro_id}_trad_{i}"
                    livros_info[livro_id]['chunks_trad'].append(req_id)

                    requests.append({
                        "request": {
                            "contents": [{"role": "user", "parts": [{"text": criar_prompt(
                                LivroRequest(req_id, titulo, idioma, 'traducao', chunk, i)
                            )}]}],
                            "generationConfig": {"temperature": 0.3, "maxOutputTokens": 8192}
                        },
                        "metadata": {"id": req_id, "livro_id": livro_id, "tipo": "traducao", "chunk": i}
                    })

            # Request de notas
            req_id = f"{livro_id}_notas"
            requests.append({
                "request": {
                    "contents": [{"role": "user", "parts": [{"text": criar_prompt(
                        LivroRequest(req_id, titulo, idioma, 'notas', texto[:6000])
                    )}]}],
                    "generationConfig": {"temperature": 0.3, "maxOutputTokens": 2048}
                },
                "metadata": {"id": req_id, "livro_id": livro_id, "tipo": "notas"}
            })

        except Exception as e:
            print(f"Erro em {folder}: {e}")

    return requests, livros_info

def salvar_batch_jsonl(requests: List[Dict], output_path: Path):
    """Salva requests em formato JSONL."""
    with open(output_path, 'w', encoding='utf-8') as f:
        for req in requests:
            f.write(json.dumps(req, ensure_ascii=False) + '\n')

# =============================================================================
# VERTEX AI BATCH
# =============================================================================

def upload_to_gcs(local_path: Path, gcs_path: str):
    """Upload arquivo para GCS."""
    cmd = f'gsutil cp "{local_path}" "gs://{GCS_BUCKET}/{gcs_path}"'
    print(f"Upload: {local_path} -> gs://{GCS_BUCKET}/{gcs_path}")
    os.system(cmd)

def download_from_gcs(gcs_path: str, local_path: Path):
    """Download arquivo do GCS."""
    cmd = f'gsutil cp "gs://{GCS_BUCKET}/{gcs_path}" "{local_path}"'
    os.system(cmd)

def submit_batch_job(input_gcs: str, output_gcs: str) -> str:
    """Submete job de batch prediction."""

    job_config = {
        "displayName": f"kdp-batch-{datetime.now():%Y%m%d-%H%M%S}",
        "model": f"publishers/google/models/{MODEL}",
        "inputConfig": {
            "instancesFormat": "jsonl",
            "gcsSource": {"uris": [f"gs://{GCS_BUCKET}/{input_gcs}"]}
        },
        "outputConfig": {
            "predictionsFormat": "jsonl",
            "gcsDestination": {"outputUriPrefix": f"gs://{GCS_BUCKET}/{output_gcs}"}
        }
    }

    config_path = BATCH_DIR / "job_config.json"
    with open(config_path, 'w') as f:
        json.dump(job_config, f, indent=2)

    print("\nSubmetendo batch job...")
    print(f"Input: gs://{GCS_BUCKET}/{input_gcs}")
    print(f"Output: gs://{GCS_BUCKET}/{output_gcs}")

    # Usa gcloud para submeter
    cmd = f'''gcloud ai batch-prediction-jobs create \
        --region={LOCATION} \
        --display-name="{job_config['displayName']}" \
        --model="{job_config['model']}" \
        --input-format=jsonl \
        --output-format=jsonl \
        --gcs-source-uris="gs://{GCS_BUCKET}/{input_gcs}" \
        --gcs-destination-output-uri-prefix="gs://{GCS_BUCKET}/{output_gcs}" \
        --project={PROJECT_ID}'''

    print("\nComando:")
    print(cmd)

    return job_config['displayName']

def check_job_status(job_name: str = None):
    """Verifica status do job."""
    cmd = f'gcloud ai batch-prediction-jobs list --region={LOCATION} --project={PROJECT_ID} --limit=5'
    os.system(cmd)

# =============================================================================
# PROCESSAMENTO DOS RESULTADOS
# =============================================================================

def processar_resultados(livros_info_path: Path, results_dir: Path):
    """Processa resultados do batch e gera DOCX."""
    print("Processando resultados...")

    # Carrega info dos livros
    with open(livros_info_path, 'r') as f:
        livros_info = json.load(f)

    # Carrega resultados
    resultados = {}
    for jsonl_file in results_dir.glob("*.jsonl"):
        with open(jsonl_file, 'r') as f:
            for line in f:
                try:
                    data = json.loads(line)
                    req_id = data.get('metadata', {}).get('id')
                    if req_id and 'response' in data:
                        text = data['response'].get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '')
                        resultados[req_id] = text
                except Exception:
                    pass

    print(f"Resultados carregados: {len(resultados)}")

    # Processa cada livro
    for livro_id, info in livros_info.items():
        titulo = info['titulo']
        print(f"  Processando: {titulo}")

        # Junta chunks OCR
        texto_ocr = []
        for req_id in info.get('chunks_ocr', []):
            if req_id in resultados:
                texto_ocr.append(resultados[req_id])

        # Junta chunks tradução (ou usa OCR se já era PT)
        if info.get('chunks_trad'):
            texto_final = []
            for req_id in info['chunks_trad']:
                if req_id in resultados:
                    texto_final.append(resultados[req_id])
            texto = '\n\n'.join(texto_final) if texto_final else '\n\n'.join(texto_ocr)
        else:
            texto = '\n\n'.join(texto_ocr)

        # Notas
        notas = []
        notas_id = f"{livro_id}_notas"
        if notas_id in resultados:
            for linha in resultados[notas_id].split('\n'):
                if '|' in linha:
                    partes = linha.split('|', 1)
                    if len(partes) == 2:
                        termo = partes[0].strip().strip('- ')
                        explicacao = partes[1].strip()
                        if termo and explicacao:
                            notas.append((termo, explicacao))

        # Cria DOCX
        if texto:
            criar_docx(texto, titulo, "Autor", notas[:15], DOCX_DIR / f"{titulo}.docx")

            # Salva TXT
            with open(TRANSLATED_DIR / f"{titulo}_pt.txt", 'w', encoding='utf-8') as f:
                f.write(texto)

def criar_docx(texto: str, titulo: str, autor: str, notas: List[Tuple[str, str]], output_path: Path):
    """Cria DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(24)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.add_run(autor)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        for para in texto.split('\n\n'):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if notas:
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run("NOTAS")
            run.bold = True

            for i, (termo, explicacao) in enumerate(notas, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{i}. {termo}: ")
                run.bold = True
                run.font.size = Pt(9)
                run = p.add_run(explicacao)
                run.font.size = Pt(9)

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Erro DOCX: {e}")
        return False

# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description="Pipeline Vertex AI Batch")
    parser.add_argument('--prepare', action='store_true', help='Prepara requests')
    parser.add_argument('--submit', action='store_true', help='Submete batch job')
    parser.add_argument('--status', action='store_true', help='Verifica status')
    parser.add_argument('--process', action='store_true', help='Processa resultados')
    parser.add_argument('--limit', type=int, default=0, help='Limite de livros')
    args = parser.parse_args()

    if args.prepare:
        print("="*60)
        print("PREPARANDO BATCH REQUESTS")
        print("="*60)

        requests, livros_info = preparar_batch_requests(args.limit)

        print(f"\nLivros: {len(livros_info)}")
        print(f"Requests totais: {len(requests)}")

        # Salva JSONL
        jsonl_path = BATCH_DIR / "requests.jsonl"
        salvar_batch_jsonl(requests, jsonl_path)
        print(f"Salvo: {jsonl_path}")

        # Salva info dos livros
        info_path = BATCH_DIR / "livros_info.json"
        with open(info_path, 'w') as f:
            json.dump(livros_info, f, indent=2, ensure_ascii=False)
        print(f"Info: {info_path}")

        # Estatísticas por idioma
        por_idioma = {}
        for info in livros_info.values():
            lang = info['idioma']
            por_idioma[lang] = por_idioma.get(lang, 0) + 1
        print(f"Por idioma: {por_idioma}")

        print(f"\nPróximo passo: python pipeline_batch.py --submit")

    elif args.submit:
        print("="*60)
        print("SUBMETENDO BATCH JOB")
        print("="*60)

        jsonl_path = BATCH_DIR / "requests.jsonl"
        if not jsonl_path.exists():
            print("ERRO: Execute --prepare primeiro!")
            return

        # Upload para GCS
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        gcs_input = f"batch/input_{timestamp}.jsonl"
        gcs_output = f"batch/output_{timestamp}/"

        upload_to_gcs(jsonl_path, gcs_input)

        # Submete job
        job_name = submit_batch_job(gcs_input, gcs_output)

        # Salva info do job
        job_info = {
            'job_name': job_name,
            'gcs_input': gcs_input,
            'gcs_output': gcs_output,
            'timestamp': timestamp
        }
        with open(BATCH_DIR / "job_info.json", 'w') as f:
            json.dump(job_info, f, indent=2)

        print(f"\nJob submetido: {job_name}")
        print(f"Verifique status: python pipeline_batch.py --status")

    elif args.status:
        check_job_status()

    elif args.process:
        print("="*60)
        print("PROCESSANDO RESULTADOS")
        print("="*60)

        # Carrega info do job
        job_info_path = BATCH_DIR / "job_info.json"
        if not job_info_path.exists():
            print("ERRO: Nenhum job encontrado!")
            return

        with open(job_info_path, 'r') as f:
            job_info = json.load(f)

        # Baixa resultados do GCS
        results_dir = BATCH_DIR / "results"
        results_dir.mkdir(exist_ok=True)

        print(f"Baixando resultados de gs://{GCS_BUCKET}/{job_info['gcs_output']}...")
        os.system(f'gsutil -m cp "gs://{GCS_BUCKET}/{job_info["gcs_output"]}*" "{results_dir}/"')

        # Processa
        livros_info_path = BATCH_DIR / "livros_info.json"
        processar_resultados(livros_info_path, results_dir)

        print("\n" + "="*60)
        print("PROCESSAMENTO COMPLETO!")
        print(f"DOCX: {DOCX_DIR}")
        print(f"TXT: {TRANSLATED_DIR}")
        print("="*60)

    else:
        parser.print_help()
        print("\nFluxo completo:")
        print("  1. python pipeline_batch.py --prepare --limit 100")
        print("  2. python pipeline_batch.py --submit")
        print("  3. python pipeline_batch.py --status  (aguarde SUCCEEDED)")
        print("  4. python pipeline_batch.py --process")

if __name__ == "__main__":
    main()

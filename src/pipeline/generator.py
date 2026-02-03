#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
MÓDULO 3: GENERATOR (com Template + Notas de Rodapé)
====================================================
Gera DOCX usando template.docx com notas explicativas.

- Usa template externo (templates/template.docx)
- Gera notas de rodapé para termos históricos/técnicos
- Usa Gemini para identificar termos

Uso:
    python -m src.pipeline.generator --input books/txt/pt --output books/docx/pt
"""

import os
import re
import json
import asyncio
import aiohttp
import argparse
import zipfile
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
from io import BytesIO
from lxml import etree
from concurrent.futures import ThreadPoolExecutor

# Configuração
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-3-pro-preview:generateContent"
TEMPLATE_PATH = Path("templates/template.docx")

# Namespaces DOCX
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class DocxGenerator:
    """Gerador de DOCX com template e notas de rodapé."""

    def __init__(
        self,
        input_dir: Path,
        output_dir: Path,
        template_path: Path = TEMPLATE_PATH,
        api_key: str = "",
        max_workers: int = 4,
        max_notes: int = 20
    ):
        self.input_dir = input_dir
        self.output_dir = output_dir
        self.template_path = template_path
        self.api_key = api_key or GEMINI_API_KEY
        self.max_workers = max_workers
        self.max_notes = max_notes
        self.generated = 0
        self.failed = 0
        self.template_bytes = None

        # Carrega template
        if self.template_path.exists():
            self.template_bytes = self.template_path.read_bytes()
            print(f"Template carregado: {self.template_path}")
        else:
            print(f"AVISO: Template não encontrado: {self.template_path}")

    def extract_metadata(self, content: str, filename: str) -> Tuple[str, str]:
        """Extrai título e autor."""
        lines = content.split('\n')[:30]
        title = filename.replace('_pt', '').replace('.txt', '')
        author = "Autor Desconhecido"

        for line in lines:
            line = line.strip()
            if line.lower().startswith('title:'):
                title = line.split(':', 1)[1].strip()
            elif line.lower().startswith('author:'):
                author = line.split(':', 1)[1].strip()

        # Tenta extrair do nome do arquivo (Autor_Titulo)
        if '_' in filename:
            parts = filename.replace('_pt.txt', '').replace('.txt', '').split('_', 1)
            if len(parts) == 2:
                author = parts[0].strip()
                title = parts[1].strip()

        return title, author

    async def get_explanatory_notes(self, text: str) -> List[Tuple[str, str]]:
        """Usa Gemini para identificar termos que precisam explicação."""
        if not self.api_key:
            return []

        prompt = f"""Analise este texto e identifique até {self.max_notes} termos que precisam de notas explicativas para leitores brasileiros.

Inclua:
- Termos históricos e datas importantes
- Conceitos filosóficos ou técnicos
- Referências culturais estrangeiras
- Nomes de pessoas importantes
- Lugares históricos
- Termos em latim/francês/etc

NÃO inclua:
- Palavras comuns
- Nomes de personagens fictícios óbvios
- Termos já explicados no texto

Formato de resposta (um por linha):
termo|explicação breve em português (máximo 25 palavras)

Texto (primeiros 8000 caracteres):
{text[:8000]}"""

        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{GEMINI_URL}?key={self.api_key}",
                    json={
                        "contents": [{"parts": [{"text": prompt}]}],
                        "generationConfig": {"temperature": 0.3, "maxOutputTokens": 2000}
                    },
                    timeout=aiohttp.ClientTimeout(total=60)
                ) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        result = data.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")

                        notes = []
                        for line in result.split('\n'):
                            if '|' in line:
                                parts = line.split('|', 1)
                                if len(parts) == 2:
                                    termo = parts[0].strip().strip('- *"\'')
                                    explicacao = parts[1].strip()
                                    if termo and explicacao and 2 < len(termo) < 60:
                                        notes.append((termo, explicacao))

                        return notes[:self.max_notes]
        except Exception as e:
            print(f"    Erro Gemini: {e}")

        return []

    def create_docx_with_template(
        self,
        content: str,
        output_path: Path,
        title: str,
        author: str,
        notes: List[Tuple[str, str]]
    ) -> bool:
        """Cria DOCX usando template com notas de rodapé.

        Estrutura do template preservada:
        - Páginas iniciais (intro, citações) - PRESERVAR
        - INSERIR CONTEÚDO DO LIVRO AQUI
        - Página final (copyright, livraria) - PRESERVAR
        """
        try:
            if not self.template_bytes:
                return self.create_docx_simple(content, output_path, title, author)

            # Constrói XML do conteúdo e notas
            content_xml, footnotes_xml, notas_usadas = self._build_content_xml(content, notes)

            template_zip = zipfile.ZipFile(BytesIO(self.template_bytes))
            output = BytesIO()

            with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as docx:
                has_footnotes = 'word/footnotes.xml' in template_zip.namelist()

                for item in template_zip.namelist():
                    if item == 'word/document.xml':
                        # Modifica document.xml para injetar conteúdo
                        doc_xml = template_zip.read(item).decode('utf-8')

                        # Encontra o parágrafo que contém "Copyright" (início do back matter)
                        copyright_match = re.search(r'<w:p[^>]*>(?:(?!</w:p>).)*Copyright(?:(?!</w:p>).)*</w:p>', doc_xml, re.DOTALL)

                        if copyright_match:
                            # Injeta conteúdo ANTES do parágrafo de Copyright
                            insert_pos = copyright_match.start()
                            page_break = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:br w:type="page"/></w:r></w:p>\n'
                            doc_xml = doc_xml[:insert_pos] + page_break + content_xml + '\n' + page_break + doc_xml[insert_pos:]
                        else:
                            # Fallback: procura sectPr
                            sectpr_match = re.search(r'(<w:sectPr[^>]*>)', doc_xml)
                            if sectpr_match:
                                insert_pos = sectpr_match.start()
                                page_break = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:br w:type="page"/></w:r></w:p>\n'
                                doc_xml = doc_xml[:insert_pos] + page_break + content_xml + '\n' + doc_xml[insert_pos:]
                            else:
                                doc_xml = doc_xml.replace('</w:body>', content_xml + '\n</w:body>')

                        docx.writestr(item, doc_xml.encode('utf-8'))

                    elif item == 'word/footnotes.xml':
                        # Substitui footnotes existente
                        docx.writestr(item, footnotes_xml)
                        has_footnotes = True

                    elif item == '[Content_Types].xml':
                        ct_xml = template_zip.read(item).decode('utf-8')
                        if 'footnotes.xml' not in ct_xml:
                            ct_xml = ct_xml.replace(
                                '</Types>',
                                '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>\n</Types>'
                            )
                        docx.writestr(item, ct_xml.encode('utf-8'))

                    elif item == 'word/_rels/document.xml.rels':
                        rels_xml = template_zip.read(item).decode('utf-8')
                        if 'footnotes.xml' not in rels_xml:
                            ids = re.findall(r'rId(\d+)', rels_xml)
                            next_id = max(int(i) for i in ids) + 1 if ids else 10
                            new_rel = f'<Relationship Id="rId{next_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>\n'
                            rels_xml = rels_xml.replace('</Relationships>', new_rel + '</Relationships>')
                        docx.writestr(item, rels_xml.encode('utf-8'))

                    else:
                        docx.writestr(item, template_zip.read(item))

                # Adiciona footnotes.xml se não existia
                if not has_footnotes:
                    docx.writestr('word/footnotes.xml', footnotes_xml)

            template_zip.close()

            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_bytes(output.getvalue())
            return True

        except Exception as e:
            print(f"    Erro template: {e}")
            import traceback
            traceback.print_exc()
            return self.create_docx_simple(content, output_path, title, author)

    @staticmethod
    def _escape_xml(text: str) -> str:
        """Escapa caracteres especiais XML."""
        return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&apos;'))

    def _build_content_xml(self, content: str, notes: List[Tuple[str, str]]) -> Tuple[str, str, List]:
        """Constrói XML do conteúdo e notas de rodapé."""
        paragraphs = content.split('\n\n')
        nota_id = 1
        notas_usadas = []
        termos_usados = set()

        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        content_parts = []

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Detectar capítulos
            is_chapter = (
                para.upper().startswith('CHAPTER') or
                para.upper().startswith('CAPÍTULO') or
                re.match(r'^[IVXLC]+\.?\s', para)
            )

            if is_chapter:
                content_parts.append(f'''<w:p xmlns:w="{W_NS}"><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>
<w:r><w:t>{self._escape_xml(para)}</w:t></w:r></w:p>''')
                continue

            # Parágrafo normal - procura termos para notas
            para_lower = para.lower()
            found_term = None

            for termo, explicacao in notes:
                if termo.lower() in termos_usados:
                    continue
                if termo.lower() in para_lower:
                    found_term = (termo, explicacao)
                    break

            if found_term:
                termo, explicacao = found_term
                pattern = re.compile(re.escape(termo), re.IGNORECASE)
                match = pattern.search(para)

                if match:
                    before = para[:match.end()]
                    after = para[match.end():]

                    content_parts.append(f'''<w:p xmlns:w="{W_NS}"><w:pPr><w:pStyle w:val="Normal"/><w:jc w:val="both"/></w:pPr>
<w:r><w:t xml:space="preserve">{self._escape_xml(before)}</w:t></w:r>
<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteReference w:id="{nota_id}"/></w:r>
<w:r><w:t xml:space="preserve">{self._escape_xml(after)}</w:t></w:r></w:p>''')

                    notas_usadas.append((termo, explicacao, nota_id))
                    termos_usados.add(termo.lower())
                    nota_id += 1
                else:
                    content_parts.append(f'''<w:p xmlns:w="{W_NS}"><w:pPr><w:pStyle w:val="Normal"/><w:jc w:val="both"/></w:pPr>
<w:r><w:t xml:space="preserve">{self._escape_xml(para)}</w:t></w:r></w:p>''')
            else:
                content_parts.append(f'''<w:p xmlns:w="{W_NS}"><w:pPr><w:pStyle w:val="Normal"/><w:jc w:val="both"/></w:pPr>
<w:r><w:t xml:space="preserve">{self._escape_xml(para)}</w:t></w:r></w:p>''')

        content_xml = '\n'.join(content_parts)

        # Footnotes XML
        fn_parts = [f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="{W_NS}">
<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>
<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>''']

        for termo, explicacao, nid in notas_usadas:
            fn_parts.append(f'''<w:footnote w:id="{nid}">
<w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>
<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteRef/></w:r>
<w:r><w:t xml:space="preserve"> {self._escape_xml(termo)}: {self._escape_xml(explicacao)}</w:t></w:r></w:p>
</w:footnote>''')

        fn_parts.append('</w:footnotes>')
        footnotes_xml = '\n'.join(fn_parts)

        return content_xml, footnotes_xml, notas_usadas

    def _add_title(self, body, title: str, author: str):
        """Adiciona página de título."""
        # Título
        p = etree.SubElement(body, f'{W}p')
        pPr = etree.SubElement(p, f'{W}pPr')
        jc = etree.SubElement(pPr, f'{W}jc')
        jc.set(f'{W}val', 'center')
        spacing = etree.SubElement(pPr, f'{W}spacing')
        spacing.set(f'{W}before', '2000')

        run = etree.SubElement(p, f'{W}r')
        rPr = etree.SubElement(run, f'{W}rPr')
        etree.SubElement(rPr, f'{W}b')
        sz = etree.SubElement(rPr, f'{W}sz')
        sz.set(f'{W}val', '56')
        t = etree.SubElement(run, f'{W}t')
        t.text = title

        # Autor
        p = etree.SubElement(body, f'{W}p')
        pPr = etree.SubElement(p, f'{W}pPr')
        jc = etree.SubElement(pPr, f'{W}jc')
        jc.set(f'{W}val', 'center')
        spacing = etree.SubElement(pPr, f'{W}spacing')
        spacing.set(f'{W}before', '500')

        run = etree.SubElement(p, f'{W}r')
        rPr = etree.SubElement(run, f'{W}rPr')
        etree.SubElement(rPr, f'{W}i')
        sz = etree.SubElement(rPr, f'{W}sz')
        sz.set(f'{W}val', '32')
        t = etree.SubElement(run, f'{W}t')
        t.text = author

    def _add_page_break(self, body):
        """Adiciona quebra de página."""
        p = etree.SubElement(body, f'{W}p')
        run = etree.SubElement(p, f'{W}r')
        br = etree.SubElement(run, f'{W}br')
        br.set(f'{W}type', 'page')

    def _add_content_with_notes(self, body, content: str, notes: List[Tuple[str, str]]) -> List[Tuple[str, str, int]]:
        """Adiciona conteúdo com referências de notas."""
        notas_usadas = []
        nota_id = 1
        termos_usados = set()

        paragraphs = content.split('\n\n')

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            p = etree.SubElement(body, f'{W}p')
            pPr = etree.SubElement(p, f'{W}pPr')
            jc = etree.SubElement(pPr, f'{W}jc')
            jc.set(f'{W}val', 'both')

            # Detectar capítulos
            is_chapter = (
                para.upper().startswith('CHAPTER') or
                para.upper().startswith('CAPÍTULO') or
                re.match(r'^[IVXLC]+\.?\s', para)
            )

            if is_chapter:
                pStyle = etree.SubElement(pPr, f'{W}pStyle')
                pStyle.set(f'{W}val', 'Heading1')

            # Processa parágrafo procurando termos para notas
            para_lower = para.lower()
            segments = [(para, None)]  # (texto, nota_id ou None)

            for termo, explicacao in notes:
                if termo.lower() in termos_usados:
                    continue
                if termo.lower() in para_lower:
                    new_segments = []
                    for seg_text, seg_note in segments:
                        if seg_note is not None or termo.lower() not in seg_text.lower():
                            new_segments.append((seg_text, seg_note))
                        else:
                            # Encontra primeira ocorrência
                            pattern = re.compile(re.escape(termo), re.IGNORECASE)
                            match = pattern.search(seg_text)
                            if match:
                                before = seg_text[:match.end()]
                                after = seg_text[match.end():]
                                new_segments.append((before, nota_id))
                                if after:
                                    new_segments.append((after, None))
                                notas_usadas.append((termo, explicacao, nota_id))
                                termos_usados.add(termo.lower())
                                nota_id += 1
                            else:
                                new_segments.append((seg_text, seg_note))
                    segments = new_segments
                    break  # Só uma nota por parágrafo

            # Adiciona runs ao parágrafo
            for seg_text, seg_note in segments:
                if seg_text:
                    run = etree.SubElement(p, f'{W}r')
                    t = etree.SubElement(run, f'{W}t')
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    t.text = seg_text

                    if seg_note is not None:
                        # Adiciona referência de nota
                        note_run = etree.SubElement(p, f'{W}r')
                        rPr = etree.SubElement(note_run, f'{W}rPr')
                        va = etree.SubElement(rPr, f'{W}vertAlign')
                        va.set(f'{W}val', 'superscript')
                        fnRef = etree.SubElement(note_run, f'{W}footnoteReference')
                        fnRef.set(f'{W}id', str(seg_note))

        return notas_usadas

    def _create_footnotes_xml(self, notas: List[Tuple[str, str, int]]) -> bytes:
        """Cria XML das notas de rodapé."""
        footnotes = etree.Element(f'{W}footnotes', nsmap={'w': NSMAP['w']})

        # Notas especiais (separator, continuationSeparator)
        for sid, stype in [('-1', 'separator'), ('0', 'continuationSeparator')]:
            fn = etree.SubElement(footnotes, f'{W}footnote')
            fn.set(f'{W}type', stype)
            fn.set(f'{W}id', sid)
            p = etree.SubElement(fn, f'{W}p')
            r = etree.SubElement(p, f'{W}r')
            etree.SubElement(r, f'{W}{stype}')

        # Notas reais
        for termo, explicacao, nid in notas:
            fn = etree.SubElement(footnotes, f'{W}footnote')
            fn.set(f'{W}id', str(nid))
            p = etree.SubElement(fn, f'{W}p')

            # Referência
            r1 = etree.SubElement(p, f'{W}r')
            rPr = etree.SubElement(r1, f'{W}rPr')
            va = etree.SubElement(rPr, f'{W}vertAlign')
            va.set(f'{W}val', 'superscript')
            etree.SubElement(r1, f'{W}footnoteRef')

            # Texto da nota
            r2 = etree.SubElement(p, f'{W}r')
            t = etree.SubElement(r2, f'{W}t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = f" {termo}: {explicacao}"

        return etree.tostring(footnotes, xml_declaration=True, encoding='UTF-8')

    def _update_content_types(self, docx, template_zip):
        """Atualiza Content_Types.xml para incluir footnotes."""
        try:
            ct = template_zip.read('[Content_Types].xml').decode('utf-8')
            if 'footnotes.xml' not in ct:
                ct = ct.replace('</Types>',
                    '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/></Types>')
                docx.writestr('[Content_Types].xml', ct.encode('utf-8'))
        except:
            pass

    def _update_relationships(self, docx, template_zip):
        """Atualiza relationships para incluir footnotes."""
        try:
            rels = template_zip.read('word/_rels/document.xml.rels').decode('utf-8')
            if 'footnotes.xml' not in rels:
                # Encontra próximo rId
                ids = re.findall(r'rId(\d+)', rels)
                next_id = max(int(i) for i in ids) + 1 if ids else 2
                new_rel = f'<Relationship Id="rId{next_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                rels = rels.replace('</Relationships>', f'{new_rel}</Relationships>')
                docx.writestr('word/_rels/document.xml.rels', rels.encode('utf-8'))
        except:
            pass

    def create_docx_simple(self, content: str, output_path: Path, title: str, author: str) -> bool:
        """Fallback: cria DOCX simples sem template."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Título
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(24)

            # Autor
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(author)
            run.italic = True

            doc.add_page_break()

            # Conteúdo
            for para in content.split('\n\n'):
                para = para.strip()
                if para:
                    p = doc.add_paragraph(para)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            return True

        except Exception as e:
            print(f"    Erro fallback: {e}")
            return False

    async def process_file_async(self, filepath: Path) -> Optional[Dict]:
        """Processa um arquivo com notas assíncronas."""
        try:
            content = filepath.read_text(encoding='utf-8', errors='ignore')

            if len(content) < 500:
                return None

            title, author = self.extract_metadata(content, filepath.name)
            print(f"[...] {title[:50]}")

            # Busca notas explicativas
            notes = await self.get_explanatory_notes(content)
            print(f"      {len(notes)} notas identificadas")

            # Cria DOCX
            output_name = filepath.stem + ".docx"
            output_path = self.output_dir / output_name

            success = self.create_docx_with_template(content, output_path, title, author, notes)

            if success:
                print(f"[OK] {output_name}")
                return {
                    "source": str(filepath),
                    "output": str(output_path),
                    "title": title,
                    "author": author,
                    "notes": len(notes)
                }

        except Exception as e:
            print(f"[ERRO] {filepath.name}: {e}")

        return None

    def process_file(self, filepath: Path) -> Optional[Dict]:
        """Wrapper síncrono."""
        return asyncio.run(self.process_file_async(filepath))

    async def run_async(self, limit: int = 0) -> List[Dict]:
        """Executa geração assíncrona."""
        print("=" * 60)
        print("GENERATOR - Template + Notas de Rodapé")
        print("=" * 60)
        print(f"Entrada: {self.input_dir}")
        print(f"Saída: {self.output_dir}")
        print(f"Template: {self.template_path}")
        print()

        self.output_dir.mkdir(parents=True, exist_ok=True)

        files = list(self.input_dir.glob("*.txt"))
        if limit > 0:
            files = files[:limit]

        print(f"Arquivos: {len(files)}")
        print()

        results = []
        for i, f in enumerate(files):
            print(f"[{i+1}/{len(files)}]", end=" ")
            result = await self.process_file_async(f)
            if result:
                results.append(result)
                self.generated += 1
            else:
                self.failed += 1

            # Rate limit para Gemini
            if self.api_key:
                await asyncio.sleep(1)

        print()
        print("=" * 60)
        print(f"Gerados: {self.generated}")
        print(f"Falhas: {self.failed}")
        print("=" * 60)

        return results

    def run(self, limit: int = 0) -> List[Dict]:
        """Executa geração."""
        return asyncio.run(self.run_async(limit))


def main():
    parser = argparse.ArgumentParser(description="Gerar DOCX com template e notas")
    parser.add_argument("--input", default="books/txt/pt", help="Pasta de entrada")
    parser.add_argument("--output", default="books/docx/pt", help="Pasta de saída")
    parser.add_argument("--template", default="templates/template.docx", help="Template DOCX")
    parser.add_argument("--limit", type=int, default=0, help="Limite de arquivos")
    parser.add_argument("--notes", type=int, default=20, help="Máximo de notas por livro")
    parser.add_argument("--api-key", help="Gemini API Key")

    args = parser.parse_args()

    generator = DocxGenerator(
        Path(args.input),
        Path(args.output),
        Path(args.template),
        args.api_key or "",
        max_notes=args.notes
    )
    generator.run(args.limit)


if __name__ == "__main__":
    main()

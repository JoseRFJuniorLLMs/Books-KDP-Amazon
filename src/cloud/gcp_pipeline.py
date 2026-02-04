# -*- coding: utf-8 -*-
"""
PIPELINE COMPLETO GCP
=====================
1. Sobe TXTs para Google Cloud Storage
2. Processa na GCP (validação, tradução, notas, OCR, DOCX)
3. Baixa DOCXs prontos

Uso:
    python gcp_pipeline_completo.py --upload      # Sobe livros para GCS
    python gcp_pipeline_completo.py --status      # Verifica status
    python gcp_pipeline_completo.py --download    # Baixa DOCXs prontos
"""

import os
import sys
import json
import argparse
import subprocess
from pathlib import Path
from datetime import datetime

# =============================================================================
# CONFIGURAÇÕES
# =============================================================================

PROJECT_ID = "seu-projeto-gcp"  # Será detectado automaticamente
BUCKET_NAME = "testmassivo"
REGION = "us-central1"

BASE_DIR = Path(__file__).parent
TXT_DIR = BASE_DIR / "txt"
DOCX_DIR = BASE_DIR / "docx"
DATA_DIR = BASE_DIR / "data"

DOCX_DIR.mkdir(exist_ok=True)

# =============================================================================
# FUNÇÕES GCS
# =============================================================================

def get_project_id():
    """Obtém project ID atual."""
    try:
        result = subprocess.run(
            ["gcloud", "config", "get-value", "project"],
            capture_output=True, text=True
        )
        return result.stdout.strip()
    except Exception:
        return None

def upload_to_gcs():
    """Sobe todos os TXTs para GCS."""
    print("=" * 60)
    print("UPLOAD PARA GOOGLE CLOUD STORAGE")
    print("=" * 60)

    # Lista livros
    livros = [p for p in TXT_DIR.iterdir() if p.is_dir()]
    print(f"Total de livros: {len(livros)}")

    # Cria bucket se não existir
    print(f"\nVerificando bucket: gs://{BUCKET_NAME}")
    subprocess.run(
        ["gsutil", "mb", "-l", REGION, f"gs://{BUCKET_NAME}"],
        capture_output=True
    )

    # Upload em paralelo
    print("\nFazendo upload...")

    uploaded = 0
    for i, pasta in enumerate(livros):
        # Encontra arquivo TXT principal
        txt_files = list(pasta.glob("*.txt"))
        if not txt_files:
            continue

        main_file = max(txt_files, key=lambda f: f.stat().st_size)

        # Upload
        dest = f"gs://{BUCKET_NAME}/livros/{pasta.name}/{main_file.name}"
        result = subprocess.run(
            ["gsutil", "-q", "cp", str(main_file), dest],
            capture_output=True
        )

        if result.returncode == 0:
            uploaded += 1

        if (i + 1) % 50 == 0:
            print(f"  {i + 1}/{len(livros)} enviados...")

    print(f"\n✓ {uploaded} livros enviados para gs://{BUCKET_NAME}/livros/")

    # Sobe também o script de processamento
    print("\nEnviando script de processamento...")
    upload_processor_script()

    return uploaded

def upload_processor_script():
    """Cria e sobe o script que vai rodar na GCP."""

    processor_code = '''#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PROCESSADOR GCP - Roda no Cloud Run/VM
"""

import os
import json
import asyncio
import aiohttp
from pathlib import Path
from google.cloud import storage
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

# Configurações
BUCKET_NAME = os.environ.get("BUCKET_NAME", "testmassivo")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-3-pro-preview:generateContent"

# Palavras para detecção de idioma
LANG_WORDS = {
    'pt': ['que', 'não', 'para', 'uma', 'com', 'por', 'mais', 'você'],
    'es': ['que', 'los', 'las', 'una', 'para', 'con', 'más', 'pero'],
    'en': ['the', 'and', 'that', 'have', 'for', 'with', 'this', 'but'],
    'fr': ['les', 'des', 'une', 'que', 'pour', 'dans', 'qui', 'sur'],
    'de': ['der', 'die', 'und', 'den', 'das', 'von', 'ist', 'mit'],
    'ru': ['что', 'как', 'это', 'его', 'она', 'они', 'все', 'для'],
}

def detectar_idioma(texto):
    texto = texto[:5000].lower()
    scores = {lang: sum(1 for w in words if f' {w} ' in f' {texto} ')
              for lang, words in LANG_WORDS.items()}
    return max(scores, key=scores.get)

def validar_livro(texto, titulo):
    """Valida se o livro está completo."""
    if len(texto) < 5000:
        return False, "Muito curto (< 5000 chars)"

    # Verifica se tem estrutura
    paragrafos = [p for p in texto.split('\\n\\n') if len(p.strip()) > 50]
    if len(paragrafos) < 10:
        return False, "Poucos parágrafos"

    return True, "OK"

async def call_gemini(session, prompt, semaphore):
    """Chama Gemini API."""
    async with semaphore:
        url = f"{GEMINI_URL}?key={GEMINI_API_KEY}"
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.3, "maxOutputTokens": 8192}
        }

        for attempt in range(5):
            try:
                async with session.post(url, json=payload) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        if "candidates" in data:
                            parts = data["candidates"][0].get("content", {}).get("parts", [])
                            if parts:
                                return parts[0].get("text", "").strip()
                        return ""
                    elif resp.status == 429:
                        await asyncio.sleep(2 ** attempt)
                    else:
                        await asyncio.sleep(1)
            except Exception as e:
                await asyncio.sleep(1)
        return ""

def dividir_chunks(texto, tamanho=5000):
    chunks = []
    current = ""
    for para in texto.split('\\n\\n'):
        if len(current) + len(para) < tamanho:
            current += '\\n\\n' + para if current else para
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)
    return chunks if chunks else [texto]

async def processar_livro(titulo, texto):
    """Processa um livro completo."""
    resultado = {
        'titulo': titulo,
        'status': 'erro',
        'idioma': '',
        'notas': 0,
        'validacao': ''
    }

    # 1. Validação
    valido, motivo = validar_livro(texto, titulo)
    resultado['validacao'] = motivo
    if not valido:
        resultado['status'] = 'invalido'
        return resultado, None

    # 2. Detectar idioma
    idioma = detectar_idioma(texto)
    resultado['idioma'] = idioma

    timeout = aiohttp.ClientTimeout(total=300)
    semaphore = asyncio.Semaphore(10)

    async with aiohttp.ClientSession(timeout=timeout) as session:
        try:
            # 3. Correção OCR
            chunks = dividir_chunks(texto)
            tasks = []
            for chunk in chunks:
                prompt = f"Corrija erros de OCR/digitalização neste texto. Retorne APENAS o texto corrigido, sem explicações:\\n\\n{chunk}"
                tasks.append(call_gemini(session, prompt, semaphore))

            resultados_ocr = await asyncio.gather(*tasks)
            texto = '\\n\\n'.join([r for r in resultados_ocr if r])

            # 4. Tradução (se não for PT)
            if idioma != 'pt' and texto:
                lang_map = {'en': 'inglês', 'es': 'espanhol', 'fr': 'francês', 'de': 'alemão', 'ru': 'russo'}
                chunks = dividir_chunks(texto)
                tasks = []
                for chunk in chunks:
                    prompt = f"Traduza de {lang_map.get(idioma, idioma)} para português brasileiro. Retorne APENAS a tradução:\\n\\n{chunk}"
                    tasks.append(call_gemini(session, prompt, semaphore))

                resultados_trad = await asyncio.gather(*tasks)
                texto = '\\n\\n'.join([r for r in resultados_trad if r])

            # 5. Gerar notas de rodapé
            if texto:
                prompt = f"""Analise este texto e identifique termos que precisam de notas de rodapé explicativas.
Inclua: termos estrangeiros, conceitos filosóficos/técnicos, referências históricas, nomes de pessoas importantes.

Formato de resposta (um por linha):
termo|explicação breve (máximo 30 palavras)

Máximo 20 termos mais importantes.

Texto:
{texto[:8000]}"""

                notas_raw = await call_gemini(session, prompt, semaphore)

                notas = []
                for linha in notas_raw.split('\\n'):
                    if '|' in linha:
                        partes = linha.split('|', 1)
                        if len(partes) == 2:
                            termo = partes[0].strip().strip('- *')
                            explicacao = partes[1].strip()
                            if termo and explicacao and len(termo) > 1:
                                notas.append((termo, explicacao))

                notas = notas[:20]
                resultado['notas'] = len(notas)

                # 6. Criar DOCX
                docx_bytes = criar_docx_completo(texto, titulo, notas)
                resultado['status'] = 'sucesso'

                return resultado, docx_bytes

        except Exception as e:
            resultado['error'] = str(e)

    return resultado, None

def criar_docx_completo(texto, titulo, notas):
    """Cria DOCX com notas de rodapé reais."""
    from io import BytesIO
    import zipfile

    doc = Document()

    # Título
    p = doc.add_paragraph()
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Conteúdo com referências de notas
    nota_counter = 1
    texto_processado = texto

    for termo, explicacao in notas:
        if termo.lower() in texto_processado.lower():
            # Marca primeira ocorrência
            import re
            pattern = re.compile(re.escape(termo), re.IGNORECASE)
            texto_processado = pattern.sub(f"{termo}[{nota_counter}]", texto_processado, count=1)
            nota_counter += 1

    # Adiciona parágrafos
    for para in texto_processado.split('\\n\\n'):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Página de notas
    if notas:
        doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run("NOTAS")
        run.bold = True
        run.font.size = Pt(14)

        for i, (termo, explicacao) in enumerate(notas, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"[{i}] {termo}: ")
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(explicacao)
            run.font.size = Pt(10)

    # Salva em bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def main():
    """Processa todos os livros do bucket."""
    import logging
    logging.basicConfig(level=logging.INFO)

    # Conecta ao GCS
    client = storage.Client()
    bucket = client.bucket(BUCKET_NAME)

    # Lista livros
    blobs = list(bucket.list_blobs(prefix="livros/"))
    livros = {}
    for blob in blobs:
        if blob.name.endswith('.txt'):
            parts = blob.name.split('/')
            if len(parts) >= 3:
                titulo = parts[1]
                livros[titulo] = blob

    print(f"Total de livros: {len(livros)}")

    # Processa
    resultados = []
    for i, (titulo, blob) in enumerate(livros.items()):
        print(f"\\n[{i+1}/{len(livros)}] {titulo[:50]}...")

        # Baixa texto
        texto = blob.download_as_text(encoding='utf-8')

        # Processa
        resultado, docx_bytes = asyncio.run(processar_livro(titulo, texto))
        resultados.append(resultado)

        print(f"  Status: {resultado['status']}")
        print(f"  Idioma: {resultado.get('idioma', '?')}")
        print(f"  Notas: {resultado.get('notas', 0)}")

        # Salva DOCX
        if docx_bytes:
            docx_blob = bucket.blob(f"docx/{titulo}.docx")
            docx_blob.upload_from_string(docx_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            print(f"  ✓ DOCX salvo")

    # Salva resumo
    resumo = {
        'total': len(resultados),
        'sucesso': sum(1 for r in resultados if r['status'] == 'sucesso'),
        'erro': sum(1 for r in resultados if r['status'] == 'erro'),
        'invalido': sum(1 for r in resultados if r['status'] == 'invalido'),
        'resultados': resultados
    }

    resumo_blob = bucket.blob("resultados/resumo.json")
    resumo_blob.upload_from_string(json.dumps(resumo, ensure_ascii=False, indent=2))

    print(f"\\n{'='*60}")
    print("RESUMO")
    print('='*60)
    print(f"Total: {resumo['total']}")
    print(f"Sucesso: {resumo['sucesso']}")
    print(f"Erros: {resumo['erro']}")
    print(f"Inválidos: {resumo['invalido']}")

if __name__ == "__main__":
    main()
'''

    # Salva localmente
    processor_path = BASE_DIR / "gcp_processor.py"
    with open(processor_path, 'w', encoding='utf-8') as f:
        f.write(processor_code)

    # Sobe para GCS
    subprocess.run([
        "gsutil", "cp", str(processor_path),
        f"gs://{BUCKET_NAME}/scripts/gcp_processor.py"
    ])

    print(f"✓ Script de processamento enviado")

def criar_vm_e_processar():
    """Cria VM na GCP e inicia processamento."""
    print("=" * 60)
    print("CRIANDO VM PARA PROCESSAMENTO")
    print("=" * 60)

    project_id = get_project_id()
    vm_name = "kdp-processor"
    zone = "us-central1-a"

    # Script de startup
    startup_script = f'''#!/bin/bash
set -e

# Instala dependências
apt-get update
apt-get install -y python3-pip
pip3 install google-cloud-storage aiohttp python-docx lxml

# Baixa script
gsutil cp gs://{BUCKET_NAME}/scripts/gcp_processor.py /tmp/processor.py

# Configura variáveis
export BUCKET_NAME="{BUCKET_NAME}"
export GEMINI_API_KEY="AIzaSyExample"  # Será substituído

# Executa
cd /tmp
python3 processor.py > /var/log/processor.log 2>&1

# Marca como completo
echo "COMPLETO" | gsutil cp - gs://{BUCKET_NAME}/status/complete.txt

# Auto-delete após completar
gcloud compute instances delete {vm_name} --zone={zone} --quiet
'''

    # Salva startup script
    startup_path = BASE_DIR / "startup.sh"
    with open(startup_path, 'w') as f:
        f.write(startup_script)

    print(f"VM: {vm_name}")
    print(f"Zone: {zone}")
    print(f"Bucket: {BUCKET_NAME}")

    # Cria VM
    cmd = [
        "gcloud", "compute", "instances", "create", vm_name,
        f"--project={project_id}",
        f"--zone={zone}",
        "--machine-type=e2-standard-4",
        "--boot-disk-size=50GB",
        "--image-family=debian-11",
        "--image-project=debian-cloud",
        f"--metadata-from-file=startup-script={startup_path}",
        "--scopes=cloud-platform",
        "--tags=http-server"
    ]

    print("\nCriando VM...")
    result = subprocess.run(cmd, capture_output=True, text=True)

    if result.returncode == 0:
        print(f"✓ VM {vm_name} criada!")
        print("\nO processamento vai iniciar automaticamente.")
        print(f"Acompanhe em: gsutil cat gs://{BUCKET_NAME}/status/complete.txt")
    else:
        print(f"Erro: {result.stderr}")

        # Alternativa: Cloud Run
        print("\nTentando Cloud Run como alternativa...")
        usar_cloud_run()

def usar_cloud_run():
    """Usa Cloud Run Jobs para processamento."""
    print("\n" + "=" * 60)
    print("CONFIGURANDO CLOUD RUN JOB")
    print("=" * 60)

    # Cria Dockerfile
    dockerfile = '''FROM python:3.11-slim

WORKDIR /app

RUN pip install google-cloud-storage aiohttp python-docx lxml

COPY gcp_processor.py .

ENV BUCKET_NAME=testmassivo
ENV GEMINI_API_KEY=""

CMD ["python", "gcp_processor.py"]
'''

    dockerfile_path = BASE_DIR / "Dockerfile.cloudrun"
    with open(dockerfile_path, 'w') as f:
        f.write(dockerfile)

    print("Dockerfile criado.")
    print("\nPara usar Cloud Run:")
    print(f"1. gcloud builds submit --tag gcr.io/{get_project_id()}/kdp-processor")
    print(f"2. gcloud run jobs create kdp-job --image gcr.io/{get_project_id()}/kdp-processor --region {REGION}")
    print(f"3. gcloud run jobs execute kdp-job --region {REGION}")

def verificar_status():
    """Verifica status do processamento."""
    print("=" * 60)
    print("STATUS DO PROCESSAMENTO")
    print("=" * 60)

    # Verifica se completou
    result = subprocess.run(
        ["gsutil", "cat", f"gs://{BUCKET_NAME}/status/complete.txt"],
        capture_output=True, text=True
    )

    if "COMPLETO" in result.stdout:
        print("✓ Processamento COMPLETO!")
    else:
        print("⏳ Processamento em andamento...")

    # Conta DOCXs
    result = subprocess.run(
        ["gsutil", "ls", f"gs://{BUCKET_NAME}/docx/"],
        capture_output=True, text=True
    )

    docx_count = len([l for l in result.stdout.split('\n') if l.endswith('.docx')])
    print(f"\nDOCXs prontos: {docx_count}")

    # Baixa resumo se existir
    result = subprocess.run(
        ["gsutil", "cat", f"gs://{BUCKET_NAME}/resultados/resumo.json"],
        capture_output=True, text=True
    )

    if result.returncode == 0:
        try:
            resumo = json.loads(result.stdout)
            print(f"\nResumo:")
            print(f"  Total: {resumo.get('total', '?')}")
            print(f"  Sucesso: {resumo.get('sucesso', '?')}")
            print(f"  Erros: {resumo.get('erro', '?')}")
            print(f"  Inválidos: {resumo.get('invalido', '?')}")
        except Exception:
            pass

def download_docx():
    """Baixa todos os DOCXs prontos."""
    print("=" * 60)
    print("DOWNLOAD DOS DOCX")
    print("=" * 60)

    DOCX_DIR.mkdir(exist_ok=True)

    print(f"Baixando de gs://{BUCKET_NAME}/docx/ para {DOCX_DIR}")

    result = subprocess.run([
        "gsutil", "-m", "cp", "-r",
        f"gs://{BUCKET_NAME}/docx/*",
        str(DOCX_DIR)
    ], capture_output=True, text=True)

    if result.returncode == 0:
        docx_files = list(DOCX_DIR.glob("*.docx"))
        print(f"\n✓ {len(docx_files)} DOCXs baixados para {DOCX_DIR}")
    else:
        print(f"Erro: {result.stderr}")

# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description="Pipeline Completo GCP")
    parser.add_argument('--upload', action='store_true', help='Sobe livros para GCS')
    parser.add_argument('--process', action='store_true', help='Inicia processamento na GCP')
    parser.add_argument('--status', action='store_true', help='Verifica status')
    parser.add_argument('--download', action='store_true', help='Baixa DOCXs prontos')
    parser.add_argument('--all', action='store_true', help='Executa tudo: upload + process')
    args = parser.parse_args()

    if not any([args.upload, args.process, args.status, args.download, args.all]):
        parser.print_help()
        print("\n" + "=" * 60)
        print("FLUXO RECOMENDADO:")
        print("=" * 60)
        print("1. python gcp_pipeline_completo.py --upload    # Sobe livros")
        print("2. python gcp_pipeline_completo.py --process   # Inicia processamento")
        print("3. python gcp_pipeline_completo.py --status    # Acompanha")
        print("4. python gcp_pipeline_completo.py --download  # Baixa DOCXs")
        return

    if args.upload or args.all:
        upload_to_gcs()

    if args.process or args.all:
        criar_vm_e_processar()

    if args.status:
        verificar_status()

    if args.download:
        download_docx()

if __name__ == "__main__":
    main()

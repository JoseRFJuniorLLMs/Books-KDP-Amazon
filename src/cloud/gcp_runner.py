# -*- coding: utf-8 -*-
"""
GCP RODAR - Pipeline Simples e Direto
=====================================
Sobe, processa na nuvem, baixa pronto.

Uso:
    python gcp_rodar.py --run
"""

import os
import sys
import json
import subprocess
import time
from pathlib import Path

BUCKET = "aurorav2-484411-kdp"
REGION = "us-central1"
BASE_DIR = Path(__file__).parent
TXT_DIR = BASE_DIR / "txt"
DOCX_DIR = BASE_DIR / "docx"

# Google Cloud SDK path no Windows
GCLOUD_BIN = r"C:\Program Files (x86)\Google\Cloud SDK\google-cloud-sdk\bin"
GSUTIL = os.path.join(GCLOUD_BIN, "gsutil.cmd")
GCLOUD = os.path.join(GCLOUD_BIN, "gcloud.cmd")

# Vertex AI API Key (já configurada)
VERTEX_API_KEY = "AQ.Ab8RN6I5q1OHyDoU18jD39-8LwCZwPlXHd9SyGsPtLzmzjF_hw"

def run_cmd(cmd, show=True):
    """Executa comando e retorna output."""
    if show:
        print(f"$ {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True, shell=True)
    return result.returncode == 0, result.stdout, result.stderr

def step1_upload():
    """Passo 1: Upload dos TXTs para GCS."""
    print("\n" + "="*60)
    print("PASSO 1: UPLOAD PARA GCS")
    print("="*60)

    # Cria bucket
    run_cmd([GSUTIL, "mb", "-l", REGION, f"gs://{BUCKET}"], show=False)

    # Conta livros
    livros = [p for p in TXT_DIR.iterdir() if p.is_dir()]
    print(f"Livros para enviar: {len(livros)}")

    # Upload paralelo de toda a pasta txt
    print("\nEnviando... (isso pode levar alguns minutos)")
    ok, out, err = run_cmd([
        GSUTIL, "-m", "cp", "-r",
        str(TXT_DIR),
        f"gs://{BUCKET}/"
    ])

    if ok:
        print(f"✓ Upload completo!")
    else:
        print(f"Erro: {err}")

    return ok

def step2_criar_job():
    """Passo 2: Cria Cloud Run Job para processar."""
    print("\n" + "="*60)
    print("PASSO 2: CRIAR JOB DE PROCESSAMENTO")
    print("="*60)

    # Pega project ID
    ok, project, _ = run_cmd([GCLOUD, "config", "get-value", "project"], show=False)
    project = project.strip()
    print(f"Projeto: {project}")

    # Cria o script de processamento inline
    processor_script = f'''
import os
import json
import asyncio
import aiohttp
from google.cloud import storage
from io import BytesIO

BUCKET = "{BUCKET}"
API_KEY = "AQ.Ab8RN6I5q1OHyDoU18jD39-8LwCZwPlXHd9SyGsPtLzmzjF_hw"
API_URL = "https://aiplatform.googleapis.com/v1/publishers/google/models/gemini-2.0-flash:generateContent"

LANG_WORDS = {{
    'pt': ['que', 'não', 'para', 'uma', 'com', 'por', 'mais'],
    'es': ['que', 'los', 'las', 'una', 'para', 'con', 'más'],
    'en': ['the', 'and', 'that', 'have', 'for', 'with', 'this'],
}}

def detectar_idioma(texto):
    texto = texto[:3000].lower()
    scores = {{lang: sum(1 for w in words if f' {{w}} ' in texto)
              for lang, words in LANG_WORDS.items()}}
    return max(scores, key=scores.get) if scores else 'en'

async def chamar_gemini(session, prompt, sem):
    async with sem:
        url = f"{{API_URL}}?key={{API_KEY}}"
        payload = {{"contents": [{{"role": "user", "parts": [{{"text": prompt}}]}}],
                   "generationConfig": {{"temperature": 0.3, "maxOutputTokens": 8192}}}}
        for _ in range(3):
            try:
                async with session.post(url, json=payload) as r:
                    if r.status == 200:
                        d = await r.json()
                        return d.get("candidates", [{{}}])[0].get("content", {{}}).get("parts", [{{}}])[0].get("text", "")
                    await asyncio.sleep(2)
            except:
                await asyncio.sleep(1)
        return ""

def chunks(texto, tam=4000):
    result, cur = [], ""
    for p in texto.split("\\n\\n"):
        if len(cur) + len(p) < tam:
            cur += "\\n\\n" + p if cur else p
        else:
            if cur: result.append(cur)
            cur = p
    if cur: result.append(cur)
    return result or [texto]

async def processar(titulo, texto):
    idioma = detectar_idioma(texto)
    sem = asyncio.Semaphore(10)

    async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=300)) as s:
        # OCR
        tasks = [chamar_gemini(s, f"Corrija erros de OCR. Retorne só o texto:\\n{{c}}", sem) for c in chunks(texto)]
        texto = "\\n\\n".join([r for r in await asyncio.gather(*tasks) if r])

        # Tradução
        if idioma != 'pt' and texto:
            lang = {{'en': 'inglês', 'es': 'espanhol'}}.get(idioma, idioma)
            tasks = [chamar_gemini(s, f"Traduza de {{lang}} para português:\\n{{c}}", sem) for c in chunks(texto)]
            texto = "\\n\\n".join([r for r in await asyncio.gather(*tasks) if r])

        # Notas
        notas = []
        if texto:
            r = await chamar_gemini(s, f"Liste termos para notas de rodapé (termo|explicação):\\n{{texto[:5000]}}", sem)
            for l in r.split("\\n"):
                if "|" in l:
                    p = l.split("|", 1)
                    if len(p) == 2 and p[0].strip():
                        notas.append((p[0].strip(), p[1].strip()))

        return texto, notas[:15], idioma

def criar_docx(texto, titulo, notas):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run(titulo)
    r.bold, r.font.size = True, Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for para in texto.split("\\n\\n"):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if notas:
        doc.add_page_break()
        p = doc.add_paragraph()
        r = p.add_run("NOTAS")
        r.bold = True
        for i, (t, e) in enumerate(notas, 1):
            p = doc.add_paragraph()
            r = p.add_run(f"[{{i}}] {{t}}: ")
            r.bold, r.font.size = True, Pt(10)
            r = p.add_run(e)
            r.font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def main():
    client = storage.Client()
    bucket = client.bucket(BUCKET)

    # Lista livros
    blobs = {{b.name.split("/")[1]: b for b in bucket.list_blobs(prefix="txt/")
             if b.name.endswith(".txt") and "/" in b.name}}

    print(f"Livros: {{len(blobs)}}")
    results = []

    for i, (titulo, blob) in enumerate(blobs.items()):
        print(f"[{{i+1}}/{{len(blobs)}}] {{titulo[:40]}}")
        try:
            texto = blob.download_as_text()
            if len(texto) < 3000:
                print("  Skip: muito curto")
                continue

            texto_proc, notas, idioma = asyncio.run(processar(titulo, texto))
            print(f"  Idioma: {{idioma}}, Notas: {{len(notas)}}")

            if texto_proc:
                docx = criar_docx(texto_proc, titulo, notas)
                bucket.blob(f"docx/{{titulo}}.docx").upload_from_string(
                    docx, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                print(f"  ✓ Salvo")
                results.append({{"titulo": titulo, "status": "ok", "notas": len(notas)}})
            else:
                results.append({{"titulo": titulo, "status": "erro"}})
        except Exception as e:
            print(f"  Erro: {{e}}")
            results.append({{"titulo": titulo, "status": "erro", "msg": str(e)}})

    # Salva resumo
    bucket.blob("resultado.json").upload_from_string(json.dumps(results, ensure_ascii=False))
    bucket.blob("COMPLETO.txt").upload_from_string("OK")
    print(f"\\n✓ Completo! {{sum(1 for r in results if r['status']=='ok')}}/{{len(results)}} sucesso")

if __name__ == "__main__":
    main()
'''

    # Salva script
    script_path = BASE_DIR / "cloud_processor.py"
    with open(script_path, 'w', encoding='utf-8') as f:
        f.write(processor_script)

    # Cria requirements
    req_path = BASE_DIR / "requirements_cloud.txt"
    with open(req_path, 'w') as f:
        f.write("google-cloud-storage\naiohttp\npython-docx\n")

    # Cria Dockerfile
    dockerfile = '''FROM python:3.11-slim
WORKDIR /app
COPY requirements_cloud.txt .
RUN pip install --no-cache-dir -r requirements_cloud.txt
COPY cloud_processor.py .
CMD ["python", "cloud_processor.py"]
'''
    docker_path = BASE_DIR / "Dockerfile"
    with open(docker_path, 'w') as f:
        f.write(dockerfile)

    print("\nArquivos criados:")
    print(f"  - {script_path}")
    print(f"  - {req_path}")
    print(f"  - {docker_path}")

    print("\n" + "-"*60)
    print("PRÓXIMOS PASSOS:")
    print("-"*60)
    print(f"""
# Build da imagem
gcloud builds submit --tag gcr.io/{project}/kdp-processor

# Criar job (API key já inclusa no código)
gcloud run jobs create kdp-processor --image gcr.io/{project}/kdp-processor --region {REGION} --memory 4Gi --cpu 2 --task-timeout 3600

# Executar
gcloud run jobs execute kdp-processor --region {REGION}

# Acompanhar
gcloud run jobs executions list --job kdp-processor --region {REGION}

# Quando terminar, baixar
gsutil -m cp -r gs://{BUCKET}/docx/* {DOCX_DIR}/
""")

    return True

def step3_verificar():
    """Verifica status e baixa resultados."""
    print("\n" + "="*60)
    print("VERIFICANDO STATUS")
    print("="*60)

    # Verifica se completou
    ok, out, _ = run_cmd([GSUTIL, "cat", f"gs://{BUCKET}/COMPLETO.txt"], show=False)

    if ok and "OK" in out:
        print("✓ Processamento COMPLETO!")

        # Conta DOCXs
        ok, out, _ = run_cmd([GSUTIL, "ls", f"gs://{BUCKET}/docx/"], show=False)
        docx_count = len([l for l in out.split('\n') if '.docx' in l])
        print(f"DOCXs prontos: {docx_count}")

        # Baixa
        print(f"\nBaixando para {DOCX_DIR}...")
        DOCX_DIR.mkdir(exist_ok=True)
        run_cmd([GSUTIL, "-m", "cp", "-r", f"gs://{BUCKET}/docx/*", str(DOCX_DIR)])

        local_count = len(list(DOCX_DIR.glob("*.docx")))
        print(f"✓ {local_count} DOCXs baixados!")
    else:
        print("⏳ Ainda processando...")

        # Mostra progresso
        ok, out, _ = run_cmd([GSUTIL, "ls", f"gs://{BUCKET}/docx/"], show=False)
        if ok:
            docx_count = len([l for l in out.split('\n') if '.docx' in l])
            print(f"DOCXs processados até agora: {docx_count}")

def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--upload', action='store_true', help='Sobe livros')
    parser.add_argument('--setup', action='store_true', help='Cria job')
    parser.add_argument('--status', action='store_true', help='Verifica e baixa')
    parser.add_argument('--run', action='store_true', help='Executa tudo')
    args = parser.parse_args()

    if args.run:
        step1_upload()
        step2_criar_job()
        print("\n" + "="*60)
        print("Siga as instruções acima para iniciar o processamento!")
        print("="*60)
    elif args.upload:
        step1_upload()
    elif args.setup:
        step2_criar_job()
    elif args.status:
        step3_verificar()
    else:
        parser.print_help()

if __name__ == "__main__":
    main()

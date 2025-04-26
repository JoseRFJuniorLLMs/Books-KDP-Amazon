# -*- coding: utf-8 -*-
# --- Using Google's Gemini API (gemini-1.5-pro) ---
# --- VERSÃO v11: Refatorada (Otimiza DOCX Gen, Mantém Funcionalidade v10) ---

# Standard Python Libraries
import sys
import os
import re
import logging
import time
import shutil
import traceback
import glob
import smtplib
import ssl
from email.message import EmailMessage
import subprocess
from typing import Tuple, Optional, List, Dict, Set, Any # For type hinting

# Third-party Libraries
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from tqdm import tqdm
import google.generativeai as genai

# === SETUP LOGGING ===
log_dir = "logs"
if not os.path.exists(log_dir): os.makedirs(log_dir)
log_filepath = os.path.join(log_dir, "book_processor_multi_author_mem_v11.log") # Log file versioned
PROCESSED_LOG_FILE = os.path.join(log_dir, "processed_books.log")
TRANSLATED_LOG_FILE = os.path.join(log_dir, "translated_books.log")

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(module)s:%(lineno)d - %(funcName)s - %(message)s',
    handlers=[ logging.FileHandler(log_filepath, encoding='utf-8'), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# === CARREGA VARIÁVEIS DE AMBIENTE ===
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
EMAIL_SENDER_ADDRESS = os.getenv("EMAIL_SENDER_ADDRESS")
EMAIL_SENDER_APP_PASSWORD = os.getenv("EMAIL_SENDER_APP_PASSWORD")
EMAIL_RECIPIENT_ADDRESS = os.getenv("EMAIL_RECIPIENT_ADDRESS", "web2ajax@gmail.com")
EMAIL_SMTP_SERVER = os.getenv("EMAIL_SMTP_SERVER", "smtp.gmail.com")
EMAIL_SMTP_PORT = int(os.getenv("EMAIL_SMTP_PORT", 587))

# === CONFIGURAÇÕES ===

# -- Diretórios Base --
BASE_INPUT_TXT_DIR = "txt"
BASE_OUTPUT_DOCX_DIR = "docx"

# -- Nomes de Arquivos Base --
TEMPLATE_DOCX = "Estrutura.docx"

# -- Nomes Base dos Arquivos de Saída --
FINAL_DOCX_CLEAN_BASENAME = "Livro_Final_Formatado_Sem_Notas.docx" # Para Tradutor
FINAL_DOCX_WITH_NOTES_BASENAME = "Livro_Final_Com_Notas_No_Fim.docx" # Final
TRANSLATED_DOCX_SUFFIX = "-A0.docx" # Sufixo para traduzido

# -- Nomes para Filtro de Input (Arquivos TXT não gerados mais) --
FINAL_NUMBERED_TXT_BASENAME_FILTER = "Livro_Final_Com_Notas_Numeros.txt"
NOTES_TXT_FILE_BASENAME_FILTER = "notas_rodape.txt"

# -- Configurações da API e Processamento --
MODEL_NAME = "gemini-1.5-pro"
MAX_CHUNK_TOKENS = 1500
MAX_OUTPUT_TOKENS = 8192
TEMPERATURE = 0.5

# -- Configurações para o Script Tradutor --
PATH_TO_TRANSLATOR_SCRIPT = "script_tradutor_hibrido.py"
NUM_WORDS_TO_TRANSLATE = 100

# -- Estilos e Padrões --
NORMAL_STYLE_NAME = "Normal"
CHAPTER_STYLE_NAME = "Heading 1" # <<< Adicionado (Usado em apply_formatting)
CHAPTER_PATTERNS = [ r'^\s*Capítulo \w+', r'^\s*CAPÍTULO \w+', r'^\s*Capítulo \d+', r'^\s*CHAPTER \w+', r'^\s*Chapter \d+', r'^\s*LIVRO \w+', r'^\s*PARTE \w+']
PAGE_BREAK_MARKER = "===QUEBRA_DE_PAGINA==="
AI_FAILURE_MARKER = "*** FALHA NA IA - TEXTO ORIGINAL ABAIXO ***"
FORMATTING_ERROR_MARKER = "*** ERRO DE FORMATAÇÃO - TEXTO ORIGINAL ABAIXO ***"

# -- Padrões Regex para Notas --
# Captura o marcador completo [NOTA_...][CONTEUDO_NOTA:...]
FOOTNOTE_MARKER_PATTERN = re.compile(r'(\[NOTA_(?:IDIOMA|CITACAO|NOME|TERMO):[^\]]+?\]\s*\[CONTEUDO_NOTA:[^\]]*?\])', re.IGNORECASE)
# Extrai apenas o conteúdo de dentro de [CONTEUDO_NOTA:...]
FOOTNOTE_CONTENT_EXTRACT_PATTERN = re.compile(r'\[CONTEUDO_NOTA:([^\]]*?)\]', re.IGNORECASE | re.DOTALL) # DOTALL para multiline
# Encontra os placeholders inseridos temporariamente
PLACEHOLDER_REGEX = re.compile(r'(__AI_NOTE_PLACEHOLDER_(\d+)__)')

# --- Tipos de Dados para Type Hinting ---
StatsTuple = Tuple[float, int, int, int] # latency, p_tokens, o_tokens, t_tokens
ApiResultTuple = Tuple[Optional[str], float, int, int, int] # text | None, latency, p_tokens, o_tokens, t_tokens

# --- Validações e Setup ---
if not GOOGLE_API_KEY: logger.error("FATAL: GOOGLE_API_KEY não encontrada."); sys.exit(1)
email_configured = bool(EMAIL_SENDER_ADDRESS and EMAIL_SENDER_APP_PASSWORD and EMAIL_RECIPIENT_ADDRESS)
if not email_configured: logger.warning("AVISO: Config e-mail incompleta. Notificação desativada.")
else: logger.info(f"Config e-mail OK: De '{EMAIL_SENDER_ADDRESS}' Para '{EMAIL_RECIPIENT_ADDRESS}'.")
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    safety_settings_lenient = {'HATE': 'BLOCK_NONE', 'HARASSMENT': 'BLOCK_NONE', 'SEXUAL': 'BLOCK_NONE', 'DANGEROUS': 'BLOCK_NONE'}
    generation_config = genai.GenerationConfig(temperature=TEMPERATURE, max_output_tokens=MAX_OUTPUT_TOKENS)
    gemini_model = genai.GenerativeModel(MODEL_NAME, safety_settings=safety_settings_lenient, generation_config=generation_config)
    logger.info(f"Modelo Gemini '{MODEL_NAME}' inicializado.")
except Exception as e: logger.error(f"FATAL: Falha inicializar Gemini: {e}"); logger.error(traceback.format_exc()); sys.exit(1)

# --- Funções Auxiliares ---

def count_tokens_approx(text: Optional[str]) -> int:
    """Estima a contagem de tokens (aproximadamente 3 chars/token)."""
    if not text: return 0
    return len(text) // 3

def create_chunks(text: str, max_tokens: int, author_name: str = "N/A", book_name: str = "N/A") -> List[str]:
    """Divide o texto em chunks, subdividindo parágrafos grandes. (Lógica inalterada)"""
    log_prefix = f"[{author_name}/{book_name}]"; logger.info(f"{log_prefix} Iniciando chunking (Max: {max_tokens} tokens)...")
    chunks = []; current_chunk = ""; current_chunk_tokens = 0
    paragraphs = text.split("\n\n")
    logger.debug(f"{log_prefix} Texto dividido em {len(paragraphs)} blocos iniciais.")

    for i, paragraph_text in enumerate(paragraphs):
        if not paragraph_text.strip():
            if chunks and chunks[-1].strip() and not chunks[-1].endswith("\n\n"): chunks[-1] += "\n\n"
            continue

        paragraph_tokens = count_tokens_approx(paragraph_text);
        tokens_with_separator = paragraph_tokens + (count_tokens_approx("\n\n") if current_chunk else 0)

        # Combinar ou iniciar novo chunk
        if current_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if current_chunk else ""; current_chunk += separator + paragraph_text; current_chunk_tokens = count_tokens_approx(current_chunk)
        else:
            if current_chunk: chunks.append(current_chunk); logger.debug(f"{log_prefix} Chunk {len(chunks)} add (limite). Tokens: {count_tokens_approx(chunks[-1])}")
            current_chunk = paragraph_text; current_chunk_tokens = paragraph_tokens

            # Se o novo chunk (o parágrafo sozinho) já excede, precisa subdividir
            if paragraph_tokens > max_tokens:
                logger.warning(f"{log_prefix} Parágrafo {i+1} ({paragraph_tokens} tk) > limite {max_tokens}. SUBDIVIDINDO.")
                current_chunk = ""; current_chunk_tokens = 0 # Reset current chunk as it's being replaced by sub-chunks
                sub_chunks_added_count = 0
                # Try splitting by sentence, then by line
                sentences = re.split(r'(?<=[.!?])\s+', paragraph_text);
                if len(sentences) <= 1: sentences = paragraph_text.split('\n')
                logger.debug(f"{log_prefix} -> Subdividindo em {len(sentences)} sentenças/linhas.")

                current_sub_chunk = ""; current_sub_chunk_tokens = 0
                for sentence_num, sentence in enumerate(sentences):
                    sentence_clean = sentence.strip();
                    if not sentence_clean: continue

                    sentence_tokens = count_tokens_approx(sentence);
                    tokens_with_sub_separator = sentence_tokens + (count_tokens_approx("\n") if current_sub_chunk else 0)

                    if current_sub_chunk_tokens + tokens_with_sub_separator <= max_tokens:
                        sub_separator = "\n" if current_sub_chunk else ""; current_sub_chunk += sub_separator + sentence; current_sub_chunk_tokens = count_tokens_approx(current_sub_chunk)
                    else:
                        # Salva o sub-chunk anterior se existir
                        if current_sub_chunk: chunks.append(current_sub_chunk); sub_chunks_added_count += 1; logger.debug(f"{log_prefix} Sub-chunk {len(chunks)} add (limite sub). Tokens: {count_tokens_approx(chunks[-1])}")
                        # Se a própria sentença for muito grande, adiciona-a sozinha (pode falhar na API)
                        if sentence_tokens > max_tokens:
                            logger.warning(f"{log_prefix} -> Sentença/Linha {sentence_num+1} ({sentence_tokens} tk) > limite. Adicionada separadamente.")
                            chunks.append(sentence); sub_chunks_added_count += 1
                            current_sub_chunk = ""; current_sub_chunk_tokens = 0 # Reset sub-chunk
                        else: # Inicia novo sub-chunk com a sentença atual
                             current_sub_chunk = sentence; current_sub_chunk_tokens = sentence_tokens

                # Adiciona o último sub-chunk
                if current_sub_chunk: chunks.append(current_sub_chunk); sub_chunks_added_count += 1; logger.debug(f"{log_prefix} Sub-chunk final {len(chunks)} add. Tokens: {count_tokens_approx(chunks[-1])}")

                if sub_chunks_added_count == 0: # Se não conseguiu subdividir
                    logger.warning(f"{log_prefix} Parágrafo {i+1} não subdividido apesar de exceder limite. Adicionando original.")
                    chunks.append(paragraph_text)

                # Reseta o chunk principal pois foi substituído pela subdivisão
                current_chunk = ""; current_chunk_tokens = 0

    # Adiciona o último chunk que sobrou
    if current_chunk: chunks.append(current_chunk); logger.debug(f"{log_prefix} Chunk final {len(chunks)} add. Tokens: {count_tokens_approx(chunks[-1])}")

    # --- Merge pass ---
    if len(chunks) < 2: # No need to merge if less than 2 chunks
         logger.info(f"{log_prefix} Chunking finalizado. {len(chunks)} chunks.")
         return chunks

    logger.info(f"{log_prefix} Tentando merge de {len(chunks)} chunks...")
    merged_chunks = []; temp_chunk = ""; temp_chunk_tokens = 0
    for i, chunk in enumerate(chunks):
        chunk_tokens = count_tokens_approx(chunk);
        tokens_with_separator = chunk_tokens + (count_tokens_approx("\n\n") if temp_chunk else 0)

        if temp_chunk_tokens + tokens_with_separator <= max_tokens:
            separator = "\n\n" if temp_chunk else ""; temp_chunk += separator + chunk; temp_chunk_tokens = count_tokens_approx(temp_chunk)
        else:
            if temp_chunk: merged_chunks.append(temp_chunk) # Salva o chunk anterior que estava cheio
            temp_chunk = chunk; temp_chunk_tokens = chunk_tokens # Inicia o novo

    if temp_chunk: merged_chunks.append(temp_chunk) # Salva o último

    final_chunk_count = len(merged_chunks)
    if final_chunk_count < len(chunks): logger.info(f"{log_prefix} Merge concluído: {len(chunks)} -> {final_chunk_count} chunks.")
    else: logger.info(f"{log_prefix} Merge não reduziu o número de chunks.")
    logger.info(f"{log_prefix} Chunking finalizado. {final_chunk_count} chunks.")
    return merged_chunks

def _call_gemini_api(model: genai.GenerativeModel, prompt_text: str, chunk_for_log: str, author_name: str = "N/A", book_name: str = "N/A") -> ApiResultTuple:
    """Chama API Gemini com retries e retorna ApiResultTuple."""
    log_prefix = f"[{author_name}/{book_name}]"; max_retries = 5; base_wait_time = 5
    log_chunk_preview = chunk_for_log[:100].replace('\n', '\\n') + '...'
    default_return: ApiResultTuple = (None, 0, 0, 0, 0)

    for attempt in range(max_retries):
        start_time = time.time(); response = None; result_text: Optional[str] = None
        latency = 0; prompt_tokens = 0; output_tokens = 0; total_tokens = 0
        logger.info(f"{log_prefix} Chamando API (Tentativa {attempt + 1}/{max_retries}). Chunk: '{log_chunk_preview}'")
        try:
            # Use generation_config from the model object
            response = model.generate_content(prompt_text)
            latency = time.time() - start_time

            # --- Processamento da Resposta ---
            # 1. Usage Metadata
            try:
                if hasattr(response, 'usage_metadata'):
                    usage = response.usage_metadata
                    prompt_tokens = getattr(usage, 'prompt_token_count', 0)
                    output_tokens = getattr(usage, 'candidates_token_count', 0)
                    total_tokens = getattr(usage, 'total_token_count', prompt_tokens + output_tokens)
                    logger.debug(f"{log_prefix} API OK ({latency:.2f}s). Tokens: P{prompt_tokens}+O{output_tokens}=T{total_tokens}")
                else:
                    logger.debug(f"{log_prefix} API OK ({latency:.2f}s). Usage metadata não encontrada.")
            except Exception as e_usage:
                logger.warning(f"{log_prefix} Erro ao processar usage_metadata: {e_usage}")

            # 2. Prompt Feedback (Bloqueio)
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback and \
               hasattr(response.prompt_feedback, 'block_reason') and response.prompt_feedback.block_reason:
                block_reason = response.prompt_feedback.block_reason.name
                logger.error(f"{log_prefix} API BLOQUEOU PROMPT (Tentativa {attempt+1}): {block_reason}. Lat:{latency:.2f}s.")
                return default_return # Não adianta tentar de novo

            # 3. Candidatos e Texto
            if not response.candidates:
                logger.error(f"{log_prefix} API retornou SEM CANDIDATOS (Tentativa {attempt+1}). Lat:{latency:.2f}s.")
            else:
                try:
                    candidate = response.candidates[0]
                    finish_reason = getattr(candidate.finish_reason, 'name', "UNKNOWN")

                    # Extrair texto
                    if hasattr(candidate, 'content') and hasattr(candidate.content, 'parts'):
                        text_parts = [part.text for part in candidate.content.parts if hasattr(part, 'text')]
                        result_text = "".join(text_parts).strip() if text_parts else None
                    elif hasattr(response, 'text') and response.text: # Fallback
                        result_text = response.text.strip()

                    # Verificar resultado
                    if finish_reason == "STOP" and result_text is not None:
                        logger.debug(f"{log_prefix} API retornou texto com sucesso. Finish: {finish_reason}. Lat: {latency:.2f}s")
                        return (result_text, latency, prompt_tokens, output_tokens, total_tokens) # SUCESSO
                    else:
                        safety_ratings_str = "N/A"
                        if candidate.safety_ratings:
                             safety_ratings_str = '; '.join([f"{r.category.name}: {r.probability.name}" for r in candidate.safety_ratings])
                        logger.warning(f"{log_prefix} API não OK ou texto vazio (Tentativa {attempt+1}). Finish:{finish_reason}. Safety:{safety_ratings_str}. Lat:{latency:.2f}s.")
                        if finish_reason == "MAX_TOKENS": logger.warning(f"{log_prefix} -> Resposta truncada por MAX_OUTPUT_TOKENS.")
                        # Continua para retry se não for bloqueio explícito

                except Exception as e_details:
                    logger.error(f"{log_prefix} Erro ao extrair resposta API (Tentativa {attempt+1}): {e_details}. Lat:{latency:.2f}s.")
                    logger.debug(traceback.format_exc()) # Debug traceback

            # --- Fim Processamento Resposta ---

        except Exception as e:
            latency = time.time() - start_time
            logger.warning(f"{log_prefix} Erro durante chamada API ({model.model_name}) (Tentativa {attempt + 1}): {e}. Lat:{latency:.2f}s")
            logger.debug(traceback.format_exc()) # Debug traceback

            # Lógica de Espera Exponencial com Jitter
            if attempt < max_retries - 1:
                wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                if "RESOURCE_EXHAUSTED" in str(e) or "429" in str(e):
                    base_wait_time = max(15, base_wait_time) # Aumenta espera base para erros de cota
                    wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
                    logger.warning(f"{log_prefix} Erro de cota detectado. Aumentando espera base para {base_wait_time}s.")
                logger.info(f"{log_prefix} Tentando API novamente em {wait_time:.2f}s...")
                time.sleep(wait_time)
            else: # Última tentativa falhou
                logger.error(f"{log_prefix} Falha final na API após {max_retries} tentativas (erro na chamada): '{log_chunk_preview}'")
                return default_return

        # Se chegou aqui sem retornar sucesso, tenta de novo (se não for a última tentativa)
        if attempt < max_retries - 1 and result_text is None : # Verifica result_text p/ evitar retry desnecessário se houve erro após obter texto
             wait_time = base_wait_time * (2 ** attempt) + (os.urandom(1)[0] / 255.0 * base_wait_time)
             logger.info(f"{log_prefix} Resposta não OK, tentando API novamente em {wait_time:.2f}s...")
             time.sleep(wait_time)
        elif result_text is not None: # Caso estranho: obteve texto mas não retornou sucesso (ex: finish_reason!=STOP)
            logger.warning(f"{log_prefix} Obtido texto mas finish_reason não era STOP. Retornando texto parcial/possivelmente incompleto.")
            return (result_text, latency, prompt_tokens, output_tokens, total_tokens)


    logger.error(f"{log_prefix} Loop de tentativas da API concluído sem sucesso explícito: '{log_chunk_preview}'")
    return default_return


# --- Funções dos Passos de Processamento ---

def run_correction_pass(model: genai.GenerativeModel, input_txt_path: str, author_name: str, book_name: str) -> Tuple[bool, Optional[str], StatsTuple]:
    """
    Executa o Passo 1: Corrige texto via API.
    Retorna: (success_bool, corrected_text_str | None, accumulated_stats)
    """
    log_prefix = f"[{author_name}/{book_name}]"; logger.info(f"{log_prefix} --- Iniciando Passo 1: Correção ---")
    accumulated_stats: List[float | int] = [0.0, 0, 0, 0] # latency, p, o, t
    default_return: Tuple[bool, Optional[str], StatsTuple] = (False, None, (0.0, 0, 0, 0))

    try:
        with open(input_txt_path, "r", encoding="utf-8") as f: texto_original = f.read()
        logger.info(f"{log_prefix} Arquivo '{os.path.basename(input_txt_path)}' lido ({len(texto_original)} chars).")
    except Exception as e: logger.error(f"{log_prefix} FATAL ao ler entrada '{input_txt_path}': {e}"); return default_return

    text_chunks = create_chunks(texto_original, MAX_CHUNK_TOKENS, author_name, book_name)
    if not text_chunks: logger.error(f"{log_prefix} Nenhum chunk gerado. Abortando Passo 1."); return default_return

    logger.info(f"{log_prefix} Iniciando chamadas API para CORREÇÃO de {len(text_chunks)} chunks...")
    corrected_text_list = []; processed_chunks_count = 0; failed_chunks_count = 0

    # --- Prompt de Correção (Exemplo, adaptar conforme necessidade) ---
    ocr_errors_examples = """
* **Troca de letras similares:** 'rn' vs 'm', 'c' vs 'e', 't' vs 'f', 'l' vs 'i', 'I' vs 'l', 'O' vs '0', 'S' vs '5', 'B' vs '8'.
* **Hífens:** Indevidos ou ausentes.
* **Espaços:** Ausentes ou extras.
* **Pontuação/Acentuação:** Incorreta ou ausente.
* **Letras:** Duplicadas ou ausentes.
* **Maiúsculas/Minúsculas:** Inconsistentes.
* **Ruído/Quebras:** Caracteres estranhos, quebras de linha indevidas.
    """
    base_correction_prompt = f"""Você é um editor literário proficiente em português do Brasil. Sua tarefa é CORRIGIR e REFORMATAR o fragmento de texto a seguir, que pertence a um livro do autor {author_name} chamado "{book_name}". O texto pode conter erros de OCR, digitação e gramática.
SIGA RIGOROSAMENTE:
1.  **Correção Profunda:** Corrija TODOS os erros gramaticais, ortográficos, de pontuação, acentuação e concordância (Português do Brasil). FOQUE em erros comuns de OCR como: {ocr_errors_examples}
2.  **Estilo e Fidelidade:** Mantenha o estilo e tom do autor. NÃO altere o significado. Mantenha a estrutura de parágrafos (separados por \\n\\n).
3.  **Sem Adições/Remoções:** NÃO omita conteúdo. NÃO adicione introduções, resumos ou suas próprias ideias. SEJA ESTRITAMENTE FIEL AO CONTEÚDO.
4.  **Marcadores:** Se houver marcadores de capítulo/quebra (ex: 'Capítulo X', '***'), MANTENHA-OS no início do parágrafo onde estão. Se houver '{PAGE_BREAK_MARKER}', MANTENHA-O em sua própria linha.
5.  **Saída:** Retorne APENAS o texto corrigido e formatado. Use parágrafos separados por \\n\\n. NÃO use Markdown ou qualquer outra formatação especial. NÃO inclua comentários sobre o que fez.

Texto do fragmento para processar:
\"\"\"
{{chunk}}
\"\"\"
Texto Corrigido:"""
    # --- Fim Prompt ---

    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P1: Corrigindo", unit="chunk", leave=False)
    for i, chunk in progress_bar:
        # Ajusta contexto para primeiro chunk se necessário
        context_start = "Você está formatando o início do livro." if i == 0 else "Você está continuando a formatação de um texto de livro existente."
        # Monta prompt final (substitui {chunk} e {context_start} se o prompt base os usar)
        # Nota: O prompt base acima já inclui autor/livro, então só precisamos do chunk.
        current_prompt = base_correction_prompt.format(chunk=chunk) # Adaptar se prompt usar context_start

        corrected_chunk, latency, p_tokens, o_tokens, t_tokens = _call_gemini_api(
            model, current_prompt, chunk, author_name, book_name
        )
        # Acumula stats (usando OR 0 para evitar TypeError se API retornar None em stats)
        accumulated_stats[0] += (latency or 0.0)
        accumulated_stats[1] += (p_tokens or 0)
        accumulated_stats[2] += (o_tokens or 0)
        accumulated_stats[3] += (t_tokens or 0)

        if corrected_chunk is not None:
            corrected_text_list.append(corrected_chunk)
            processed_chunks_count += 1
        else:
            logger.warning(f"{log_prefix} Chunk {i+1} falhou na CORREÇÃO (API). Usando fallback.")
            fallback_text = f"{AI_FAILURE_MARKER}\n\n{chunk}"
            corrected_text_list.append(fallback_text)
            failed_chunks_count += 1

    full_corrected_text = "\n\n".join(corrected_text_list)
    final_stats: StatsTuple = (accumulated_stats[0], accumulated_stats[1], accumulated_stats[2], accumulated_stats[3])

    logger.info(f"{log_prefix} --- Passo 1 concluído. Chunks OK: {processed_chunks_count}, Falhas: {failed_chunks_count} ---")
    logger.info(f"{log_prefix} Stats API P1: Lat:{final_stats[0]:.2f}s, Toks:{final_stats[3]}(P:{final_stats[1]},O:{final_stats[2]})")

    success = failed_chunks_count == 0 # Considera sucesso apenas se nenhum chunk falhou na API
    return (success, full_corrected_text if success else None, final_stats) # Retorna None para texto se houve falha


def run_footnote_id_pass(model: genai.GenerativeModel, corrected_text_content: str, author_name: str, book_name: str) -> Tuple[bool, Optional[str], StatsTuple]:
    """
    Executa o Passo 2: Identifica notas no texto já corrigido.
    Retorna: (success_bool, marked_text_str | None, accumulated_stats)
    """
    log_prefix = f"[{author_name}/{book_name}]"; logger.info(f"{log_prefix} --- Iniciando Passo 2: ID Notas ---")
    accumulated_stats: List[float | int] = [0.0, 0, 0, 0] # latency, p, o, t
    default_return: Tuple[bool, Optional[str], StatsTuple] = (False, None, (0.0, 0, 0, 0))

    if not corrected_text_content:
        logger.error(f"{log_prefix} Input NULO para Passo 2."); return default_return

    text_chunks = create_chunks(corrected_text_content, MAX_CHUNK_TOKENS, author_name, book_name)
    if not text_chunks: logger.error(f"{log_prefix} Nenhum chunk gerado (P2). Abortando Passo 2."); return default_return

    logger.info(f"{log_prefix} Iniciando API ID Notas em {len(text_chunks)} chunks...")
    marked_list = []; ok_chunks = 0; fail_chunks = 0

    # --- Prompt de Identificação de Notas (Exemplo, adaptar conforme necessidade) ---
    footnote_id_prompt = f"""Você é um assistente de edição para o livro "{book_name}" do autor {author_name}. Sua tarefa é analisar o fragmento de texto A SEGUIR, que JÁ FOI CORRIGIDO, e APENAS inserir marcadores para potenciais notas de rodapé.
REGRAS IMPORTANTES:
1.  **NÃO ALTERE O TEXTO CORRIGIDO.** Apenas insira os marcadores.
2.  **MARCADORES:** Insira `[NOTA_TIPO:Referência][CONTEUDO_NOTA:Explicação]` APENAS para:
    * **Idioma Estrangeiro (raro):** `[NOTA_IDIOMA:palavra_original][CONTEUDO_NOTA:Tradução ou breve explicação]`
    * **Citações/Referências:** `[NOTA_CITACAO:Texto citado ou ref][CONTEUDO_NOTA:Referência completa ou fonte]` (Ex: (Autor, Ano))
    * **Nomes Próprios (contexto essencial):** `[NOTA_NOME:Nome Mencionado][CONTEUDO_NOTA:Breve ID (datas, relevância)]` (Use com MODERAÇÃO)
    * **Termos Técnicos/Jargão (essencial):** `[NOTA_TERMO:Termo Técnico][CONTEUDO_NOTA:Definição concisa]` (Use com MUITA MODERAÇÃO)
3.  **FORMATO:** Use EXATAMENTE o formato. NÃO adicione espaços entre os marcadores.
4.  **CRITÉRIO:** Seja CONSERVADOR. Notas apenas se úteis e provavelmente desconhecidas. É MELHOR NÃO ADICIONAR do que adicionar demais.
5.  **NÃO INVENTE CONTEÚDO:** Use traduções óbvias, refs diretas ou contexto mínimo. Se não souber, NÃO insira a nota.
6.  **SAÍDA:** Retorne APENAS o texto original com os marcadores inseridos. Mantenha parágrafos (\\n\\n). NÃO adicione comentários.

Texto JÁ CORRIGIDO para analisar:
\"\"\"
{{chunk}}
\"\"\"
Texto com Marcadores (se houver):"""
    # --- Fim Prompt ---

    progress_bar = tqdm(enumerate(text_chunks), total=len(text_chunks), desc=f"{log_prefix} P2: Notas", unit="chunk", leave=False)
    for i, chunk in progress_bar:
        current_prompt = footnote_id_prompt.format(chunk=chunk)
        marked_chunk, latency, p_tokens, o_tokens, t_tokens = _call_gemini_api(
            model, current_prompt, chunk, author_name, book_name
        )
        # Acumula stats
        accumulated_stats[0] += (latency or 0.0)
        accumulated_stats[1] += (p_tokens or 0)
        accumulated_stats[2] += (o_tokens or 0)
        accumulated_stats[3] += (t_tokens or 0)

        if marked_chunk is not None:
            marked_list.append(marked_chunk)
            ok_chunks += 1
            if "[NOTA_" in marked_chunk: logger.debug(f"{log_prefix} Chunk {i+1}: Marcadores de nota encontrados/inseridos.")
            else: logger.debug(f"{log_prefix} Chunk {i+1}: Nenhum marcador de nota inserido.")
        else:
            logger.warning(f"{log_prefix} Chunk {i+1} falhou na ID Notas (API). Usando texto original sem marcadores.")
            marked_list.append(chunk) # Fallback usa o chunk original (corrigido)
            fail_chunks += 1

    full_marked_text = "\n\n".join(marked_list)
    final_stats: StatsTuple = (accumulated_stats[0], accumulated_stats[1], accumulated_stats[2], accumulated_stats[3])

    logger.info(f"{log_prefix} --- Passo 2 concluído. Chunks OK: {ok_chunks}, Falhas: {fail_chunks} ---")
    logger.info(f"{log_prefix} Stats API P2: Lat:{final_stats[0]:.2f}s, Toks:{final_stats[3]}(P:{final_stats[1]},O:{final_stats[2]})")

    success = fail_chunks == 0 # Considera sucesso mesmo que não encontre notas, mas API não pode falhar
    return (success, full_marked_text if success else None, final_stats) # Retorna None para texto se API falhou


def apply_formatting_to_doc(doc: Document, text_content: str,
                            normal_style_name: str, chapter_style_name: str, chapter_patterns: List[str],
                            author_name: str, book_name: str,
                            note_markers_map: Optional[Dict[str, str]] = None) -> Tuple[Document, int]:
    """
    Adiciona text_content formatado ao objeto doc, substituindo placeholders por [Num] (SEM negrito).
    Retorna o objeto doc modificado e o número de placeholders substituídos.
    """
    log_prefix = f"[{author_name}/{book_name}][apply_formatting]"
    replacements_made = 0
    if note_markers_map is None: note_markers_map = {}

    # Get styles safely
    normal_style: Optional[Any] = None
    chapter_style: Optional[Any] = None
    try:
        if normal_style_name in doc.styles: normal_style = doc.styles[normal_style_name]
        else: logger.warning(f"{log_prefix} Estilo Normal '{normal_style_name}' não encontrado no doc.")
        if chapter_style_name in doc.styles: chapter_style = doc.styles[chapter_style_name]
        else: logger.warning(f"{log_prefix} Estilo Capítulo '{chapter_style_name}' não encontrado no doc.")
    except Exception as e_style: logger.error(f"{log_prefix} Erro ao acessar estilos: {e_style}.")

    chapter_regex = re.compile('|'.join(chapter_patterns), re.IGNORECASE)

    parts = text_content.split(PAGE_BREAK_MARKER)
    # Check if document already has content before adding initial page break
    # has_initial_content = any(p.text.strip() or p._element.xpath('.//w:drawing') for p in doc.paragraphs) # Check text or images

    for part_index, part in enumerate(parts):
        part_clean = part.strip()

        # Add page break logic (Needs refinement if appending to complex templates)
        if part_index > 0:
             last_p = doc.paragraphs[-1] if doc.paragraphs else None
             # Add break if last paragraph exists and is not already effectively a break
             if last_p and (last_p.text.strip() or not any(run.text and '\f' in run.text for run in last_p.runs)):
                  logger.debug(f"{log_prefix} Adicionando quebra de página antes da parte {part_index+1}")
                  try:
                      doc.add_page_break()
                  except Exception as e_pb:
                       logger.warning(f"{log_prefix} Falha ao adicionar quebra de página: {e_pb}")

        if not part_clean: continue

        paragraphs_in_part = part_clean.split("\n\n")
        for para_idx, paragraph_text in enumerate(paragraphs_in_part):
            paragraph_text_clean = paragraph_text.strip()
            if not paragraph_text_clean:
                 # Add empty paragraph only if the previous one wasn't empty
                 if doc.paragraphs and doc.paragraphs[-1].text.strip():
                     p_empty = doc.add_paragraph()
                     if normal_style:
                         try: p_empty.style = normal_style
                         except Exception as e_style_empty: logger.debug(f"Err style empty para: {e_style_empty}")
                 continue

            # Determine paragraph type
            is_ai_failure = paragraph_text_clean.startswith(AI_FAILURE_MARKER)
            is_formatting_error = paragraph_text_clean.startswith(FORMATTING_ERROR_MARKER)
            is_chapter = not is_ai_failure and not is_formatting_error and chapter_regex.match(paragraph_text_clean) is not None

            p = doc.add_paragraph()
            # Apply base style/alignment
            try:
                if is_chapter and chapter_style: p.style = chapter_style
                elif is_chapter: p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Fallback alignment
                elif normal_style: p.style = normal_style
                else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Fallback alignment
            except Exception as e_style_para: logger.warning(f"{log_prefix} Erro aplicar estilo base parágrafo {para_idx+1}: {e_style_para}")


            # --- Process Text with Placeholders ---
            last_end = 0
            for match in PLACEHOLDER_REGEX.finditer(paragraph_text_clean):
                start, end = match.span()
                placeholder_found = match.group(1) # __AI_NOTE_PLACEHOLDER_N__
                note_number_str = match.group(2)   # N

                # Add text before placeholder
                if start > last_end:
                    run_text = p.add_run(paragraph_text_clean[last_end:start])
                    # Apply error formatting if needed
                    if is_ai_failure or is_formatting_error:
                         run_text.italic = True; run_text.font.color.rgb = RGBColor(0xFF, 0, 0)

                # Add the note marker [N] (without bolding here)
                if placeholder_found in note_markers_map:
                     #replacement_num = note_markers_map[placeholder_found] # Should just be the number string "N"
                     p.add_run(f"[{note_number_str}]") # Just add [N] plainly
                     replacements_made += 1
                else:
                    logger.warning(f"{log_prefix} Placeholder '{placeholder_found}' não encontrado no mapa! Inserindo placeholder original.")
                    p.add_run(placeholder_found) # Add the raw placeholder if not found

                last_end = end

            # Add remaining text after the last placeholder (or the whole text if no placeholders)
            if last_end < len(paragraph_text_clean):
                run_text = p.add_run(paragraph_text_clean[last_end:])
                # Apply error formatting if needed
                if is_ai_failure or is_formatting_error:
                    run_text.italic = True; run_text.font.color.rgb = RGBColor(0xFF, 0, 0)

            # Apply special chapter formatting (e.g., font) after content is added
            if is_chapter and not chapter_style: # Apply manual font only if style wasn't applied
                try:
                    # Example: Apply specific font/size to the first run (usually the whole title)
                    if p.runs:
                         first_run = p.runs[0]
                         # first_run.font.name = 'Times New Roman' # Keep default or set specific
                         first_run.font.size = Pt(16) # Example size
                         first_run.bold = True
                except Exception as e_font:
                    logger.warning(f"{log_prefix} Erro ao aplicar formatação manual capítulo: {e_font}")

    #logger.debug(f"{log_prefix} Formatação aplicada. {replacements_made} placeholders inseridos (sem formatação).")
    return doc, replacements_made


def extract_notes_and_prepare_text(marked_text_string: str, author_name: str, book_name: str) -> Tuple[Optional[str], Dict[str, str], List[str]]:
    """
    Extrai notas do texto marcado, cria texto com placeholders e mapa de substituição.
    Retorna: (text_with_placeholders | None, note_markers_map, notes_content_list)
    """
    log_prefix = f"[{author_name}/{book_name}][extract_notes]"
    notes_content_list: List[str] = []
    note_markers_map: Dict[str, str] = {} # placeholder -> note_number_str
    text_with_placeholders = ""
    note_number_counter = 1
    last_match_end = 0

    try:
        logger.info(f"{log_prefix} Extraindo notas e criando placeholders...")
        # Iterate through matches of the *full* marker pattern
        for match in FOOTNOTE_MARKER_PATTERN.finditer(marked_text_string):
            full_marker = match.group(1) # The whole [NOTA_...][CONTEUDO_NOTA:...]
            # Use the *content* extraction pattern on the full marker found
            content_match = FOOTNOTE_CONTENT_EXTRACT_PATTERN.search(full_marker)
            note_content = content_match.group(1).strip() if content_match else ""

            # Append text segment before the current match
            text_with_placeholders += marked_text_string[last_match_end:match.start()]

            if note_content:
                notes_content_list.append(note_content)
                placeholder = f"__AI_NOTE_PLACEHOLDER_{note_number_counter}__"
                note_markers_map[placeholder] = str(note_number_counter) # Map placeholder to number string
                text_with_placeholders += placeholder
                note_number_counter += 1
            else:
                logger.warning(f"{log_prefix} Nota com conteúdo vazio removida: '{full_marker[:60]}...'")
                # If content is empty, we remove the marker pair entirely by not adding a placeholder

            last_match_end = match.end()

        # Append any remaining text after the last match
        text_with_placeholders += marked_text_string[last_match_end:]
        logger.info(f"{log_prefix} Extraídas {len(notes_content_list)} notas válidas.")
        return text_with_placeholders, note_markers_map, notes_content_list

    except Exception as e_extract:
        logger.error(f"{log_prefix} ERRO durante extração de notas: {e_extract}")
        logger.error(traceback.format_exc())
        return None, {}, []


def bold_footnote_markers_in_doc(doc: Document, author_name: str, book_name: str) -> int:
    """Encontra marcadores [N] no documento e aplica negrito."""
    log_prefix = f"[{author_name}/{book_name}][bold_markers]"
    count = 0
    marker_regex = re.compile(r'(\[\d+\])') # Regex to find [N]
    try:
        logger.info(f"{log_prefix} Aplicando negrito aos marcadores de nota [N]...")
        for para in doc.paragraphs:
            if '[' not in para.text: continue # Quick check

            inline_runs = list(para.runs) # Work on a copy if modifying list while iterating
            para.clear() # Clear existing runs

            current_text = "".join(run.text for run in inline_runs)
            last_end = 0
            for match in marker_regex.finditer(current_text):
                 start, end = match.span()
                 marker_text = match.group(1)

                 # Add text before marker (preserving original run formatting if possible - complex, skipping for now)
                 if start > last_end:
                      para.add_run(current_text[last_end:start])

                 # Add bold marker
                 run_marker = para.add_run(marker_text)
                 run_marker.bold = True
                 count += 1

                 last_end = end

            # Add remaining text
            if last_end < len(current_text):
                 para.add_run(current_text[last_end:])

    except Exception as e_bold:
        logger.error(f"{log_prefix} Erro ao aplicar negrito aos marcadores: {e_bold}")
        logger.error(traceback.format_exc())
    logger.info(f"{log_prefix} Negrito aplicado a {count} marcadores.")
    return count


def append_notes_section(doc: Document, notes_content_list: List[str], normal_style_name: str, author_name: str, book_name: str):
    """Adiciona a seção de notas formatada ao final do documento."""
    log_prefix = f"[{author_name}/{book_name}][append_notes]"
    if not notes_content_list: return # Do nothing if no notes

    logger.info(f"{log_prefix} Adicionando seção de {len(notes_content_list)} notas ao DOCX...")
    try:
        # Ensure some space before the notes section
        if doc.paragraphs and doc.paragraphs[-1].text.strip():
            doc.add_paragraph() # Add an empty paragraph for spacing

        # Notes Title
        try:
            # Use Heading 1 if available, otherwise simple bold paragraph
            heading_style_found = False
            if 'Heading 1' in doc.styles:
                 try:
                     doc.add_heading("Notas", level=1); heading_style_found = True
                 except Exception as e_h1: logger.warning(f"{log_prefix} Erro add_heading level 1: {e_h1}")
            if not heading_style_found:
                 p_heading = doc.add_paragraph("Notas")
                 if p_heading.runs: p_heading.runs[0].bold = True
                 logger.warning(f"{log_prefix} Estilo 'Heading 1' não aplicável/encontrado, usando parágrafo negrito.")
        except Exception as e_head:
            logger.error(f"{log_prefix} Erro ao adicionar título 'Notas': {e_head}")

        # Get Normal style for note paragraphs
        normal_style_notes = None
        try:
            if normal_style_name in doc.styles: normal_style_notes = doc.styles[normal_style_name]
        except Exception as e_style: logger.debug(f"{log_prefix} Erro buscar estilo normal p/ notas: {e_style}")

        # Add each note
        for i, note_text in enumerate(notes_content_list):
            p_note = doc.add_paragraph()
            try: # Apply style first if available
                if normal_style_notes: p_note.style = normal_style_notes
            except Exception as e_style_note: logger.debug(f"Err style nota {i+1}: {e_style_note}")

            run_num = p_note.add_run(f"{i+1}. ")
            run_num.bold = True
            p_note.add_run(note_text)

        logger.info(f"{log_prefix} Seção de notas adicionada com sucesso.")

    except Exception as e_append:
        logger.error(f"{log_prefix} ERRO CRÍTICO ao adicionar seção de notas: {e_append}")
        logger.error(traceback.format_exc())


def generate_final_docx_outputs(marked_text_string: str, template_path: str,
                                output_path_clean: str, output_path_with_notes: str,
                                author_name: str, book_name: str) -> bool:
    """
    Passo 3 Consolidado: Gera ambos os DOCX finais (limpo e com notas).
    Retorna: bool (sucesso geral)
    """
    log_prefix = f"[{author_name}/{book_name}][generate_docx]"
    logger.info(f"{log_prefix} --- Iniciando Passo 3: Geração DOCX Finais ---")

    # 1. Extrair Notas e Preparar Texto com Placeholders
    text_with_placeholders, note_markers_map, notes_content_list = extract_notes_and_prepare_text(
        marked_text_string, author_name, book_name
    )
    if text_with_placeholders is None: # Erro na extração
        return False

    # Se não houver notas, apenas copia o resultado da correção original (que não temos mais aqui!)
    # Precisamos gerar o doc limpo mesmo sem notas.
    # if not notes_content_list:
    #     logger.info(f"{log_prefix} Nenhuma nota encontrada. Gerando apenas DOCX limpo.")
    #     # TODO: Handle this case - needs original corrected text string from step 1 ideally
    #     # For now, proceed to generate the clean doc from the text_with_placeholders (which is just the corrected text)
    #     pass # Let the process continue

    # 2. Carregar Template e Aplicar Formatação Inicial (para DOCX Limpo)
    doc_object: Optional[Document] = None
    try:
        if not os.path.exists(template_path):
             logger.error(f"{log_prefix} FATAL: Template '{template_path}' não encontrado."); return False
        doc_object = Document(template_path)
        logger.info(f"{log_prefix} Template '{os.path.basename(template_path)}' carregado.")

        # Aplica texto COM placeholders, que serão formatados como [N] sem negrito
        doc_object, replacements_count = apply_formatting_to_doc(
            doc_object, text_with_placeholders,
            NORMAL_STYLE_NAME, CHAPTER_STYLE_NAME, CHAPTER_PATTERNS,
            author_name, book_name,
            note_markers_map # Passa o mapa placeholder -> number_str
        )
        logger.info(f"{log_prefix} Texto principal formatado no DOCX (com {replacements_count} marcadores [N] simples).")
        # Verifica se o número de substituições bate com o esperado
        if replacements_count != len(note_markers_map):
             logger.warning(f"{log_prefix} Discrepância: {len(note_markers_map)} notas mapeadas, {replacements_count} marcadores [N] inseridos.")

    except Exception as e_load_format:
        logger.error(f"{log_prefix} ERRO ao carregar template ou aplicar formatação inicial: {e_load_format}")
        logger.error(traceback.format_exc())
        return False

    # 3. Salvar DOCX Limpo (para Tradutor)
    try:
        output_dir = os.path.dirname(output_path_clean)
        if not os.path.exists(output_dir): os.makedirs(output_dir, exist_ok=True) # Garante diretório

        # Backup
        if os.path.exists(output_path_clean):
             backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
             backup_clean_path = os.path.join(output_dir, f"backup_{os.path.splitext(os.path.basename(output_path_clean))[0]}_{backup_timestamp}.docx")
             try: shutil.copy2(output_path_clean, backup_clean_path); logger.info(f"{log_prefix} Backup DOCX limpo criado.")
             except Exception as e_bkp: logger.warning(f"{log_prefix} Falha backup DOCX limpo: {e_bkp}")

        logger.info(f"{log_prefix} Salvando DOCX limpo (p/ tradutor): {os.path.basename(output_path_clean)}")
        doc_object.save(output_path_clean)
    except Exception as e_save_clean:
        logger.error(f"{log_prefix} ERRO CRÍTICO ao salvar DOCX limpo: {e_save_clean}")
        logger.error(traceback.format_exc())
        return False # Falha crítica

    # --- Modificações para DOCX Final com Notas ---

    # Se não havia notas, o trabalho acabou, o arquivo limpo já serve como "com notas"
    if not notes_content_list:
        logger.info(f"{log_prefix} Nenhuma nota para processar. Copiando DOCX limpo para final 'com notas'.")
        try:
            if os.path.exists(output_path_with_notes): os.remove(output_path_with_notes)
            shutil.copy2(output_path_clean, output_path_with_notes)
            logger.info(f"{log_prefix} Cópia final salva: {os.path.basename(output_path_with_notes)}")
            return True
        except Exception as e_copy:
            logger.error(f"{log_prefix} ERRO ao copiar arquivo final (sem notas): {e_copy}")
            return False


    # 4. Aplicar Negrito aos Marcadores [N] no *mesmo* objeto doc
    bold_count = bold_footnote_markers_in_doc(doc_object, author_name, book_name)
    if bold_count != replacements_count:
         logger.warning(f"{log_prefix} Discrepância: {replacements_count} marcadores inseridos, {bold_count} formatados em negrito.")

    # 5. Adicionar Seção de Notas ao Final do *mesmo* objeto doc
    append_notes_section(doc_object, notes_content_list, NORMAL_STYLE_NAME, author_name, book_name)

    # 6. Salvar DOCX Final com Notas
    try:
        output_dir_notes = os.path.dirname(output_path_with_notes)
        if not os.path.exists(output_dir_notes): os.makedirs(output_dir_notes, exist_ok=True) # Garante diretório

        # Backup
        if os.path.exists(output_path_with_notes):
             backup_timestamp = time.strftime("%Y%m%d_%H%M%S")
             backup_notes_path = os.path.join(output_dir_notes, f"backup_{os.path.splitext(os.path.basename(output_path_with_notes))[0]}_{backup_timestamp}.docx")
             try: shutil.copy2(output_path_with_notes, backup_notes_path); logger.info(f"{log_prefix} Backup DOCX com notas criado.")
             except Exception as e_bkp: logger.warning(f"{log_prefix} Falha backup DOCX com notas: {e_bkp}")

        logger.info(f"{log_prefix} Salvando DOCX final com notas: {os.path.basename(output_path_with_notes)}")
        doc_object.save(output_path_with_notes)
    except Exception as e_save_notes:
        logger.error(f"{log_prefix} ERRO CRÍTICO ao salvar DOCX com notas: {e_save_notes}")
        logger.error(traceback.format_exc())
        return False

    logger.info(f"{log_prefix} --- Passo 3 concluído (Geração DOCX Finais). ---")
    return True


# --- Funções para Gerenciar Logs de Processados ---
def load_processed_files(filepath: str) -> Set[str]:
    """Carrega identificadores de arquivos do log de correções."""
    processed: Set[str] = set()
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            processed.update(line.strip() for line in f if line.strip())
        logger.info(f"Carregados {len(processed)} registros CORREÇÕES de '{os.path.basename(filepath)}'.")
    except FileNotFoundError: logger.info(f"Log correções '{os.path.basename(filepath)}' não encontrado.")
    except Exception as e: logger.error(f"Erro carregar log correções '{filepath}': {e}")
    return processed

def log_processed_file(filepath: str, file_identifier: str):
    """Registra um arquivo no log de correções."""
    try:
        with open(filepath, 'a', encoding='utf-8') as f: f.write(f"{file_identifier}\n")
    except Exception as e: logger.error(f"Erro registrar '{file_identifier}' log correções '{filepath}': {e}")

def load_translated_files(filepath: str) -> Set[str]:
    """Carrega identificadores de arquivos do log de traduções."""
    processed: Set[str] = set()
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            processed.update(line.strip() for line in f if line.strip())
        logger.info(f"Carregados {len(processed)} registros TRADUÇÕES de '{os.path.basename(filepath)}'.")
    except FileNotFoundError: logger.info(f"Log traduções '{os.path.basename(filepath)}' não encontrado.")
    except Exception as e: logger.error(f"Erro carregar log traduções '{filepath}': {e}")
    return processed

def log_translated_file(filepath: str, file_identifier: str):
    """Registra um arquivo no log de traduções."""
    try:
        with open(filepath, 'a', encoding='utf-8') as f: f.write(f"{file_identifier}\n")
        logger.debug(f"Registrado '{file_identifier}' como TRADUZIDO em '{os.path.basename(filepath)}'.")
    except Exception as e: logger.error(f"Erro registrar '{file_identifier}' log traduções '{filepath}': {e}")

# --- FUNÇÃO DE ENVIO DE E-MAIL ---
def send_completion_email(sender_email: str, sender_password: str, recipient_email: str, smtp_server: str, smtp_port: int,
                          processed_correction: int, skipped_correction: int, failed_correction: int,
                          processed_translation: int, skipped_translation: int, failed_translation: int,
                          total_duration_seconds: float,
                          main_log_path: str, processed_log_path: str, translated_log_path: str,
                          total_correction_latency_secs: float, total_correction_tokens: int,
                          total_footnote_latency_secs: float, total_footnote_tokens: int,
                          avg_correction_time_secs: float, total_correction_time_secs: float,
                          avg_translation_time_secs: float, total_translation_time_secs: float,
                          processed_correction_books: List[str], skipped_correction_books: List[str], failed_correction_books: List[str],
                          processed_translation_books: List[str], skipped_translation_books: List[str], failed_translation_books: List[str]
                          ):
    """Envia e-mail de resumo (lógica inalterada)."""
    global email_configured
    if not email_configured: logger.warning("Envio e-mail desativado."); return
    logger.info(f"Preparando e-mail para {recipient_email}...")
    subject = "Script Processador Livros v11 (Correção+Tradução) - Concluído" # Version subject
    body = f"""Olá,\n\nO script v11 (Refatorado) concluiu a execução.\n\nResumo Geral:\n{'-'*50}\n"""
    body += f"- Tempo Total: {total_duration_seconds:.2f}s ({total_duration_seconds/60:.2f}m)\n"
    body += f"\nResumo Correção/Notas DOCX:\n{'-'*50}\n" # Updated title
    body += f"- Corrigidos OK: {processed_correction}\n- Pulados: {skipped_correction}\n- Falhas: {failed_correction}\n"
    if processed_correction > 0:
        body += f"- Tempo Total Correção/Notas (OK): {total_correction_time_secs:.2f}s ({total_correction_time_secs/60:.2f}m)\n"
        body += f"- Tempo Médio Correção/Notas: {avg_correction_time_secs:.2f}s\n"
        body += f"- API P1 (Correção): Lat {total_correction_latency_secs:.2f}s / Tokens {total_correction_tokens}\n"
        body += f"- API P2 (Notas ID): Lat {total_footnote_latency_secs:.2f}s / Tokens {total_footnote_tokens}\n"
    body += f"\nResumo Tradução:\n{'-'*50}\n"
    body += f"- Traduzidos OK: {processed_translation}\n- Pulados: {skipped_translation}\n- Falhas: {failed_translation}\n"
    if processed_translation > 0:
        body += f"- Tempo Total Tradução (OK): {total_translation_time_secs:.2f}s ({total_translation_time_secs/60:.2f}m)\n"
        body += f"- Tempo Médio Tradução: {avg_translation_time_secs:.2f}s\n"
    body += f"\nDetalhes por Livro:\n{'-'*50}\n"
    max_list_items = 20
    def format_book_list(label: str, book_list: List[str]) -> str:
        if not book_list: return ""
        list_str = f"\n{label} ({len(book_list)}):\n - "
        list_str += "\n - ".join(book_list[:max_list_items])
        if len(book_list) > max_list_items: list_str += f"\n - ... (e mais {len(book_list) - max_list_items} - ver logs)"
        list_str += "\n"
        return list_str
    body += format_book_list("Correção/Notas OK", processed_correction_books) # Updated label
    body += format_book_list("Pulados Correção/Notas", skipped_correction_books) # Updated label
    body += format_book_list("Falha Correção/Notas", failed_correction_books) # Updated label
    body += format_book_list("Traduzidos OK", processed_translation_books)
    body += format_book_list("Pulados Tradução", skipped_translation_books)
    body += format_book_list("Falha Tradução", failed_translation_books)
    body += f"\n{'-'*50}\nLogs:\n- Detalhado: {os.path.abspath(main_log_path)}\n- Correções OK: {os.path.abspath(processed_log_path)}\n- Traduções OK: {os.path.abspath(translated_log_path)}\n\nAtenciosamente,\nScript Processador v11"
    message = EmailMessage(); message['Subject'] = subject; message['From'] = sender_email; message['To'] = recipient_email; message.set_content(body)
    context = ssl.create_default_context()
    server = None
    try:
        logger.info(f"Conectando SMTP: {smtp_server}:{smtp_port}...")
        if smtp_port == 465: server = smtplib.SMTP_SSL(smtp_server, smtp_port, context=context, timeout=30); server.login(sender_email, sender_password)
        else: server = smtplib.SMTP(smtp_server, smtp_port, timeout=30); server.ehlo(); server.starttls(context=context); server.ehlo(); server.login(sender_email, sender_password)
        logger.info("Enviando e-mail resumo..."); server.send_message(message); logger.info(f"✅ E-mail enviado para {recipient_email}.")
    except Exception as e: logger.error(f"ERRO ao enviar e-mail: {e}"); logger.debug(traceback.format_exc())
    finally:
        if server:
            try: server.quit()
            except Exception: pass

# --- FUNÇÃO PRINCIPAL (main - Refatorada v11) ---
def main():
    start_time_main = time.time()
    logger.info("========================================================")
    logger.info(f"Iniciando Processador v11 (Refatorado: DOCX Otimizado) - {time.strftime('%Y-%m-%d %H:%M:%S')}")
    logger.info(f"Diretório Entrada: {BASE_INPUT_TXT_DIR}, Saída DOCX: {BASE_OUTPUT_DOCX_DIR}")
    logger.info(f"Template: {TEMPLATE_DOCX}, Script Tradutor: {PATH_TO_TRANSLATOR_SCRIPT}")
    logger.info(f"Log Correção: {PROCESSED_LOG_FILE}, Log Tradução: {TRANSLATED_LOG_FILE}")
    logger.info("========================================================")

    processed_files_set = load_processed_files(PROCESSED_LOG_FILE)
    translated_files_set = load_translated_files(TRANSLATED_LOG_FILE)

    if not os.path.isdir(BASE_INPUT_TXT_DIR): logger.error(f"FATAL: Diretório '{BASE_INPUT_TXT_DIR}' não encontrado!"); return
    try: author_folders = sorted([f for f in os.listdir(BASE_INPUT_TXT_DIR) if os.path.isdir(os.path.join(BASE_INPUT_TXT_DIR, f))])
    except Exception as e: logger.error(f"FATAL: Erro listar autores: {e}"); return
    if not author_folders: logger.warning(f"Nenhuma pasta de autor encontrada."); return
    logger.info(f"Autores encontrados ({len(author_folders)}): {', '.join(author_folders)}")

    # --- Inicializa Contadores e Acumuladores ---
    stats = {
        "corr_ok": 0, "corr_skip": 0, "corr_fail": 0,
        "trans_ok": 0, "trans_skip": 0, "trans_fail": 0,
        "lat_p1": 0.0, "p_p1": 0, "o_p1": 0, "t_p1": 0,
        "lat_p2": 0.0, "p_p2": 0, "o_p2": 0, "t_p2": 0,
        "corr_times": [], "trans_times": [],
        "lists": {"corr_ok": [], "corr_skip": [], "corr_fail": [],
                  "trans_ok": [], "trans_skip": [], "trans_fail": []}
    }

    # === LOOP PRINCIPAL: AUTOR ===
    for author_name in author_folders:
        author_input_dir = os.path.join(BASE_INPUT_TXT_DIR, author_name)
        logger.info(f"--- Verificando Autor: {author_name} ---")
        try:
            search_pattern = os.path.join(author_input_dir, '**', '*.txt') # Busca recursiva
            input_txt_files_found = sorted(glob.glob(search_pattern, recursive=True))
            # Filtra arquivos TXT que não parecem ser intermediários/log
            input_txt_files = [
                f for f in input_txt_files_found if not (
                    os.path.basename(f).endswith(FINAL_NUMBERED_TXT_BASENAME_FILTER) or
                    os.path.basename(f).endswith(NOTES_TXT_FILE_BASENAME_FILTER) or
                    os.path.basename(f).startswith("backup_")
                )
            ]
        except Exception as e: logger.error(f"[{author_name}] Erro buscar .txt: {e}"); continue
        if not input_txt_files: logger.debug(f"[{author_name}] Nenhum .txt válido encontrado para processar."); continue
        logger.info(f"[{author_name}] Processando {len(input_txt_files)} arquivos .txt válidos.")

        # === LOOP INTERNO: LIVRO ===
        for input_txt_path in input_txt_files:
            # --- Definição de Paths e Identificadores ---
            try:
                relative_path = os.path.relpath(input_txt_path, BASE_INPUT_TXT_DIR)
                file_identifier = relative_path.replace('\\', '/') # Usar como ID único
                log_prefix_book = f"[{file_identifier}]"
                path_parts = file_identifier.split('/')
                author_name_from_path = path_parts[0]
                book_subpath_parts = path_parts[1:-1]
                book_filename = path_parts[-1]
                base_book_name = os.path.splitext(book_filename)[0]

                # Diretório de saída DOCX (considerando subpastas dentro do autor)
                book_subdir_rel = os.path.join(*book_subpath_parts) if book_subpath_parts else ""
                author_output_docx_book_dir = os.path.join(BASE_OUTPUT_DOCX_DIR, author_name_from_path, book_subdir_rel)
                os.makedirs(author_output_docx_book_dir, exist_ok=True) # Garante diretório

                # Caminhos dos arquivos DOCX de saída
                output_path_clean_docx = os.path.join(author_output_docx_book_dir, f"{base_book_name}_{FINAL_DOCX_CLEAN_BASENAME}")
                output_path_notes_docx = os.path.join(author_output_docx_book_dir, f"{base_book_name}_{FINAL_DOCX_WITH_NOTES_BASENAME}")
                translated_docx_path = os.path.join(author_output_docx_book_dir, f"{base_book_name}_{FINAL_DOCX_CLEAN_BASENAME.replace('.docx', TRANSLATED_DOCX_SUFFIX)}")

            except Exception as e_path:
                logger.error(f"Erro fatal processando caminhos para '{input_txt_path}': {e_path}")
                logger.error(traceback.format_exc())
                stats["corr_fail"] += 1; stats["lists"]["corr_fail"].append(f"{file_identifier} (Erro Path)")
                continue # Pula para o próximo livro

            logger.info(f"--------------------------------------------------------")
            logger.info(f"{log_prefix_book} Processando Livro...")

            correction_notes_step_success = False
            translator_input_path: Optional[str] = None # Caminho do arquivo limpo para tradutor

            # --- Etapa 1: Correção + Notas (Verifica Log) ---
            if file_identifier in processed_files_set:
                logger.info(f"{log_prefix_book} CORREÇÃO/NOTAS DOCX já concluída (log). Pulando etapa.")
                stats["corr_skip"] += 1; stats["lists"]["corr_skip"].append(file_identifier)
                correction_notes_step_success = True
                # Assume que o arquivo limpo existe para o tradutor
                translator_input_path = output_path_clean_docx if os.path.exists(output_path_clean_docx) else None
                if translator_input_path is None: logger.warning(f"{log_prefix_book} Log de correção OK, mas arquivo DOCX limpo não encontrado: {output_path_clean_docx}")

            else:
                logger.info(f"{log_prefix_book} Iniciando Correção/Notas DOCX...")
                book_start_time = time.time()
                # --- Executa Passos Refatorados ---
                corrected_text: Optional[str] = None
                marked_text: Optional[str] = None
                book_success = False
                book_stats_p1: StatsTuple = (0.0,0,0,0)
                book_stats_p2: StatsTuple = (0.0,0,0,0)

                try:
                    # Passo 1: Correção -> Retorna (bool, str|None, stats)
                    p1_ok, corrected_text, book_stats_p1 = run_correction_pass(
                        gemini_model, input_txt_path, author_name_from_path, base_book_name
                    )
                    if not p1_ok or corrected_text is None: raise ValueError("Falha no Passo 1 (Correção)")

                    # Passo 2: ID Notas -> Retorna (bool, str|None, stats)
                    p2_ok, marked_text, book_stats_p2 = run_footnote_id_pass(
                        gemini_model, corrected_text, author_name_from_path, base_book_name
                    )
                    # Mesmo que p2_ok seja False (nenhuma nota encontrada), precisamos do marked_text (que será igual ao corrected_text)
                    if marked_text is None: raise ValueError("Falha no Passo 2 (ID Notas) - Texto marcado ausente")

                    # Passo 3: Geração DOCX -> Retorna (bool)
                    p3_ok = generate_final_docx_outputs(
                        marked_text, TEMPLATE_DOCX,
                        output_path_clean_docx, output_path_notes_docx,
                        author_name_from_path, base_book_name
                    )
                    if not p3_ok: raise ValueError("Falha no Passo 3 (Geração DOCX Finais)")

                    book_success = True # Se chegou aqui, tudo OK

                except Exception as e_steps:
                    logger.error(f"{log_prefix_book} Erro durante Passos 1-3: {e_steps}")
                    logger.debug(traceback.format_exc())
                    book_success = False

                book_end_time = time.time()
                book_total_time = book_end_time - book_start_time

                # --- Atualiza Stats de Correção/Notas ---
                if book_success:
                    logger.info(f"✅ {log_prefix_book} Etapa CORREÇÃO/NOTAS DOCX SUCESSO em {book_total_time:.2f} seg.")
                    log_processed_file(PROCESSED_LOG_FILE, file_identifier); processed_files_set.add(file_identifier)
                    stats["corr_ok"] += 1; stats["lists"]["corr_ok"].append(file_identifier); stats["corr_times"].append(book_total_time)
                    # Acumula stats das APIs
                    stats["lat_p1"] += book_stats_p1[0]; stats["p_p1"] += book_stats_p1[1]; stats["o_p1"] += book_stats_p1[2]; stats["t_p1"] += book_stats_p1[3]
                    stats["lat_p2"] += book_stats_p2[0]; stats["p_p2"] += book_stats_p2[1]; stats["o_p2"] += book_stats_p2[2]; stats["t_p2"] += book_stats_p2[3]
                    correction_notes_step_success = True
                    translator_input_path = output_path_clean_docx # Define input para tradutor
                else:
                    logger.warning(f"⚠️ {log_prefix_book} Etapa CORREÇÃO/NOTAS DOCX FALHAS em {book_total_time:.2f} seg.")
                    stats["corr_fail"] += 1; stats["lists"]["corr_fail"].append(file_identifier)
                    correction_notes_step_success = False
                    translator_input_path = None

            # --- Etapa 2: Tradução (Se Correção/Notas OK) ---
            if correction_notes_step_success:
                if translator_input_path is None or not os.path.exists(translator_input_path):
                     logger.warning(f"{log_prefix_book} Input DOCX '{os.path.basename(output_path_clean_docx)}' tradução não encontrado ou inválido. Pulando tradução.");
                     # Considera falha na tradução se o input não existe
                     if translator_input_path is not None: # Só conta como falha se deveria existir mas não existe
                          stats["trans_fail"] += 1; stats["lists"]["trans_fail"].append(f"{file_identifier} (Input DOCX ausente)")
                elif file_identifier in translated_files_set:
                    logger.info(f"{log_prefix_book} TRADUÇÃO já feita (log). Pulando.");
                    stats["trans_skip"] += 1; stats["lists"]["trans_skip"].append(file_identifier)
                else:
                    logger.info(f"{log_prefix_book} >>> Iniciando TRADUÇÃO HÍBRIDA...")
                    translation_start_time = time.time()
                    translation_success = False
                    if not os.path.exists(PATH_TO_TRANSLATOR_SCRIPT):
                         logger.error(f"{log_prefix_book} ERRO CRÍTICO: Script tradutor '{PATH_TO_TRANSLATOR_SCRIPT}' não encontrado.");
                    else:
                        try:
                            command = [ sys.executable or 'python', # Usa o executável python atual ou 'python'
                                        PATH_TO_TRANSLATOR_SCRIPT,
                                        '--input', translator_input_path,
                                        '--output', translated_docx_path,
                                        '--words', str(NUM_WORDS_TO_TRANSLATE) ]
                            logger.info(f"{log_prefix_book} Executando: {' '.join(command)}")
                            # Usar timeout para evitar bloqueios indefinidos
                            result = subprocess.run(command, capture_output=True, text=True, encoding='utf-8', check=False, timeout=600) # Timeout 10 min
                            translation_end_time = time.time(); translation_total_time = translation_end_time - translation_start_time
                            if result.returncode == 0:
                                logger.info(f"✅ {log_prefix_book} TRADUÇÃO HÍBRIDA SUCESSO em {translation_total_time:.2f} seg.")
                                log_translated_file(TRANSLATED_LOG_FILE, file_identifier); translated_files_set.add(file_identifier)
                                stats["trans_ok"] += 1; stats["lists"]["trans_ok"].append(file_identifier); stats["trans_times"].append(translation_total_time)
                                if result.stdout: logger.debug(f"{log_prefix_book} Saída tradutor:\n{result.stdout}")
                                translation_success = True
                            else:
                                logger.error(f"❌ {log_prefix_book} TRADUÇÃO HÍBRIDA FALHOU (código: {result.returncode}) em {translation_total_time:.2f} seg.")
                                if result.stderr: logger.error(f"{log_prefix_book} Erro tradutor:\n{result.stderr}")
                                else: logger.error(f"{log_prefix_book} Tradutor não reportou erro específico no stderr.")
                        except subprocess.TimeoutExpired:
                             logger.error(f"❌ {log_prefix_book} TRADUÇÃO HÍBRIDA TIMEOUT (após 600s).")
                             translation_total_time = time.time() - translation_start_time # Tempo até timeout
                        except Exception as e_translate_sub:
                             logger.error(f"{log_prefix_book} Erro CRÍTICO subprocesso tradução: {e_translate_sub}"); logger.error(traceback.format_exc());
                             translation_total_time = time.time() - translation_start_time # Tempo até erro

                    if not translation_success:
                        stats["trans_fail"] += 1; stats["lists"]["trans_fail"].append(f"{file_identifier} (Erro script tradutor)")
            # --- Fim Tradução ---
            logger.info(f"{log_prefix_book} --- Fim processamento livro ---")
            # --- Fim Loop Livros ---
        logger.info(f"--- Concluída verificação Autor: {author_name} ---")
    # --- Fim Loop Autores ---

    end_time_main = time.time(); total_time_main = end_time_main - start_time_main
    total_corr_time_ok = sum(stats["corr_times"]); avg_corr_time_ok = total_corr_time_ok / len(stats["corr_times"]) if stats["corr_times"] else 0
    total_trans_time_ok = sum(stats["trans_times"]); avg_trans_time_ok = total_trans_time_ok / len(stats["trans_times"]) if stats["trans_times"] else 0

    # --- Resumo Final Logging ---
    logger.info("===================== RESUMO FINAL (v11) =====================")
    logger.info(f"Tempo total geral: {total_time_main:.2f} seg ({total_time_main/60:.2f} min).")
    logger.info("--- Resumo Etapa de Correção/Notas DOCX ---")
    logger.info(f"Livros Processados OK: {stats['corr_ok']}")
    logger.info(f"Livros Pulados (já feitos): {stats['corr_skip']}")
    logger.info(f"Livros com Falha: {stats['corr_fail']}")
    if stats["corr_times"]: logger.info(f"Tempo Total Correção/Notas (livros OK): {total_corr_time_ok:.2f} seg ({total_corr_time_ok/60:.2f} min)"); logger.info(f"Tempo Médio por Livro OK: {avg_corr_time_ok:.2f} seg")
    logger.info(f"API Correção (P1) - Latência Total: {stats['lat_p1']:.2f}s / Tokens Totais: {stats['t_p1']} (P: {stats['p_p1']}, O: {stats['o_p1']})")
    logger.info(f"API Notas ID (P2) - Latência Total: {stats['lat_p2']:.2f}s / Tokens Totais: {stats['t_p2']} (P: {stats['p_p2']}, O: {stats['o_p2']})")
    logger.info("--- Resumo Etapa de Tradução ---")
    logger.info(f"Livros Traduzidos OK: {stats['trans_ok']}")
    logger.info(f"Livros Pulados (já traduzidos): {stats['trans_skip']}")
    logger.info(f"Livros com Falha: {stats['trans_fail']}")
    if stats["trans_times"]: logger.info(f"Tempo Total Tradução (livros OK): {total_trans_time_ok:.2f} seg ({total_trans_time_ok/60:.2f} min)"); logger.info(f"Tempo Médio por Tradução OK: {avg_trans_time_ok:.2f} seg")
    logger.info("--- Logs ---"); logger.info(f"Log detalhado: {os.path.abspath(log_filepath)}"); logger.info(f"Log de correções OK: {os.path.abspath(PROCESSED_LOG_FILE)}"); logger.info(f"Log de traduções OK: {os.path.abspath(TRANSLATED_LOG_FILE)}")
    logger.info("--- Arquivos Gerados (Estrutura Exemplo v11) ---")
    logger.info(f"  - DOCX Limpo (p/ Tradutor): {BASE_OUTPUT_DOCX_DIR}/<Autor>/<Subpasta>/<Livro>_{FINAL_DOCX_CLEAN_BASENAME}")
    logger.info(f"  - DOCX Com Notas no Fim:    {BASE_OUTPUT_DOCX_DIR}/<Autor>/<Subpasta>/<Livro>_{FINAL_DOCX_WITH_NOTES_BASENAME}")
    logger.info(f"  - DOCX Traduzido: {BASE_OUTPUT_DOCX_DIR}/<Autor>/<Subpasta>/<Livro>_{FINAL_DOCX_CLEAN_BASENAME.replace('.docx', TRANSLATED_DOCX_SUFFIX)}")
    logger.info("========================================================")

    # === Envio de E-mail FINAL ===
    if email_configured:
        send_completion_email(
            sender_email=EMAIL_SENDER_ADDRESS, sender_password=EMAIL_SENDER_APP_PASSWORD, recipient_email=EMAIL_RECIPIENT_ADDRESS,
            smtp_server=EMAIL_SMTP_SERVER, smtp_port=EMAIL_SMTP_PORT,
            processed_correction=stats["corr_ok"], skipped_correction=stats["corr_skip"], failed_correction=stats["corr_fail"],
            processed_translation=stats["trans_ok"], skipped_translation=stats["trans_skip"], failed_translation=stats["trans_fail"],
            total_duration_seconds=total_time_main, main_log_path=log_filepath, processed_log_path=PROCESSED_LOG_FILE, translated_log_path=TRANSLATED_LOG_FILE,
            total_correction_latency_secs=stats["lat_p1"], total_correction_tokens=stats["t_p1"],
            total_footnote_latency_secs=stats["lat_p2"], total_footnote_tokens=stats["t_p2"],
            avg_correction_time_secs=avg_corr_time_ok, total_correction_time_secs=total_corr_time_ok,
            avg_translation_time_secs=avg_trans_time_ok, total_translation_time_secs=total_trans_time_ok,
            processed_correction_books=stats["lists"]["corr_ok"], skipped_correction_books=stats["lists"]["corr_skip"], failed_correction_books=stats["lists"]["corr_fail"],
            processed_translation_books=stats["lists"]["trans_ok"], skipped_translation_books=stats["lists"]["trans_skip"], failed_translation_books=stats["lists"]["trans_fail"]
        )
    else: logger.info("Envio de e-mail de resumo final pulado (configuração ausente).")

# --- Ponto de Entrada ---
if __name__ == "__main__":
    try: main()
    except KeyboardInterrupt: logger.warning("\nProcesso interrompido manualmente (Ctrl+C).")
    except Exception as e_main: logger.critical(f"FATAL Erro na execução principal (main): {e_main}"); logger.critical(traceback.format_exc())